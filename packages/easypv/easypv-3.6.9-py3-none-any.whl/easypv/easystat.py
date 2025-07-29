#!/usr/bin/env python
# coding: utf-8
# 开发人：蔡权周

############################################################################################################################
#导入基本模块
############################################################################################################################

#系统基本模块
import os
import sys
import io
import re
import math
import random
import string
import itertools
import argparse
import sqlite3
import zipfile
import shutil
import json
import threading
import datetime
import hashlib
import queue
import time
import ast
import requests
import tempfile
import platform
import subprocess
from datetime import datetime, timedelta
from collections import defaultdict, Counter
import re

#禁止独立运行
#from easypv import run_apps

# 数据处理与分析
import numpy as np
import pandas as pd
from pandas.core.groupby.groupby import GroupBy
from pandas.core.base import NoNewAttributesMixin
      
try:
    from scipy.stats import ranksums
    import scipy.stats as st
    import scipy.stats as stats
    from scipy.stats import ranksums, iqr, chi2_contingency, fisher_exact
    from scipy.special import factorial
except:
    print('scipy模块未正确导入，部分功能可能不正常。')
    
try:
    from sklearn.preprocessing import StandardScaler, MinMaxScaler
    from sklearn.neighbors import NearestNeighbors
    from sklearn.metrics import roc_auc_score
except:
    print('sklearn模块未正确导入，部分功能可能不正常。')
    
try:   
    import statsmodels.api as sm
    from statsmodels.formula.api import ols, logit, poisson, mnlogit
    from statsmodels.tools import add_constant
except:
    print('statsmodels模块未正确导入，部分功能可能不正常。')


# 数据可视化
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from matplotlib import cm
from matplotlib.text import Text
from matplotlib.transforms import Affine2D
import matplotlib.font_manager as fm
import seaborn as sns

# 文档处理
from docx import Document

# 数据库
from sqlalchemy import create_engine, text


# Excel 处理
from openpyxl import load_workbook

# Tkinter GUI
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog, scrolledtext
from tkinter.scrolledtext import ScrolledText
from tkinter.ttk import Treeview, Style, Checkbutton
from tkinter.messagebox import showinfo
from tkinter import Tk, filedialog, Listbox, Button, END, Scrollbar, Frame, Label, StringVar, IntVar, Checkbutton, messagebox,Menu
from tkinter import Label, Entry, Button, StringVar, OptionMenu, LabelFrame




# XML 处理
import xml.etree.ElementTree as ET
import uuid
from random import choices
from string import digits


# 并发处理
try:
    from concurrent.futures import ThreadPoolExecutor, as_completed
    import queue
    import concurrent.futures
    from Crypto.Cipher import AES
    from Crypto.Protocol.KDF import PBKDF2
    from Crypto.Util.Padding import pad, unpad
except:
    print('并发模块为正确导入，部分功能可能不正常。')



############################################################################################################################
#预先配置的函数
############################################################################################################################
class AAA_01_ORI():
    pass
    
def EasyInf():
    """
    软件版本信息。
    """
    inf = {
        '软件名称': '易析数据分析工具',
        '版本号': '2.3.3',
        '功能介绍': '用于表单数据统计和分析。',
        'PID': 'MDRDSLF006',
        '分组': '药物警戒',
        '依赖': 'pandas,numpy,scipy,statsmodels,sklearn,matplotlib,seaborn,python-docx,sqlalchemy,openpyxl',
        '资源库版本':'20250210'    
    }
    return inf
   
  
  
def ORI_get_font():
    # 获取系统信息
    system = platform.system().lower()

    if system == 'windows':
        # Windows 可用 SimSun（宋体），但推荐手动安装思源黑体
        my_font_ch = 'SimSun'  # 或手动安装 Noto Sans CJK SC
    elif system == 'linux':
        # 国产 Linux（UOS/麒麟）通常自带思源黑体
        my_font_ch = 'Noto Sans CJK SC'  # 或 'Source Han Sans SC'
    else:
        # Mac 或其他系统
        my_font_ch = 'PingFang SC'  # 苹方（Mac 自带，仅限 Apple 生态使用）
    print(f"选择的中文字体: {my_font_ch}")
    return my_font_ch
            
def ORI_get_data_path(filename):
    """
    获取数据文件的绝对路径，用于开始获取sor.zip等文件的路径。
    """
    # 获取当前包的安装目录
    package_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建数据文件的路径
    data_path = os.path.join(package_dir, 'data', filename)   
    return data_path

    
def ORI_csdir_and_peizhidir():
    """
    设置 csdir 和 peizhidir：
    1. 检查 setting.xml 文件是否存在：
       - 如果不存在，弹出警示框并让用户选择工作目录，创建 setting.xml 文件。
       - 如果存在，读取 work_dir 作为 csdir。
    2. 检查 csdir 目录下是否存在 '资源库' 目录：
       - 如果存在，返回 csdir 和 peizhidir（资源库路径）。
       - 如果不存在，尝试创建 '资源库' 目录，并解压 sor.zip（zip 文件）到资源库中。
       - 如果 sor.zip 不存在，提示未找到资源库文件，但仍返回 peizhidir 路径。
    返回：
       - csdir: 工作目录路径。
       - peizhidir: 资源库路径。
    """
    # 获取当前脚本所在的目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    setting_file = os.path.join(script_dir, "setting.xml")

    # 初始化 tkinter 主窗口（仅一次）
    root_tk = tk.Tk()
    root_tk.withdraw()  # 隐藏主窗口

    # 1. 设置 csdir
    if not os.path.exists(setting_file):
        # 弹出警示框
        messagebox.showinfo("提示", "未找到配置文件，请选择工作目录。")

        # 弹出对话框让用户选择工作目录
        selected_dir = filedialog.askdirectory(title="请选择工作目录")
        if not selected_dir:
            raise ValueError("未选择工作目录，程序退出。")

        # 创建 setting.xml 文件
        root = ET.Element("settings")
        python_command = ET.SubElement(root, "python_command")
        python_command.text = "python"
        work_dir = ET.SubElement(root, "work_dir")
        work_dir.text = selected_dir
        check_interval = ET.SubElement(root, "check_interval")
        check_interval.text = "60000"

        # 保存 setting.xml 文件
        tree = ET.ElementTree(root)
        tree.write(setting_file, encoding="UTF-8", xml_declaration=True)

    # 读取 setting.xml 文件
    tree = ET.parse(setting_file)
    root = tree.getroot()
    csdir = root.find("work_dir").text

    # 2. 设置 peizhidir
    peizhidir = os.path.join(csdir, "资源库")

    # 检查资源库是否存在
    if not os.path.exists(peizhidir):
        # 弹出提示框
        messagebox.showinfo("提示", "未找到资源库，将自动生成资源库。")

        # 创建资源库目录
        os.makedirs(peizhidir, exist_ok=True)

        # 获取程序目录下的 sor.zip（实际是 zip 文件）
        #def_zip_path = os.path.join(script_dir, "sor.zip")
        def_zip_path=sor_path
        
        # 检查 sor.zip 是否存在
        if os.path.exists(def_zip_path):
            try:
                # 解压 sor.zip 到资源库目录，处理中文文件名
                with zipfile.ZipFile(def_zip_path, 'r') as zip_ref:
                    for file_info in zip_ref.infolist():
                        # 处理中文文件名
                        file_name = file_info.filename.encode('cp437').decode('gbk')  # 从 cp437 转换为 gbk
                        target_path = os.path.join(peizhidir, file_name)
                        # 确保目标目录存在
                        os.makedirs(os.path.dirname(target_path), exist_ok=True)
                        # 解压文件
                        with open(target_path, 'wb') as target_file:
                            target_file.write(zip_ref.read(file_info.filename))
                print(f"资源库配置完成: {peizhidir}")
            except Exception as e:
                print(f"解压资源库文件失败: {e}")
        else:
            print("未找到资源库文件。")

    # 销毁 tkinter 主窗口
    root_tk.destroy()

    return csdir, peizhidir
    

############################################################################################################################
#全局变量
############################################################################################################################
global version_now
global usergroup
global setting_cfg
global csdir
global peizhidir
global biaozhun
global global_dfs
global psur
global mytitle
global my_font_ch
global_dfs = {}
biaozhun={}
version_now=EasyInf()
# 全局变量，用于存储所有保存的DataFrame
global_dfs = {}
version_now=version_now['版本号']
db_ver_i=EasyInf()['资源库版本']
csdir,peizhidir = ORI_csdir_and_peizhidir()
psur=0

try:
    sor_path = ORI_get_data_path('sor.zip')
    #print(f"sor.zip path: {sor_path}")
except FileNotFoundError as e:
    print(e)  
    
print(f"工作目录设置为: {csdir}")
print(f"资源库目录设置为: {peizhidir}")
my_font_ch=ORI_get_font()
psur=0  
mytitle='易析数据分析工具 EFS_'+version_now
    

default_data = pd.DataFrame({  
            '条目': ['名称', '版本', '用途说明'],  
            '信息': ['易析数据统计分析工具', version_now, '供数据统计使用。']}) 

############################################################################################################################
#品种自定义清洗或配置函数
############################################################################################################################
class AAA_02_PSUR:
    pass
def PSUR_read_files_and_clean():
    """智能导入：打开文件并执行清洗。"""
    data=SMALL_read_and_merge_files()
    try:
        if '报告表编码' in data.columns and  '不良反应过程描述' in data.columns and  '-持有人-品名-证号列' not in data.columns:
           print('监测到这是一份药品不良反应的报告列表，正在尝试进行数据清洗...') 
           PSUR_yaopin(data) 
        elif '报告编码' in data.columns and  '器械故障表现' in data.columns and  '-持有人-品名-证号列' not in data.columns:
           print('监测到这是一份医疗器械不良事件的报告列表，正在尝试进行数据清洗...')  
           PSUR_qixie(data)        
    except:
        print('未执行任何清洗操作。')
        PROGRAM_display_df_in_treeview(data,0,0)
def PSUR_DRUG_merge_dfs(df1, df2,methon='完全'):
    """
    药品预制-药品信息标准化（含厂家-调用历史知识库）
    """
    # 1. 将 df2 的列改名
    df2 = df2.rename(columns={
        '注册证备案证': '-注册证备案证',
        '通用名称': '通用名称（历史匹配规整）',
        '剂型': '剂型（历史匹配规整）',
        '-注册人备案人': '-注册人备案人（历史匹配规整）'
    })

    # 2. 筛选 df2 的列
    if methon=='完全':
        df2 = df2[['-注册证备案证', '通用名称（历史匹配规整）', '药品分类', '最主要的一个活性成分', '剂型（历史匹配规整）', '-注册人备案人（历史匹配规整）']]
        # 3. 删除 df1 中 -注册证备案证 列的 "国药准字"
        df1['-注册证备案证'] = df1['-注册证备案证'].str.replace('国药准字', '', regex=False)
        # 4. 左连接合并 df1 和 df2
        result = pd.merge(df1, df2, on='-注册证备案证', how='left')
    elif methon=='精简':
               df2 = df2[['通用名称（历史匹配规整）', '药品分类', '最主要的一个活性成分', '剂型（历史匹配规整）']].drop_duplicates("通用名称（历史匹配规整）")
               df2.rename(columns={'通用名称（历史匹配规整）': "通用名称"}, inplace=True)
               result = pd.merge(df1, df2, on='通用名称', how='left')  
    return result


 
def PSUR_qixie(data):
    """
    器械的智能导入模式。
    """
    global psur
    print('程序正在启动...')
    print('正在载入，请稍候...')
    if psur!=0:
        for i in range(1000000000000):   
            a=input("该模式目前不可用。") 
        
    if isinstance(data, pd.DataFrame):
        psur_data=data      
    else:
        psur_data=SMALL_read_and_merge_files()
    
    methon='普通模式' 
    guize3=SQL_gettable(memory_db,"easy_器械规整-SOC-关键词")  
    guize3=guize3[guize3['适用范围'].isin(['无源', '有源', '体外诊断试剂'])]
    guize3['结果列']='通用故障表现'
    if '是否已清洗' not in psur_data.columns:
        print('●正在执行基础清洗...')
        psur_data=CLEAN_replay_operations(psur_data,biaozhun["器械清洗"]) 
        print('正在添加敏感词标记...')
        try:
            psur_data["医院填写的信息"]=psur_data["器械故障表现"].astype(str)+psur_data["伤害表现"].astype(str)+psur_data["使用过程"].astype(str)+psur_data["事件原因分析描述"].astype(str)+psur_data["初步处置情况"].astype(str)#器械故障表现|伤害表现|使用过程|事件原因分析描述|初步处置情况
        except:
            psur_data["医院填写的信息"]=psur_data["器械故障表现"].astype(str)+psur_data["伤害表现"].astype(str)+psur_data["使用过程"].astype(str)
        kx = pd.read_excel(os.path.join(peizhidir,"easy_WarningParameters.xlsx"), header=0, sheet_name=0).reset_index(drop=True)


        k4=kx["权重"][3] #高度关注关键字（一级）        
        k4_values=kx["值"][3] #高度关注关键字（一级） 值    
        
        k5=kx["权重"][4] #高度关注关键字（二级）        
        k5_values=kx["值"][4] #高度关注关键字（二级） 值
        k5_values=k4_values+"|"+k5_values
        psur_data['敏感词']=''
        
        # 定义一个函数来检查匹配并更新列
        def update_keywords(row, patterns):
            matched_keywords = [pattern.pattern for pattern in patterns if pattern.search(row["医院填写的信息"])]
            if matched_keywords:
                return ";".join(matched_keywords)
            return ""
            
        k5_set = k5_values.split('|')        

        compiled_patterns_k5 = [re.compile(pattern) for pattern in k5_set]        
        psur_data["敏感词"] = psur_data.apply(lambda row: update_keywords(row, compiled_patterns_k5), axis=1)
        
        print('正在添加通用关键词标记...')
        psur_data=CLEAN_easystat(psur_data,guize3,'加关键词')

        
        try:
            print('正在添加产品分类信息...')
            psur_data=pd.merge(psur_data,biaozhun['器械品种库'], on="产品名称", how='left') 
            psur_data['产品大类']=psur_data['产品大类'].fillna('-未分类-')
            psur_data['一级类别']=psur_data['一级类别'].fillna('-未分类-')
            psur_data['二级类别']=psur_data['二级类别'].fillna('-未分类-')
        except:
            print('产品分类信息添加失败...')
            

        psur_data['是否已清洗']='是'
        print('●数据清洗完成。')          
    else:
        print('●检测到导入的数据前期已清洗过，不再执行数据清洗。')        

    PROGRAM_display_df_in_treeview(psur_data,'psur',psur_data) 



def PSUR_yaopin(data):
    """
    药品的智能导入模式。
    """
    global psur
    print('程序正在启动...')
    print('正在载入，请稍候...')
    if psur!=0:
        for i in range(1000000000000):   
            a=input("该模式目前不可用。") 
        
    if isinstance(data, pd.DataFrame):
        psur_data=data      
    else:
        psur_data=SMALL_read_and_merge_files()
    
    if ('不良反应-术语' in psur_data.columns) and  ('不良反应-code' in psur_data.columns) :
        methon='品种评价'
    else:
        methon='普通模式' 
    
        
    if '是否已清洗' not in psur_data.columns:
        print('●正在执行基础清洗...')
        if methon=='品种评价':
            psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品清洗品种评价"])
            try: 
                psur_data["关联性评价汇总"]="("+psur_data["报告单位评价"].astype(str)+"("+psur_data["县评价"].astype(str)+"("+psur_data["市评价"].astype(str)+"("+psur_data["省评价"].astype(str)+"("+psur_data["国家评价"].astype(str)+")"
                psur_data["关联性评价汇总"]=psur_data["关联性评价汇总"].str.replace("(nan","",regex=False)
                psur_data["关联性评价汇总"]=psur_data["关联性评价汇总"].str.replace("nan)","",regex=False)
                psur_data["关联性评价汇总"]=psur_data["关联性评价汇总"].str.replace("nan","",regex=False)
                psur_data['关联性评价汇总'] = psur_data["关联性评价汇总"].str.extract('.*\((.*)\).*', expand=False)
            except:
                print('关联性评价整合失败。') 
            psur_data=psur_data.fillna(0)
           
        else:
            psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品清洗"]) 

        print('●正在执行添加药品额外信息...') 
        try:       
            psur_data=PSUR_DRUG_merge_dfs(psur_data, biaozhun["药品信息库"],'精简')
        except:
            print('药品信息添加失败，原因可能是药品信息库没有正确配置。')
        
        print(len(psur_data))

        
        print('●正在执行扩行...')
        psur_data=CLEAN_expand_rows(psur_data, ';', ["不良反应名称（规整）"])
        psur_data=CLEAN_expand_rows(psur_data, '；', ["不良反应名称（规整）"])
        print('●正在执行PT规整...')
        #psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品PT清洗"])
        psur_data=SMALL_merge_dataframes(psur_data, "不良反应名称（规整）",biaozhun["药品不良反应名称规整-AI"][['不良反应名称AI','AI初步规整PT']], df2_col='不良反应名称AI') 
        
        print(len(psur_data))


        del psur_data['不良反应名称AI']
        psur_data['AI初步规整PT'] = psur_data['AI初步规整PT'].fillna(psur_data['不良反应名称（规整）'])
        del psur_data['不良反应名称（规整）']

        
        #psur_data.rename(columns={'通用名称（历史匹配规整）': "通用名称"}, inplace=True)
        print('●再次执行扩行...')
        

        psur_data=CLEAN_expand_rows(psur_data, '；', ['AI初步规整PT'])
        psur_data=CLEAN_expand_rows(psur_data, ';', ['AI初步规整PT'])

        print('●正在执行SOC初步映射...')
        pss=biaozhun["药品不良反应名称规整-AI"][(biaozhun["药品不良反应名称规整-AI"]['主SOC']=='Y')].copy() 
        pss=pss[["AI初步规整PT", 'code','PT','HLT','HLGT','SOC']].drop_duplicates("AI初步规整PT")

        psur_data=pd.merge(psur_data,pss, on="AI初步规整PT", how='left') 
        psur_data['code']=psur_data['code'].astype(str)

        psur_data.rename(columns={'AI初步规整PT': "不良反应名称（规整）"}, inplace=True)
        print('●正在执行SOC进一步映射...')
        try:
            if isinstance(biaozhun["meddra"], int):
                biaozhun["meddra"]=pd.read_excel(os.path.join(peizhidir,'share_easy_adrmdr_药品规整-SOC-Meddra库.xlsx')).reset_index(drop=True)  
            psur_data['PTx']=psur_data['PT'].copy()
            psur_data['HLTx']=psur_data['HLT'].copy()
            psur_data['HLGTx']=psur_data['HLGT'].copy()
            psur_data['SOCx']=psur_data['SOC'].copy()

            psur_data = psur_data.drop(columns=['code', 'PT', 'HLT', 'HLGT', 'SOC'])
            
            psur_data=SMALL_merge_dataframes(psur_data, "不良反应名称（规整）",biaozhun["meddra"], df2_col='PT') 
            psur_data['PT'] = psur_data['PT'].fillna(psur_data['PTx'])
            psur_data['HLT'] = psur_data['HLT'].fillna(psur_data['HLTx'])
            psur_data['HLGT'] = psur_data['HLGT'].fillna(psur_data['HLGTx'])
            psur_data['SOC'] = psur_data['SOC'].fillna(psur_data['SOCx'])
            del psur_data['PTx']
            del psur_data['HLTx']
            del psur_data['HLGTx']
            del psur_data['SOCx']
            
            print('●正在执行SOC进一步映射成功。')
        except:
            print('●正在执行SOC进一步映射失败...')
        print('●正在执行用量单位规整...')
        try:   
            psur_data['用量（规整后）']= psur_data['用量'].copy()
            psur_data['用量单位（规整后）']= psur_data['用量单位'].copy()
            psur_data.loc[psur_data['用量单位（规整后）'].isin(['g', '克']), '用量（规整后）'] *= 1000  
            psur_data.loc[psur_data['用量单位（规整后）'].isin(['g', '克','毫克']), '用量单位（规整后）'] = 'mg'
            psur_data['用法用量']=psur_data['用量（规整后）'].astype(str)+psur_data['用量单位（规整后）'].astype(str)+' '+psur_data['用法-日'].astype(str)+'日'+psur_data['用法-次'].astype(str)+'次'
        except:
            print('●未成功对用法用量进行规整。')   
        print('●正在执行重点关注标注...')
        psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品清洗-扩行后"].reset_index(drop=True) ) 

        psur_data['疑似新的ADR']='未载入说明书'  
        psur_data['SOC']=psur_data['SOC'].fillna('其他（未规整）')     
        psur_data.loc[(psur_data['重点关注']=='重点关注'),'重点关注ADR'] = psur_data.loc[psur_data['重点关注']=='重点关注','不良反应名称（规整）']
        psur_data['是否已清洗']='是'
        print('●数据清洗完成。')          
    else:
        print('●检测到导入的数据前期已清洗过，不再执行数据清洗。')        

    PROGRAM_display_df_in_treeview(psur_data,'psur',psur_data) 


def PSUR_get_guize2(ori_owercount_easyread):
    """切换通用规则使用"""
    global memory_db
    print('为了方便测试，该模式下，关键词库实时载入。')
    guize3=SQL_gettable(memory_db,"easy_器械规整-SOC-关键词")  
    auto_guize3=guize3[["适用范围列","适用范围"]].drop_duplicates("适用范围")
    print(auto_guize3)

    se = tk.Toplevel()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(se)
    se.title('表现归类')
    sw_se = se.winfo_screenwidth()
    #得到屏幕宽度
    sh_se = se.winfo_screenheight()
    #得到屏幕高度
    ww_se = 450
    wh_se = 500
    #窗口宽高为100  #
    x_se = (sw_se-ww_se) / 2
    y_se = (sh_se-wh_se) / 2
    se.geometry("%dx%d+%d+%d" %(ww_se,wh_se,x_se,y_se)) 
    import_se55=tk.Label(se,text="请选择关键词查找位置：")
    import_se55.pack()    
    
    # 创建一个Frame来放置Treeview和Scrollbar
    frame = tk.Frame(se)
    frame.pack(fill='both', expand=True)
 
    # 设置Treeview的样式
    style = Style()
    style.configure("Treeview", rowheight=25)
 
    # Treeview组件，用于显示df的列名，并允许多选
    treeview = Treeview(frame, columns=("Column",), show='headings', selectmode='extended')
    treeview.heading("#0", text="选择列", anchor='w')
    treeview.heading("#1", text="Column", anchor='w')
 
    # 填充Treeview
    for col in ori_owercount_easyread.columns:
        treeview.insert("", "end", values=(col,), iid=col, tags=("selectable",))
 
    # Scrollbar组件，用于Treeview的滚动
    scrollbar = Scrollbar(frame, orient="vertical", command=treeview.yview)
    treeview.configure(yscrollcommand=scrollbar.set)
 
    scrollbar.pack(side="right", fill="y")
    treeview.pack(side="left", fill="both", expand=True)
    
    
    
    import_se2=tk.Label(se,text="请选择您所需要的通用规则关键字：")
    import_se2.pack()
    def xt11set(*arg):
        comvalue.set(comboxlist.get())
    comvalue = StringVar()
    comboxlist = ttk.Combobox(se, width=14, height=30, state="readonly", textvariable=comvalue )  # 初始化
    comboxlist["values"] = auto_guize3["适用范围"].tolist()
    comboxlist.current(0)  # 选择第一个
    comboxlist.bind("<<ComboboxSelected>>", xt11set)  # 绑定事件,(下拉列表框被选中时，绑定XT11SET函数)
    comboxlist.pack()


    labFrame_Button_se=tk.LabelFrame(se)
    btn_se=tk.Button(labFrame_Button_se,text="确定",width=10,command=lambda:get_guize2(guize3,comvalue.get(),treeview.selection()))
    btn_se.pack(side=tk.LEFT,padx=1,pady=1)
    labFrame_Button_se.pack()    
    
    def get_guize2(guize3,covn,selected_items):
        
        msdd=guize3.loc[guize3["适用范围"].str.contains(covn, na=False)].copy().reset_index(drop=True)
        if not selected_items:
            print("未选择任何列！将使用配置表文件预设的查找位置:",msdd["查找位置"][0])
            ori_owercount_easyread2=ori_owercount_easyread.loc[ori_owercount_easyread[msdd["适用范围列"][0]].str.contains(msdd["适用范围"][0], na=False)].copy().reset_index(drop=True)
            ori_owercount_easyread2[msdd["结果列"][0]]=''
            ori_owercount_easyread2=CLEAN_easystat(ori_owercount_easyread2.copy(),msdd,'加关键词')

        else:
            selected_columns = [treeview.item(item, "values")[0] for item in selected_items]
            ori_owercount_easyread2=ori_owercount_easyread.loc[ori_owercount_easyread[msdd["适用范围列"][0]].str.contains(msdd["适用范围"][0], na=False)].copy().reset_index(drop=True)
            ori_owercount_easyread2[msdd["结果列"][0]]=''
            ori_owercount_easyread2=CLEAN_easystat(ori_owercount_easyread2.copy(),msdd,'加关键词',selected_columns)
        print("原始报告数量：",len(ori_owercount_easyread2))
        ori2=CLEAN_expand_rows(ori_owercount_easyread2,";",[msdd["结果列"][0]])
        result=TOOLS_create_pivot_tool(ori2,[[msdd["结果列"][0]], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']])
        result['报表类型']=result['报表类型'].str.replace('grouped','group_sep')
        PROGRAM_display_df_in_treeview(result,0,ori_owercount_easyread2) 





def PSUR_check_adr_in_word(df): 
    '''药品载入说明书进行比对 ''' 
    # 弹出文件选择对话框，让用户选择一个Word文档  
    root = tk.Tk()  
    root.withdraw()  # 隐藏tkinter的主窗口  
    file_path = filedialog.askopenfilename(  
        title="选择Word文档",  
        filetypes=[("Word files", "*.docx")]  
    )  
    root.destroy()  # 销毁tkinter窗口  
      
    if not file_path:  
        return df  # 如果用户取消选择文件，则直接返回原df  
      
    # 读取Word文档内容  
    doc = Document(file_path)  
    word_content = ''  
      
    # 提取段落文本  
    for paragraph in doc.paragraphs:  
        word_content += paragraph.text + ' '  
      
    # 提取表格中的文本  
    for table in doc.tables:  
        for row in table.rows:  
            for cell in row.cells:  
                word_content += cell.text + ' '  
                
    df['疑似新的ADR']=""  
    # 遍历df的每一行  
    for index, row in df.iterrows():  
        adr = row['不良反应名称（规整）']  
        if adr not in word_content:  
            df.at[index, '疑似新的ADR'] = adr  
      
    PROGRAM_display_df_in_treeview(df,'psur',0)  


def PSUR_get_new_GUI(df):
    """主函数：生成新的不良反应检测 GUI 并处理数据"""
    df['-注册证备案证'] = df['-注册证备案证'].str.replace('国药准字', '', regex=False)
    
    def filter_adverse_effects(df, column_name, n):
        """过滤不良反应"""
        result_rows = []
        for index, row in df.iterrows():
            adverse_effects_str = row['不良反应']
            matches = re.findall(r'([^（]+)（(\d+)）', adverse_effects_str)
            filtered_effects = [(symptom, int(count)) for symptom, count in matches if int(count) >= n]
            if filtered_effects:
                new_adverse_effects_str = '、'.join(f"{symptom}（{count}）" for symptom, count in filtered_effects)
                new_adverse_effects_str = new_adverse_effects_str.replace("、、", '、')
                new_row = row.copy()
                new_row['不良反应'] = new_adverse_effects_str
                result_rows.append(new_row)
        return pd.DataFrame(result_rows)

    def drug(data, n2=1):
        """药品不良反应名称统计"""
        data = data.drop_duplicates("报告表编码")
        rm = str(Counter(SMALL_get_list("use(不良反应名称（规整）).file", data, 1000))).replace("Counter({", "{")
        rm = rm.replace("})", "}")
        import ast
        user_dict = ast.literal_eval(rm)
        df = SMALL_easyreadT(pd.DataFrame([user_dict]))
        df = df[(df['详细描述T'] != '######################################################################')]
        df = df[(df['详细描述T'].astype(int) >= n2)]
        df = df.rename(columns={"逐条查看": "不良反应（规整）"})
        return df

    def analyze_products(df1, df2, biaozhun_drug_info, method):
        """分析产品并返回结果 DataFrame"""
        def is_new_adverse_reaction(row):
            """检查是否为新的不良反应"""
            matched_rows = df2[df2['批准文号'] == row['-注册证备案证']] if method == '证号' else df2[df2['通用名称'] == row['-产品名称']]
            if not matched_rows.empty and str(row['不良反应名称（规整）']) not in matched_rows['不良反应'].str.cat(sep='|'):
                return row['不良反应名称（规整）']
            return None

        def is_non_new_adverse_reaction(row):
            """检查是否为疑似非新的不良反应"""
            matched_rows = df2[df2['批准文号'] == row['-注册证备案证']] if method == '证号' else df2[df2['通用名称'] == row['-产品名称']]
            if not matched_rows.empty and str(row['不良反应名称（规整）']) in matched_rows['不良反应'].str.cat(sep='|'):
                return row['不良反应名称（规整）']
            return None

        def is_key_adverse_reaction(row):
            """检查是否为重点关注的不良反应"""
            adverse_effect = str(row['不良反应名称（规整）']).strip()  # 确保是字符串并去除空白
            if not adverse_effect:  # 如果为空，直接返回 None
                return None
            for key in biaozhun['药品重点关注库']['重点关注库']:
                if key in adverse_effect:  # 检查关键词是否在不良反应名称中
                    return adverse_effect
            return None

        def extract_keyword(adverse_effect):
            """根据不良反应名称提取关键词"""
            KEYWORDS = ['疹','痒','肝','肾','白细胞','血小板','粒细胞','红细胞']
            for word in KEYWORDS:
                if word in adverse_effect:
                    return word
            return adverse_effect  # 如果没有匹配的关键词，返回原始的不良反应名称

        def not_in_ai_adverse_reaction(row):
            """检查是否不在不良反应AI中"""
            adverse_effect = str(row['不良反应名称（规整）']).strip()
            if not adverse_effect:  # 如果为空，直接返回 None
                return None
            
            # 提取关键词
            keyword = extract_keyword(adverse_effect)
            
            # 检查 biaozhun_drug_info 中是否有相关药品
            if method == '证号':
                matched_rows = biaozhun_drug_info[biaozhun_drug_info['注册证备案证'] == row['-注册证备案证']]
            else:
                matched_rows = biaozhun_drug_info[biaozhun_drug_info['通用名称'] == row['-产品名称']]
            
            if matched_rows.empty:  # 如果没有相关药品
                return '标准库内无相关药品无法比对'
            
            # 如果有相关药品，检查是否在不良反应AI中
            if keyword not in matched_rows['不良反应AI'].str.cat(sep='|'):
                return adverse_effect
            return None

        def not_in_expected_adverse_reaction(row):
            """检查是否不在预期不良反应（历史数据）中"""
            adverse_effect = str(row['不良反应名称（规整）']).strip()
            if not adverse_effect:  # 如果为空，直接返回 None
                return None
            # 检查 biaozhun_drug_info 中是否有相关药品
            if method == '证号':
                matched_rows = biaozhun_drug_info[biaozhun_drug_info['注册证备案证'] == row['-注册证备案证']]
            else:
                matched_rows = biaozhun_drug_info[biaozhun_drug_info['通用名称'] == row['-产品名称']]
            
            if matched_rows.empty:  # 如果没有相关药品
                return '标准库内无相关药品无法比对'
            # 如果有相关药品，检查是否在预期不良反应（历史数据）中
            if adverse_effect not in matched_rows['预期不良反应（历史数据）'].str.cat(sep='|'):
                return adverse_effect
            return None

        def app_3(d):
            """评分函数"""
            return 3 if any(value >= 3 for value in d.values()) else 0

        def app_2(d):
            """评分函数"""
            return 2 if any(value >= 1 for value in d.values()) else 0

        # 应用函数到 df1 的每一行
        df1['疑似新的不良反应'] = df1.apply(is_new_adverse_reaction, axis=1)
        df1['疑似非新的不良反应'] = df1.apply(is_non_new_adverse_reaction, axis=1)
        df1['重点关注的不良反应'] = df1.apply(is_key_adverse_reaction, axis=1)
        df1['不在不良反应AI中的不良反应'] = df1.apply(not_in_ai_adverse_reaction, axis=1)
        df1['不在预期不良反应（历史数据）中的不良反应'] = df1.apply(not_in_expected_adverse_reaction, axis=1)

        # 调用 TOOLS_create_pivot_tool 函数
        if method == '证号':
            df3 = TOOLS_create_pivot_tool(df1, [['-注册人备案人', '产品类别', '-产品名称', '-注册证备案证'], ['-伤害'], ['报告编码'], ['nunique'], {'不良反应名称（规整）': 'SMALL_count_mode', '疑似新的不良反应': 'SMALL_count_mode', '疑似非新的不良反应': 'SMALL_count_mode', '重点关注的不良反应': 'SMALL_count_mode', '不在不良反应AI中的不良反应': 'SMALL_count_mode', '不在预期不良反应（历史数据）中的不良反应': 'SMALL_count_mode'}, ['报告编码合计']])
        elif method == '品种':
            df3 = TOOLS_create_pivot_tool(df1, [['产品类别', '-产品名称'], ['-伤害'], ['报告编码'], ['nunique'], {'不良反应名称（规整）': 'SMALL_count_mode', '疑似新的不良反应': 'SMALL_count_mode', '疑似非新的不良反应': 'SMALL_count_mode', '重点关注的不良反应': 'SMALL_count_mode', '不在不良反应AI中的不良反应': 'SMALL_count_mode', '不在预期不良反应（历史数据）中的不良反应': 'SMALL_count_mode'}, ['报告编码合计']])

        # 应用评分函数
        df3['新的评分'] = df3['疑似新的不良反应'].apply(app_3)
        df3['重点关注评分'] = df3['重点关注的不良反应'].apply(app_2)
        df3['总评分'] = df3['新的评分'] + df3['重点关注评分']

        # 排序
        df3 = df3.sort_values(by=['总评分', '报告编码合计'], ascending=[False, False]).reset_index(drop=True)
        return df3
    
    # 创建主窗口
    root = tk.Toplevel()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)
    root.title("不良反应检测（集合法）")
    root.configure(bg="white")  # 设置背景颜色为白色

    # 获取屏幕尺寸并居中窗口
    sw = root.winfo_screenwidth()
    sh = root.winfo_screenheight()
    ww = 450  # 窗口宽度
    wh = 200  # 窗口高度
    x = (sw - ww) // 2
    y = (sh - wh) // 2
    root.geometry(f"{ww}x{wh}+{x}+{y}")

    # 创建标签和下拉菜单
    label_column = tk.Label(root, text="请选择对象：", bg="white", font=("Arial", 12))
    label_column.grid(row=0, column=0, padx=10, pady=10)

    column_var = tk.StringVar(root, value="证号")  # 预设值为"证号"
    column_options = ['证号', '品种', '页面']  # 使用预设的选项
    column_menu = tk.OptionMenu(root, column_var, *column_options)
    column_menu.config(bg="white", font=("Arial", 10))
    column_menu.grid(row=0, column=1, padx=10, pady=10)

    # 创建标签和输入框
    label_number = tk.Label(root, text="请输入灵敏度：", bg="white", font=("Arial", 12))
    label_number.grid(row=1, column=0, padx=10, pady=5)

    number_var = tk.StringVar(value="1")
    number_entry = tk.Entry(root, textvariable=number_var, font=("Arial", 10))
    number_entry.grid(row=1, column=1, padx=10, pady=5)

    # 创建进度条
    progress = ttk.Progressbar(root, orient="horizontal", length=300, mode="determinate")
    progress.grid(row=2, column=0, columnspan=2, pady=10)

    # 创建确定按钮
    def on_ok_click():
        selected_column = column_var.get()
        input_number = number_var.get()

        # 简单的输入验证
        try:
            input_number = int(input_number)
        except ValueError:
            messagebox.showerror("输入错误", "请输入一个有效的数字。")
            return

        # 禁用按钮，防止重复点击
        ok_button.config(state=tk.DISABLED)

        # 启动线程处理数据
        def process_data():
            progress["value"] = 0
            root.update_idletasks()

            # 模拟数据处理过程
            total_steps = 10  # 假设总共有10个步骤
            for step in range(total_steps):
                time.sleep(0.5)  # 模拟每个步骤的耗时
                progress["value"] = (step + 1) * (100 / total_steps)
                root.update_idletasks()  # 更新界面

            # 处理数据
            if selected_column == "证号":
                df2 = TOOLS_create_pivot_tool(df[(df['报告类型-新的'] != "新的")], [['-注册证备案证'], [], ['不良反应名称（规整）'], ['SMALL_count_mode'], '', ''])
                df2.rename(columns={"-注册证备案证": "批准文号", "不良反应名称（规整）": "不良反应"}, inplace=True)
                df2 = SMALL_expand_dict_like_columns(df2)
                df2 = filter_adverse_effects(df2, '批准文号', input_number)
                df2['不良反应'] = df2['不良反应'].astype(str)
                result_df = analyze_products(df, df2, biaozhun['药品信息库'], selected_column)
            elif selected_column == "品种":
                # 对 biaozhun['药品信息库'] 按通用名称去重
                biaozhun_drug_info_unique = biaozhun['药品信息库'].drop_duplicates(subset=['通用名称'])
                df2 = TOOLS_create_pivot_tool(df[(df['报告类型-新的'] != "新的")], [['-产品名称'], [], ['不良反应名称（规整）'], ['SMALL_count_mode'], '', ''])
                df2.rename(columns={"-产品名称": "通用名称", "不良反应名称（规整）": "不良反应"}, inplace=True)
                df2 = SMALL_expand_dict_like_columns(df2)
                df2 = filter_adverse_effects(df2, '通用名称', input_number)
                df2['不良反应'] = df2['不良反应'].astype(str)
                result_df = analyze_products(df, df2, biaozhun_drug_info_unique, selected_column)
            elif selected_column == "页面":
                data_new = df.loc[df["报告类型-新的"].str.contains("新", na=False)].copy()
                data_old = df.loc[~df["报告类型-新的"].str.contains("新", na=False)].copy()
                list_new = drug(data_new, input_number)
                list_old = drug(data_old)
                new = ""
                old = ""
                for idc, colc in list_old.iterrows():
                    if "分隔符" not in colc["条目"]:
                        kde = "'" + str(colc["条目"]) + "':" + str(colc["详细描述T"]) + ","
                        old = old + kde
                for idc, colc in list_new.iterrows():
                    if str(colc["条目"]).strip() not in old and "分隔符" not in str(colc["条目"]):
                        kde = "'" + str(colc["条目"]) + "':" + str(colc["详细描述T"]) + ","
                        new = new + kde
                old = "{" + old + "}"
                new = "{" + new + "}"
                allon = "\n可能是新的不良反应：\n\n" + new + "\n\n\n可能不是新的不良反应：\n\n" + old
                PROGRAM_display_content_in_textbox(allon)
                result_df = None

            # 更新进度条
            progress["value"] = 100
            root.update_idletasks()

            # 显示结果
            if result_df is not None:
                PROGRAM_display_df_in_treeview(result_df, 0, df)

            # 启用按钮
            ok_button.config(state=tk.NORMAL)

            # 关闭窗口
            root.destroy()

        # 启动线程
        threading.Thread(target=process_data).start()

    ok_button = ttk.Button(root, text="确定", command=on_ok_click)
    ok_button.grid(row=3, column=0, columnspan=2, pady=10)

    # 启动Tkinter的事件循环
    root.mainloop()

def PSUR_keti_GUI(data):
      """日期预警功能"""
      
      def STAT_countx(x):
            """所有成分关键字计数,返回一个字典""" 
            return x.value_counts().to_dict()
            
      def STAT_countpx(x,y):
            """特定成分关键字计数,返回一个数值""" 
            return len(x[(x==y)])#.values      

      def STAT_get_mean(df):
            """返回平均值""" 
            return round(df.value_counts().mean(),2)
            
      def STAT_get_std(df):
            """返回标准差""" 
            return round(df.value_counts().std(ddof=1),2)
            

      def df_findrisk(df,target):      
            """预警模块,针对批号、月份、季度""" 
            if target=="产品批号":
                  return STAT_find_risk(df[(df["产品类别"]!="有源")],["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"],"注册证编号/曾用注册证编号",target)
            else:
                  return STAT_find_risk(df,["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"],"注册证编号/曾用注册证编号",target)
                  
      def STAT_find_risk_old(df,cols_list,main_col,target):      
                  """评分及预警模块,cols_list为所要引入的列（列表形式），main_col统计对象列（关键字），target为月份、季度或者批号等""" 
                  df=df.drop_duplicates(["报告编码"]).reset_index(drop=True)
                  dfx_findrisk1=df.groupby(cols_list).agg(
                        证号总数量=(main_col,"count"),      
                        包含元素个数=(target,"nunique"),
                        历史数据=(target,STAT_countx),            
                        均值=(target,STAT_get_mean),
                        标准差=(target,STAT_get_std),                  
                        ).reset_index()
                                    
                  cols_list2=cols_list.copy()
                  cols_list2.append(target)
                  dfx_findrisk2=df.groupby(cols_list2).agg(
                        计数=(target,"count"),
                        严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
                        死亡数量=("伤害",lambda x: STAT_countpx(x.values,"死亡")),      
                        单位个数=("单位名称","nunique"),      
                        单位列表=("单位名称",STAT_countx),                                          
                        ).reset_index()                        

                  dfx_findrisk=pd.merge(dfx_findrisk2,dfx_findrisk1,on=cols_list,how="left")#.reset_index()      
                              
                  return dfx_findrisk
            
      def STAT_find_risk(df, cols_list, main_col, target):      
            """评分及预警模块,cols_list为所要引入的列（列表形式），main_col统计对象列（关键字），target为月份、季度或者批号等""" 
            df = df.drop_duplicates(["报告编码"]).reset_index(drop=True)
            
            # 第一个聚合操作
            dfx_findrisk1 = df.groupby(cols_list).agg({
                  main_col: "count",             # 证号总数量
                  target: ["nunique", STAT_countx, STAT_get_mean, STAT_get_std]  # 包含元素个数,历史数据,均值,标准差
            }).reset_index()
            
            # 重命名列以匹配原始逻辑
            dfx_findrisk1.columns = cols_list + [
                  "证号总数量",
                  "包含元素个数",
                  "历史数据",
                  "均值",
                  "标准差"
            ]
            
            # 第二个聚合操作（修正后的写法）
            cols_list2 = cols_list.copy()
            cols_list2.append(target)
            
            # 方法1：直接聚合（推荐 pandas 1.0+）
            dfx_findrisk2 = df.groupby(cols_list2).agg({
                  target: "count",  # 计数
                  "伤害": [
                        ("严重伤害数", lambda x: STAT_countpx(x.values, "严重伤害")),
                        ("死亡数量", lambda x: STAT_countpx(x.values, "死亡"))
                  ],
                  "单位名称": [
                        ("单位个数", "nunique"),
                        ("单位列表", STAT_countx)
                  ]
            }).reset_index()
            
            # 调整列名（因为多级索引）
            dfx_findrisk2.columns = [
                  *cols_list2,
                  "计数",
                  "严重伤害数",
                  "死亡数量",
                  "单位个数",
                  "单位列表"
            ]
            
            # 合并两个结果
            dfx_findrisk = pd.merge(dfx_findrisk2, dfx_findrisk1, on=cols_list, how="left")
            
            return dfx_findrisk
            
        
      def keti(timex,time_windows,time_base,data0,kx):
            print(time_windows,time_base)
            
            if time_base=='发生日期':
                  time_base='事件发生日期'
                  time_base_m='事件发生月份'
                  time_base_q='事件发生季度'
            if time_base=='报告日期':
                  time_base='报告日期'
                  time_base_m='报告月份'
                  time_base_q='报告季度'                  
                                                                  
                  
            data0['规整后品类']=data0['产品名称']
            data0['报告日期'] = pd.to_datetime(data0['报告日期'], format='%Y-%m-%d', errors='coerce')       
            data0['事件发生日期'] = pd.to_datetime(data0['事件发生日期'], format='%Y-%m-%d', errors='coerce')       
            data0["伤害与评价"]=data0["伤害"]+data0["持有人报告状态"]                        
            data0["报告月份"] = data0["报告日期"].dt.to_period("M").astype(str)      
            data0["报告季度"] = data0["报告日期"].dt.to_period("Q").astype(str)      
            data0["报告年份"] = data0["报告日期"].dt.to_period("Y").astype(str)      #      品种评价            
            data0["事件发生月份"] = data0["事件发生日期"].dt.to_period("M").astype(str)                  
            data0["事件发生季度"] = data0["事件发生日期"].dt.to_period("Q").astype(str)                        
            data0["事件发生年份"] = data0["事件发生日期"].dt.to_period("Y").astype(str)
            
      
                  
                                    
            k1=kx["权重"][0] #严重比
            k2=kx["权重"][1] #单位数量            
            k3=kx["权重"][2] #批号或型号集中度权重
            k4=kx["权重"][3] #高度关注关键字（一级）            
            k4_values=kx["值"][3] #高度关注关键字（一级） 值      
            
            k5=kx["权重"][4] #高度关注关键字（二级）            
            k5_values=kx["值"][4] #高度关注关键字（二级） 值
            
            k6=kx["权重"][5] #高风险品种            
            k6_values=kx["值"][5] 

            k7=kx["权重"][6] #低风险品种      
            k7_values=kx["值"][6] 
            
            k8=kx["权重"][7] #低风险问题      
            k8_values=kx["值"][7] 
      
                                    
            lastdayfrom = pd.to_datetime(timex)
            data2=data0.copy().set_index(time_base)
            data2=data2.sort_index()
            data2['高度关注关键字']=''
            data2['二级敏感词']=''
            data2['高风险品种']=0
            data2['低风险品种']=0
            data2['低风险问题']=0
            if ini["模式"]=="器械":            
                  data2["关键字查找列"]=data2["器械故障表现"].astype(str)+data2["伤害表现"].astype(str)+data2["使用过程"].astype(str)+data2["事件原因分析描述"].astype(str)+data2["初步处置情况"].astype(str)#器械故障表现|伤害表现|使用过程|事件原因分析描述|初步处置情况
            else:
                  data2["关键字查找列"]=data2["器械故障表现"].astype(str)
            
            # 定义一个函数来检查匹配并更新列
            def update_keywords(row, patterns):
                  matched_keywords = [pattern.pattern for pattern in patterns if pattern.search(row["关键字查找列"])]
                  if matched_keywords:
                        return ";".join(matched_keywords)
                  return ""
                  
            k4_set = k4_values.split('|')
            k5_set = k5_values.split('|')
                        
            compiled_patterns_k4 = [re.compile(pattern) for pattern in k4_set]
            compiled_patterns_k5 = [re.compile(pattern) for pattern in k5_set]            
            data2["高度关注关键字"] = data2.apply(lambda row: update_keywords(row, compiled_patterns_k4), axis=1)
            data2["二级敏感词"] = data2.apply(lambda row: update_keywords(row, compiled_patterns_k5), axis=1)

            data2.loc[data2["关键字查找列"].str.contains(k8_values, na=False), "低风险问题"]  = 1      
            data2.loc[data2["产品名称"].str.contains(k6_values, na=False), "高风险品种"]  = 1      
            data2.loc[data2["关键字查找列"].str.contains(k7_values, na=False), "低风险问题"]  = 1      
            
            #月度数据
            if time_windows=='月份窗口':                              
                  data30 = data2.loc[lastdayfrom - pd.Timedelta(days=30):lastdayfrom].reset_index()

            #季度数据
            if time_windows=='季度窗口':      
                  data30 = data2.loc[lastdayfrom - pd.Timedelta(days=90):lastdayfrom].reset_index()
                  
            #增加对使用所有数据的兼容性
            if time_windows=='所有数据':
                  data30=data2.copy()
            
            
            
            #对于月度窗口或全部数据窗口所对应的全部数据：
            if time_windows!='季度窗口':      
                  timex_date = pd.to_datetime(timex)
                  start_date = (timex_date.replace(day=1) - pd.DateOffset(months=12)).date()
                  end_date = (timex_date.replace(day=1) - timedelta(days=1)).date()
                  data365 = data2.loc[start_date:end_date]
                  #data365=data2.loc[lastdayfrom - pd.Timedelta(days=365):lastdayfrom].reset_index()
            #季度窗口的情况
            else:
                  # 将timex转换为Pandas的datetime对象
                  timex_date = pd.to_datetime(timex)
                  # 将timex_date转换为季度频率的Period对象
                  current_quarter_period = pd.Period(timex_date, freq='Q')
                  # 获取前一个季度的Period对象
                  previous_quarter_period = current_quarter_period - 1
                  # 将前一个季度的Period对象转换为时间戳，并指定为季度末的时间点
                  # 注意：这里我们不需要再减去一天，因为to_timestamp('M', how='end')直接给我们季度的最后一天
                  previous_quarter_end_timestamp = previous_quarter_period.to_timestamp('M', how='end')
                  # 如果只需要日期部分，可以使用.date()方法
                  end_date = previous_quarter_end_timestamp.date()
                  start_date = end_date - timedelta(days=365)
                  data365 = data2.loc[start_date:end_date]
                  
            #当前时间窗数据评分
            df301=data30.groupby(["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"]).agg(
                  证号计数=("报告编码","nunique"),
                  批号个数=("产品批号","nunique"),
                  批号列表=("产品批号",STAT_countx),      
                  型号个数=("型号","nunique"),
                  型号列表=("型号",STAT_countx),            
                  规格个数=("规格","nunique"),      
                  规格列表=("规格",STAT_countx),            
                  ).sort_values(by="证号计数", ascending=[False], na_position="last").reset_index()      

            df302=data30.drop_duplicates(["报告编码"]).groupby(["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"]).agg(
                  严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
                  死亡数量=("伤害",lambda x: STAT_countpx(x.values,"死亡")),      
                  单位个数=("单位名称","nunique"),      
                  单位列表=("单位名称",STAT_countx),                              
                  待评价数=("持有人报告状态",lambda x: STAT_countpx(x.values,"待评价")),
                  严重伤害待评价数=("伤害与评价",lambda x: STAT_countpx(x.values,"严重伤害待评价")),
                  高度关注关键字=("高度关注关键字","sum"),      
                  二级敏感词=("二级敏感词","sum"),
                  高风险品种=("高风险品种","sum"),                  
                  低风险品种=("低风险品种","sum"),                  
                  低风险问题=("低风险问题","sum"),                  
                  ).reset_index()      
            
            df30=pd.merge(df301,  df302,on=["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"], how="left")      


            
      
            df30xinghao=data30.groupby(["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","型号"]).agg(
                  型号计数=("报告编码","nunique"),      
                  ).sort_values(by="型号计数", ascending=[False], na_position="last").reset_index()            
            df30xinghao=df30xinghao.drop_duplicates("注册证编号/曾用注册证编号")            
                  
            df30wu=data30.drop_duplicates(["报告编码"]).groupby(["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","产品批号"]).agg(
                  批号计数=("报告编码","nunique"),      
                  严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
                  ).sort_values(by="批号计数", ascending=[False], na_position="last").reset_index()

                              
                  
            df30wu["风险评分-影响"]=0            
            df30wu["评分说明"]=""            
            df30wu.loc[((df30wu["批号计数"]>=3)&(df30wu["严重伤害数"]>=1)&(df30wu["产品类别"]!="有源"))|((df30wu["批号计数"]>=5)&(df30wu["产品类别"]!="有源")), "风险评分-影响"] = df30wu["风险评分-影响"]+3      
            df30wu.loc[(df30wu["风险评分-影响"]>=3), "评分说明"] = df30wu["评分说明"]+"●符合预警无源规则+3;"      
            

            
            df30wu=df30wu.sort_values(by="风险评分-影响", ascending=[False], na_position="last").reset_index(drop=True)      
            df30wu=df30wu.drop_duplicates("注册证编号/曾用注册证编号")
            df30xinghao=df30xinghao[["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","型号","型号计数"]]      
            df30wu=df30wu[["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","产品批号","批号计数","风险评分-影响","评分说明"]]
            df30=pd.merge(df30, df30xinghao, on=["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"], how="left")
      
            df30=pd.merge(df30, df30wu, on=["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"], how="left")            

            #TABLE_tree_Level_2(df30,1,data30)            
            #符合省中心规则打分（因为是针对证号，按有源标准）
            df30.loc[((df30["证号计数"]>=3)&(df30["严重伤害数"]>=1)&(df30["产品类别"]=="有源"))|((df30["证号计数"]>=5)&(df30["产品类别"]=="有源")), "风险评分-影响"] = df30["风险评分-影响"]+3      
            df30.loc[(df30["风险评分-影响"]>=3)&(df30["产品类别"]=="有源"), "评分说明"] = df30["评分说明"]+"●符合预警有源规则+3;"      


                        
            #针对死亡
            df30.loc[(df30["死亡数量"]>=1), "风险评分-影响"] = df30["风险评分-影响"]+50      
            df30.loc[(df30["风险评分-影响"]>=10), "评分说明"] = df30["评分说明"]+"存在死亡报告;"      
            
            #严重比评分
            fen_yanzhong=round(k1*(df30["严重伤害数"]/df30["证号计数"]),2)
            df30["风险评分-影响"] = df30["风险评分-影响"]+      fen_yanzhong
            df30["评分说明"] = df30["评分说明"]+"严重比评分"+fen_yanzhong.astype(str)+";"                  
            
            #报告单位数评分
            fen_danwei=round(k2*(np.log(df30["单位个数"])),2)
            df30["风险评分-影响"] = df30["风险评分-影响"]+      fen_danwei
            df30["评分说明"] = df30["评分说明"]+"报告单位评分"+fen_danwei.astype(str)+";"                        
            
            #批号型号集中度评分
            df30.loc[(df30["产品类别"]=="有源")&(df30["证号计数"]>=3), "风险评分-影响"] = df30["风险评分-影响"]+k3*df30["型号计数"]/df30["证号计数"]                  
            df30.loc[(df30["产品类别"]=="有源")&(df30["证号计数"]>=3), "评分说明"] = df30["评分说明"]+"型号集中度评分"+(round(k3*df30["型号计数"]/df30["证号计数"],2)).astype(str)+";"      
            df30.loc[(df30["产品类别"]!="有源")&(df30["证号计数"]>=3), "风险评分-影响"] = df30["风险评分-影响"]+k3*df30["批号计数"]/df30["证号计数"]                  
            df30.loc[(df30["产品类别"]!="有源")&(df30["证号计数"]>=3), "评分说明"]  = df30["评分说明"]+"批号集中度评分"+(round(k3*df30["批号计数"]/df30["证号计数"],2)).astype(str)+";"                  

            #高度关注关键字（一级）
            df30.loc[(df30["高度关注关键字"]!=''), "风险评分-影响"]  = df30["风险评分-影响"]+k4
            df30.loc[(df30["高度关注关键字"]!=''), "评分说明"] = df30["评分说明"]+"●含有高度关注关键字评分"+str(k4)+"；"                                                      

            #二级敏感词
            df30.loc[(df30["二级敏感词"]!=''), "风险评分-影响"]  = df30["风险评分-影响"]+k5
            df30.loc[(df30["二级敏感词"]!=''), "评分说明"] = df30["评分说明"]+"含有二级敏感词评分"+str(k5)+"；"            
            
            #高风险品种
            df30.loc[(df30["高风险品种"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k6
            df30.loc[(df30["高风险品种"]>=1), "评分说明"] = df30["评分说明"]+"高风险品种"+str(k6)+"；"      
            
            #低风险品种
            df30.loc[(df30["低风险品种"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k7
            df30.loc[(df30["低风险品种"]>=1), "评分说明"] = df30["评分说明"]+"高风险品种"+str(k7)+"；"      
            
            
            #低风险问题
            df30.loc[(df30["低风险问题"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k8
            df30.loc[(df30["低风险问题"]>=1), "评分说明"] = df30["评分说明"]+"减分项评分"+str(k8)+"；"      
            

            
                        
            #历史比较（月份或季度）
            if time_windows=='月份窗口' or time_windows=='所有数据':      
                  df365month=df_findrisk(data365,time_base_m)
            if time_windows=='季度窗口':      
                  df365month=df_findrisk(data365,time_base_q)                  
            df365month=df365month.drop_duplicates("注册证编号/曾用注册证编号")      
            df365month=df365month[["注册证编号/曾用注册证编号",'历史数据',"均值","标准差"]]
            df30=pd.merge(df30, df365month, on=["注册证编号/曾用注册证编号"], how="left")      
            df30["均值"]=round(df30["均值"],2)      
            df30["标准差"]=round(df30["标准差"],2)
            df30["风险评分-历史"]=1
            df30["mfc"]=""

            #增加对使用所有数据的兼容性
            if time_windows!='所有数据':
                  
                  df30.loc[(df30["证号计数"]>=3), "风险评分-历史"]  = df30["风险评分-历史"]+2
                  df30.loc[(df30["证号计数"]>=3),  "mfc"] = "数量超过3例；"            
                  
                  
                  df30.loc[((df30["证号计数"]>df30["均值"])&(df30["风险评分-历史"]>=3)&(df30["标准差"].astype(str)!="nan")), "风险评分-历史"]  = df30["风险评分-历史"]+1
                  df30.loc[((df30["证号计数"]>df30["均值"])&(df30["风险评分-历史"]>=3)&(df30["标准差"].astype(str)!="nan")), "mfc"] = "月份计数超过历史均值"+df30["均值"].astype(str)+"；"      
                        
                  df30.loc[(df30["证号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-历史"]>=4), "风险评分-历史"]  = df30["风险评分-历史"]+2
                  df30.loc[(df30["证号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-历史"]>=4), "mfc"] = "超过历史均值一个标准差；"                  
                  
                        
                  df30.loc[(df30["证号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-历史"]>=6), "风险评分-历史"]  = df30["风险评分-历史"]+2
                  df30.loc[(df30["证号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-历史"]>=6), "mfc"] = "超过历史均值两个标准差；"                        

            df30["评分说明"]=df30["评分说明"]+"●●证号数量："+df30["证号计数"].astype(str)+";"+ df30["mfc"]      
            del df30["mfc"]
            df30=df30.rename(columns={"均值": "历史均值","标准差": "历史标准差","历史数据": "历史时间数据"})
            

            #历史比较（批号）
            df365month=df_findrisk(data365,"产品批号")
            df365month=df365month.drop_duplicates("注册证编号/曾用注册证编号")      
            df365month=df365month[["注册证编号/曾用注册证编号","历史数据","均值","标准差"]]
            df30=pd.merge(df30, df365month, on=["注册证编号/曾用注册证编号"], how="left")      
      
            df30["风险评分-批号"]=1
            df30.loc[(df30["产品类别"]!="有源"), "评分说明"] =df30["评分说明"]+"●●高峰批号数量："+df30["批号计数"].astype(str)+";"
            
            #增加对使用所有数据的兼容性
            if time_windows!='所有数据':
                  df30.loc[(df30["批号计数"]>=3), "风险评分-批号"]  = df30["风险评分-批号"]+2
                  df30.loc[(df30["批号计数"]>=3),  "mfc"] = "数量超过3例；"            
                  
                  
                  df30.loc[((df30["批号计数"]>df30["均值"])&(df30["风险评分-批号"]>=3)&(df30["标准差"].astype(str)!="nan")), "风险评分-批号"]  = df30["风险评分-批号"]+1
                  df30.loc[((df30["批号计数"]>df30["均值"])&(df30["风险评分-批号"]>=3)&(df30["标准差"].astype(str)!="nan")), "mfc"] = "高峰批号计数超过历史均值"+df30["均值"].astype(str)+"；"      
                        
                  df30.loc[(df30["批号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-批号"]>=4), "风险评分-批号"]  = df30["风险评分-批号"]+2
                  df30.loc[(df30["批号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-批号"]>=4), "mfc"] = "高峰批号超过历史均值一个标准差；"                  
                  
                        
                  df30.loc[(df30["批号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-批号"]>=6), "风险评分-批号"]  = df30["风险评分-批号"]+2
                  df30.loc[(df30["批号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-批号"]>=6), "mfc"] = "高峰批号超过历史均值两个标准差；"            
            
                  
            df30=df30.rename(columns={"均值": "历史批号均值","标准差": "历史批号标准差","历史数据": "历史批号数据"})

            
            df30["风险评分-影响"]=round(df30["风险评分-影响"],2)
            df30["风险评分-历史"]=round(df30["风险评分-历史"],2)
            df30["风险评分-批号"]=round(df30["风险评分-批号"],2)
            
            df30["总体评分"]=df30["风险评分-影响"].copy()
            df30["关注建议"]=""
            df30.loc[(df30["风险评分-影响"]>=3),  "关注建议"]=df30["关注建议"]+"●建议关注(影响范围)；" 
            df30.loc[(df30["风险评分-历史"]>=5),  "关注建议"]=df30["关注建议"]+"●建议关注(当期数量异常)；"
            df30.loc[(df30["风险评分-批号"]>=3),  "关注建议"]=df30["关注建议"]+"●建议关注(高峰批号数量异常)。"            
            df30.loc[(df30["风险评分-历史"]>=df30["风险评分-批号"]),  "总体评分"]=df30["风险评分-影响"]*df30["风险评分-历史"]
            df30.loc[(df30["风险评分-历史"]<df30["风险评分-批号"]),  "总体评分"]=df30["风险评分-影响"]*df30["风险评分-批号"]

            df30["总体评分"]=round(df30["总体评分"],2)            
            df30["评分说明"]=df30["关注建议"]      +df30["评分说明"]            
            df30=df30.sort_values(by=["总体评分","风险评分-影响"], ascending=[False,False], na_position="last").reset_index(drop=True)
            
                  
                  
            df30=df30[["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号","总体评分","风险评分-影响","风险评分-历史","风险评分-批号","评分说明","证号计数","严重伤害数","死亡数量","单位个数","单位列表","批号个数","批号列表","型号个数","型号列表","规格个数","规格列表","待评价数","严重伤害待评价数","高度关注关键字","二级敏感词",'历史时间数据',"历史均值","历史标准差",'历史批号数据',"历史批号均值","历史批号标准差","型号","型号计数","产品批号","批号计数"]]
            df30["报表类型"]="{'grouped':['上市许可持有人名称','产品类别','产品名称','注册证编号/曾用注册证编号']}"

            PROGRAM_display_df_in_treeview(df30,1,data30)
            pass                  
      

      if 1==1:
            ini={}
            ini['模式']='器械'
            if ini["模式"]=="药品":
                  kx = pd.read_excel(os.path.join(peizhidir,"easy_WarningParameters.xlsx"), header=0, sheet_name="药品").reset_index(drop=True)
            if ini["模式"]=="器械":
                  kx = pd.read_excel(os.path.join(peizhidir,"easy_WarningParameters.xlsx"), header=0, sheet_name="器械").reset_index(drop=True)      
            if ini["模式"]=="化妆品":
                  kx = pd.read_excel(os.path.join(peizhidir,"（范例）预警参数.xlsx"), header=0, sheet_name="化妆品").reset_index(drop=True)

      se = tk.Toplevel()
      bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(se)

      se.title('风险预警')
      sw_se = se.winfo_screenwidth()
    #得到屏幕宽度
      sh_se = se.winfo_screenheight()
    #得到屏幕高度
      ww_se = 250
      wh_se = 250
    #窗口宽高为100
      x_se = (sw_se-ww_se) / 2
      y_se = (sh_se-wh_se) / 2
      se.geometry("%dx%d+%d+%d" %(ww_se,wh_se,x_se,y_se)) 

      import_se=tk.Label(se,text="预警日期：")
      import_se.grid(row=1, column=0, sticky="w", padx=10, pady=10)
      import_se_entry=tk.Entry(se, width = 30)
      import datetime
      import_se_entry.insert(0,datetime.date.today())
      import_se_entry.grid(row=1, column=1, sticky="w")

      # 创建第一个下拉框：请选择时间窗
      time_window_var = tk.StringVar(value="月份窗口")  # 设置默认选项
      time_window_menu_options = ["月份窗口",'季度窗口','所有数据']  # 定义选项列表
      time_window_label = tk.Label(se, text="请选择时间窗：")
      time_window_label.grid(row=2, column=0, sticky="w", padx=10, pady=10)
      time_window_optionmenu = tk.OptionMenu(se, time_window_var, *time_window_menu_options)
      time_window_optionmenu.grid(row=2, column=1, sticky="w", padx=10, pady=10)
       
      # 创建第二个下拉框：请选择时间依据
      time_basis_var = tk.StringVar(value="报告日期")  # 设置默认选项
      time_basis_menu_options = ["报告日期", "发生日期"]  # 定义选项列表
      time_basis_label = tk.Label(se, text="请选择时间依据：")
      time_basis_label.grid(row=3, column=0, sticky="w", padx=10, pady=10)
      time_basis_optionmenu = tk.OptionMenu(se, time_basis_var, *time_basis_menu_options)
      time_basis_optionmenu.grid(row=3, column=1, sticky="w", padx=10, pady=10)

      # 创建第3个下拉框：请选择规则
      workbook = load_workbook(os.path.join(peizhidir,"easy_WarningParameters.xlsx"), read_only=True)
      sheet_names = workbook.sheetnames
      kx_basis_var = tk.StringVar(value="器械")  # 设置默认选项
      kx_basis_menu_options = sheet_names  # 定义选项列表
      kx_basis_label = tk.Label(se, text="请选择预警参数：")
      kx_basis_label.grid(row=4, column=0, sticky="w", padx=10, pady=10)
      kx_basis_optionmenu = tk.OptionMenu(se, kx_basis_var, *kx_basis_menu_options)
      kx_basis_optionmenu.grid(row=4, column=1, sticky="w", padx=10, pady=10)


      
      btn_se=tk.Button(se,text="确定",width=10,command=lambda:PROGRAM_display_df_in_treeview(keti(import_se_entry.get(),time_window_var.get(),time_basis_var.get(),data,pd.read_excel(os.path.join(peizhidir,"easy_WarningParameters.xlsx"), header=0, sheet_name=kx_basis_var.get()).reset_index(drop=True)),1,data))
      btn_se.grid(row=5, column=1, sticky="w", padx=10, pady=10)

      
      pass
      


############################################################################################################################
#定制的和个性化设置的自定义函数
############################################################################################################################
class AAA_03_Setting():
    pass
def SETTING_style_UI(window):
    # ==================== 浅灰现代扁平化配色方案 ====================
    style = ttk.Style(window)
    style.theme_use('clam')  # 使用clam主题作为基础

    # 浅灰色配色方案 (现代化风格)
    bg_color = "#DCDAD5"       # 主背景色 (浅灰白)
    frame_color = "#E0E0E0"    # 框架色 (稍深灰)
    button_color = "#A0A0A0"   # 按钮基础色 (中灰)
    text_color = "#212121"     # 主文本色 (深灰近黑)
    highlight_color = "#E0E0E0" # 高亮色 (浅灰)
    accent_color = "#1976D2"   # 强调色 (Material Blue)
    border_color = "#BDBDBD"   # 边框色
    
    style.configure('TLabel', background=bg_color, foreground=text_color)  # 再强制覆盖  
    
    # ==================== 控件样式配置 ====================
    # 基础样式 (会应用到所有标准tkinter控件)

    # 框架样式
    style.configure('TFrame', 
                   background=bg_color,
                   relief="flat")
    
    # 标签样式 (包括标准tk.Label)
    window.option_add('*Label*background', frame_color)
    window.option_add('*Label*foreground', text_color)
    window.option_add('*Label*font', (my_font_ch, 10))
    
    # 标签框架样式
    style.configure('TLabelFrame', 
                   background=bg_color,
                   foreground=accent_color,
                   borderwidth=1,
                   relief="flat",
                   padding=(10, 5))
    style.configure('TLabelFrame.Label', 
                   background=bg_color,
                   foreground=accent_color)
    # 同时设置tk.Label的默认选项
    window.option_add('*Label.background', bg_color)
    window.option_add('*Label.foreground', text_color)
    window.option_add('*Label.font', (my_font_ch, 10))
    
    style.configure('TButton', 
                   font=(my_font_ch, 10),
                   padding="3 2",
                   background='#E0E0E0',
                   foreground='#2E4053',
                   borderwidth=1,
                   relief="groove",
                   width=0,
                   focuscolor='',          # 新增：移除焦点颜色
                   highlightthickness=0)   # 新增：移除焦点边框

    style.map('TButton',
              background=[('active', '#D0D0D0'),
                         ('pressed', '#C0C0C0')],
              foreground=[('active', '#2E4053'),
                         ('pressed', '#2E4053')],
              focuscolor=[('!focus', '')]  # 新增：确保无焦点状态
    )
        
    # 输入框样式
    style.configure('TEntry',
                   fieldbackground="white",
                   foreground=text_color,
                   borderwidth=1,
                   insertcolor=text_color,
                   relief="flat",
                   padding=5)
    style.map('TEntry',
              fieldbackground=[('focus', 'white')],
              bordercolor=[('focus', accent_color)])
    
    # 树形视图样式 (保持白底黑字)
    style.configure('Treeview',
                   font=(my_font_ch, 10),
                   rowheight=28,
                   background="white",
                   fieldbackground="white",
                   foreground=text_color,
                   bordercolor=border_color,
                   borderwidth=0)
    style.configure('Treeview.Heading',
                   font=(my_font_ch, 10, 'bold'),
                   background=frame_color,
                   foreground=text_color,
                   relief="flat",
                   padding=5)
    style.map('Treeview',
              background=[('selected', '#E3F2FD')],  # 浅蓝色选中背景
              foreground=[('selected', text_color)])
    
    # 滚动条样式
    style.configure('Vertical.TScrollbar',
                   background=frame_color,
                   arrowcolor=text_color,
                   bordercolor=bg_color,
                   gripcount=0,
                   arrowsize=12)
    style.map('Vertical.TScrollbar',
              background=[('active', '#BDBDBD')])
    
    # 设置窗口全局样式
    window.configure(background=bg_color)
    
    return bg_color, frame_color, button_color, text_color, highlight_color



def SETTING_create_menu(data):  
    """创建菜单栏，包含File、Edit和Help三个菜单项。"""  

    win=data["windows"]
    win_progressbar=data["win_progressbar"]
    ori_owercount_easyread=data["ori_owercount_easyread"]
    ori=data["ori"]   
    datacols=ori_owercount_easyread.columns.tolist() 
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(win)
    menu_bar = tk.Menu(win, 
                      background=frame_color,  # 菜单背景色
                      foreground=text_color,   # 菜单文本色
                      activebackground=button_color,  # 悬停背景色
                      activeforeground="white",       # 悬停文本色
                      borderwidth=0,                  # 无边框
                      relief="flat")  
    
    
    win.config(menu=menu_bar)

    file_menu = tk.Menu(menu_bar, 
                        background=frame_color,  # 菜单项背景色
                        foreground=text_color,   # 菜单项文本色
                        activebackground=button_color,  # 悬停背景色
                        activeforeground="white",       # 悬停文本色
                        tearoff=0)                       # 禁止菜单项被拖出

    menu_bar.add_cascade(label="文件", menu=file_menu)  
    file_menu.add_command(label="载入文件", command=lambda:PROGRAM_display_df_in_treeview(SMALL_read_and_merge_files(),0,0))     
    file_menu.add_separator()  
 
    file_menu.add_command(label="导出文件", command=lambda:SMALL_save_dict(ori_owercount_easyread))    
     
    file_menu.add_separator()    
    file_menu.add_command(label="从数据库载入", command=lambda:SQL_create_query_gui(ori_owercount_easyread)) 
    file_menu.add_command(label="另存为数据库", command=lambda:SQL_df_to_sqlite_db_with_gui(ori_owercount_easyread))  
    file_menu.add_command(label="批量转为数据库", command=lambda:SQL_excels_to_db())  
    file_menu.add_command(label="添加到数据库", command=lambda:SQL_update_sqlite_db_with_df(ori_owercount_easyread))     
    file_menu.add_command(label="数据库去重", command=lambda:SQL_database_deduplication_tool())          
    file_menu.add_command(label="数据库分组透视", command=lambda:SQL_PivotTool(filedialog.askopenfilename(filetypes=[("DB Files", "*.db"), ("All Files", "*.*")])).create_pivot_tool_gui())          

    file_menu.add_separator()           
    file_menu.add_command(label="文件转字典", command=lambda:PROGRAM_display_content_in_textbox(str(ori_owercount_easyread.to_dict(orient='list')))) 
    #file_menu.add_command(label="转化字典列", command=lambda:PROGRAM_display_df_in_treeview(SMALL_expand_dict_like_columns(ori_owercount_easyread),0,0))   
    file_menu.add_command(label="字典筛选", command=lambda:PROGRAM_display_df_in_treeview(Small_FilterDataWithGUI(ori_owercount_easyread).run(),0,0))   

    
    file_menu.add_separator()  
    file_menu.add_command(label="追加合并表格", command=lambda:PROGRAM_merge_dataframes(ori_owercount_easyread))  
    file_menu.add_command(label="更新表格", command=lambda:CLEAN_DataFrameUpdateApp(win,ori_owercount_easyread.copy(),pd.read_excel(filedialog.askopenfilename(filetypes=[("XLSX Files", "*.xlsx"), ("All Files", "*.*")])) ))  


    file_menu.add_separator()   
    file_menu.add_command(label="导出文本文件", command=lambda:SMALL_save_df_as_txt_files(ori_owercount_easyread,filedialog.askdirectory()))
    file_menu.add_command(label="切换视图", command=lambda:PROGRAM_display_df_in_treeview(SMALL_easyreadT(ori_owercount_easyread),0,0))  
    file_menu.add_separator()   
    file_menu.add_command(label="问题和建议", command=lambda:showinfo(title="联系我们", message="如有任何问题或建议，请联系蔡老师，411703730（微信或QQ）。"))  
    file_menu.add_separator() 
    file_menu.add_command(label="退出", command=lambda:exit())  


    nomal_menu = tk.Menu(menu_bar, 
                        background=frame_color,  # 菜单项背景色
                        foreground=text_color,   # 菜单项文本色
                        activebackground=button_color,  # 悬停背景色
                        activeforeground="white",       # 悬停文本色
                        tearoff=0)                       # 禁止菜单项被拖出
 
    menu_bar.add_cascade(label="预制清洗", menu=nomal_menu)
    

    
    #加载器械清洗菜单和相关清洗任务
    
    nomal_menu.add_command(label="自定义标准清洗(字典任务)", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),SMALL_read_and_merge_files()),0,0)) 
    nomal_menu.add_command(label="自定义标准清洗（赋关键词）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),SMALL_read_and_merge_files(),'赋关键词'),0,0)) 
    nomal_menu.add_command(label="自定义标准清洗（加关键词）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),SMALL_read_and_merge_files(),'加关键词'),0,0)) 
    nomal_menu.add_separator()
    nomal_menu.add_command(label="数据编码", command=lambda:PROGRAM_display_df_in_treeview(RWS_data_encoding_gui(ori_owercount_easyread.copy()).run(),0,0)) 
    nomal_menu.add_command(label="检测每列空值", command=lambda:PROGRAM_display_df_in_treeview(RWS_summarize_missing_values(ori_owercount_easyread),0,ori_owercount_easyread))   
    nomal_menu.add_separator()
    nomal_menu.add_command(label="选列脱敏", command=lambda:CLEAN_DataMaskingApp(win,ori_owercount_easyread.copy()))  
    nomal_menu.add_command(label="脱敏恢复", command=lambda:CLEAN_restore_data(ori_owercount_easyread))     

    if  '报告编码' in datacols and '器械故障表现' in datacols:
        nomal_menu.add_command(label="医疗器械报告一键脱敏", command=lambda:CLEAN_data_masking_qixie(ori_owercount_easyread.copy()))     


    
    
    nomal_menu.add_separator()  
    nomal_menu.add_command(label="行选择", command=lambda:CLEAN_select_rows_interactively(ori_owercount_easyread.copy()))   



    nomal_menu.add_separator()
    if biaozhun!={} and '器械故障表现' in datacols:
        nomal_menu.add_command(label="○器械预制-基础清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["器械清洗"]),0,0)) 
        if "-监测机构"  in datacols:
            nomal_menu.add_command(label="○器械预制-表现归类（关键词法）", command=lambda:PSUR_get_guize2(ori_owercount_easyread)) 


    if biaozhun!={} and '用药开始时间' in datacols:
        nomal_menu.add_separator()
        nomal_menu.add_command(label="●药品预制-基础清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["药品清洗"]),0,0))   
        if "不良反应名称（规整）"  in datacols:
            nomal_menu.add_command(label="● 药品预制-PT标准化清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["药品PT清洗"]),0,0))   
            nomal_menu.add_command(label="● 药品预制-赋SOC（PT扩行-关键词法-标准库）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；', ["不良反应名称（规整）"]),biaozhun["药品关键词"],'赋关键词'),0,0)) 
            nomal_menu.add_command(label="● 药品预制-加SOC（PT扩行-MedDRA法）", command=lambda:PROGRAM_display_df_in_treeview(SMALL_merge_dataframes(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；', ["不良反应名称（规整）"]), "不良反应名称（规整）", os.path.join(peizhidir,'share_easy_adrmdr_药品规整-SOC-Meddra库.xlsx'), df2_col='PT'),0,0)) 
            nomal_menu.add_command(label="● 药品预制-加SOC（PT扩行-调用历史知识库）", command=lambda:PROGRAM_display_df_in_treeview(SMALL_merge_dataframes(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；', ["不良反应名称（规整）"]), "不良反应名称（规整）", biaozhun["药品不良反应名称规整-AI"] , df2_col='不良反应名称AI'),0,0)) 
            nomal_menu.add_command(label="● 药品预制-药品信息标准化（含厂家-调用历史知识库）", command=lambda:PROGRAM_display_df_in_treeview(PSUR_DRUG_merge_dfs(ori_owercount_easyread.copy(), biaozhun["药品信息库"]),0,0)) 
            nomal_menu.add_command(label="● 药品预制-药品信息标准化（不含厂家-调用历史知识库）", command=lambda:PROGRAM_display_df_in_treeview(PSUR_DRUG_merge_dfs(ori_owercount_easyread.copy(), biaozhun["药品信息库"],'精简'),0,0)) 


 
        nomal_menu.add_separator()

        
        nomal_menu.add_command(label="以药品定制模式打开（执行规整）", command=lambda:PSUR_yaopin(ori_owercount_easyread.copy()))  
        nomal_menu.add_command(label="以药品定制模式打开（不执行规整）", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread.copy(),'psur',0)) 
    if biaozhun!={} and '器械故障表现' in datacols:
        nomal_menu.add_command(label="以器械定制模式打开（执行规整）", command=lambda:PSUR_qixie(ori_owercount_easyread.copy()))  
        nomal_menu.add_command(label="以器械定制模式打开（不执行规整）", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread.copy(),'psur',0)) 
         
    if '---器械数据规整---' in datacols or '---药品数据规整---' in datacols: 
        
        nomal_menu2 = tk.Menu(menu_bar, 
                            background=frame_color,  # 菜单项背景色
                            foreground=text_color,   # 菜单项文本色
                            activebackground=button_color,  # 悬停背景色
                            activeforeground="white",       # 悬停文本色
                            tearoff=0)       
        
        menu_bar.add_cascade(label="预制统计", menu=nomal_menu2)

        if 1==1:
            nomal_menu2.add_command(label="报告年份", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['报告年份'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="事件发生年份", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['事件发生年份'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_separator()
            
            nomal_menu2.add_command(label="-注册人备案人", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-产品类别", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['产品类别'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-产品名称", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[["产品类别",'-产品名称'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 


            nomal_menu2.add_command(label="-证号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-批号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-批号'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-型号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-型号'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-规格", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-规格'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 

            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="-性别", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-性别'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-年龄段", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['年龄段'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 


        if '---药品数据规整---' in datacols:
            nomal_menu2.add_command(label="-时隔", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['不良反应发生时间减用药开始时间'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="原患疾病（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),";",["原患疾病"]),[['原患疾病'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="用药原因（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),";",["用药原因"]),[['用药原因'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="不良反应（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),"；",["不良反应名称（规整）"]),[['不良反应名称（规整）'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            if '-表现归类(关键词法)'  in datacols:
                nomal_menu2.add_command(label="累及器官系统（关键词法）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),"；",["不良反应名称（规整）"]),[['-表现归类(关键词法)'], ['-伤害'], ['报告编码'], ['nunique'], {'不良反应名称（规整）': 'SMALL_count_mode','报告表编码': 'count'}, ['报告表编码']]),0,ori)) 
            
            nomal_menu2.add_separator()
        if '-监测机构'in datacols and '-单位名称' in datacols:
            nomal_menu2.add_command(label="-监测机构（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode', '-单位名称': 'SMALL_count_mode', '报告超时': sum, '新的加严重的': sum}, ['报告编码合计',["报告编码严重",'新的加严重的','报告超时'],['报告编码合计','报告编码合计','报告编码合计']]]),0,ori)) 
            nomal_menu2.add_command(label="-报告单位（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构', '-单位名称'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode',  '报告超时': sum, '新的加严重的': sum} ,['报告编码合计',["报告编码严重",'新的加严重的','报告超时'],['报告编码合计','报告编码合计','报告编码合计']]]),0,ori)) 
        if '关联性评价汇总'in datacols:
            nomal_menu2.add_command(label="关联性评价汇总", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['关联性评价汇总'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 

        if '药品分类'in datacols:
            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="药品分类", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.fillna('-未填写-'),[['药品分类'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="药品剂型", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.fillna('-未填写-'),[['剂型（历史匹配规整）'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="主要活性成分", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.fillna('-未填写-'),[['最主要的一个活性成分'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
   



          
    if '---器械数据规整---' in datacols or '---药品数据规整---' in datacols: 
        
        
        nomal_menu3 = tk.Menu(menu_bar, 
                            background=frame_color,  # 菜单项背景色
                            foreground=text_color,   # 菜单项文本色
                            activebackground=button_color,  # 悬停背景色
                            activeforeground="white",       # 悬停文本色
                            tearoff=0)                       # 禁止菜单项被拖出 
            

        menu_bar.add_cascade(label="风险监测", menu=nomal_menu3)
        if '器械故障表现' in datacols and ('报告编码' in datacols or '报告表编码'):
            nomal_menu3.add_separator()        
            nomal_menu3.add_command(label="风险预警", command=lambda:PSUR_keti_GUI(ori_owercount_easyread)) 

      
        if '---药品数据规整---' in datacols:
            nomal_menu3.add_separator()   
            nomal_menu3.add_command(label="新的不良反应（集合法）", command=lambda:PROGRAM_display_df_in_treeview(PSUR_get_new_GUI(ori_owercount_easyread),0,0))  
 
      
    stat_menu = tk.Menu(menu_bar, 
                        background=frame_color,  # 菜单项背景色
                        foreground=text_color,   # 菜单项文本色
                        activebackground=button_color,  # 悬停背景色
                        activeforeground="white",       # 悬停文本色
                        tearoff=0)                       # 禁止菜单项被拖出
 
    menu_bar.add_cascade(label="统计工具", menu=stat_menu) 
    stat_menu.add_command(label="ROR和PRR计算(df)", command=lambda:TOOLS_ROR_from_df_with_gui(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="ROR和PRR计算(db)", command=lambda:TOOLS_ROR_from_DB_GUI(ori_owercount_easyread))  

    stat_menu.add_separator()
    stat_menu.add_command(label="批量透视(df)", command=lambda:TOOLS_stat_all_gui(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="趋势分析(df)", command=lambda:TOOLS_trend_analysis_GUI(ori_owercount_easyread.copy()))        
    stat_menu.add_separator()   
    stat_menu.add_command(label="描述性统计(依据原始数据)", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread.describe(include='all').T,0,0))     
    stat_menu.add_command(label="秩和检验(依据统计表)", command=lambda:TOOLS_rank_sum_test(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="卡方检验(依据统计表)", command=lambda:TOOLS_drug_reaction_CH2_create_gui(ori_owercount_easyread.copy()))
    stat_menu.add_separator()   
    #stat_menu.add_command(label="倾向性评分(依据编码的数据)", command=lambda:RWS_PropensityScoreAnalysisGUI(ori_owercount_easyread.copy()).run())  
    stat_menu.add_command(label="多因素回归(依据编码的数据)", command=lambda:RWS_MultiFactorRegressionGUI(ori_owercount_easyread.copy()).run())  


    ai_menu = tk.Menu(menu_bar, 
                        background=frame_color,  # 菜单项背景色
                        foreground=text_color,   # 菜单项文本色
                        activebackground=button_color,  # 悬停背景色
                        activeforeground="white",       # 悬停文本色
                        tearoff=0)                       # 禁止菜单项被拖出 
    menu_bar.add_cascade(label='智能辅助', menu=ai_menu) 
    ai_menu.add_command(label="中文分词(df)", command=lambda:LLM_fenci(ori_owercount_easyread.copy())) 
    ai_menu.add_command(label="大模型API设置", command=lambda:LLM_edit_xml(os.path.join(csdir,'tok.xml')))  
    ai_menu.add_separator() 
    ai_menu.add_command(label="大模型辅助规整（逐行模式）", command=lambda:LLM_Send_Row(ori_owercount_easyread.copy(),win))  
    ai_menu.add_command(label="大模型辅助分析（整表模式）", command=lambda:LLM_Send_Form(ori_owercount_easyread.copy(),win))  
    ai_menu.add_separator() 
    ai_menu.add_command(label="大模型辅助数据库操作", command=lambda:LLM_SQL(os.path.join(csdir,'tok.xml'))) 
    ai_menu.add_separator() 
    ai_menu.add_command(label="大模型辅助数据统计", command=lambda:LLM_DF_Analyzer(os.path.join(csdir,'tok.xml'), ori_owercount_easyread.copy())) 
    ai_menu.add_command(label="大模型辅助数据可视化", command=lambda:LLM_DF_PLT(os.path.join(csdir,'tok.xml'), ori_owercount_easyread.copy())) 
    ai_menu.add_separator() 
    if  '报告编码' in datacols and '器械故障表现' in datacols:
        ai_menu.add_command(label="大模型辅助自动化规整和统计可视化（器械）", command=lambda:LLM_qixie(ori_owercount_easyread.copy()))  



def SETTING_get_width():  
    column_widths = {  
        "评分说明": 800,  
        "该单位喜好上报的品种统计": 200,  
        "报告编码": 200,  
        "产品名称": 200,  
        "上报机构描述": 200,  
        "持有人处理描述": 200,  
        "该注册证编号/曾用注册证编号报告数量": 200,  
        "通用名称": 200,  
        "该批准文号报告数量": 200,  
        "上市许可持有人名称": 200,  
        "注册证编号/曾用注册证编号": 140,  
        "监测机构": 140,  
        "报告月份": 140,  
        "报告季度": 140,  
        "单位列表": 140,  
        "单位名称": 140,  
        "管理类别": 40,  
        "报告日期": 100,          
        "报告人": 50,         
        "报告表编码一般": 100,  
        "报告表编码严重": 100,      
        "报告表编码新的一般": 100,  
        "报告表编码新的严重": 100,  
        "报告表编码其他": 100,  
        "报告表编码合计": 100,  
        "报告编码一般": 100,  
        "报告编码严重": 100,      
        "报告编码新的一般": 100,  
        "报告编码新的严重": 100,  
        "报告编码其他": 100,  
        "报告编码合计": 100,  
        "一般": 40,  
        "严重": 40,      
        "新的一般": 55,  
        "新的严重": 55,    
        "合计": 50,          
        "信息":1000,                          
        "#0":40,
        "发生地":40,      
        "联系人":40, 
        "产品类别":40,
        "产品批号":60,      
        "型号":60, 
        "规格":60,         
        "-批号":60,      
        "-型号":60, 
        "-规格":60,         
        "曾用注册证编号上报":40, 
        "联系电话":40  
        
    }  
    return column_widths
    
def SETTING_ReportType(s_dict,methon_treeview,ori):
    #先处理报表类型。
    
    report_type=eval(str(s_dict["报表类型"]))


    #等于的情况
    if "grouped" in s_dict["报表类型"]:
        data_s=ori.copy()
        #print(data_s)
        for i in report_type["grouped"]:
            mask=data_s[i].astype(str)==str(s_dict[i])
            data_s =data_s[mask].copy() # bao括的
            #print('check:',i,data_s)

        PROGRAM_display_df_in_treeview(data_s,methon_treeview,data_s)
        return 0      
    
    #包含的情况    
    elif "group_sep" in s_dict["报表类型"]:
        data_s=ori.copy()        
        for i in report_type["group_sep"]:
            #print(i)
            #escaped_value = re.escape(str(s_dict[i]))
             
            mask=ori[i].str.contains(str(s_dict[i]),na=False)
            
            data_s =ori[mask].copy() # bao括的
        PROGRAM_display_df_in_treeview(data_s,methon_treeview,data_s)
        return 0
                   
    elif "fenci" in s_dict["报表类型"]:
        data_s=ori.copy()
        for i in report_type["fenci"]:
            mask=data_s[i].str.contains(s_dict['关键词'],na=False)   
            data_s =data_s[mask].copy() # bao括的
        PROGRAM_display_df_in_treeview(data_s,methon_treeview,data_s)
        return 0      

    elif "group_nan" in s_dict["报表类型"]:
        data_s=ori.copy()
        for i in report_type["group_nan"]:
            #print(data_s.columns)
            filtered_df = ori[data_s[i].isnull()]
        PROGRAM_display_df_in_treeview(filtered_df,methon_treeview,filtered_df)
        return 0

    elif "SQL" in s_dict["报表类型"]:
        results = [] 
        selfconn = sqlite3.connect(s_dict["数据库文件"])
        selfcursor = selfconn.cursor()
        
        # 选择所有列（或指定需要的列）
        select_columns = '*'  # 或者指定列名，例如：'"列1", "列2", "列3"'
        
        # 构建查询条件
        conditions = []
        params = []
        for col in report_type["SQL"]:
            if col in s_dict:  # 确保列名在 s_dict 中存在
                value = s_dict[col]
                conditions.append(f'"{col}" = ?')
                params.append(value)
        
        # 将条件用 AND 连接
        where_clause = ' AND '.join(conditions)
        
        # 构建完整的查询语句
        query = f'SELECT {select_columns} FROM table1'
        if where_clause:  # 如果有条件，添加 WHERE 子句
            query += f' WHERE {where_clause}'
        #print("Generated SQL Query:", query)  # 打印生成的 SQL 查询，便于调试
        #print("Query Parameters:", params)  # 打印查询参数，便于调试
        
        # 执行查询
        try:
            selfcursor.execute(query, params)
            # 获取所有符合条件的行
            rows = selfcursor.fetchall()
            #print("Query Results (Rows):", rows)  # 打印查询结果，便于调试
            
            # 获取列名
            column_names = [description[0] for description in selfcursor.description]
            #print("Column Names:", column_names)  # 打印列名，便于调试
            
            # 将结果存储到列表中
            for row in rows:
                results.append(dict(zip(column_names, row)))  # 将每行数据转换为字典并存储
            
            # 将结果列表转换为 DataFrame
            result_df = pd.DataFrame(results)
            #print("Result DataFrame:")  # 打印 DataFrame，便于调试
            #print(result_df)
            selfcursor.close()
            selfconn.close()
            # 显示结果
            PROGRAM_display_df_in_treeview(result_df, 0, 0)
        except sqlite3.Error as e:
            print("SQLite Error:", e)  # 捕获并打印 SQLite 错误

        
        return 0

def SETTING_improve_ui_style(root):
    """
    对指定的Tk根窗口应用现代扁平化主题，采用蓝色和灰色配色方案
    仅影响当前传入的窗口，不影响其他窗口
    
    参数:
        root: 要应用样式的Tkinter根窗口对象
    """
    from tkinter import ttk
    import tkinter.font as tkFont
    
    # 为当前窗口创建专属样式对象
    style = ttk.Style(root)  # 关键修改：将样式绑定到特定窗口
    
    # 设置当前窗口的字体
    default_font = tkFont.Font(root, name="TkDefaultFont", exists=True) or \
                  tkFont.Font(root, name="TkDefaultFont", family="Microsoft YaHei", size=9)
    default_font.configure(family="Microsoft YaHei", size=9)
    
    text_font = tkFont.Font(root, name="TkTextFont", exists=True) or \
               tkFont.Font(root, name="TkTextFont", family="Microsoft YaHei", size=9)
    text_font.configure(family="Microsoft YaHei", size=9)
    
    fixed_font = tkFont.Font(root, name="TkFixedFont", exists=True) or \
                tkFont.Font(root, name="TkFixedFont", family="Consolas", size=9)
    fixed_font.configure(family="Consolas", size=9)
    
    # 使用clam主题作为基础（仅影响当前窗口）
    style.theme_use('clam')
    
    # 定义配色方案
    bg_color = "#f5f5f5"
    frame_color = "#e0e0e0"
    button_normal = "#4a7a8c"
    button_hover = "#5a8a9c"
    button_pressed = "#3a6a7c"
    button_disabled = "#a0a0a0"
    text_color = "#333333"
    highlight_color = "#e1f5fe"
    white = "#FFFFFF"
    
    # 仅配置当前窗口的背景
    root.configure(bg=bg_color)
    
    # 为当前窗口创建专属样式名称
    style_name = f"{str(root)}_style"
    
    # 配置样式（使用专属样式名称）
    style.configure(f"{style_name}.TFrame", background=bg_color)
    style.configure(f"{style_name}.TLabel", 
                  background=bg_color,
                  foreground=text_color,
                  padding=2)
    style.configure(f"{style_name}.TButton",
                  background=button_normal,
                  foreground=white,
                  font=("Microsoft YaHei", 9),
                  padding=6,
                  relief="flat",
                  borderwidth=0,
                  focuscolor='',
                  focusthickness=0)
    
    style.map(f"{style_name}.TButton",
            background=[('active', button_hover),
                       ('pressed', button_pressed),
                       ('disabled', button_disabled)],
            foreground=[('active', white),
                       ('pressed', white),
                       ('disabled', '#e0e0e0')])
    
    style.configure(f"{style_name}.TEntry",
                  fieldbackground=white,
                  foreground=text_color,
                  insertcolor=text_color,
                  padding=5,
                  relief="flat")
    
    style.map(f"{style_name}.TEntry",
             bordercolor=[('focus', button_normal),
                         ('hover', button_hover),
                         ('!focus', frame_color)])
    
    style.configure(f"{style_name}.TCombobox",
                  fieldbackground=white,
                  foreground=text_color,
                  arrowsize=12,
                  padding=5,
                  relief="flat")
    
    style.map(f"{style_name}.TCombobox",
             bordercolor=[('focus', button_normal),
                         ('hover', button_hover),
                         ('!focus', frame_color)])
    
    style.configure(f"{style_name}.Treeview",
                  background=white,
                  foreground=text_color,
                  rowheight=25,
                  fieldbackground=white,
                  borderwidth=0,
                  padding=5)
    
    style.configure(f"{style_name}.Treeview.Heading",
                  background=frame_color,
                  foreground="#444444",
                  padding=5,
                  font=("Microsoft YaHei", 9, "bold"),
                  relief="flat")
    
    style.map(f"{style_name}.Treeview.Heading",
             background=[('active', button_hover)],
             foreground=[('active', white)])
    
    style.map(f"{style_name}.Treeview",
             background=[('selected', highlight_color)],
             foreground=[('selected', text_color)])
    
    style.configure(f"{style_name}.Horizontal.TScrollbar",
                  background=frame_color,
                  troughcolor=bg_color,
                  arrowsize=12,
                  gripcount=0,
                  width=12,
                  padding=0,
                  relief="flat")
    
    style.configure(f"{style_name}.Vertical.TScrollbar",
                  background=frame_color,
                  troughcolor=bg_color,
                  arrowsize=12,
                  gripcount=0,
                  width=12,
                  padding=0,
                  relief="flat")
    
    style.configure(f"{style_name}.TLabelFrame",
                  background=bg_color,
                  foreground="#444444",
                  borderwidth=2,
                  relief="groove")
    
    style.configure(f"{style_name}.TSeparator",
                  background=frame_color)
    
    # 返回修改后的样式对象以便进一步自定义
    return style


############################################################################################################################
#大模型函数
############################################################################################################################
class AAA_04_LLM:
    pass

 
def LLM_qixie(dfs):
    """
    医疗器械分析智能体
    """
    class MedicalDataNormalizer:
        def __init__(self):
            self.root = tk.Tk()
            bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
            self.root.title("医疗器械不良事件报告智能规整系统 V1.3")
            
            sw = self.root.winfo_screenwidth()    
            sh = self.root.winfo_screenheight()    
            ww = 800  # 窗口宽度    
            wh = 700  # 窗口高度    
            x = (sw - ww) // 2    
            y = (sh - wh) // 2    
            self.root.geometry(f"{ww}x{wh}+{x}+{y}")  
            
            self.token_file = os.path.join(csdir, 'tok.xml')
            self.queue = queue.Queue()
            self.setup_ui()
            self.load_available_models()
            try:
                self.df=dfs
                self.df['报告编码']=self.df['报告编码'].astype(str)
                self.log(f"✅ 成功读取数据。| 行数: {len(self.df)}")
            except Exception as e:
                self.log(f"❌ 文件读取失败: {str(e)}")
                raise
 
        def setup_ui(self):
            """初始化界面"""
            def display_df():
                PROGRAM_display_df_in_treeview(self.df,0,0) 
            main_frame = ttk.Frame(self.root, padding=10)
            main_frame.pack(fill=tk.BOTH, expand=True)
            
            def display_gz():
                if '故障表现(规整)' not in self.df.columns:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '器械故障表现', '事件发生日期')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '故障表现(规整)', '事件发生日期')

            def display_gzbw():
                if '故障部位(规整)' not in self.df.columns:
                    print('规整未完成。')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '故障部位(规整)', '事件发生日期')

            def display_sh():
                if '伤害表现(规整)' not in self.df.columns:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '伤害表现', '事件发生日期')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '伤害表现(规整)', '事件发生日期')

            # 参数设置
            param_frame = ttk.LabelFrame(main_frame, text="处理参数", padding=10)
            param_frame.pack(fill=tk.X, pady=5)
            
            # 第一行参数
            row1_frame = ttk.Frame(param_frame)
            row1_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row1_frame, text="大模型选择:").grid(row=0, column=0, sticky=tk.W)
            self.model_var = tk.StringVar()
            self.model_combobox = ttk.Combobox(row1_frame, textvariable=self.model_var, width=20, state="readonly")
            self.model_combobox.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            ttk.Label(row1_frame, text="每批处理量:").grid(row=0, column=2, sticky=tk.W)
            self.batch_size = ttk.Entry(row1_frame, width=10)
            self.batch_size.insert(0, "10")
            self.batch_size.grid(row=0, column=3, sticky=tk.W, padx=5)
            
            ttk.Label(row1_frame, text="请求间隔(秒):").grid(row=0, column=4, sticky=tk.W)
            self.wait_time = ttk.Entry(row1_frame, width=10)
            self.wait_time.insert(0, "3")
            self.wait_time.grid(row=0, column=5, sticky=tk.W, padx=5)
            
            # 第二行参数
            row2_frame = ttk.Frame(param_frame)
            row2_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row2_frame, text="术语集最大数量:").grid(row=0, column=0, sticky=tk.W)
            self.max_terms = ttk.Entry(row2_frame, width=10)
            self.max_terms.insert(0, "15")
            self.max_terms.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            ttk.Label(row2_frame, text="选择模式:").grid(row=0, column=2, sticky=tk.W)
            self.mode_var = tk.StringVar()
            self.mode_combobox = ttk.Combobox(row2_frame, textvariable=self.mode_var, 
                                            values=["单品类自建术语集","单品类IMDRF术语集","多品类IMDRF术语集"], width=20)
            self.mode_combobox.current(0)
            self.mode_combobox.grid(row=0, column=3, sticky=tk.W, padx=5)
            
            # 第三行参数 - 匹配模式
            row3_frame = ttk.Frame(param_frame)
            row3_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row3_frame, text="匹配模式:").grid(row=0, column=0, sticky=tk.W)
            self.match_mode_var = tk.StringVar()
            self.match_mode_combobox = ttk.Combobox(row3_frame, textvariable=self.match_mode_var, 
                                                  values=["一对一", "一对多"], width=10)
            self.match_mode_combobox.current(0)
            self.match_mode_combobox.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            # 日志区域
            log_frame = ttk.LabelFrame(main_frame, text="处理日志", padding=10)
            log_frame.pack(fill=tk.BOTH, expand=True)
            self.log_area = scrolledtext.ScrolledText(
                log_frame, height=20, width=100, 
                font=('Consolas', 9), wrap=tk.WORD
            )
            self.log_area.pack(fill=tk.BOTH, expand=True)
            
            # 进度条
            self.progress = ttk.Progressbar(
                main_frame, orient="horizontal", 
                mode="determinate"
            )
            self.progress.pack(fill=tk.X, pady=10)
            
            # 按钮区域 - 恢复所有功能按钮
            btn_frame = ttk.Frame(main_frame)
            btn_frame.pack(fill=tk.X, pady=5)
            
            # 第一行按钮
            btn_row1 = ttk.Frame(btn_frame)
            btn_row1.pack(fill=tk.X)
            
            ttk.Button(
                btn_row1, text="开始处理", 
                command=self.start_processing
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="查看数据", 
                command=display_df
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="故障分析", 
                command=display_gz
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="部位分析", 
                command=display_gzbw
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="伤害分析", 
                command=display_sh
            ).pack(side=tk.LEFT, padx=5)
            
            # 第二行按钮
            btn_row2 = ttk.Frame(btn_frame)
            btn_row2.pack(fill=tk.X, pady=5)
            
            ttk.Button(
                btn_row2, text="信号检测", 
                command=lambda:TOOLS_ROR_from_df_with_gui(self.df)
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row2, text="趋势分析", 
                command=lambda:TOOLS_trend_analysis_GUI(self.df)
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row2, text="退出", 
                command=self.root.quit
            ).pack(side=tk.RIGHT, padx=5)
          
          
          
        #def confirm_send_columns(self, default_columns=['产品名称', '使用过程', '器械故障表现', '伤害表现']):
        def confirm_send_columns(self, default_columns=['产品名称', '使用过程', '器械故障表现', '伤害表现']):
            """弹出窗口让用户选择要发送给API的列（报告编码为必选）"""
            class ColumnSelector(tk.Toplevel):
                def __init__(self, parent, all_columns, default_columns):
                    super().__init__(parent)
                    self.title("选择发送给API的列")
                    self.parent = parent
                    self.result = None
                    
                    
                    sw = self.winfo_screenwidth()    
                    sh = self.winfo_screenheight()    
                    ww = 500  # 窗口宽度    
                    wh = 800  # 窗口高度    
                    x = (sw - ww) // 2    
                    y = (sh - wh) // 2    
                    self.geometry(f"{ww}x{wh}+{x}+{y}")  
                    
                    # 窗口设置
                    #self.geometry("800x400")
                    #self.resizable(True, True)
                    
                    # 主框架
                    main_frame = ttk.Frame(self)
                    main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                    
                    # 说明标签
                    ttk.Label(main_frame, 
                             text="请选择要发送给API的列（报告编码将自动包含）").pack(pady=(0, 10))
                    
                    # 创建Treeview和滚动条
                    tree_frame = ttk.Frame(main_frame)
                    tree_frame.pack(fill=tk.BOTH, expand=True)
                    
                    self.tree = ttk.Treeview(tree_frame, columns=('select', 'name'), 
                                            show='headings', selectmode='none', height=15)
                    
                    # 设置列
                    self.tree.heading('select', text='选择')
                    self.tree.heading('name', text='列名')
                    self.tree.column('select', width=80, anchor='center')
                    self.tree.column('name', width=400, anchor='w')
                    
                    # 滚动条
                    vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
                    hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
                    self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
                    
                    # 布局
                    self.tree.grid(row=0, column=0, sticky='nsew')
                    vsb.grid(row=0, column=1, sticky='ns')
                    hsb.grid(row=1, column=0, sticky='ew')
                    
                    tree_frame.grid_rowconfigure(0, weight=1)
                    tree_frame.grid_columnconfigure(0, weight=1)
                    
                    # 存储选择状态
                    self.selection_vars = {}
                    
                    # 添加数据
                    for col in all_columns:
                        var = tk.BooleanVar(value=col in default_columns)
                        self.selection_vars[col] = var
                        self.tree.insert('', 'end', values=('☑' if var.get() else '☐', col))
                    
                    # 绑定点击事件
                    self.tree.bind('<Button-1>', self.on_click)
                    
                    # 按钮框架
                    btn_frame = ttk.Frame(main_frame)
                    btn_frame.pack(pady=(10, 0))
                    
                    ttk.Button(btn_frame, text="全选", command=self.select_all).pack(side=tk.LEFT, padx=5)
                    ttk.Button(btn_frame, text="全不选", command=self.select_none).pack(side=tk.LEFT, padx=5)
                    ttk.Button(btn_frame, text="反选", command=self.toggle_all).pack(side=tk.LEFT, padx=5)
                    ttk.Button(btn_frame, text="确认", command=self.on_confirm).pack(side=tk.RIGHT, padx=5)
                
                def on_click(self, event):
                    """处理复选框点击"""
                    region = self.tree.identify("region", event.x, event.y)
                    if region == "cell":
                        column = self.tree.identify_column(event.x)
                        item = self.tree.identify_row(event.y)
                        
                        if column == "#1":  # 只在选择列点击时切换
                            col_name = self.tree.item(item, 'values')[1]
                            current = self.selection_vars[col_name].get()
                            self.selection_vars[col_name].set(not current)
                            self.tree.item(item, values=('☑' if not current else '☐', col_name))
                
                def select_all(self):
                    """全选"""
                    for item in self.tree.get_children():
                        col_name = self.tree.item(item, 'values')[1]
                        self.selection_vars[col_name].set(True)
                        self.tree.item(item, values=('☑', col_name))
                
                def select_none(self):
                    """全不选"""
                    for item in self.tree.get_children():
                        col_name = self.tree.item(item, 'values')[1]
                        self.selection_vars[col_name].set(False)
                        self.tree.item(item, values=('☐', col_name))
                
                def toggle_all(self):
                    """反选"""
                    for item in self.tree.get_children():
                        col_name = self.tree.item(item, 'values')[1]
                        new_state = not self.selection_vars[col_name].get()
                        self.selection_vars[col_name].set(new_state)
                        self.tree.item(item, values=('☑' if new_state else '☐', col_name))
                
                def on_confirm(self):
                    """确认选择"""
                    self.result = [col for col, var in self.selection_vars.items() if var.get()]
                    self.destroy()
            
            # 获取可用列（排除报告编码）
            available_columns = [col for col in self.df.columns if col != '报告编码']
            
            # 创建并显示选择窗口
            selector = ColumnSelector(self.root, available_columns, default_columns)
            selector.transient(self.root)
            selector.grab_set()
            selector.wait_window()
            
            # 返回结果（确保包含报告编码）
            if selector.result is not None:
                return ['报告编码'] + selector.result
            return ['报告编码'] + default_columns  # 如果用户取消，返回默认值
            
            
            
            
        def load_available_models(self):
            """从tok.xml加载可用模型"""
            try:
                tree = ET.parse(self.token_file)
                root = tree.getroot()
                
                models = [model.get('name') for model in root.findall("model")]
                
                if not models:
                    raise ValueError("tok.xml中没有找到任何模型配置")
                    
                self.model_combobox['values'] = models
                self.model_var.set(models[0])
                self.log(f"✅ 加载可用模型: {', '.join(models)}")
                
            except Exception as e:
                self.log(f"❌ 加载模型列表失败: {str(e)}")
                self.model_combobox['values'] = ["默认模型"]
                self.model_var.set("默认模型")

        def log(self, message):
            """记录带时间戳的日志"""
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.queue.put(("log", f"[{timestamp}] {message}\n"))

        def start_processing(self):
            """启动处理流程"""
            self.progress['value'] = 0
            self.progress['maximum'] = len(self.df)
            threading.Thread(
                target=self.process_data_pipeline, 
                daemon=True
            ).start()
            self.process_queue()

        def process_data_pipeline(self):
            """主处理流程"""
            try:
                # 1. 一次性生成三个术语集
                self.log("\n🔧 开始生成所有术语集")
                fault_terms, harm_terms, part_terms = self.generate_all_terms()
                
                if not all([fault_terms, harm_terms, part_terms]):
                    raise ValueError("术语集生成不完整")
                
                # 2. 批量规整处理
                self.log("\n⚙️ 开始批量规整数据")
                self.batch_normalize_data(fault_terms, harm_terms, part_terms)
                
                # 3. 保存结果
                output_path = os.path.join(os.path.dirname(__file__), "规整后.xlsx")
                self.log(f"\n✅ 处理完成！")
                self.queue.put(("complete", output_path))
                
            except Exception as e:
                self.log(f"\n❌ 处理失败: {str(e)}")
                self.queue.put(("error", str(e)))

        def generate_all_terms(self):
            """一次性生成三个术语集并确认"""
            try:
                max_terms = int(self.max_terms.get()) if self.max_terms.get() else 15
                self.df.reset_index(inplace=True)
                fault_raw = self.df['器械故障表现'].value_counts().head(100)
                harm_raw = self.df['伤害表现'].value_counts().head(100)
                part_raw = self.df['器械故障表现'].value_counts().head(100)
                name_raw = self.df['产品名称'].value_counts().head(1)
                mode_select=self.mode_combobox.get()
                print('模式:',mode_select)
                if mode_select=='多品类IMDRF术语集':
                    name_x='多个品类的医疗器械'
                else:
                    name_x=self.df['产品名称'][0]
                imdfr_guzhang=["患者器械相互作用问题（生物相容性、排异反应、形状尺寸不当、过敏、患者感染等）",
                "包装或运输问题（包装损坏、运输损坏等）",
                "化学问题（异味、结块、沉淀、PH异常、显色异常、凝血、溶血等）",
                "材料完整性问题（破裂、断裂、粘合失效、材料变形、材料穿孔等）",
                "机械问题（堵塞、非材料问题产生的渗漏、设备机械故障、机械异响等）",
                "光学问题（散焦、可见光透视异常等）",
                "电气或电子特性问题（电池、充电、电源、放电等）",
                "准确度、精确度或校准问题（读数不准确、过高过低、过快过慢、精确度不足、假阳性、假阴性、校准失效、矫正不足等）",
                "输出问题（显示信息有错误、图像有问题、能量/辐射输出问题等）",
                "温度问题（过热、过冷等）",
                "计算机软件问题（计算机系统、程序、网络、数据丢失、时间不准等）",
                "连接问题（连接松动、连接断开、松动脱落、连接不充分、连接受阻等）",
                "通信或传输问题（与器械读取用于解释或测量的信号故障相关的问题）",
                "输注或流量问题（与器械未能按预期输送或抽吸液体或气体相关的问题）",
                "激活、定位或分离问题（激活失败、器械错位、难以进入、难以分离或过早分离等）",
                "保护措施问题（无报警、误报警、低声音报警、延迟报警、保护措施失效、自检失败等）",
                "相容性问题（组件、附件、系统间不兼容）",
                "污染或去污问题（与器械、其表面或包装材料中发现存在任何非预期异物相关的问题、微生物污染等）",
                "环境相容性问题（环境噪声、环境温湿度、供水、供电等）",
                "安装相关问题（组装错误、配置错误等）",
                "标签、使用说明书或培训问题（与器械标标签、使用说明书、培训和维护文件或指南相关的问题）",
                "器械使用问题（与未能按照制造商的建议或公认的最佳实践处理、维修或操作器械相关的问题）",
                '设计缺陷问题（器械设计缺陷）',
                "其他（不属于以上任何归类的）"]
                prompt2='且故障表现术语严格在以下不良事件术语集选取（不能改变任何文字和符号，也要保留括号和括号里的内容）：'+str(imdfr_guzhang)
                if mode_select=='单品类自建术语集':
                    prompt2='' 
                # 构造组合提示词，明确指定最大数量
                prompt = (
                    f"这是一份{name_x}所发生的不良事件的情况，请结合该产品的机构特点，请分别从以下三类描述中归纳标准术语（每类不超过{max_terms}个，并且最后一个为包括编号的'其他'：\n"
                    "=== 故障表现 ===\n" + fault_raw.to_string() + "\n\n"
                    "=== 伤害表现 ===\n" + harm_raw.to_string() + "\n\n"
                    "=== 故障部位 ===\n" + part_raw.to_string() + "\n\n"
                    "格式要求：\n"
                    f"1. 故障表现术语（最多{max_terms}个，要合并类似但不同表述的故障表现，每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次{prompt2}）：\n"
                    "  1.1 术语1\n"
                    "  1.2 术语2\n"
                    "  ...\n"
                    f"2. 伤害表现术语（最多{max_terms}个，要合并类似但不同表述的伤害表现（注意区分故障表现），每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次）：\n"
                    "  2.1 术语1\n"
                    "  ...\n"
                    f"3. 故障部位术语（最多{max_terms}个，要合并类似但不同表述的故障部位，注意列出的是部位或部件，每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次）：\n"
                    "  3.1 术语1\n"
                    "  ..."               
                )
                
                self.log("\n发送术语集生成请求...")
                response = self.call_api(
                    prompt,
                    f"你是一个医疗设备和医疗耗材的管理专家，请从三类描述中分别提取不超过{max_terms}个标准术语"
                )
                self.log(f"大模型返回:\n{response}")
                
                # 解析三类术语，并确保不超过max_terms
                fault_terms = self.extract_terms(response, "1. 故障表现术语")[:max_terms]
                harm_terms = self.extract_terms(response, "2. 伤害表现术语")[:max_terms]
                part_terms = self.extract_terms(response, "3. 故障部位术语")[:max_terms]
                
                # 用户确认
                confirmed_terms = self.confirm_all_terms(
                    fault_terms, harm_terms, part_terms
                )
                
                # 再次确保不超过max_terms
                confirmed_terms = (
                    confirmed_terms[0][:max_terms],
                    confirmed_terms[1][:max_terms],
                    confirmed_terms[2][:max_terms]
                )
                
                if not all(confirmed_terms):
                    raise ValueError("用户取消了术语确认")
                    
                return confirmed_terms
                
            except Exception as e:
                self.log(f"术语集生成失败: {str(e)}")
                raise

        def extract_terms(self, response, section_title):
            """从响应中提取特定类别的术语"""
            terms = []
            in_section = False
            max_terms = int(self.max_terms.get()) if self.max_terms.get() else 15
            
            for line in response.split('\n'):
                line = line.strip()
                if line.startswith(section_title):
                    in_section = True
                    continue
                elif in_section and line and line[0].isdigit():
                    term = line.split('. ')[1] if '. ' in line else line
                    terms.append(term.strip())
                elif in_section and not line:
                    break  # 遇到空行结束当前section
                    
            return terms

        def confirm_all_terms(self, fault_terms, harm_terms, part_terms):
            """三栏式术语确认窗口"""
            confirmed = {}
            
            def on_confirm():
                nonlocal confirmed
                confirmed['fault'] = fault_text.get("1.0", tk.END).split('\n')
                confirmed['harm'] = harm_text.get("1.0", tk.END).split('\n')
                confirmed['part'] = part_text.get("1.0", tk.END).split('\n')
                confirm_win.destroy()
            
            confirm_win = tk.Toplevel(self.root)
            bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(confirm_win)

            confirm_win.title("术语集确认")
            
            sw = confirm_win.winfo_screenwidth()    
            sh =confirm_win.winfo_screenheight()    
            ww = 1200  # 窗口宽度    
            wh = 600  # 窗口高度    
            x = (sw - ww) // 2    
            y = (sh - wh) // 2    
            confirm_win.geometry(f"{ww}x{wh}+{x}+{y}")  
            
            # 三栏布局
            frame = ttk.Frame(confirm_win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            # 故障表现术语
            ttk.Label(frame, text="故障表现术语").grid(row=0, column=0)
            fault_text = scrolledtext.ScrolledText(frame, width=80, height=35)
            fault_text.insert(tk.END, '\n'.join(fault_terms))
            fault_text.grid(row=1, column=0, padx=5, pady=5)
            
            # 伤害表现术语
            ttk.Label(frame, text="伤害表现术语").grid(row=0, column=1)
            harm_text = scrolledtext.ScrolledText(frame, width=30, height=35)
            harm_text.insert(tk.END, '\n'.join(harm_terms))
            harm_text.grid(row=1, column=1, padx=5, pady=5)
            
            # 故障部位术语
            ttk.Label(frame, text="故障部位术语").grid(row=0, column=2)
            part_text = scrolledtext.ScrolledText(frame, width=30, height=35)
            part_text.insert(tk.END, '\n'.join(part_terms))
            part_text.grid(row=1, column=2, padx=5, pady=5)
            
            # 确认按钮
            ttk.Button(
                confirm_win, text="确认所有术语", 
                command=on_confirm
            ).pack(pady=10)
            
            # 等待窗口关闭
            confirm_win.transient(self.root)
            confirm_win.grab_set()
            confirm_win.wait_window(confirm_win)
            
            # 清理空行和空值
            return (
                [t.strip() for t in confirmed.get('fault', []) if t.strip()],
                [t.strip() for t in confirmed.get('harm', []) if t.strip()],
                [t.strip() for t in confirmed.get('part', []) if t.strip()]
            )

        def batch_normalize_data(self, fault_terms, harm_terms, part_terms):
            """批量规整三类数据（支持一对一和一对多模式）"""
            try:
                batch_size = int(self.batch_size.get())
                wait_time = int(self.wait_time.get())
                is_one_to_many = self.match_mode_combobox.get() == "一对多"
                
                # 让用户选择要发送给API的列
                send_columns = self.confirm_send_columns()
                self.log(f"将发送以下列给API: {', '.join(send_columns)}")
                
                for i in range(0, len(self.df), batch_size):
                    batch = self.df.iloc[i:i+batch_size]
                    prompt_lines = []
                    
                    # 为每行构建描述
                    for _, row in batch.iterrows():
                        # 构建每行的数据字典
                        row_data = {col: str(row[col]) for col in send_columns}
                        
                        # 构建基本描述
                        base_desc = f"这是一份{row_data.get('产品名称', '')}所发生的不良事件的报告"
                        
                        # 收集额外参考信息
                        additional_info = []
                        for col, value in row_data.items():
                            if col not in ['报告编码', '产品名称'] and value:
                                additional_info.append(f"{col}: '{value}'")
                        
                        # 组合完整描述
                        if additional_info:
                            base_desc += f"，其中可额外参考的信息为 {'; '.join(additional_info)}"
                        
                        prompt_lines.append(f"{base_desc}。编码: {row_data['报告编码']}")

                    # 构造提示词
                    prompt = (
                        "请严格按以下要求分类（必须返回JSON格式）：\n"
                        "=== 标准术语 ===\n"
                        "1. 故障表现:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(fault_terms)) + "\n"
                        "2. 伤害表现:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(harm_terms)) + "\n"
                        "3. 故障部位:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(part_terms)) + "\n\n"
                        "=== 待分类数据 ===\n" +
                        '\n'.join(prompt_lines) + "\n\n"
                        "返回格式要求：\n"
                        "```json\n"
                        "{\n"
                        '  "结果": [\n'
                    )
                    
                    if is_one_to_many:
                        prompt += (
                            '    {"编码": "1253705092023000050", "故障": [1,2], "伤害": [2,3], "部位": [3]},\n'
                            '    {"编码": "1253718172023000029", "故障": [4], "伤害": [5], "部位": [6]}\n'
                            "  ]\n"
                            "}\n"
                            "```\n"
                            "注意：\n"
                            "1. 必须使用双引号\n"
                            "2. 编码必须与输入完全一致\n"
                            "3. 数字对应术语编号\n"
                            "4. 可以返回多个编号，用数组表示,不允许返回空数组"
                        )
                    else:
                        prompt += (
                            '    {"编码": "1253705092023000050", "故障": 1, "伤害": 2, "部位": 3},\n'
                            '    {"编码": "1253718172023000029", "故障": 4, "伤害": 5, "部位": 6}\n'
                            "  ]\n"
                            "}\n"
                            "```\n"
                            "注意：\n"
                            "1. 必须使用双引号\n"
                            "2. 编码必须与输入完全一致\n"
                            "3. 数字对应术语编号\n"
                            "4. 每个字段只返回一个整数"
                        )
                    
                    self.log(f"\n处理行 {i+1}-{i+len(batch)} | 发送请求...")
                    print(prompt)
                    response = self.call_api(
                        prompt,
                        "你必须是严格的JSON格式生成器，只返回指定格式的JSON数据"
                    )
                    self.log(f"API返回原始结果:\n{response}")
                    
                    # 解析JSON
                    try:
                        import json
                        response = response.strip()
                        if response.startswith('```json'):
                            response = response[7:].strip()
                        if response.endswith('```'):
                            response = response[:-3].strip()
                        
                        result_data = json.loads(response)
                        code_mapping = {}
                        
                        for item in result_data.get('结果', []):
                            code = item['编码']
                            
                            # 处理故障表现
                            if is_one_to_many:
                                fault_indices = [min(int(f)-1, len(fault_terms)-1) for f in item.get('故障', [])]
                                harm_indices = [min(int(h)-1, len(harm_terms)-1) for h in item.get('伤害', [])]
                                part_indices = [min(int(p)-1, len(part_terms)-1) for p in item.get('部位', [])]
                            else:
                                fault_indices = [min(int(item.get('故障', 1))-1, len(fault_terms)-1)]
                                harm_indices = [min(int(item.get('伤害', 1))-1, len(harm_terms)-1)]
                                part_indices = [min(int(item.get('部位', 1))-1, len(part_terms)-1)]
                            
                            code_mapping[code] = (fault_indices, harm_indices, part_indices)
                            
                    except Exception as e:
                        self.log(f"❌ JSON解析失败: {str(e)}")
                        self.log(f"原始响应内容: {response}")
                        code_mapping = {}
                    
                    # 填充结果
                    for _, row in batch.iterrows():
                        code = str(row['报告编码']).strip()
                        if code in code_mapping:
                            fault_indices, harm_indices, part_indices = code_mapping[code]
                            
                            # 处理故障表现
                            if fault_indices:
                                self.df.at[row.name, '故障表现(规整)'] = ";".join([fault_terms[i] for i in fault_indices if 0 <= i < len(fault_terms)])
                            else:
                                self.df.at[row.name, '故障表现(规整)'] = fault_terms[-1]  # 其他
                                
                            # 处理伤害表现
                            if harm_indices:
                                self.df.at[row.name, '伤害表现(规整)'] = ";".join([harm_terms[i] for i in harm_indices if 0 <= i < len(harm_terms)])
                            else:
                                self.df.at[row.name, '伤害表现(规整)'] = harm_terms[-1]  # 其他
                                
                            # 处理故障部位
                            if part_indices:
                                self.df.at[row.name, '故障部位(规整)'] = ";".join([part_terms[i] for i in part_indices if 0 <= i < len(part_terms)])
                            else:
                                self.df.at[row.name, '故障部位(规整)'] = part_terms[-1]  # 其他
                        else:
                            self.log(f"⚠️ 未找到编码 {code} 的分类结果")
                            self.df.at[row.name, '故障表现(规整)'] = fault_terms[-1]
                            self.df.at[row.name, '伤害表现(规整)'] = harm_terms[-1]
                            self.df.at[row.name, '故障部位(规整)'] = part_terms[-1]
                    
                    try:
                        self.df = self.df.drop(columns=['level_0', 'index'])
                    except:
                        pass
                    
                    self.queue.put(("progress", len(batch)))
                    time.sleep(wait_time)
                    
            except Exception as e:
                self.log(f"批量规整失败: {str(e)}")
                raise

        def call_api(self, prompt, system_prompt=None):
            """API调用"""
            try:
                tree = ET.parse(self.token_file)
                root = tree.getroot()
                selected_model = self.model_var.get()
                
                # 修改为直接从根节点查找model节点
                model = None
                for m in root.findall("model"):
                    if m.get('name') == selected_model:
                        model = m
                        break
                        
                if not model:
                    raise ValueError(f"模型配置未找到: {selected_model}")
                    
                url = model.find("url").text
                api_key = model.find("api_key").text
                model_name = model.find("model").text           
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                
                messages = [{"role": "user", "content": prompt}]
                if system_prompt:
                    messages.insert(0, {"role": "system", "content": system_prompt})
                    
                data = {
                    "model": model.find("model").text,
                    "messages": messages,
                    "temperature": 0.3,
                    "max_tokens": 8000,
                }
                response = requests.post(
                    url, headers=headers, json=data, 
                    timeout=30
                )
                response.raise_for_status()
                
                # 更健壮的响应处理
                response_json = response.json()
                if 'choices' not in response_json or len(response_json['choices']) == 0:
                    raise ValueError("API返回无效响应: 无choices字段")
                    
                content = response_json['choices'][0]['message']['content']
                
                # 清理响应内容 - 移除可能的Markdown代码块标记
                if content.startswith('```json') and content.endswith('```'):
                    content = content[7:-3].strip()
                elif content.startswith('```') and content.endswith('```'):
                    content = content[3:-3].strip()
                    
                return content
                
            except Exception as e:
                raise Exception(f"API调用失败: {str(e)}")

        def process_queue(self):
            """处理消息队列"""
            try:
                while True:
                    try:
                        msg_type, *msg_data = self.queue.get_nowait()
                        if msg_type == "log":
                            self.log_area.insert(tk.END, msg_data[0])
                            self.log_area.see(tk.END)
                        elif msg_type == "progress":
                            self.progress.step(msg_data[0])
                        elif msg_type == "complete":
                            print("完成")
                            DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号' ,'故障表现(规整)','事件发生日期')
                        elif msg_type == "error":
                            messagebox.showerror("错误", msg_data[0])
                    except queue.Empty:
                        break
            finally:
                self.root.after(100, self.process_queue)

    try:
        app = MedicalDataNormalizer()
        app.root.mainloop()
    except Exception as e:
        print(f"程序初始化失败: {str(e)}")

        
def LLM_qixie_old(dfs):
    """
    医疗器械分析智能体
    """
    class MedicalDataNormalizer:
        def __init__(self):
            self.root = tk.Tk()
            bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
            self.root.title("医疗器械不良事件报告智能规整系统 V1.2")
            
            sw = self.root.winfo_screenwidth()    
            sh = self.root.winfo_screenheight()    
            ww = 800  # 窗口宽度    
            wh = 700  # 窗口高度    
            x = (sw - ww) // 2    
            y = (sh - wh) // 2    
            self.root.geometry(f"{ww}x{wh}+{x}+{y}")  
            
            self.token_file = os.path.join(csdir, 'tok.xml')
            self.queue = queue.Queue()
            self.setup_ui()
            self.load_available_models()
            try:
                self.df=dfs
                self.df['报告编码']=self.df['报告编码'].astype(str)
                self.log(f"✅ 成功读取数据。| 行数: {len(self.df)}")
            except Exception as e:
                self.log(f"❌ 文件读取失败: {str(e)}")
                raise
 
        def setup_ui(self):
            """初始化界面"""
            def display_df():
                PROGRAM_display_df_in_treeview(self.df,0,0) 
            main_frame = ttk.Frame(self.root, padding=10)
            main_frame.pack(fill=tk.BOTH, expand=True)
            
            def display_gz():
                if '故障表现(规整)' not in self.df.columns:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '器械故障表现', '事件发生日期')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '故障表现(规整)', '事件发生日期')

            def display_gzbw():
                if '故障部位(规整)' not in self.df.columns:
                    print('规整未完成。')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '故障部位(规整)', '事件发生日期')

            def display_sh():
                if '伤害表现(规整)' not in self.df.columns:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '伤害表现', '事件发生日期')
                else:
                    DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号', '伤害表现(规整)', '事件发生日期')

            # 参数设置
            param_frame = ttk.LabelFrame(main_frame, text="处理参数", padding=10)
            param_frame.pack(fill=tk.X, pady=5)
            
            # 第一行参数
            row1_frame = ttk.Frame(param_frame)
            row1_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row1_frame, text="大模型选择:").grid(row=0, column=0, sticky=tk.W)
            self.model_var = tk.StringVar()
            self.model_combobox = ttk.Combobox(row1_frame, textvariable=self.model_var, width=20, state="readonly")
            self.model_combobox.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            ttk.Label(row1_frame, text="每批处理量:").grid(row=0, column=2, sticky=tk.W)
            self.batch_size = ttk.Entry(row1_frame, width=10)
            self.batch_size.insert(0, "10")
            self.batch_size.grid(row=0, column=3, sticky=tk.W, padx=5)
            
            ttk.Label(row1_frame, text="请求间隔(秒):").grid(row=0, column=4, sticky=tk.W)
            self.wait_time = ttk.Entry(row1_frame, width=10)
            self.wait_time.insert(0, "3")
            self.wait_time.grid(row=0, column=5, sticky=tk.W, padx=5)
            
            # 第二行参数
            row2_frame = ttk.Frame(param_frame)
            row2_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row2_frame, text="术语集最大数量:").grid(row=0, column=0, sticky=tk.W)
            self.max_terms = ttk.Entry(row2_frame, width=10)
            self.max_terms.insert(0, "15")
            self.max_terms.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            ttk.Label(row2_frame, text="选择模式:").grid(row=0, column=2, sticky=tk.W)
            self.mode_var = tk.StringVar()
            self.mode_combobox = ttk.Combobox(row2_frame, textvariable=self.mode_var, 
                                            values=["单品类自建术语集","单品类IMDRF术语集","多品类IMDRF术语集"], width=20)
            self.mode_combobox.current(0)
            self.mode_combobox.grid(row=0, column=3, sticky=tk.W, padx=5)
            
            # 第三行参数 - 匹配模式
            row3_frame = ttk.Frame(param_frame)
            row3_frame.pack(fill=tk.X, pady=5)
            
            ttk.Label(row3_frame, text="匹配模式:").grid(row=0, column=0, sticky=tk.W)
            self.match_mode_var = tk.StringVar()
            self.match_mode_combobox = ttk.Combobox(row3_frame, textvariable=self.match_mode_var, 
                                                  values=["一对一", "一对多"], width=10)
            self.match_mode_combobox.current(0)
            self.match_mode_combobox.grid(row=0, column=1, sticky=tk.W, padx=5)
            
            # 日志区域
            log_frame = ttk.LabelFrame(main_frame, text="处理日志", padding=10)
            log_frame.pack(fill=tk.BOTH, expand=True)
            self.log_area = scrolledtext.ScrolledText(
                log_frame, height=20, width=100, 
                font=('Consolas', 9), wrap=tk.WORD
            )
            self.log_area.pack(fill=tk.BOTH, expand=True)
            
            # 进度条
            self.progress = ttk.Progressbar(
                main_frame, orient="horizontal", 
                mode="determinate"
            )
            self.progress.pack(fill=tk.X, pady=10)
            
            # 按钮区域 - 恢复所有功能按钮
            btn_frame = ttk.Frame(main_frame)
            btn_frame.pack(fill=tk.X, pady=5)
            
            # 第一行按钮
            btn_row1 = ttk.Frame(btn_frame)
            btn_row1.pack(fill=tk.X)
            
            ttk.Button(
                btn_row1, text="开始处理", 
                command=self.start_processing
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="查看数据", 
                command=display_df
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="故障分析", 
                command=display_gz
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="部位分析", 
                command=display_gzbw
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row1, text="伤害分析", 
                command=display_sh
            ).pack(side=tk.LEFT, padx=5)
            
            # 第二行按钮
            btn_row2 = ttk.Frame(btn_frame)
            btn_row2.pack(fill=tk.X, pady=5)
            
            ttk.Button(
                btn_row2, text="信号检测", 
                command=lambda:TOOLS_ROR_from_df_with_gui(self.df)
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row2, text="趋势分析", 
                command=lambda:TOOLS_trend_analysis_GUI(self.df)
            ).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(
                btn_row2, text="退出", 
                command=self.root.quit
            ).pack(side=tk.RIGHT, padx=5)
          
        def load_available_models(self):
            """从tok.xml加载可用模型"""
            try:
                tree = ET.parse(self.token_file)
                root = tree.getroot()
                
                models = [model.get('name') for model in root.findall("model")]
                
                if not models:
                    raise ValueError("tok.xml中没有找到任何模型配置")
                    
                self.model_combobox['values'] = models
                self.model_var.set(models[0])
                self.log(f"✅ 加载可用模型: {', '.join(models)}")
                
            except Exception as e:
                self.log(f"❌ 加载模型列表失败: {str(e)}")
                self.model_combobox['values'] = ["默认模型"]
                self.model_var.set("默认模型")

        def log(self, message):
            """记录带时间戳的日志"""
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.queue.put(("log", f"[{timestamp}] {message}\n"))

        def start_processing(self):
            """启动处理流程"""
            self.progress['value'] = 0
            self.progress['maximum'] = len(self.df)
            threading.Thread(
                target=self.process_data_pipeline, 
                daemon=True
            ).start()
            self.process_queue()

        def process_data_pipeline(self):
            """主处理流程"""
            try:
                # 1. 一次性生成三个术语集
                self.log("\n🔧 开始生成所有术语集")
                fault_terms, harm_terms, part_terms = self.generate_all_terms()
                
                if not all([fault_terms, harm_terms, part_terms]):
                    raise ValueError("术语集生成不完整")
                
                # 2. 批量规整处理
                self.log("\n⚙️ 开始批量规整数据")
                self.batch_normalize_data(fault_terms, harm_terms, part_terms)
                
                # 3. 保存结果
                output_path = os.path.join(os.path.dirname(__file__), "规整后.xlsx")
                self.log(f"\n✅ 处理完成！")
                self.queue.put(("complete", output_path))
                
            except Exception as e:
                self.log(f"\n❌ 处理失败: {str(e)}")
                self.queue.put(("error", str(e)))

        def generate_all_terms(self):
            """一次性生成三个术语集并确认"""
            try:
                max_terms = int(self.max_terms.get()) if self.max_terms.get() else 15
                self.df.reset_index(inplace=True)
                # 生成原始术语集，限制为max_terms的2倍以便大模型筛选
                fault_raw = self.df['器械故障表现'].value_counts().head(max_terms*2)
                harm_raw = self.df['伤害表现'].value_counts().head(max_terms*2)
                part_raw = self.df['器械故障表现'].value_counts().head(max_terms*2)
                name_raw = self.df['产品名称'].value_counts().head(1)
                mode_select=self.mode_combobox.get()
                print('模式:',mode_select)
                if mode_select=='多品类IMDRF术语集':
                    name_x='多个品类的医疗器械'
                else:
                    name_x=self.df['产品名称'][0]
                imdfr_guzhang=["患者器械相互作用问题（生物相容性、排异反应、形状尺寸不当、过敏、患者感染等）",
                "包装或运输问题（包装损坏、运输损坏等）",
                "化学问题（异味、结块、沉淀、PH异常、显色异常、凝血、溶血等）",
                "材料完整性问题（破裂、断裂、粘合失效、材料变形、材料穿孔等）",
                "机械问题（堵塞、非材料问题产生的渗漏、设备机械故障、机械异响等）",
                "光学问题（散焦、可见光透视异常等）",
                "电气或电子特性问题（电池、充电、电源、放电等）",
                "准确度、精确度或校准问题（读数不准确、过高过低、过快过慢、精确度不足、假阳性、假阴性、校准失效、矫正不足等）",
                "输出问题（显示信息有错误、图像有问题、能量/辐射输出问题等）",
                "温度问题（过热、过冷等）",
                "计算机软件问题（计算机系统、程序、网络、数据丢失、时间不准等）",
                "连接问题（连接松动、连接断开、松动脱落、连接不充分、连接受阻等）",
                "通信或传输问题（与器械读取用于解释或测量的信号故障相关的问题）",
                "输注或流量问题（与器械未能按预期输送或抽吸液体或气体相关的问题）",
                "激活、定位或分离问题（激活失败、器械错位、难以进入、难以分离或过早分离等）",
                "保护措施问题（无报警、误报警、低声音报警、延迟报警、保护措施失效、自检失败等）",
                "相容性问题（组件、附件、系统间不兼容）",
                "污染或去污问题（与器械、其表面或包装材料中发现存在任何非预期异物相关的问题、微生物污染等）",
                "环境相容性问题（环境噪声、环境温湿度、供水、供电等）",
                "安装相关问题（组装错误、配置错误等）",
                "标签、使用说明书或培训问题（与器械标标签、使用说明书、培训和维护文件或指南相关的问题）",
                "器械使用问题（与未能按照制造商的建议或公认的最佳实践处理、维修或操作器械相关的问题）",
                '设计缺陷问题（器械设计缺陷）',
                "其他（不属于以上任何归类的）"]
                prompt2='且故障表现术语严格在以下不良事件术语集选取（不能改变任何文字和符号，也要保留括号和括号里的内容）：'+str(imdfr_guzhang)
                if mode_select=='单品类自建术语集':
                    prompt2='' 
                # 构造组合提示词，明确指定最大数量
                prompt = (
                    f"这是一份{name_x}所发生的不良事件的情况，请结合该产品的机构特点，请分别从以下三类描述中归纳标准术语（每类不超过{max_terms}个，并且最后一个为包括编号的'其他'：\n"
                    "=== 故障表现 ===\n" + fault_raw.to_string() + "\n\n"
                    "=== 伤害表现 ===\n" + harm_raw.to_string() + "\n\n"
                    "=== 故障部位 ===\n" + part_raw.to_string() + "\n\n"
                    "格式要求：\n"
                    f"1. 故障表现术语（最多{max_terms}个，要合并类似但不同表述的故障表现，每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次{prompt2}）：\n"
                    "  1.1 术语1\n"
                    "  1.2 术语2\n"
                    "  ...\n"
                    f"2. 伤害表现术语（最多{max_terms}个，要合并类似但不同表述的伤害表现（注意区分故障表现），每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次）：\n"
                    "  2.1 术语1\n"
                    "  ...\n"
                    f"3. 故障部位术语（最多{max_terms}个，要合并类似但不同表述的故障部位，注意列出的是部位或部件，每个术语尽量简短，不能有重复，从频次多到少排列，但不用列出频次）：\n"
                    "  3.1 术语1\n"
                    "  ..."               
                )
                
                self.log("\n发送术语集生成请求...")
                response = self.call_api(
                    prompt,
                    f"你是一个医疗设备和医疗耗材的管理专家，请从三类描述中分别提取不超过{max_terms}个标准术语"
                )
                self.log(f"大模型返回:\n{response}")
                
                # 解析三类术语，并确保不超过max_terms
                fault_terms = self.extract_terms(response, "1. 故障表现术语")[:max_terms]
                harm_terms = self.extract_terms(response, "2. 伤害表现术语")[:max_terms]
                part_terms = self.extract_terms(response, "3. 故障部位术语")[:max_terms]
                
                # 用户确认
                confirmed_terms = self.confirm_all_terms(
                    fault_terms, harm_terms, part_terms
                )
                
                # 再次确保不超过max_terms
                confirmed_terms = (
                    confirmed_terms[0][:max_terms],
                    confirmed_terms[1][:max_terms],
                    confirmed_terms[2][:max_terms]
                )
                
                if not all(confirmed_terms):
                    raise ValueError("用户取消了术语确认")
                    
                return confirmed_terms
                
            except Exception as e:
                self.log(f"术语集生成失败: {str(e)}")
                raise

        def extract_terms(self, response, section_title):
            """从响应中提取特定类别的术语"""
            terms = []
            in_section = False
            max_terms = int(self.max_terms.get()) if self.max_terms.get() else 15
            
            for line in response.split('\n'):
                line = line.strip()
                if line.startswith(section_title):
                    in_section = True
                    continue
                elif in_section and line and line[0].isdigit():
                    term = line.split('. ')[1] if '. ' in line else line
                    terms.append(term.strip())
                elif in_section and not line:
                    break  # 遇到空行结束当前section
                    
            return terms

        def confirm_all_terms(self, fault_terms, harm_terms, part_terms):
            """三栏式术语确认窗口"""
            confirmed = {}
            
            def on_confirm():
                nonlocal confirmed
                confirmed['fault'] = fault_text.get("1.0", tk.END).split('\n')
                confirmed['harm'] = harm_text.get("1.0", tk.END).split('\n')
                confirmed['part'] = part_text.get("1.0", tk.END).split('\n')
                confirm_win.destroy()
            
            confirm_win = tk.Toplevel(self.root)
            bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(confirm_win)

            confirm_win.title("术语集确认")
            
            sw = confirm_win.winfo_screenwidth()    
            sh =confirm_win.winfo_screenheight()    
            ww = 1200  # 窗口宽度    
            wh = 600  # 窗口高度    
            x = (sw - ww) // 2    
            y = (sh - wh) // 2    
            confirm_win.geometry(f"{ww}x{wh}+{x}+{y}")  
            
            # 三栏布局
            frame = ttk.Frame(confirm_win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            # 故障表现术语
            ttk.Label(frame, text="故障表现术语").grid(row=0, column=0)
            fault_text = scrolledtext.ScrolledText(frame, width=80, height=35)
            fault_text.insert(tk.END, '\n'.join(fault_terms))
            fault_text.grid(row=1, column=0, padx=5, pady=5)
            
            # 伤害表现术语
            ttk.Label(frame, text="伤害表现术语").grid(row=0, column=1)
            harm_text = scrolledtext.ScrolledText(frame, width=30, height=35)
            harm_text.insert(tk.END, '\n'.join(harm_terms))
            harm_text.grid(row=1, column=1, padx=5, pady=5)
            
            # 故障部位术语
            ttk.Label(frame, text="故障部位术语").grid(row=0, column=2)
            part_text = scrolledtext.ScrolledText(frame, width=30, height=35)
            part_text.insert(tk.END, '\n'.join(part_terms))
            part_text.grid(row=1, column=2, padx=5, pady=5)
            
            # 确认按钮
            ttk.Button(
                confirm_win, text="确认所有术语", 
                command=on_confirm
            ).pack(pady=10)
            
            # 等待窗口关闭
            confirm_win.transient(self.root)
            confirm_win.grab_set()
            confirm_win.wait_window(confirm_win)
            
            # 清理空行和空值
            return (
                [t.strip() for t in confirmed.get('fault', []) if t.strip()],
                [t.strip() for t in confirmed.get('harm', []) if t.strip()],
                [t.strip() for t in confirmed.get('part', []) if t.strip()]
            )

        def batch_normalize_data(self, fault_terms, harm_terms, part_terms):
            """批量规整三类数据（支持一对一和一对多模式）"""
            try:
                print('匹配模式：',self.match_mode_combobox.get())
                batch_size = int(self.batch_size.get())
                wait_time = int(self.wait_time.get())
                is_one_to_many = self.match_mode_combobox.get() == "一对多"
                
                for i in range(0, len(self.df), batch_size):
                    batch = self.df.iloc[i:i+batch_size]
                    
                    # 构造提示词
                    prompt = (
                        "请严格按以下要求分类（必须返回JSON格式）：\n"
                        "=== 标准术语 ===\n"
                        "1. 故障表现:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(fault_terms)) + "\n"
                        "2. 伤害表现:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(harm_terms)) + "\n"
                        "3. 故障部位:\n" + '\n'.join(f"{idx+1}.{term}" for idx, term in enumerate(part_terms)) + "\n\n"
                        "=== 待分类数据 ===\n" +
                        '\n'.join(
                            f"这是一份{row['产品名称']}所发生的不良事件的报告，其中可额外参考的信息为{row['使用过程']}。编码:{row['报告编码']} 故障:'{row['器械故障表现']}' | 伤害:'{row['伤害表现']}'"
                            for _, row in batch.iterrows()
                        ) + "\n\n"
                        "返回格式要求：\n"
                        "```json\n"
                        "{\n"
                        '  "结果": [\n'
                    )
                    
                    if is_one_to_many:
                        prompt += (
                            '    {"编码": "1253705092023000050", "故障": [1,2], "伤害": [2,3], "部位": [3]},\n'
                            '    {"编码": "1253718172023000029", "故障": [4], "伤害": [5], "部位": [6]}\n'
                            "  ]\n"
                            "}\n"
                            "```\n"
                            "注意：\n"
                            "1. 必须使用双引号\n"
                            "2. 编码必须与输入完全一致\n"
                            "3. 数字对应术语编号\n"
                            "4. 可以返回多个编号，用数组表示,不允许返回空数组"
                        )
                    else:
                        prompt += (
                            '    {"编码": "1253705092023000050", "故障": 1, "伤害": 2, "部位": 3},\n'
                            '    {"编码": "1253718172023000029", "故障": 4, "伤害": 5, "部位": 6}\n'
                            "  ]\n"
                            "}\n"
                            "```\n"
                            "注意：\n"
                            "1. 必须使用双引号\n"
                            "2. 编码必须与输入完全一致\n"
                            "3. 数字对应术语编号\n"
                            "4. 每个字段只返回一个整数"
                        )
                    
                    self.log(f"\n处理行 {i+1}-{i+len(batch)} | 发送请求...")
                    response = self.call_api(
                        prompt,
                        "你必须是严格的JSON格式生成器，只返回指定格式的JSON数据"
                    )
                    self.log(f"API返回原始结果:\n{response}")
                    
                    # 解析JSON
                    try:
                        import json
                        response = response.strip()
                        if response.startswith('```json'):
                            response = response[7:].strip()
                        if response.endswith('```'):
                            response = response[:-3].strip()
                        
                        result_data = json.loads(response)
                        code_mapping = {}
                        
                        for item in result_data.get('结果', []):
                            code = item['编码']
                            
                            # 处理故障表现
                            if is_one_to_many:
                                fault_indices = [min(int(f)-1, len(fault_terms)-1) for f in item.get('故障', [])]
                                harm_indices = [min(int(h)-1, len(harm_terms)-1) for h in item.get('伤害', [])]
                                part_indices = [min(int(p)-1, len(part_terms)-1) for p in item.get('部位', [])]
                            else:
                                fault_indices = [min(int(item.get('故障', 1))-1, len(fault_terms)-1)]
                                harm_indices = [min(int(item.get('伤害', 1))-1, len(harm_terms)-1)]
                                part_indices = [min(int(item.get('部位', 1))-1, len(part_terms)-1)]
                            
                            code_mapping[code] = (fault_indices, harm_indices, part_indices)
                            
                    except Exception as e:
                        self.log(f"❌ JSON解析失败: {str(e)}")
                        self.log(f"原始响应内容: {response}")
                        code_mapping = {}
                    
                    # 填充结果
                    for _, row in batch.iterrows():
                        code = str(row['报告编码']).strip()
                        if code in code_mapping:
                            fault_indices, harm_indices, part_indices = code_mapping[code]
                            
                            # 处理故障表现
                            if fault_indices:
                                self.df.at[row.name, '故障表现(规整)'] = ";".join([fault_terms[i] for i in fault_indices if 0 <= i < len(fault_terms)])
                            else:
                                self.df.at[row.name, '故障表现(规整)'] = fault_terms[-1]  # 其他
                                
                            # 处理伤害表现
                            if harm_indices:
                                self.df.at[row.name, '伤害表现(规整)'] = ";".join([harm_terms[i] for i in harm_indices if 0 <= i < len(harm_terms)])
                            else:
                                self.df.at[row.name, '伤害表现(规整)'] = harm_terms[-1]  # 其他
                                
                            # 处理故障部位
                            if part_indices:
                                self.df.at[row.name, '故障部位(规整)'] = ";".join([part_terms[i] for i in part_indices if 0 <= i < len(part_terms)])
                            else:
                                self.df.at[row.name, '故障部位(规整)'] = part_terms[-1]  # 其他
                        else:
                            self.log(f"⚠️ 未找到编码 {code} 的分类结果")
                            self.df.at[row.name, '故障表现(规整)'] = fault_terms[-1]
                            self.df.at[row.name, '伤害表现(规整)'] = harm_terms[-1]
                            self.df.at[row.name, '故障部位(规整)'] = part_terms[-1]
                    
                    try:
                        self.df = self.df.drop(columns=['level_0', 'index'])
                    except:
                        pass
                    self.queue.put(("progress", len(batch)))
                    time.sleep(wait_time)
                    
            except Exception as e:
                self.log(f"批量规整失败: {str(e)}")
                raise

        def call_api(self, prompt, system_prompt=None):
            """API调用"""
            try:
                tree = ET.parse(self.token_file)
                root = tree.getroot()
                selected_model = self.model_var.get()
                
                # 修改为直接从根节点查找model节点
                model = None
                for m in root.findall("model"):
                    if m.get('name') == selected_model:
                        model = m
                        break
                        
                if not model:
                    raise ValueError(f"模型配置未找到: {selected_model}")
                    
                url = model.find("url").text
                api_key = model.find("api_key").text
                model_name = model.find("model").text           
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                
                messages = [{"role": "user", "content": prompt}]
                if system_prompt:
                    messages.insert(0, {"role": "system", "content": system_prompt})
                    
                data = {
                    "model": model.find("model").text,
                    "messages": messages,
                    "temperature": 0.3,
                    "max_tokens": 8000,
                }
                response = requests.post(
                    url, headers=headers, json=data, 
                    timeout=30
                )
                response.raise_for_status()
                
                # 更健壮的响应处理
                response_json = response.json()
                if 'choices' not in response_json or len(response_json['choices']) == 0:
                    raise ValueError("API返回无效响应: 无choices字段")
                    
                content = response_json['choices'][0]['message']['content']
                
                # 清理响应内容 - 移除可能的Markdown代码块标记
                if content.startswith('```json') and content.endswith('```'):
                    content = content[7:-3].strip()
                elif content.startswith('```') and content.endswith('```'):
                    content = content[3:-3].strip()
                    
                return content
                
            except Exception as e:
                raise Exception(f"API调用失败: {str(e)}")

        def process_queue(self):
            """处理消息队列"""
            try:
                while True:
                    try:
                        msg_type, *msg_data = self.queue.get_nowait()
                        if msg_type == "log":
                            self.log_area.insert(tk.END, msg_data[0])
                            self.log_area.see(tk.END)
                        elif msg_type == "progress":
                            self.progress.step(msg_data[0])
                        elif msg_type == "complete":
                            print("完成")
                            DRAW_show_analysis(self.df,'注册证编号/曾用注册证编号' ,'故障表现(规整)','事件发生日期')
                        elif msg_type == "error":
                            messagebox.showerror("错误", msg_data[0])
                    except queue.Empty:
                        break
            finally:
                self.root.after(100, self.process_queue)

    try:
        app = MedicalDataNormalizer()
        app.root.mainloop()
    except Exception as e:
        print(f"程序初始化失败: {str(e)}")

       
def LLM_edit_xml(file_path):
    """
    编辑XML文件的函数，提供统一的GUI界面进行操作。
    支持新增、删除和修改model节点，样式仅对当前窗口有效。
    """
    def initialize_xml_file(file_path):
        """初始化XML文件结构"""
        root = ET.Element("data")

        # 创建 DeepSeek 模型部分
        deepseek_model = ET.SubElement(root, "model", name="DeepSeek")
        ET.SubElement(deepseek_model, "url").text = "https://api.deepseek.com/chat/completions"
        ET.SubElement(deepseek_model, "model").text = "deepseek-chat"
        ET.SubElement(deepseek_model, "api_key").text = "your_api_key_here"

        # 创建 Kimi 模型部分
        kimi_model = ET.SubElement(root, "model", name="Kimi")
        ET.SubElement(kimi_model, "url").text = "https://api.moonshot.cn/v1/chat/completions"
        ET.SubElement(kimi_model, "model").text = "moonshot-v1-8k"
        ET.SubElement(kimi_model, "api_key").text = "your_api_key_here"

        # 生成 XML 文件
        tree = ET.ElementTree(root)
        tree.write(file_path, encoding="utf-8", xml_declaration=True)
        
    def save_xml():
        """保存XML文件"""
        tree.write(file_path, encoding="utf-8", xml_declaration=True)

    # 尝试加载XML文件
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
    except FileNotFoundError:
        messagebox.showinfo("提示", "文件未找到，将创建新配置文件！")
        initialize_xml_file(file_path)
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError:
        messagebox.showerror("错误", "XML文件格式错误！")
        return

    # 创建主窗口
    window = tk.Tk()
    window.title("API管理器 v1.1")

    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(window)
    # ==================== 布局设置 ====================
    # 窗口大小和位置
    window_width = 950
    window_height = 750
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = (screen_width - window_width) // 2
    y = (screen_height - window_height) // 2
    window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    window.minsize(800, 650)

    # 配置网格布局权重
    window.grid_rowconfigure(2, weight=1)
    window.grid_rowconfigure(4, weight=1)
    window.grid_columnconfigure(0, weight=1)

    # ==================== 控件创建 ====================
    # 1. 输入框框架
    input_frame = ttk.LabelFrame(
        window, 
        text=" 模型信息 ", 
        padding=(15, 10)
    )
    input_frame.grid(row=0, column=0, columnspan=4, padx=15, pady=(15, 10), sticky="ew")
    
    # 配置输入框列权重
    for i in range(4):
        input_frame.grid_columnconfigure(i, weight=1)

    # 输入框标签
    ttk.Label(input_frame, text="名称*").grid(row=0, column=0, padx=5, pady=2, sticky="w")
    ttk.Label(input_frame, text="API URL*").grid(row=0, column=1, padx=5, pady=2, sticky="w")
    ttk.Label(input_frame, text="模型*").grid(row=0, column=2, padx=5, pady=2, sticky="w")
    ttk.Label(input_frame, text="API Key*").grid(row=0, column=3, padx=5, pady=2, sticky="w")

    # 输入框
    name_entry = ttk.Entry(input_frame)
    name_entry.grid(row=1, column=0, padx=5, pady=5, sticky="ew")
    url_entry = ttk.Entry(input_frame)
    url_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
    model_entry = ttk.Entry(input_frame)
    model_entry.grid(row=1, column=2, padx=5, pady=5, sticky="ew")
    api_key_entry = ttk.Entry(input_frame, show="*")
    api_key_entry.grid(row=1, column=3, padx=5, pady=5, sticky="ew")

    # 2. 模型列表框架
    tree_frame = ttk.LabelFrame(
        window, 
        text=" 模型列表 ", 
        padding=(15, 10)
    )
    tree_frame.grid(row=2, column=0, columnspan=4, padx=15, pady=10, sticky="nsew")
    tree_frame.grid_columnconfigure(0, weight=1)
    tree_frame.grid_rowconfigure(0, weight=1)
    
    # 树形视图
    columns = ("name", "url", "model", "api_key")
    tree_view = ttk.Treeview(
        tree_frame, 
        columns=columns, 
        show="headings",
        selectmode="browse"
    )
    
    # 设置列
    tree_view.heading("name", text="名称", anchor="w")
    tree_view.heading("url", text="API URL", anchor="w")
    tree_view.heading("model", text="模型", anchor="w")
    tree_view.heading("api_key", text="API Key", anchor="w")
    
    # 列宽度
    tree_view.column("name", width=180, anchor="w", stretch=False)
    tree_view.column("url", width=280, anchor="w")
    tree_view.column("model", width=180, anchor="w")
    tree_view.column("api_key", width=200, anchor="w")
    
    # 滚动条
    vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree_view.yview)
    hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree_view.xview)
    tree_view.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
    
    # 布局
    tree_view.grid(row=0, column=0, sticky="nsew")
    vsb.grid(row=0, column=1, sticky="ns")
    hsb.grid(row=1, column=0, sticky="ew")

    # 填充数据
    for model in root.findall("model"):
        name = model.get("name")
        url = model.find("url").text if model.find("url") is not None else ""
        model_name = model.find("model").text if model.find("model") is not None else ""
        api_key = model.find("api_key").text if model.find("api_key") is not None else ""
        tree_view.insert("", "end", values=(name, url, model_name, api_key))

    # 3. 操作按钮框架
    button_frame = ttk.Frame(window)
    button_frame.grid(row=3, column=0, columnspan=4, pady=(5, 15), sticky="ew")
    button_frame.grid_columnconfigure((0, 1, 2, 3), weight=1, uniform="buttons")

    # 按钮
    add_button = ttk.Button(
        button_frame, 
        text="新增模型", 
        command=lambda: add_model()
    )
    delete_button = ttk.Button(
        button_frame, 
        text="删除模型", 
        command=lambda: delete_model()
    )
    modify_button = ttk.Button(
        button_frame, 
        text="修改模型", 
        command=lambda: modify_model()
    )
    test_button = ttk.Button(
        button_frame, 
        text="测试模型", 
        command=lambda: test_model()
    )

    add_button.grid(row=0, column=0, padx=10)
    delete_button.grid(row=0, column=1, padx=10)
    modify_button.grid(row=0, column=2, padx=10)
    test_button.grid(row=0, column=3, padx=10)

    # 4. 测试框架
    test_frame = ttk.LabelFrame(
        window, 
        text=" 模型测试 ", 
        padding=(15, 10)
    )
    test_frame.grid(row=4, column=0, columnspan=4, padx=15, pady=(0, 15), sticky="nsew")
    test_frame.grid_columnconfigure(0, weight=1)
    test_frame.grid_rowconfigure(1, weight=1)

    # 测试提示词
    ttk.Label(
        test_frame, 
        text="测试提示词:"
    ).grid(row=0, column=0, padx=5, pady=5, sticky="w")

    prompt_text = tk.Text(
        test_frame, 
        height=8, 
        wrap=tk.WORD, 
        font=('Microsoft YaHei', 9),
        bg="white",
        fg=text_color,
        padx=5,
        pady=5,
        insertbackground=text_color,
        selectbackground=highlight_color
    )
    prompt_text.grid(row=1, column=0, columnspan=3, padx=5, pady=5, sticky="nsew")
    prompt_text.insert(tk.END, '什么是药物警戒和医疗器械警戒，有什么不同？')

    # 测试结果标签
    result_label = ttk.Label(
        test_frame,
        text="测试结果将显示在控制台"
    )
    result_label.grid(row=2, column=0, padx=5, pady=(10, 5), sticky="w")

    # 滚动条
    scrollbar = ttk.Scrollbar(test_frame, command=prompt_text.yview)
    scrollbar.grid(row=1, column=3, sticky="ns")
    prompt_text.config(yscrollcommand=scrollbar.set)

    # ==================== 功能函数 ====================
    def on_double_click(event):
        """双击事件：填充选中模型到输入框"""
        selected_item = tree_view.selection()
        if not selected_item:
            return
        item = selected_item[0]
        values = tree_view.item(item, "values")
        name_entry.delete(0, tk.END)
        name_entry.insert(0, values[0])
        url_entry.delete(0, tk.END)
        url_entry.insert(0, values[1])
        model_entry.delete(0, tk.END)
        model_entry.insert(0, values[2])
        api_key_entry.delete(0, tk.END)
        api_key_entry.insert(0, values[3])

    tree_view.bind("<Double-1>", on_double_click)

    def add_model():
        """新增模型"""
        name = name_entry.get().strip()
        url = url_entry.get().strip()
        model_name = model_entry.get().strip()
        api_key = api_key_entry.get().strip()
        
        if not all([name, url, model_name, api_key]):
            messagebox.showerror("错误", "所有带*字段必须填写！", parent=window)
            return
        
        # 检查名称是否已存在
        for item in tree_view.get_children():
            if tree_view.item(item, "values")[0] == name:
                messagebox.showerror("错误", f"模型名称 '{name}' 已存在！", parent=window)
                return
        
        # 添加到XML和Treeview
        new_model = ET.SubElement(root, "model", name=name)
        ET.SubElement(new_model, "url").text = url
        ET.SubElement(new_model, "model").text = model_name
        ET.SubElement(new_model, "api_key").text = api_key
        tree_view.insert("", "end", values=(name, url, model_name, api_key))
        
        # 清空输入框并保存
        name_entry.delete(0, tk.END)
        url_entry.delete(0, tk.END)
        model_entry.delete(0, tk.END)
        api_key_entry.delete(0, tk.END)
        save_xml()
        messagebox.showinfo("成功", f"模型 '{name}' 添加成功！", parent=window)

    def delete_model():
        """删除模型"""
        selected_item = tree_view.selection()
        if not selected_item:
            messagebox.showerror("错误", "请先选择要删除的模型！", parent=window)
            return
        
        item = selected_item[0]
        name = tree_view.item(item, "values")[0]
        
        if not messagebox.askyesno("确认", f"确定要删除模型 '{name}' 吗？", parent=window):
            return
        
        # 从XML和Treeview中删除
        for model in root.findall("model"):
            if model.get("name") == name:
                root.remove(model)
                break
        tree_view.delete(item)
        save_xml()
        messagebox.showinfo("成功", f"模型 '{name}' 已删除！", parent=window)

    def modify_model():
        """修改模型"""
        selected_item = tree_view.selection()
        if not selected_item:
            messagebox.showerror("错误", "请先选择要修改的模型！", parent=window)
            return
        
        old_name = tree_view.item(selected_item[0], "values")[0]
        name = name_entry.get().strip()
        url = url_entry.get().strip()
        model_name = model_entry.get().strip()
        api_key = api_key_entry.get().strip()
        
        if not all([name, url, model_name, api_key]):
            messagebox.showerror("错误", "所有带*字段必须填写！", parent=window)
            return
        
        # 检查名称是否冲突（如果修改了名称）
        if name != old_name:
            for item in tree_view.get_children():
                if tree_view.item(item, "values")[0] == name:
                    messagebox.showerror("错误", f"模型名称 '{name}' 已存在！", parent=window)
                    return
        
        # 更新XML和Treeview
        for model in root.findall("model"):
            if model.get("name") == old_name:
                model.set("name", name)
                model.find("url").text = url
                model.find("model").text = model_name
                model.find("api_key").text = api_key
                break
        
        tree_view.item(selected_item[0], values=(name, url, model_name, api_key))
        save_xml()
        messagebox.showinfo("成功", f"模型 '{name}' 修改成功！", parent=window)


    def test_model():
        selected_item = tree_view.selection()
        if not selected_item:
            messagebox.showerror("错误", "请先选择一个model！")
            return
        prompt = prompt_text.get("1.0", tk.END).strip()  # 获取多行文本框的内容
        if not prompt:
            messagebox.showerror("错误", "提示词不能为空！")
            return

        # 获取选中的模型信息
        values = tree_view.item(selected_item, "values")
        name = values[0]
        url = values[1]
        model_name = values[2]
        api_key = values[3]

        # 检查 URL 和 API Key 是否有效
        if not url or not api_key:
            messagebox.showerror("错误", "URL 或 API Key 无效！")
            return

        # 构造请求数据（完全保持旧代码格式）
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model_name,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,  # 可调整生成的随机性
            "max_tokens": 8000,  # 可调整生成的最大令牌数
            "top_p": 0.9,  # 可调整生成的多样性
            "frequency_penalty": 1.0,  # 频率惩罚
            "presence_penalty": 1.0  # 存在惩罚
        }

        # 发送请求到百度文心 API（保持旧代码逻辑）
        try:
            response = requests.post(url, headers=headers, json=data)
            response.raise_for_status()  # 检查请求是否成功
            result = response.json()  # 获取返回的 JSON 数据

            # 提取生成的文本（保持旧代码逻辑）
            if "choices" in result and result["choices"]:
                generated_text = result["choices"][0]["message"]["content"].strip()
                print(f"模型 {model_name} 的测试结果：")
                print(generated_text)
                # 添加结果显示在消息框中（旧代码没有，这是改进点）
                PROGRAM_display_content_in_textbox(generated_text)
            else:
                messagebox.showerror("错误", "未获取到有效的响应内容！")
                print("未获取到有效的响应内容！")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("错误", f"请求失败：{e}")
            print(f"请求失败：{e}")
        except Exception as e:
            messagebox.showerror("错误", f"处理返回结果时出错：{e}")
            print(f"处理返回结果时出错：{e}")



    # 启动主循环
    window.mainloop()


class LLM_DF_Analyzer:
    """
    大模型辅助数据统计（当前表单）。
    """
    def __init__(self, xmlpath, df):
        # 创建临时文件保存传入的DataFrame
        self.temp_file_path = os.path.join(tempfile.gettempdir(), f"temp_analyze_data_{os.getpid()}.csv")
        df.to_csv(self.temp_file_path, index=False, encoding='utf-8')
        print(f"原始数据文件已保存到: {self.temp_file_path}")
        
        # 创建结果文件路径
        self.result_file_path = os.path.join(tempfile.gettempdir(), f"temp_analyze_result_{os.getpid()}.csv")
        print(f"分析结果将保存到: {self.result_file_path}")
        
        print("\n正在使用大模型理解用户需求...")
        self.root = tk.Toplevel()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
        self.root.title("智能数据分析")
        self.df = df
        self.tok_xml_path = xmlpath
        self.models = None
        self.selected_model = None
        self.generated_code = ""
        
        # 设置窗口布局
        self.center_window(800, 700)
        self.create_widgets()
        self.load_models()
        
        # 确保窗口关闭时删除临时文件
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def center_window(self, width, height):
        """将窗口居中显示"""
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def on_close(self):
        """窗口关闭时清理资源"""
        try:
            for filepath in [self.temp_file_path, self.result_file_path]:
                if os.path.exists(filepath):
                    os.unlink(filepath)
                    print(f"已删除临时文件: {filepath}")
        except Exception as e:
            print(f"删除临时文件时出错: {e}")
        finally:
            self.root.destroy()

    def create_widgets(self):
        """创建GUI组件"""
        # 主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 模型选择组件
        model_frame = ttk.Frame(main_frame)
        model_frame.pack(fill=tk.X, pady=5)
        ttk.Label(model_frame, text="选择大模型:").pack(side=tk.LEFT, padx=5)
        self.model_combobox = ttk.Combobox(model_frame, state="readonly")
        self.model_combobox.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # 提示词输入
        prompt_frame = ttk.Frame(main_frame)
        prompt_frame.pack(fill=tk.BOTH, pady=5)
        ttk.Label(prompt_frame, text="数据分析需求描述:").pack(anchor=tk.W, padx=5)
        self.prompt_text = tk.Text(prompt_frame, height=8)
        self.prompt_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 代码生成按钮
        self.generate_btn = ttk.Button(main_frame, text='分析数据需求', command=self.generate_code)
        self.generate_btn.pack(fill=tk.X, pady=5)

        # 模型响应显示
        response_frame = ttk.Frame(main_frame)
        response_frame.pack(fill=tk.BOTH, pady=5)
        ttk.Label(response_frame, text="模型响应:").pack(anchor=tk.W, padx=5)
        self.response_text = tk.Text(response_frame, height=10)
        self.response_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 操作按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)


        self.execute_btn2 = ttk.Button(btn_frame, text="查看代码", command=lambda:PROGRAM_display_content_in_textbox(str(self.generated_code)))
        self.execute_btn2.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.execute_btn = ttk.Button(btn_frame, text="执行分析", command=self.execute_code)
        self.execute_btn.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)


    def load_models(self):
        """从XML文件加载大模型配置"""
        try:
            tree = ET.parse(self.tok_xml_path)
            root = tree.getroot()
            self.models = {}
            for model in root.findall('model'):
                model_info = {
                    'name': model.get('name'),
                    'url': model.find('url').text,
                    'model': model.find('model').text,
                    'api_key': model.find('api_key').text
                }
                self.models[model_info['name']] = model_info
            self.model_combobox['values'] = list(self.models.keys())
            if self.models:
                self.model_combobox.current(0)
        except Exception as e:
            messagebox.showerror("加载模型配置失败", f"错误信息：{str(e)}")

    def generate_code(self):
        """调用大模型生成数据分析代码"""
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        if not prompt:
            messagebox.showerror("错误", "请输入数据分析需求描述！")
            return
        
        selected_model_name = self.model_combobox.get()
        if not selected_model_name or selected_model_name not in self.models:
            messagebox.showerror("错误", "请选择有效的大模型！")
            return
        
        selected_model = self.models[selected_model_name]
        
        # 构造包含数据信息的提示词
        columns_info = (
            f"数据文件路径: {self.temp_file_path}\n"
            f"列名: {', '.join(self.df.columns)}\n"
            f"数据类型:\n{self.df.dtypes.to_string()}\n"
            f"前3行数据示例:\n{self.df.head(3).to_string()}"
        )
        
        full_prompt = (
            f"请根据以下数据和需求生成Python数据分析代码:\n\n"
            f"数据信息:\n{columns_info}\n\n"
            f"代码要求:\n"
            f"1. 使用pd.read_csv(r'{self.temp_file_path}')读取数据\n"
            f"2. 进行必要的数据清洗和分析\n"
            f"3. 最终结果保存到变量result_df中\n"
            f"4. 使用result_df.to_csv(r'{self.result_file_path}', index=False)保存结果\n"
            f"5. 返回完整可执行代码(包含import语句)\n"
            f"6. 不要使用图形绘制相关代码\n"
            f"7. 确保代码不要重复导入和重复读取数据\n\n"
            f"数据分析需求: {prompt}"
        )
        
        print("\n正在获取解决方案...")

        try:
            headers = {
                "Authorization": f"Bearer {selected_model['api_key']}",
                "Content-Type": "application/json"
            }
            data = {
                "model": selected_model['model'],
                "messages": [{"role": "user", "content": full_prompt}],
                "temperature": 0.3,
                "max_tokens": 8000
            }
            
            print("\n正在发送请求到大模型...")
            start_time = time.time()
            
            response = requests.post(
                selected_model['url'],
                headers=headers,
                json=data,
                timeout=60
            )
            response.raise_for_status()
            
            result = response.json()
            elapsed_time = time.time() - start_time
            print(f"收到响应 (耗时: {elapsed_time:.2f}秒)")
            
            if "choices" not in result or not result["choices"]:
                raise ValueError("无效的模型响应结构")
            
            raw_code = result["choices"][0]["message"]["content"]
            self.process_response(raw_code)
            
        except requests.exceptions.RequestException as e:
            messagebox.showerror("请求失败", f"网络错误: {str(e)}")
        except Exception as e:
            messagebox.showerror("处理错误", f"错误信息: {str(e)}")

    def process_response(self, raw_response):
        """处理大模型返回的原始响应"""
        print("\n收到大模型原始响应.")
        
        # 尝试提取代码块
        code_pattern = re.compile(r"```(?:python)?\n(.*?)```", re.DOTALL)
        match = code_pattern.search(raw_response)
        
        if match:
            try:
                self.generated_code = match.group(1).strip()
                
                # 确保代码中包含必要部分
                required_parts = [
                    f"pd.read_csv(r'{self.temp_file_path}')",
                    "result_df.to_csv"
                ]
                
                for part in required_parts:
                    if part not in self.generated_code:
                        raise ValueError(f"生成的代码缺少必要部分: {part}")
                
                self.response_text.delete("1.0", tk.END)
                self.response_text.insert("1.0", "✅ 分析方案已生成！")
            except Exception as e:
                error_msg = f"处理代码时出错: {str(e)}"
                messagebox.showerror("处理错误", error_msg)
                self.response_text.delete("1.0", tk.END)
                self.response_text.insert("1.0", f"错误: {error_msg}\n\n原始响应:\n{raw_response}")
        else:
            messagebox.showwarning("警告", "未检测到标准代码块，请检查响应内容")
            self.response_text.delete("1.0", tk.END)
            self.response_text.insert("1.0", raw_response)

    def execute_code(self):
        """执行生成的代码，生成结果DataFrame并打印"""
        if not self.generated_code:
            messagebox.showerror("错误", "请先生成有效代码！")
            return
        
        try:
            # 准备执行环境
            namespace = {
                'pd': pd,
                'np': np,
                'df': self.df
            }
            
            print("\n执行数据分析方案...")
            print("="*50)
            
            # 在受限环境中执行代码
            exec(self.generated_code, namespace)
            
            # 检查结果文件是否生成
            if os.path.exists(self.result_file_path):
                result_df = pd.read_csv(self.result_file_path)
                print("\n数据分析完成!")
                print("="*50)
                print("\n分析结果DataFrame:")
                print("="*50)
                PROGRAM_display_df_in_treeview(result_df,0,0)
                print("\n结果已保存到:", self.result_file_path)
            else:
                print("警告: 结果文件未生成")
            
        except Exception as e:
            error_msg = f"分析执行失败: {str(e)}"
            messagebox.showerror("执行错误", error_msg)
            print(error_msg)
            import traceback
            traceback.print_exc()
            
            
            
class LLM_DF_PLT:
    """
    大模型辅助数据可视化（当前表单）。
    """
    def __init__(self, xmlpath, df):
        # 创建临时文件保存传入的DataFrame（Windows兼容方式）
        self.temp_file_path = os.path.join(tempfile.gettempdir(), f"temp_plot_data_{os.getpid()}.csv")
        df.to_csv(self.temp_file_path, index=False, encoding='utf-8')
        print(f"临时数据文件已保存到: {self.temp_file_path}")
        
        # 打印列名信息
        #print("\n发送给大模型的列名信息:")
        #print("-" * 40)
        #print(f"列名: {', '.join(df.columns)}")
        #print("\n各列数据类型:")
        #print(df.dtypes.to_string())
        #print("-" * 40)
        print("\n正在使用大模型理解用户需求...")
        self.root = tk.Toplevel()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)

        self.root.title("智能数据可视化")
        self.df = df
        self.tok_xml_path = xmlpath
        self.models = None
        self.selected_model = None
        self.generated_code = ""
        self.current_figure = None
        
        # 初始化Matplotlib字体设置
        plt.rcParams['font.sans-serif'] = [my_font_ch]
        plt.rcParams['axes.unicode_minus'] = False

        # 设置窗口布局
        self.center_window(800, 500)
        self.create_widgets()
        self.load_models()
        
        # 确保窗口关闭时删除临时文件
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def center_window(self, width, height):
        """将窗口居中显示"""
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def on_close(self):
        """窗口关闭时清理资源"""
        try:
            if hasattr(self, 'temp_file_path') and os.path.exists(self.temp_file_path):
                os.unlink(self.temp_file_path)
                #print(f"已删除临时文件: {self.temp_file_path}")
        except Exception as e:
            print(f"删除临时文件时出错: {e}")
        finally:
            if plt.get_fignums():
                plt.close('all')
            self.root.destroy()

    def create_widgets(self):
        """创建GUI组件"""
        # 主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 模型选择组件
        model_frame = ttk.Frame(main_frame)
        model_frame.pack(fill=tk.X, pady=5)
        ttk.Label(model_frame, text="选择大模型:").pack(side=tk.LEFT, padx=5)
        self.model_combobox = ttk.Combobox(model_frame, state="readonly")
        self.model_combobox.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # 提示词输入
        prompt_frame = ttk.Frame(main_frame)
        prompt_frame.pack(fill=tk.BOTH, pady=5)
        ttk.Label(prompt_frame, text="可视化需求描述:").pack(anchor=tk.W, padx=5)
        self.prompt_text = tk.Text(prompt_frame, height=8)
        self.prompt_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 代码生成按钮
        self.generate_btn = ttk.Button(main_frame, text='分析可视化需求', command=self.generate_code)
        self.generate_btn.pack(fill=tk.X, pady=5)

        # 模型响应显示
        response_frame = ttk.Frame(main_frame)
        response_frame.pack(fill=tk.BOTH, pady=5)
        ttk.Label(response_frame, text="模型响应:").pack(anchor=tk.W, padx=5)
        self.response_text = tk.Text(response_frame, height=10)
        self.response_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 生成代码显示
        code_frame = ttk.Frame(main_frame)
        #code_frame.pack(fill=tk.BOTH, pady=5)
        #ttk.Label(code_frame, text="生成的绘图代码:").pack(anchor=tk.W, padx=5)
        self.code_text = tk.Text(code_frame, height=12)
        #self.code_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 操作按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        self.execute_btn2 = ttk.Button(btn_frame, text="查看代码", command=lambda:PROGRAM_display_content_in_textbox(str(self.generated_code)))
        self.execute_btn2.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.execute_btn = ttk.Button(btn_frame, text="执行方案", command=self.execute_code)
        self.execute_btn.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        


    def load_models(self):
        """从XML文件加载大模型配置"""
        try:
            tree = ET.parse(self.tok_xml_path)
            root = tree.getroot()
            self.models = {}
            for model in root.findall('model'):
                model_info = {
                    'name': model.get('name'),
                    'url': model.find('url').text,
                    'model': model.find('model').text,
                    'api_key': model.find('api_key').text
                }
                self.models[model_info['name']] = model_info
            self.model_combobox['values'] = list(self.models.keys())
            self.model_combobox.current(0)
        except Exception as e:
            messagebox.showerror("加载模型配置失败", f"错误信息：{str(e)}")

    def generate_code(self):
        """调用大模型生成绘图代码"""
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        if not prompt:
            messagebox.showerror("错误", "请输入可视化需求描述！")
            return
        
        selected_model = self.models[self.model_combobox.get()]
        
        # 构造包含数据信息的提示词（使用原始字符串避免转义问题）
        columns_info = (
            fr"数据文件路径: {self.temp_file_path}\n"
            f"列名: {', '.join(self.df.columns)}\n"
            f"数据类型:\n{self.df.dtypes.to_string()}\n"
            f"前3行数据示例:\n{self.df.head(3).to_string()}"
        )
        
        full_prompt = (
            f"请根据以下数据和需求生成Python绘图代码:\n\n"
            f"数据信息:\n{columns_info}\n\n"
            f"代码要求:\n"
            f"1. 使用pd.read_csv(r'{self.temp_file_path}')读取数据\n"
            f"2. 正确处理中文显示(已配置plt.rcParams)\n"
            f"3. 包含标题、坐标轴标签等必要元素\n"
            f"4. 返回完整可执行代码(包含import语句)\n"
            f"5. 不要使用\\u Unicode转义字符\n\n"
            f"6. 确保代码不要重复导入和重复读取数据\n\n"         
            f"可视化需求: {prompt}"
        )
        
        #print("\n发送给大模型的完整提示词:")
        #print("-" * 60)
        #print(full_prompt)
        #print("-" * 60)
        print("\n正在获取解决方案...")

        try:
            headers = {
                "Authorization": f"Bearer {selected_model['api_key']}",
                "Content-Type": "application/json"
            }
            data = {
                "model": selected_model['model'],
                "messages": [{"role": "user", "content": full_prompt}],
                "temperature": 0.3,
                "max_tokens": 8000
            }
            
            print("\n正在发送请求到大模型...")
            start_time = time.time()
            
            response = requests.post(
                selected_model['url'],
                headers=headers,
                json=data,
                timeout=60
            )
            response.raise_for_status()
            
            result = response.json()
            elapsed_time = time.time() - start_time
            print(f"收到响应 (耗时: {elapsed_time:.2f}秒)")
            
            if "choices" not in result or not result["choices"]:
                raise ValueError("无效的模型响应结构")
            
            raw_code = result["choices"][0]["message"]["content"]
            self.process_response(raw_code)
            
        except requests.exceptions.RequestException as e:
            messagebox.showerror("请求失败", f"网络错误: {str(e)}")
        except Exception as e:
            messagebox.showerror("处理错误", f"错误信息: {str(e)}")

    def process_response(self, raw_response):
        """处理大模型返回的原始响应，修复Unicode转义问题"""
        print("\n收到大模型原始响应.")
        #print("-" * 60)
        #print(raw_response)
        #print("-" * 60)
        
        # 尝试提取代码块（处理多种代码块格式）
        code_pattern = re.compile(r"```(?:python)?\n(.*?)```", re.DOTALL)
        match = code_pattern.search(raw_response)
        
        if match:
            try:
                # 修复Unicode转义问题的新方法
                self.generated_code = self.safe_unicode_decode(match.group(1).strip())
                
                # 确保代码中包含数据读取部分（使用原始字符串）
                if f"pd.read_csv(r'{self.temp_file_path}')" not in self.generated_code:
                    self.generated_code = (
                        f"import pandas as pd\n"
                        f"import matplotlib.pyplot as plt\n\n"
                        f"# 读取数据\n"
                        f"df = pd.read_csv(r'{self.temp_file_path}')\n\n"
                        f"{self.generated_code}"
                    )
                    print("已自动添加数据读取代码")
                
                self.code_text.delete("1.0", tk.END)
                self.code_text.insert("1.0", self.generated_code)
                self.response_text.delete("1.0", tk.END)
                self.response_text.insert("1.0", "✅ 解决方案已返回！")
            except Exception as e:
                error_msg = f"处理代码时出错: {str(e)}"
                messagebox.showerror("处理错误", error_msg)
                self.response_text.delete("1.0", tk.END)
                self.response_text.insert("1.0", f"错误: {error_msg}\n\n原始响应:\n{raw_response}")
        else:
            messagebox.showwarning("警告", "未检测到标准代码块，请检查响应内容")
            self.response_text.delete("1.0", tk.END)
            self.response_text.insert("1.0", raw_response)

    def safe_unicode_decode(self, text):
        """安全处理包含Unicode转义的字符串"""
        try:
            # 方法1: 尝试直接解码
            return text.encode('utf-8').decode('unicode_escape')
        except UnicodeDecodeError:
            try:
                # 方法2: 替换无效的Unicode转义
                fixed_text = re.sub(r'\\U[0-9a-fA-F]{1,8}', lambda m: m.group(0).ljust(10, '0')[:10], text)
                return fixed_text.encode('utf-8').decode('unicode_escape')
            except Exception:
                # 方法3: 使用错误忽略模式
                return text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

    def execute_code(self):
        """执行生成的代码，增强错误处理"""
        if not self.generated_code:
            messagebox.showerror("错误", "请先生成有效代码！")
            return
        
        try:
            # 关闭所有现有的Matplotlib图形
            plt.close('all')
            
            # 创建新的图形
            fig = plt.figure(figsize=(10, 6))
            
            # 准备执行环境
            namespace = {
                'plt': plt,
                'pd': pd,
                'np': np,
                'df': self.df,
                'fig': fig
            }
            
            print("\n执行生成的方案...")
            
            # 修复代码中的Unicode问题
            cleaned_code = self.fix_unicode_in_code(self.generated_code)
            
            # 在受限环境中执行代码
            exec(cleaned_code, namespace)
            
            # 显示图形（阻塞模式）
            plt.show(block=False)
            
            # 保存当前图形引用
            self.current_figure = fig
            print("方案执行成功！")
            
        except Exception as e:
            error_msg = f"方案执行失败: {str(e)}"
            messagebox.showerror("执行错误", error_msg)
            print(error_msg)
            import traceback
            traceback.print_exc()

    def fix_unicode_in_code(self, code):
        """修复代码中的Unicode转义问题"""
        # 处理简单的Unicode转义
        code = re.sub(r'\\u[0-9a-fA-F]{1,4}', lambda m: m.group(0).ljust(6, '0')[:6], code)
        code = re.sub(r'\\U[0-9a-fA-F]{1,8}', lambda m: m.group(0).ljust(10, '0')[:10], code)
        
        # 处理其他可能的编码问题
        try:
            return code.encode('utf-8').decode('unicode_escape')
        except:
            return code.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

    def save_figure(self):
        """保存当前图形"""
        if not self.current_figure:
            messagebox.showerror("错误", "没有可保存的图形！")
            return
            
        filepath = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[
                ("PNG文件", "*.png"),
                ("JPEG文件", "*.jpg"),
                ("PDF文件", "*.pdf"),
                ("SVG文件", "*.svg"),
                ("所有文件", "*.*")
            ],
            title="保存图形"
        )
        
        if filepath:
            try:
                self.current_figure.savefig(
                    filepath, 
                    dpi=300, 
                    bbox_inches='tight',
                    facecolor='white'  # 确保背景为白色
                )
                messagebox.showinfo("保存成功", f"图形已保存至:\n{filepath}")
                print(f"图形已保存到: {filepath}")
            except Exception as e:
                messagebox.showerror("保存失败", f"保存图形时出错: {str(e)}")



class LLM_Send_Form:
    """
    大模型辅助分析（整表模式）
    """
    
    def __init__(self, df, root=None):
        self.df = df
        self.rootfsx = tk.Toplevel(root)
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.rootfsx)

        self.rootfsx.title("通用AI数据处理工具 V1.0")
        self.center_window(800, 600)
        self.token_file = os.path.join(csdir, 'tok.xml')
        self.reference_file = os.path.join(csdir, 'ai_cankao2.xml')  # 修改为 ai_cankao2.xml
        self.queue = queue.Queue()
        self.cache = {}
        self.token_1 = '未载入'
        self.default_prompt_fixed = """
        你是一名医学和药学背景的监测人员，熟悉药品不良反应监测和医疗器械不良事件监测，请根据以下要求分析输入数据并返回结果。
        """
        self.default_prompt_editable = (
            "你是一个专业的助手，负责对数据进行规整。"
            "请根据输入数据生成结构化结果。"
        )
        self.style = ttk.Style()
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure("TLabel", background="#f0f0f0", font=("Arial", 10))
        self.style.configure("TButton", font=("Arial", 10))
        self.style.configure("TEntry", font=("Arial", 10))
        self.main_frame = ttk.Frame(self.rootfsx)
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        self.left_frame = ttk.Frame(self.main_frame)
        self.left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 模型选择下拉菜单
        self.model_frame = ttk.Frame(self.left_frame)
        self.model_frame.pack(pady=10, fill=tk.X)
        self.model_label = ttk.Label(self.model_frame, text="选择调用模型：")
        self.model_label.pack(side=tk.LEFT, padx=(0, 10))
        self.model_var = tk.StringVar(self.model_frame)
        self.model_var.set("DeepSeek")  # 默认选择 DeepSeek
        self.model_option = ttk.OptionMenu(self.model_frame, self.model_var, "DeepSeek", *self.get_available_models())
        self.model_option.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.model_var.trace_add("write", lambda *args: self.load_token())

        # 列选择界面
        self.column_frame = ttk.Frame(self.left_frame)
        self.column_frame.pack(pady=10, fill=tk.BOTH, expand=True)
        self.column_label = ttk.Label(self.column_frame, text="选择需要处理的列:")
        self.column_label.pack()
        self.column_tree = ttk.Treeview(self.column_frame, columns=("列名"), show="headings")
        self.column_tree.heading("列名", text="列名")
        self.column_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.column_scrollbar = ttk.Scrollbar(self.column_frame, orient=tk.VERTICAL, command=self.column_tree.yview)
        self.column_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.column_tree.config(yscrollcommand=self.column_scrollbar.set)
        for col in self.df.columns:
            self.column_tree.insert("", tk.END, values=(col,))

        # 按钮框架
        self.button_frame = ttk.Frame(self.left_frame)
        self.button_frame.pack(pady=10, fill=tk.X)
        self.reference_button = ttk.Button(self.button_frame, text="参考实例", command=self.load_reference_file)
        self.reference_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        self.process_button = ttk.Button(self.button_frame, text="处理数据", command=self.start_processing)
        self.process_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)

        # 右侧框架（提示词和历史记录）
        self.right_frame = ttk.Frame(self.main_frame)
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=10, pady=10)
        self.prompt_label = ttk.Label(self.right_frame, text="提示词:")
        self.prompt_label.pack(pady=(10, 0))
        self.prompt_text = scrolledtext.ScrolledText(self.right_frame, height=10, width=40, font=("Arial", 10))
        self.prompt_text.pack(fill=tk.BOTH, expand=True)
        self.prompt_text.bind("<<Modified>>", self.on_prompt_modified)  # 绑定事件
        self.history_label = ttk.Label(self.right_frame, text="历史记录:")
        self.history_label.pack(pady=(10, 0))
        self.history_frame = ttk.Frame(self.right_frame)
        self.history_frame.pack(fill=tk.BOTH, expand=True)
        self.history_listbox = tk.Listbox(self.history_frame, height=5, font=("Arial", 10))
        self.history_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.history_scrollbar = ttk.Scrollbar(self.history_frame, orient=tk.VERTICAL, command=self.history_listbox.yview)
        self.history_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_listbox.config(yscrollcommand=self.history_scrollbar.set)
        self.history_listbox.bind("<<ListboxSelect>>", self.load_selected_history)
        self.load_prompt_and_history()
        self.rootfsx.after(100, self.process_queue)
    def on_prompt_modified(self, event=None):
        if self.prompt_text.edit_modified():
            self.save_prompt_and_history()
            self.prompt_text.edit_modified(False)  # 重置修改标志
    def center_window(self, width, height):
        screen_width = self.rootfsx.winfo_screenwidth()
        screen_height = self.rootfsx.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.rootfsx.geometry(f"{width}x{height}+{x}+{y}")

    def get_available_models(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        tree = ET.parse(self.token_file)
        root = tree.getroot()
        return [model.get("name") for model in root.findall("model")]

    def extract_json_part(self, text):
        start = text.find('```json')
        if start == -1:
            return None
        start += 5
        end = text.find('```', start)
        if end == -1:
            return None
        json_part = text[start:end].strip()
        return json_part

    def load_token(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        try:
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            selected_model = self.model_var.get()
            token = root.find(f"./model[@name='{selected_model}']/api_key").text
            if token:
                self.token_1 = token
        except ET.ParseError:
            self.initialize_xml_file(self.token_file)

    def initialize_xml_file(self, token_file):
        root = ET.Element("data")
        deepseek_model = ET.SubElement(root, "model", name="DeepSeek")
        ET.SubElement(deepseek_model, "url").text = "https://qianfan.baidubce.com/v2/chat/completions"
        ET.SubElement(deepseek_model, "model").text = "deepseek-v3"
        ET.SubElement(deepseek_model, "api_key").text = "token"
        kimi_model = ET.SubElement(root, "model", name="Kimi")
        ET.SubElement(kimi_model, "url").text = "https://api.moonshot.cn/v1/chat/completions"
        ET.SubElement(kimi_model, "model").text = "moonshot-v1-8k"
        ET.SubElement(kimi_model, "api_key").text = "token"
        tree = ET.ElementTree(root)
        tree.write(token_file, encoding="utf-8", xml_declaration=True)

    def load_prompt_and_history(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        try:
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            prompt = root.find("prompt")
            if prompt is not None and prompt.text:
                self.prompt_text.insert(tk.END, prompt.text)
            else:
                self.prompt_text.insert(tk.END, self.default_prompt_editable)
                self.save_prompt_and_history()
            history = root.find("history")
            if history is not None:
                for item in history.findall("item"):
                    self.history_listbox.insert(tk.END, item.text)
        except ET.ParseError:
            self.initialize_xml_file()

    def save_prompt_and_history(self):
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        tree = ET.parse(self.token_file)
        root = tree.getroot()
        prompt_element = root.find("prompt")
        if prompt_element is None:
            prompt_element = ET.SubElement(root, "prompt")
        prompt_element.text = prompt
        history_element = root.find("history")
        if history_element is None:
            history_element = ET.SubElement(root, "history")
        if prompt:
            history_items = [item.text for item in history_element.findall("item")]
            if prompt not in history_items:
                if len(history_items) >= 10:
                    history_element.remove(history_element[-1])
                new_item = ET.SubElement(history_element, "item")
                new_item.text = prompt
                new_item.tail = "\n"
        self.history_listbox.delete(0, tk.END)
        for item in history_element.findall("item"):
            self.history_listbox.insert(0, item.text)
        tree.write(self.token_file, encoding="utf-8", xml_declaration=True)

    def load_selected_history(self, event=None):
        selected_index = self.history_listbox.curselection()
        if selected_index:
            selected_prompt = self.history_listbox.get(selected_index)
            self.prompt_text.delete("1.0", tk.END)
            self.prompt_text.insert(tk.END, selected_prompt)

    def load_reference_file(self):
        self.reference_window = tk.Toplevel(self.rootfsx)
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.rootfsx)
        self.reference_window.title("参考实例")
        self.center_window_on_parent(self.reference_window, 600, 400)
        self.reference_tree = ttk.Treeview(self.reference_window, columns=("标题", "内容"), show="headings")
        self.reference_tree.heading("标题", text="标题")
        self.reference_tree.heading("内容", text="内容")
        self.reference_tree.pack(fill=tk.BOTH, expand=True)
        self.reference_tree.bind("<Double-1>", self.on_reference_double_click)
        self.reference_tree.bind("<Button-3>", self.show_reference_context_menu)
        self.load_reference_examples()

    def load_reference_examples(self):
        if not os.path.exists(self.reference_file):
            self.initialize_reference_file()
        try:
            tree = ET.parse(self.reference_file)
            root = tree.getroot()
            self.reference_examples = []
            for example in root.findall("example"):
                title = example.find("title").text
                content = example.find("content").text
                self.reference_examples.append({"title": title, "content": content})
                self.reference_tree.insert("", tk.END, values=(title, content))
        except ET.ParseError:
            self.initialize_reference_file()

    def initialize_reference_file(self):
        root = ET.Element("examples")
        tree = ET.ElementTree(root)
        tree.write(self.reference_file, encoding="utf-8", xml_declaration=True)

    def save_reference_examples(self):
        root = ET.Element("examples")
        for example in self.reference_examples:
            example_element = ET.SubElement(root, "example")
            ET.SubElement(example_element, "title").text = example["title"]
            ET.SubElement(example_element, "content").text = example["content"]
        tree = ET.ElementTree(root)
        tree.write(self.reference_file, encoding="utf-8", xml_declaration=True)

    def show_reference_context_menu(self, event):
        selected_item = self.reference_tree.identify_row(event.y)
        if selected_item:
            self.reference_tree.selection_set(selected_item)
            menu = tk.Menu(self.reference_window, tearoff=0)
            menu.add_command(label="新增", command=self.add_reference_example)
            menu.add_command(label="编辑", command=lambda: self.edit_reference_example(selected_item))
            menu.add_command(label="移除", command=lambda: self.remove_reference_example(selected_item))
            menu.post(event.x_root, event.y_root)
        else:
            menu = tk.Menu(self.reference_window, tearoff=0)
            menu.add_command(label="新增", command=self.add_reference_example)
            menu.post(event.x_root, event.y_root)

    def add_reference_example(self):
        self.edit_reference_example(None)

    def edit_reference_example(self, item):
        self.edit_window = tk.Toplevel(self.reference_window)
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.reference_window)

        self.edit_window.title("编辑参考实例")
        self.center_window_on_parent(self.edit_window, 400, 300)
        title_label = ttk.Label(self.edit_window, text="标题:")
        title_label.pack(pady=(10, 0))
        self.edit_title_entry = ttk.Entry(self.edit_window, width=40)
        self.edit_title_entry.pack(pady=5)
        content_label = ttk.Label(self.edit_window, text="内容:")
        content_label.pack(pady=(10, 0))
        self.edit_content_text = scrolledtext.ScrolledText(self.edit_window, height=10, width=40, font=("Arial", 10))
        self.edit_content_text.pack(fill=tk.BOTH, expand=True)
        if item:
            title, content = self.reference_tree.item(item, "values")
            self.edit_title_entry.insert(0, title)
            self.edit_content_text.insert(tk.END, content)
        save_button = ttk.Button(self.edit_window, text="保存", command=lambda: self.save_reference_example_edit(item))
        save_button.pack(pady=10)

    def save_reference_example_edit(self, item):
        title = self.edit_title_entry.get().strip()
        content = self.edit_content_text.get("1.0", tk.END).strip()
        if not title or not content:
            messagebox.showwarning("警告", "标题和内容不能为空")
            return
        if item:
            self.reference_tree.item(item, values=(title, content))
            index = self.reference_tree.index(item)
            self.reference_examples[index] = {"title": title, "content": content}
        else:
            self.reference_tree.insert("", tk.END, values=(title, content))
            self.reference_examples.append({"title": title, "content": content})
        self.save_reference_examples()
        self.edit_window.destroy()

    def remove_reference_example(self, item):
        index = self.reference_tree.index(item)
        self.reference_tree.delete(item)
        del self.reference_examples[index]
        self.save_reference_examples()

    def on_reference_double_click(self, event):
        selected_item = self.reference_tree.selection()
        if selected_item:
            title, content = self.reference_tree.item(selected_item, "values")
            self.prompt_text.delete("1.0", tk.END)
            self.prompt_text.insert(tk.END, content)

    def center_window_on_parent(self, window, width, height):
        parent_x = self.rootfsx.winfo_x()
        parent_y = self.rootfsx.winfo_y()
        parent_width = self.rootfsx.winfo_width()
        parent_height = self.rootfsx.winfo_height()
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        window.geometry(f"{width}x{height}+{x}+{y}")

    def start_processing(self):
        # 获取用户选择的列
        selected_columns = [self.column_tree.item(item, "values")[0] for item in self.column_tree.selection()]
        if not selected_columns:
            messagebox.showwarning("警告", "请先选择列")
            return
        
        # 禁用按钮，防止重复点击
        self.process_button.config(state=tk.DISABLED)
        
        # 启动数据处理线程
        threading.Thread(
            target=self.process_data,
            args=(selected_columns,),
            daemon=True
        ).start()

    def process_data(self, selected_columns):
        try:
            # 生成编号列
            self.df["数据链接代码"] = 'R' + (self.df.index + 1).astype(str)

            
            # 构建输入内容：将整个表格转换为字符串
            table_content = "以下是需要分析的数据表：\n"
            table_content += self.df[["数据链接代码"] + selected_columns].to_string(index=False)
            
            # 构建提示词
            user_prompt = self.prompt_text.get("1.0", tk.END).strip()
            full_prompt = self.default_prompt_fixed + user_prompt
            
            # 合并提示词和数据表内容
            input_content = f"{full_prompt}\n\n{table_content}"

            # 调用 API
            selected_model = self.model_var.get()
            response = self.call_api(self.token_1, input_content, selected_model)
            
            # 打印 API 返回结果
            result = response["choices"][0]["message"]["content"]
            PROGRAM_display_content_in_textbox(result)
            #print(f"API 返回的总体分析结果:\n{result}")
        
        except Exception as e:
            print(f"处理数据时出错: {e}")
        finally:
            # 恢复按钮状态
            self.process_button.config(state=tk.NORMAL)

    def call_api(self, api_key, content, selected_model=None):
        """
        调用 API 处理数据
        :param api_key: API 密钥
        :param content: 请求内容（包含提示词和数据表内容）
        :param selected_model: 选择的模型名称（可选，默认为当前选择的模型）
        :return: API 响应结果
        """
        if selected_model is None:
            selected_model = self.model_var.get()
        
        if not api_key:
            raise ValueError("API Key 不能为空")
        
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        
        try:
            # 加载模型配置
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            model_element = root.find(f"./model[@name='{selected_model}']")
            if model_element is None:
                raise ValueError(f"模型 {selected_model} 未配置")
            
            # 获取 API URL 和模型名称
            url = model_element.find("url").text
            model_name = model_element.find("model").text
            
            # 构建请求头
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json; charset=utf-8"
            }
            
            # 构建请求数据
            data = {
                "model": model_name,
                "messages": [
                    {"role": "system", "content": "你是一个数据分析助手，请根据提供的数据表进行分析。"},  # 系统提示词
                    {"role": "user", "content": content}  # 用户输入内容（提示词 + 数据表）
                ],
                "stream": False
            }
            
            # 发送请求
            print('发送的信息：',data)
            response = requests.post(url, headers=headers, json=data)
            
            # 检查响应状态
            if response.status_code == 200:
                return response.json()
            else:
                raise Exception(f"API 调用失败: {response.status_code}, {response.text}")
        
        except ET.ParseError as e:
            raise Exception(f"解析配置文件时出错: {str(e)}")
        except requests.RequestException as e:
            raise Exception(f"网络请求失败: {str(e)}")
        except Exception as e:
            raise Exception(f"API 调用过程中发生错误: {str(e)}")

    def process_queue(self):
        try:
            while True:
                try:
                    message = self.queue.get_nowait()
                    if message[0] == "log":
                        print(message[1])
                    elif message[0] == "error":
                        messagebox.showerror("错误", message[1])
                except queue.Empty:
                    break
        except Exception as e:
            print(f"处理队列时出错: {e}")
        finally:
            self.rootfsx.after(100, self.process_queue)



class LLM_Send_Row:
    """
    大模型辅助分析（逐行模式）
    """
    def __init__(self, df, root=None):
        self.df = df.copy()
        self.df.reset_index(inplace=True)
        self.rootfsx =tk.Toplevel(root)#tk.Tk()# root if root else tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.rootfsx)

        self.rootfsx.title("通用AI数据处理工具 V1.0")
        self.center_window(800, 600)
        self.token_file = os.path.join(csdir,'tok.xml')
        self.reference_file =os.path.join(csdir,'ai_cankao.xml') 
        #self.rootfsx.state('zoomed')
        self.queue = queue.Queue()
        self.batch_size = 20
        self.executor = ThreadPoolExecutor(max_workers=5)
        self.cache = {}
        self.token_1='未载入'
        self.default_prompt_fixed = """
        请根据以下要求分析输入数据并返回结果：
        1. **输入数据说明**：
           - 输入数据中包含用户选择的列。
           - 每行数据都有一个唯一的 `数据链接代码`，格式为 `R + 数字`（例如：R1, R2, R3）。
        2. **返回结果格式**：
           - 返回结果必须为 JSON 格式。以```json开头。
           - JSON 的键为 `数据链接代码`，值为分析结果。
        3. **注意事项**：
           - 确保每个 `数据链接代码` 对应的结果与输入数据一致。
           - 不要遗漏任何 `数据链接代码`。
        """
        self.default_prompt_editable = (
            "你是一个专业的助手，负责对数据进行规整。"
            "请根据输入数据生成结构化结果。"
        )
        self.style = ttk.Style()
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure("TLabel", background="#f0f0f0", font=("Arial", 10))
        self.style.configure("TButton", font=("Arial", 10))
        self.style.configure("TEntry", font=("Arial", 10))
        self.main_frame = ttk.Frame(self.rootfsx)
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        self.left_frame = ttk.Frame(self.main_frame)
        self.left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 模型选择下拉菜单
        self.model_frame = ttk.Frame(self.left_frame)
        self.model_frame.pack(pady=10, fill=tk.X)
        self.model_label = ttk.Label(self.model_frame, text="选择调用模型：")
        self.model_label.pack(side=tk.LEFT, padx=(0, 10))
        self.model_var = tk.StringVar(self.model_frame)
        self.model_var.set("DeepSeek")  # 默认选择 DeepSeek
        self.model_option = ttk.OptionMenu(self.model_frame, self.model_var, "DeepSeek", *self.get_available_models())
        self.model_option.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.model_var.trace_add("write", lambda *args: self.load_token())

        # 每批处理行数输入框
        self.batch_size_frame = ttk.Frame(self.left_frame)
        self.batch_size_frame.pack(pady=10, fill=tk.X)
        self.batch_size_label = ttk.Label(self.batch_size_frame, text="每批处理行数(受API限制，超过20可能会有返回结果不完整风险）:")
        self.batch_size_label.pack(side=tk.LEFT, padx=(0, 10))
        self.batch_size_entry = ttk.Entry(self.batch_size_frame, width=10)
        self.batch_size_entry.insert(0, "20")
        self.batch_size_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)



        # 并发选项下拉菜单
        self.concurrent_frame = ttk.Frame(self.left_frame)
        self.concurrent_frame.pack(pady=10, fill=tk.X)
        self.concurrent_label = ttk.Label(self.concurrent_frame, text="并发选项：")
        self.concurrent_label.pack(side=tk.LEFT, padx=(0, 10))
        # 修正并发选项初始化
        self.concurrent_var = tk.StringVar(self.concurrent_frame)
        self.concurrent_options = ["不并发", "并发"]  # 明确指定所有可用选项
        self.concurrent_var.set(self.concurrent_options[0])  # 设置默认值
        # 修正OptionMenu参数
        self.concurrent_option = ttk.OptionMenu(
            self.concurrent_frame, 
            self.concurrent_var, 
            self.concurrent_options[0],  # 默认值
            *self.concurrent_options     # 所有可用选项
        )
            
        self.concurrent_option.pack(side=tk.LEFT, fill=tk.X, expand=True)


        # 等待时间设置框（默认隐藏）
        self.wait_time_frame = ttk.Frame(self.left_frame)
        self.wait_time_label = ttk.Label(self.wait_time_frame, text="等待时间（秒）:")
        self.wait_time_label.pack(side=tk.LEFT, padx=(0, 10))
        self.wait_time_entry = ttk.Entry(self.wait_time_frame, width=10)
        self.wait_time_entry.insert(0, "3")  # 默认等待时间为3秒
        self.wait_time_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.wait_time_frame.pack_forget()  # 初始状态下隐藏

        # 监听并发选项的变化
        self.concurrent_var.trace_add("write", self.toggle_wait_time_visibility)



        # 加载已保存的 Token
        self.load_token()

        # 列选择界面
        self.column_frame = ttk.Frame(self.left_frame)
        self.column_frame.pack(pady=10, fill=tk.BOTH, expand=True)
        self.column_label = ttk.Label(self.column_frame, text="选择需要处理的列:")
        self.column_label.pack()
        self.column_tree = ttk.Treeview(self.column_frame, columns=("列名"), show="headings")
        self.column_tree.heading("列名", text="列名")
        self.column_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.column_scrollbar = ttk.Scrollbar(self.column_frame, orient=tk.VERTICAL, command=self.column_tree.yview)
        self.column_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.column_tree.config(yscrollcommand=self.column_scrollbar.set)
        for col in self.df.columns:
            self.column_tree.insert("", tk.END, values=(col,))

        # 按钮框架
        self.button_frame = ttk.Frame(self.left_frame)
        self.button_frame.pack(pady=10, fill=tk.X)
        self.reference_button = ttk.Button(self.button_frame, text="参考实例", command=self.load_reference_file)
        self.reference_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        self.process_button = ttk.Button(self.button_frame, text="处理数据", command=self.start_processing)
        self.process_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)

        # 进度条
        self.progress_frame = ttk.Frame(self.left_frame)
        self.progress_frame.pack(pady=10, fill=tk.X)
        self.progress = ttk.Progressbar(self.progress_frame, orient="horizontal", length=400, mode="determinate")
        self.progress.pack(fill=tk.X, expand=True)

        # 右侧框架（提示词和历史记录）
        self.right_frame = ttk.Frame(self.main_frame)
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=10, pady=10)
        self.prompt_label = ttk.Label(self.right_frame, text="提示词:")
        self.prompt_label.pack(pady=(10, 0))
        self.prompt_text = scrolledtext.ScrolledText(self.right_frame, height=10, width=40, font=("Arial", 10))
        self.prompt_text.pack(fill=tk.BOTH, expand=True)
        self.history_label = ttk.Label(self.right_frame, text="历史记录:")
        self.history_label.pack(pady=(10, 0))
        self.history_frame = ttk.Frame(self.right_frame)
        self.history_frame.pack(fill=tk.BOTH, expand=True)
        self.history_listbox = tk.Listbox(self.history_frame, height=5, font=("Arial", 10))
        self.history_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.history_scrollbar = ttk.Scrollbar(self.history_frame, orient=tk.VERTICAL, command=self.history_listbox.yview)
        self.history_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_listbox.config(yscrollcommand=self.history_scrollbar.set)
        self.history_listbox.bind("<<ListboxSelect>>", self.load_selected_history)
        self.load_prompt_and_history()
        self.toggle_wait_time_visibility()
        self.rootfsx.after(100, self.process_queue)
    def toggle_wait_time_visibility(self, *args):
        """
        根据并发选项的当前值，显示或隐藏等待时间设置框。
        """
        if self.concurrent_var.get() == "不并发":
            self.wait_time_frame.pack(pady=10, fill=tk.X)  # 显示等待时间设置框
        else:
            self.wait_time_frame.pack_forget()  # 隐藏等待时间设置框
    def center_window(self, width, height):
        screen_width = self.rootfsx.winfo_screenwidth()
        screen_height = self.rootfsx.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.rootfsx.geometry(f"{width}x{height}+{x}+{y}")

    def get_available_models(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        tree = ET.parse(self.token_file)
        root = tree.getroot()
        return [model.get("name") for model in root.findall("model")]

    def extract_json_part(self, text):
        start = text.find('```json')
        if start == -1:
            return None
        start += 5
        end = text.find('```', start)
        if end == -1:
            return None
        json_part = text[start:end].strip()
        return json_part

    def load_token(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        try:
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            selected_model = self.model_var.get()
            token = root.find(f"./model[@name='{selected_model}']/api_key").text
            if token:
                self.token_1=token

        except ET.ParseError:
            self.initialize_xml_file(self.token_file)

    def initialize_xml_file(self, token_file):
        root = ET.Element("data")

        # 创建 DeepSeek 模型部分
        deepseek_model = ET.SubElement(root, "model", name="DeepSeek")
        ET.SubElement(deepseek_model, "url").text = "https://qianfan.baidubce.com/v2/chat/completions"
        ET.SubElement(deepseek_model, "model").text = "deepseek-v3"
        ET.SubElement(deepseek_model, "api_key").text = "token"

        # 创建 Kimi 模型部分
        kimi_model = ET.SubElement(root, "model", name="Kimi")
        ET.SubElement(kimi_model, "url").text = "https://api.moonshot.cn/v1/chat/completions"
        ET.SubElement(kimi_model, "model").text = "moonshot-v1-8k"
        ET.SubElement(kimi_model, "api_key").text = "token"

        # 生成 XML 文件
        tree = ET.ElementTree(root)
        tree.write(token_file, encoding="utf-8", xml_declaration=True)
    def load_prompt_and_history(self):
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        try:
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            prompt = root.find("prompt")
            if prompt is not None and prompt.text:
                self.prompt_text.insert(tk.END, prompt.text)
            else:
                self.prompt_text.insert(tk.END, self.default_prompt_editable)
                self.save_prompt_and_history()
            history = root.find("history")
            if history is not None:
                for item in history.findall("item"):
                    self.history_listbox.insert(tk.END, item.text)
        except ET.ParseError:
            self.initialize_xml_file()

    def save_prompt_and_history(self):
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        tree = ET.parse(self.token_file)
        root = tree.getroot()
        prompt_element = root.find("prompt")
        if prompt_element is None:
            prompt_element = ET.SubElement(root, "prompt")
        prompt_element.text = prompt
        history_element = root.find("history")
        if history_element is None:
            history_element = ET.SubElement(root, "history")
        if prompt:
            history_items = [item.text for item in history_element.findall("item")]
            if prompt not in history_items:
                if len(history_items) >= 10:
                    history_element.remove(history_element[-1])
                new_item = ET.SubElement(history_element, "item")
                new_item.text = prompt
                new_item.tail = "\n"
        self.history_listbox.delete(0, tk.END)
        for item in history_element.findall("item"):
            self.history_listbox.insert(0, item.text)
        tree.write(self.token_file, encoding="utf-8", xml_declaration=True)

    def load_selected_history(self, event=None):
        selected_index = self.history_listbox.curselection()
        if selected_index:
            selected_prompt = self.history_listbox.get(selected_index)
            self.prompt_text.delete("1.0", tk.END)
            self.prompt_text.insert(tk.END, selected_prompt)

    def load_reference_file(self):
        self.reference_window = tk.Toplevel(self.rootfsx)
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.rootfsx)

        self.reference_window.title("参考实例")
        self.center_window_on_parent(self.reference_window, 600, 400)
        self.reference_tree = ttk.Treeview(self.reference_window, columns=("标题", "内容"), show="headings")
        self.reference_tree.heading("标题", text="标题")
        self.reference_tree.heading("内容", text="内容")
        self.reference_tree.pack(fill=tk.BOTH, expand=True)
        self.reference_tree.bind("<Double-1>", self.on_reference_double_click)
        self.reference_tree.bind("<Button-3>", self.show_reference_context_menu)
        self.load_reference_examples()

    def load_reference_examples(self):
        if not os.path.exists(self.reference_file):
            self.initialize_reference_file()
        try:
            tree = ET.parse(self.reference_file)
            root = tree.getroot()
            self.reference_examples = []
            for example in root.findall("example"):
                title = example.find("title").text
                content = example.find("content").text
                self.reference_examples.append({"title": title, "content": content})
                self.reference_tree.insert("", tk.END, values=(title, content))
        except ET.ParseError:
            self.initialize_reference_file()

    def initialize_reference_file(self):
        root = ET.Element("examples")
        tree = ET.ElementTree(root)
        tree.write(self.reference_file, encoding="utf-8", xml_declaration=True)

    def save_reference_examples(self):
        root = ET.Element("examples")
        for example in self.reference_examples:
            example_element = ET.SubElement(root, "example")
            ET.SubElement(example_element, "title").text = example["title"]
            ET.SubElement(example_element, "content").text = example["content"]
        tree = ET.ElementTree(root)
        tree.write(self.reference_file, encoding="utf-8", xml_declaration=True)

    def show_reference_context_menu(self, event):
        selected_item = self.reference_tree.identify_row(event.y)
        if selected_item:
            self.reference_tree.selection_set(selected_item)
            menu = tk.Menu(self.reference_window, tearoff=0)
            menu.add_command(label="新增", command=self.add_reference_example)
            menu.add_command(label="编辑", command=lambda: self.edit_reference_example(selected_item))
            menu.add_command(label="移除", command=lambda: self.remove_reference_example(selected_item))
            menu.post(event.x_root, event.y_root)
        else:
            menu = tk.Menu(self.reference_window, tearoff=0)
            menu.add_command(label="新增", command=self.add_reference_example)
            menu.post(event.x_root, event.y_root)

    def add_reference_example(self):
        self.edit_reference_example(None)

    def edit_reference_example(self, item):
        self.edit_window = tk.Toplevel(self.reference_window)
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.edit_window)

        self.edit_window.title("编辑参考实例")
        self.center_window_on_parent(self.edit_window, 400, 300)
        title_label = ttk.Label(self.edit_window, text="标题:")
        title_label.pack(pady=(10, 0))
        self.edit_title_entry = ttk.Entry(self.edit_window, width=40)
        self.edit_title_entry.pack(pady=5)
        content_label = ttk.Label(self.edit_window, text="内容:")
        content_label.pack(pady=(10, 0))
        self.edit_content_text = scrolledtext.ScrolledText(self.edit_window, height=10, width=40, font=("Arial", 10))
        self.edit_content_text.pack(fill=tk.BOTH, expand=True)
        if item:
            title, content = self.reference_tree.item(item, "values")
            self.edit_title_entry.insert(0, title)
            self.edit_content_text.insert(tk.END, content)
        save_button = ttk.Button(self.edit_window, text="保存", command=lambda: self.save_reference_example_edit(item))
        save_button.pack(pady=10)

    def save_reference_example_edit(self, item):
        title = self.edit_title_entry.get().strip()
        content = self.edit_content_text.get("1.0", tk.END).strip()
        if not title or not content:
            messagebox.showwarning("警告", "标题和内容不能为空")
            return
        if item:
            self.reference_tree.item(item, values=(title, content))
            index = self.reference_tree.index(item)
            self.reference_examples[index] = {"title": title, "content": content}
        else:
            self.reference_tree.insert("", tk.END, values=(title, content))
            self.reference_examples.append({"title": title, "content": content})
        self.save_reference_examples()
        self.edit_window.destroy()

    def remove_reference_example(self, item):
        index = self.reference_tree.index(item)
        self.reference_tree.delete(item)
        del self.reference_examples[index]
        self.save_reference_examples()

    def on_reference_double_click(self, event):
        selected_item = self.reference_tree.selection()
        if selected_item:
            title, content = self.reference_tree.item(selected_item, "values")
            self.prompt_text.delete("1.0", tk.END)
            self.prompt_text.insert(tk.END, content)

    def center_window_on_parent(self, window, width, height):
        parent_x = self.rootfsx.winfo_x()
        parent_y = self.rootfsx.winfo_y()
        parent_width = self.rootfsx.winfo_width()
        parent_height = self.rootfsx.winfo_height()
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        window.geometry(f"{width}x{height}+{x}+{y}")

    def start_processing(self):
        print('开始处理数据...')
        # 获取用户选择的列并存储为类的属性
        self.selected_columns = [self.column_tree.item(item, "values")[0] for item in self.column_tree.selection()]
        if not self.selected_columns:
            messagebox.showwarning("警告", "请先选择列")
            return
        try:
            self.batch_size = int(self.batch_size_entry.get().strip())
            if self.batch_size <= 0:
                raise ValueError("每批处理行数必须大于 0")
        except ValueError as e:
            messagebox.showerror("错误", f"每批处理行数无效: {e}")
            return
        
        self.save_prompt_and_history()
        self.process_button.config(state=tk.DISABLED)  # 禁用按钮
        
        # 获取当前选择的并发模式
        current_mode = self.concurrent_var.get()
        
        # 通过队列发送日志信息
        self.queue.put(("log", f"\n\n当前工作模式为：{current_mode}"))
        
        # 启动处理线程
        threading.Thread(
            target=self.process_data if current_mode == "并发" else self.process_data_sequential,
            daemon=True
        ).start()
        
    def process_data_sequential(self):
        #不并发
        try:
            selected_columns = [self.column_tree.item(item, "values")[0] for item in self.column_tree.selection()]
            self.df["数据链接代码"] = 'R' + (self.df.index + 1).astype(str)
            
            # 设置进度条最大值
            self.queue.put(("progress_max", len(self.df)))
            
            # 获取用户设置的等待时间
            try:
                wait_time = int(self.wait_time_entry.get().strip())
                if wait_time < 0:
                    raise ValueError("等待时间不能为负数")
            except ValueError as e:
                self.queue.put(("error", f"等待时间无效: {e}"))
                return
            
            result_dict = {}
            for i in range(0, len(self.df), self.batch_size):
                batch = self.df.iloc[i:i + self.batch_size]
                
                # 处理单个批次
                batch_result_dict = self.process_batch(
                    api_key=self.token_1,  # 传递API密钥
                    batch=batch,
                    start_index=i
                )
                result_dict.update(batch_result_dict)
                
                # 更新进度
                self.queue.put(("progress", i + len(batch)))
                
                # 不并发模式下，每条请求后等待用户设置的时间
                if self.concurrent_var.get() == "不并发":
                    time.sleep(wait_time)  # 使用用户设置的等待时间
            
            # 处理完成后更新数据
            self.df["分析结果"] = self.df["数据链接代码"].map(result_dict)
            self.queue.put(("display_df", self.df))
            
        except Exception as e:
            self.queue.put(("error", f"处理文件时出错: {str(e)}"))
        finally:
            # 确保发送 enable_button 消息
            self.queue.put(("enable_button",))
            print("不并发模式：执行完毕。")  # 调试日志


    def process_data(self):
        try:
            selected_columns = [self.column_tree.item(item, "values")[0] for item in self.column_tree.selection()]
            self.df["数据链接代码"] = 'R' + (self.df.index + 1).astype(str)
            
            # 设置进度条最大值
            self.queue.put(("progress_max", len(self.df)))
            
            result_dict = {}
            futures = []
            for i in range(0, len(self.df), self.batch_size):
                batch = self.df.iloc[i:i + self.batch_size]
                future = self.executor.submit(self.process_batch, self.token_1, batch, i)
                futures.append(future)
            
            # 等待所有任务完成
            for future in as_completed(futures):
                try:
                    batch_result_dict = future.result()
                    result_dict.update(batch_result_dict)
                except Exception as e:
                    print(f"处理批次时出错: {e}")
            
            # 处理完成后更新数据
            self.df["分析结果"] = self.df["数据链接代码"].map(result_dict)
            self.queue.put(("display_df", self.df))
            
        except Exception as e:
            self.queue.put(("error", f"处理文件时出错: {str(e)}"))
        finally:
            # 确保发送 enable_button 消息
            self.queue.put(("enable_button",))
            print("并发模式：执行完毕。")  # 调试日志

    def process_batch(self, api_key, batch, start_index):
        try:
            content = "请对以下内容进行规整：\n"
            for index, row in batch.iterrows():
                content += f"数据链接代码: {row['数据链接代码']}\n"
                # 使用 self.selected_columns 过滤列
                content += "\n".join([f"{col}: {row[col]}" for col in self.selected_columns]) + "\n\n"
            
            # 添加模型选择逻辑
            selected_model = self.model_var.get()
            
            response = self.call_api(api_key, content, selected_model)
            
            result = response["choices"][0]["message"]["content"]
            print(f"\n\n批次 {start_index + 1}-{start_index + len(batch)} API 返回结果:\n {result}")
            result = self.extract_json_part(result)
            result_dict = self.parse_and_fill_result(result)
            return result_dict
        except Exception as e:
            print(f"处理批次 {start_index + 1}-{start_index + len(batch)} 时出错: {e}")
            return {}

    def call_api(self, api_key, content, selected_model=None):
        """
        调用API处理数据
        :param api_key: API密钥
        :param content: 要处理的内容
        :param selected_model: 选择的模型名称（可选，默认为当前选择的模型）
        :return: API响应结果
        """
        # 如果未指定模型，使用当前选择的模型
        if selected_model is None:
            selected_model = self.model_var.get()

        # 检查API密钥
        if not api_key:
            raise ValueError("API Key 不能为空")

        # 加载模型配置
        if not os.path.exists(self.token_file):
            self.initialize_xml_file(self.token_file)
        
        try:
            tree = ET.parse(self.token_file)
            root = tree.getroot()
            
            # 查找模型配置
            model_element = root.find(f"./model[@name='{selected_model}']")
            if model_element is None:
                raise ValueError(f"模型 {selected_model} 未配置")
            
            # 获取API URL和模型名称
            url = model_element.find("url").text
            model_name = model_element.find("model").text
            
            # 准备请求头
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json; charset=utf-8"
            }
            
            # 构建提示词
            user_prompt = self.prompt_text.get("1.0", tk.END).strip()
            full_prompt = self.default_prompt_fixed + user_prompt

            # 构建请求数据
            data = {
                "model": model_name,
                "messages": [
                    {"role": "system", "content": full_prompt},
                    {"role": "user", "content": content}
                ],
                "stream": False
            }
            
            # 发送请求
            print('发送给大模型的数据：',data)
            response = requests.post(url, headers=headers, json=data)
            
            # 检查响应状态
            if response.status_code == 200:
                return response.json()
            else:
                raise Exception(f"API 调用失败: {response.status_code}, {response.text}")
                
        except ET.ParseError as e:
            raise Exception(f"解析配置文件时出错: {str(e)}")
        except requests.RequestException as e:
            raise Exception(f"网络请求失败: {str(e)}")
        except Exception as e:
            raise Exception(f"API调用过程中发生错误: {str(e)}")

    def parse_and_fill_result(self, result):
        try:
            result_dict = json.loads(result.strip().strip("```json").strip())
            return result_dict
        except Exception as e:
            self.queue.put(("log", f"解析 API 结果时出错: {e}"))
            return {}


    def process_queue(self):
        try:
            while True:
                try:
                    message = self.queue.get_nowait()
                    if message[0] == "log":
                        print(message[1])  # 控制台输出日志
                    elif message[0] == "progress_max":
                        self.progress["maximum"] = message[1]
                    elif message[0] == "progress":
                        self.progress["value"] = message[1]
                    elif message[0] == "display_df":
                        print('恢复按钮')
                        #self.process_button.config(state=tk.NORMAL)
                        PROGRAM_display_df_in_treeview(message[1], 0, 0)
  
                    elif message[0] == "enable_button":
                        self.process_button.config(state=tk.NORMAL)  # 恢复按钮状态
                        print("按钮状态已恢复")  # 调试日志
                    elif message[0] == "error":
                        messagebox.showerror("错误", message[1])
                except queue.Empty:
                    break
        except Exception as e:
            print(f"处理队列时出错: {e}")  # 捕获队列处理异常
        finally:
            self.rootfsx.after(100, self.process_queue)


class LLM_SQL:
    """
    大模型辅助SQL分析
    """
    def __init__(self, xmlpath):
        self.root = tk.Toplevel()
        
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
        self.root.title("SQL智能查询")
        self.tok_xml_path = xmlpath
        self.db_path = None
        self.columns = None
        self.models = None
        self.selected_model = None

        # 设置窗口居中
        self.center_window(800, 600)

        # 初始化 GUI 组件
        self.model_label = tk.Label(self.root, text="选择大模型:")
        self.model_label.grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.model_combobox = ttk.Combobox(self.root, state="readonly")
        self.model_combobox.grid(row=0, column=1, padx=10, pady=10, sticky="ew")

        self.db_label = tk.Label(self.root, text="选择 SQLite 数据库:")
        self.db_label.grid(row=1, column=0, padx=10, pady=10, sticky="w")
        self.db_entry = tk.Entry(self.root, width=50)
        self.db_entry.grid(row=1, column=1, padx=10, pady=10, sticky="ew")
        self.db_button = ttk.Button(self.root, text="打开", command=self.open_db)
        self.db_button.grid(row=1, column=2, padx=10, pady=10)

        self.prompt_label = tk.Label(self.root, text="输入提示词:")
        self.prompt_label.grid(row=2, column=0, padx=10, pady=10, sticky="w")
        self.prompt_text = tk.Text(self.root, height=10, width=60)
        self.prompt_text.grid(row=2, column=1, padx=10, pady=10, sticky="nsew")
        self.prompt_scroll = ttk.Scrollbar(self.root, orient="vertical", command=self.prompt_text.yview)
        self.prompt_scroll.grid(row=2, column=2, padx=(0, 10), pady=10, sticky="ns")
        self.prompt_text.configure(yscrollcommand=self.prompt_scroll.set)

        self.generate_button = ttk.Button(self.root, text="生成 SQL 查询", command=self.generate_sql)
        self.generate_button.grid(row=2, column=3, padx=10, pady=10)

        self.response_label = tk.Label(self.root, text="大模型返回的内容:")
        #self.response_label.grid(row=3, column=0, padx=10, pady=10, sticky="w")
        self.response_text = tk.Text(self.root, height=10, width=60)
        #self.response_text.grid(row=3, column=1, padx=10, pady=10, sticky="nsew")
        self.response_scroll = ttk.Scrollbar(self.root, orient="vertical", command=self.response_text.yview)
        #self.response_scroll.grid(row=3, column=2, padx=(0, 10), pady=10, sticky="ns")
        self.response_text.configure(yscrollcommand=self.response_scroll.set)

        self.sql_label = tk.Label(self.root, text="生成的 SQL 查询语句:")
        self.sql_label.grid(row=4, column=0, padx=10, pady=10, sticky="w")
        self.sql_text = tk.Text(self.root, height=10, width=60)
        self.sql_text.grid(row=4, column=1, padx=10, pady=10, sticky="nsew")
        self.sql_scroll = ttk.Scrollbar(self.root, orient="vertical", command=self.sql_text.yview)
        self.sql_scroll.grid(row=4, column=2, padx=(0, 10), pady=10, sticky="ns")
        self.sql_text.configure(yscrollcommand=self.sql_scroll.set)

        self.execute_button = ttk.Button(self.root, text="执行 SQL 查询", command=self.execute_sql)
        self.execute_button.grid(row=4, column=3, padx=10, pady=10)

        # 配置网格布局权重
        self.root.grid_rowconfigure(2, weight=1)
        self.root.grid_rowconfigure(3, weight=1)
        self.root.grid_rowconfigure(4, weight=1)
        self.root.grid_columnconfigure(1, weight=1)

        # 在 GUI 组件初始化完成后，再调用 open_tok_xml
        self.open_tok_xml()

    def center_window(self, width, height):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def parse_tok_xml(self, file_path):
        tree = ET.parse(file_path)
        root = tree.getroot()
        models = {}
        for model in root.findall('model'):
            name = model.get('name')
            url = model.find('url').text
            model_name = model.find('model').text
            api_key = model.find('api_key').text
            models[name] = {'url': url, 'model': model_name, 'api_key': api_key}
        return models

    def connect_db_and_get_columns(self, db_path):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        table_name = cursor.fetchone()[0]
        cursor.execute(f"PRAGMA table_info({table_name});")
        columns = [column[1] for column in cursor.fetchall()]
        conn.close()
        return columns

    def generate_sql_query(self, prompt, columns, model_info):
        url = model_info['url']
        model_name = model_info['model']
        api_key = model_info['api_key']
        
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model_name,
            "messages": [
                {
                    "role": "user",
                    "content": f"数据库表的列名有: {', '.join(columns)},表单名为table1。请根据以下提示生成SQL查询语句(注意要能在python内置的sqllite3中使用，不会报错）: {prompt}"
                }
            ],
            "temperature": 0.3,
            "max_tokens": 8000,
            "top_p": 0.9,
            "frequency_penalty": 1.0,
            "presence_penalty": 1.0
        }
        
        try:
            response = requests.post(url, headers=headers, json=data)
            response.raise_for_status()
            result = response.json()
            print("返回结果：",result)
            if "choices" in result and result["choices"]:
                return result["choices"][0]["message"]["content"].strip()
            else:
                messagebox.showerror("错误", "未获取到有效的响应内容！")
                return None
        except requests.exceptions.RequestException as e:
            messagebox.showerror("错误", f"请求失败：{e}")
            return None
        except Exception as e:
            messagebox.showerror("错误", f"处理返回结果时出错：{e}")
            return None

    def extract_sql_from_response(self, response):
        # 使用正则表达式提取 SQL 查询语句
        sql_pattern = re.compile(r"```sql\n(.*?)\n```", re.DOTALL)
        match = sql_pattern.search(response)
        if match:
            return match.group(1).strip()
        return None

    def execute_sql_query(self, db_path, query):
        conn = sqlite3.connect(db_path)
        try:
            # 使用 pandas 直接执行 SQL 查询并返回 DataFrame
            df = pd.read_sql_query(query, conn)
            return df
        except sqlite3.Error as e:
            messagebox.showerror("错误", f"SQL查询执行失败：{e}")
            return None
        except Exception as e:
            messagebox.showerror("错误", f"处理查询结果时出错：{e}")
            return None
        finally:
            conn.close()  # 确保连接被关闭

    def open_tok_xml(self):
        self.models = self.parse_tok_xml(self.tok_xml_path)
        self.model_combobox['values'] = list(self.models.keys())
        self.model_combobox.current(0)
    
    def open_db(self):
        self.db_path = filedialog.askopenfilename(filetypes=[("SQLite files", "*.db *.sqlite")])
        self.db_entry.delete(0, tk.END)
        self.db_entry.insert(0, self.db_path)
        self.columns = self.connect_db_and_get_columns(self.db_path)
        print("连接数据库成功", f"连接成功！表的列名为: {', '.join(self.columns)}")
        self.root.lift()
        self.root.attributes("-topmost", True)
        self.root.attributes("-topmost", False)
    
    def generate_sql(self):
        if not self.models or not self.columns:
            messagebox.showerror("错误", "请先选择 tok.xml 文件和数据库！")
            return
        
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        if not prompt:
            messagebox.showerror("错误", "请输入提示词！")
            return
        
        # 获取选择的模型
        selected_model_name = self.model_combobox.get()
        model_info = self.models[selected_model_name]
        
        # 调用大模型生成 SQL 查询
        response = self.generate_sql_query(prompt, self.columns, model_info)
        if response:
            self.response_text.delete("1.0", tk.END)
            self.response_text.insert("1.0", response)
            
            # 提取 SQL 查询语句
            sql_query = self.extract_sql_from_response(response)
            if sql_query:
                self.sql_text.delete("1.0", tk.END)
                self.sql_text.insert("1.0", sql_query)
            else:
                messagebox.showerror("错误", "未提取到有效的 SQL 查询语句！")
    
    def execute_sql(self):
        if not self.db_path:
            messagebox.showerror("错误", "请先选择数据库！")
            return
        
        sql_query = self.sql_text.get("1.0", tk.END).strip()
        if not sql_query:
            messagebox.showerror("错误", "请先生成 SQL 查询语句！")
            return
        
        result = self.execute_sql_query(self.db_path, sql_query)
        if result is not None:
            PROGRAM_display_df_in_treeview(result, 0, 0)  # 替换为你的显示数据的方法
            


def LLM_fenci(df):
    
    print(
        """
    **********************************************
    分词工具

    功能：
    1. 对指定列的内容进行分词处理。
    2. 支持使用停用词库过滤分词结果。
    3. 统计分词结果的词频，并显示匹配情况。

    使用方法：
    1. 启动程序后，选择需要分词的列。
    2. 选择停用词库（可选）。
    3. 点击“提交”按钮，程序会对选定列的内容进行分词，并统计词频。
    4. 分词结果会显示在 Treeview 中。

    配置表文件要求：
    1. 文件：`easy_ai.xlsx`。
    2. 文件格式：Excel 文件。
    3. 工作表内容：
       - 每个工作表代表一个停用词库，包含一列名为 "stopword" 的停用词。
       - 如果不需要停用词，可以选择 "无"。

    注意：
    - 如果未选择任何列，程序会提示警告。
    - 分词结果会显示在 Treeview 中，并统计词频。
    - 本功能依赖于jieba模块，如未安装，请安装。参考代码：python -m pip install jieba
    **********************************************
    """
    )
    try:
        import jieba
        import jieba.analyse
        import openpyxl
    except:
        print('jieba模块或openpyxl模块未安装。')    
        return
    def segment_text(text, stopwords):
        # 使用jieba进行分词，并去除停用词
        words = jieba.cut(text)
        filtered_words = [word for word in words if word not in stopwords and word.strip()]
        return filtered_words
    def update_df1(df1, df2):
        # 创建一个空的匹配情况列（如果还没有的话）
        if '匹配情况' not in df1.columns:
            df1['匹配情况'] = ''
        
        # 遍历 df2 中的每一行
        for value in df2['值']:
            # 检查 df1 中的关键词列是否包含 df2 中的值
            mask = df1['关键词'].str.contains(value, na=False)
            # 更新匹配情况列
            df1.loc[mask, '匹配情况'] =df1.loc[mask, '匹配情况'] +' ●'+ SMALL_get_list(value,0)[0]
  
        # 返回修改后的 df1
        return df1     
    def process_df(df1, df2, selected_columns):
        # 从df2中读取停用词
        stopwords = set(df2['stopword'].dropna().tolist())
     
        # 初始化一个空列表来存储分词结果和频次
        word_freq = []
     
        for col in selected_columns:
            for text in df1[col].astype(str).tolist():
                words = segment_text(text, stopwords)
                word_counter = Counter(words)
                word_freq.extend(word_counter.items())
     
        # 将结果转换为DataFrame
        word_freq_df = pd.DataFrame(word_freq, columns=['关键词', '频次'])
        word_freq_df = word_freq_df.groupby('关键词').sum().reset_index()
        word_freq_df = word_freq_df.sort_values(by='频次', ascending=False).reset_index(drop=True)


        try:
            dfx=SQL_gettable(memory_db,'easy_器械规整-SOC-关键词').reset_index(drop=True)
            word_freq_df['匹配情况']=""
            word_freq_df = update_df1(word_freq_df, dfx)
        except:
            pass     
        return word_freq_df



 
    # 创建主窗口

    root = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("选择分词列")
    if 1==1:
        sw = root.winfo_screenwidth()  
        sh = root.winfo_screenheight()  
        ww = 600  # 窗口宽度  
        wh = 800 # 窗口高度  
        x = (sw - ww) // 2  
        y = (sh - wh) // 2  
        root.geometry(f"{ww}x{wh}+{x}+{y}")
        
    label = ttk.Label(root, text='请选择需要分词的列和停用词库：')  
    label.pack( pady=5)   
        
    # 创建一个Frame来放置Treeview和Scrollbar
    frame = tk.Frame(root)
    frame.pack(fill='both', expand=True)
 
    # 设置Treeview的样式
    style = Style()
    style.configure("Treeview", rowheight=25)
 
    # Treeview组件，用于显示df的列名，并允许多选
    treeview = Treeview(frame, columns=("Column",), show='headings', selectmode='extended')
    treeview.heading("#0", text="选择列", anchor='w')
    treeview.heading("#1", text="Column", anchor='w')
 
    # 填充Treeview
    for col in df.columns:
        treeview.insert("", "end", values=(col,), iid=col, tags=("selectable",))
 
    # Scrollbar组件，用于Treeview的滚动
    scrollbar = Scrollbar(frame, orient="vertical", command=treeview.yview)
    treeview.configure(yscrollcommand=scrollbar.set)
 
    scrollbar.pack(side="right", fill="y")
    treeview.pack(side="left", fill="both", expand=True)
 
    # 提交按钮
    def on_submit(selected_items,use_stopwords_var):
        #print(selected_items)
        if not selected_items:
            messagebox.showwarning("警告", "未选择任何列！")
        else:
            selected_columns = [treeview.item(item, "values")[0] for item in selected_items]
            root.destroy()
            if use_stopwords_var=="":
                use_stopwords_var="无"
            df2=pd.read_excel(os.path.join(peizhidir,'easy_ai.xlsx'),sheet_name=use_stopwords_var).reset_index(drop=True)
            
            result=process_df(df, df2, selected_columns)
            
            result['报表类型']="{'fenci':" + str(selected_columns) + "}" 
            #print(result)
            PROGRAM_display_df_in_treeview(result,0,df)
            
    # 下拉框选择是否使用停用词
    use_stopwords_var = StringVar(value="无") 
    workbook = openpyxl.load_workbook(os.path.join(peizhidir,'easy_ai.xlsx'), read_only=True)
    sheet_namess = workbook.sheetnames
    dropdown = ttk.Combobox(frame, textvariable=use_stopwords_var, values=sheet_namess, state='readonly',width= 15)  
    dropdown.pack(padx=10) 

    
    submit_button = ttk.Button(root, text="提交", command=lambda:on_submit(treeview.selection(),dropdown.get()))
    submit_button.pack(pady=10)
 
    root.mainloop()
 


############################################################################################################################
#数据库函数
############################################################################################################################
class AAA_05_SQL():
    pass

# 生成加密密钥
def SQL_generate_key_from_password(password: str, salt: bytes = b'salt_value'):
    """
    从密码生成加密密钥。
    :param password: 用户定义的密码
    :param salt: 盐值（固定值或随机生成）
    :return: 加密密钥
    """
    key = PBKDF2(password.encode(), salt, dkLen=32, count=100000)
    return key

# 加密文件
def SQL_encrypt_file(input_file, output_file, key):
    """
    加密文件。
    :param input_file: 输入文件路径
    :param output_file: 输出文件路径
    :param key: 加密密钥
    """
    cipher = AES.new(key, AES.MODE_CBC)
    with open(input_file, 'rb') as f:
        data = f.read()
    padded_data = pad(data, AES.block_size)
    encrypted_data = cipher.encrypt(padded_data)
    with open(output_file, 'wb') as f:
        f.write(cipher.iv + encrypted_data)

# 解密文件到内存
def SQL_decrypt_file_to_memory(input_file, key):
    """
    解密文件到内存。
    :param input_file: 输入文件路径
    :param key: 加密密钥
    :return: 解密后的数据（字节流）
    """
    with open(input_file, 'rb') as f:
        iv = f.read(16)
        encrypted_data = f.read()
    cipher = AES.new(key, AES.MODE_CBC, iv=iv)
    decrypted_data = unpad(cipher.decrypt(encrypted_data), AES.block_size)
    return decrypted_data

# 函数1：保存所有xlsx文件的第一个单元表到db并加密
def SQL_save_xlsx_to_db_and_encrypt(password: str):
    """
    将当前目录下的所有xlsx文件的第一个单元表保存到SQLite数据库，并对数据库文件加密。
    :param password: 加密密码
    """
    global peizhidir
    # 获取当前目录下的所有xlsx文件
    xlsx_files = [f for f in os.listdir('.') if f.endswith('.xlsx')]
    if not xlsx_files:
        print("未找到任何xlsx文件！")
        return

    # 创建或连接到SQLite数据库
    db_file = filename=os.path.join(peizhidir,'packages.db')
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()

    # 遍历所有xlsx文件并保存到数据库
    for file in xlsx_files:
        table_name = os.path.splitext(file)[0]  # 去掉扩展名作为表名
        df = pd.read_excel(file, sheet_name=0)  # 只读取第一个单元表
        df.to_sql(table_name, conn, if_exists='replace', index=False)
        print(f"已保存文件 '{file}' 的第一个单元表到表 '{table_name}'")

    # 关闭数据库连接
    conn.close()

    # 对数据库文件进行加密
    key = SQL_generate_key_from_password(password)
    SQL_encrypt_file(db_file, 'packages_encrypted.db', key)
    print(f"数据库已加密并保存为 'packages_encrypted.db'")

    # 删除未加密的数据库文件
    os.remove(db_file)
    print(f"已删除未加密的数据库文件 '{db_file}'")

# 函数2：解密db文件并打印表单清单（不保存到硬盘）
def SQL_decrypt_db_and_print_tables(password: str):
    """
    解密数据库文件并打印表单清单。
    :param password: 解密密码
    """
    global peizhidir
    # 解密数据库文件到内存
    db_file = filename=os.path.join(peizhidir,'packages.db')
    key = SQL_generate_key_from_password(password)
    decrypted_data = SQL_decrypt_file_to_memory(db_file, key)

    # 将解密后的数据加载到内存中的SQLite数据库
    memory_db = sqlite3.connect(':memory:')
    with open('temp_decrypted.db', 'wb') as f:
        f.write(decrypted_data)
    disk_db = sqlite3.connect('temp_decrypted.db')
    disk_db.backup(memory_db)
    disk_db.close()
    os.remove('temp_decrypted.db')

    # 获取所有表名
    cursor = memory_db.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()

    # 打印表名
    if tables:
        print("默认资源库已连接。")
    else:
        print("默认资源库中没有表单！")

    # 返回内存数据库连接对象
    return memory_db

            
# 函数3：从内存数据库中获取指定表单并返回为DataFrame
def SQL_gettable(memory_db, table_name: str):
    """
    从内存数据库中获取指定表单并返回为DataFrame。
    如果当前目录下存在与表名对应的xlsx文件，则直接从xlsx文件中读取数据。
    :param memory_db: 内存数据库连接对象
    :param table_name: 表单名
    :return: 表单内容（DataFrame）
    """
    # 检查当前目录下是否存在与表名对应的xlsx文件
    global peizhidir


    xlsx_file=os.path.join(peizhidir, f"{table_name}.xlsx")
    #print(xlsx_file)
    if os.path.exists(xlsx_file):
        print(f"发现自定义文件 '{xlsx_file}'，忽略默认资源库，直接从文件中读取数据。")
        try:
            df = pd.read_excel(xlsx_file, sheet_name=0)  # 读取第一个单元表
            return df
        except Exception as e:
            print(f"从文件 '{xlsx_file}' 读取数据时出错：{e}")
            return None
    else:
        # 如果不存在对应的xlsx文件，则从内存数据库中获取数据
        # 将表单名用反引号括起来，避免特殊字符导致的语法错误
        query = f'SELECT * FROM `{table_name}`;'
        try:
            df = pd.read_sql_query(query, memory_db)
            return df
        except Exception as e:
            print(f"从默认资源库读取表单 '{table_name}' 时出错：{e}")
            return None


class SQL_PivotTool:
    """
    数据库分组透视
    """
    def __init__(self, db_path):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path)
        self.cursor = self.conn.cursor()
        self.table_name = self.get_first_table_name()

    def get_first_table_name(self):
        """获取数据库中的第一个表名"""
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        return self.cursor.fetchone()[0]

    def SMALL_count_mode(self, cursor, table_name, column_name, mode, group_by_columns):
        """
        对给定的 SQLite 表中的某一列进行计数，并根据 mode 参数指定的分隔符拆分复合症状并计算每个症状的计数。

        参数:
            cursor (sqlite3.Cursor): SQLite 游标对象。
            table_name (str): 表名。
            column_name (str): 列名。
            mode (str): 指定用于拆分复合症状的分隔符，可以是单个字符或多个字符。
                        如果传入 "count_all"，则使用常见符号（,；.、，；。）作为分隔符。
            group_by_columns (list): 分组列名列表。

        返回:
            dict: 格式为 {分组键: "症状1（计数1）、症状2（计数2）"} 的字典。
        """
        group_by_str = ", ".join([f"`{col}`" for col in group_by_columns])
        sql = f"""
            SELECT {group_by_str}, `{column_name}`
            FROM `{table_name}`
        """
        cursor.execute(sql)
        rows = cursor.fetchall()

        result = {}
        for row in rows:
            group_key = tuple(row[:-1])
            key = row[-1]
            if key is not None:
                if group_key not in result:
                    result[group_key] = {}
                if key in result[group_key]:
                    result[group_key][key] += 1
                else:
                    result[group_key][key] = 1

        filtered_result = {k: {k2: v2 for k2, v2 in v.items() if v2 > 0} for k, v in result.items()}

        if mode == "count_all":
            mode = r"[,；.、，；。]"

        if mode == "":
            final_result = {}
            for k, v in filtered_result.items():
                sorted_items = sorted(v.items(), key=lambda x: x[1], reverse=True)
                final_result["_".join(map(str, k))] = "、".join([f"{k2}（{v2}）" for k2, v2 in sorted_items])
            return final_result

        new_dict = {}
        for group_key, group_result in filtered_result.items():
            new_dict[group_key] = {}
            for key, value in group_result.items():
                symptoms = re.split(mode, key)
                for symptom in symptoms:
                    symptom = symptom.strip()
                    if symptom:
                        if symptom in new_dict[group_key]:
                            new_dict[group_key][symptom] += value
                        else:
                            new_dict[group_key][symptom] = value

        final_result = {}
        for k, v in new_dict.items():
            sorted_items = sorted(v.items(), key=lambda x: x[1], reverse=True)
            final_result["_".join(map(str, k))] = "、".join([f"{k2}（{v2}）" for k2, v2 in sorted_items])

        return final_result

    def create_pivot_tool(self, table_name, method):
        """
        数据透视工具的核心功能函数。

        参数:
            table_name (str): 表名。
            method (list): 包含行标签、列标签、值列和聚合方法等信息的列表。

        返回:
            pd.DataFrame: 生成的数据透视表结果。
        """
        row_labels = method[0]
        col_labels = method[1]
        value_cols = method[2]
        agg_methods = method[3]
        text_content = method[4]
        all_ratio = method[5]

        group_by_columns = row_labels + col_labels
        group_by_str = ", ".join([f"`{col}`" for col in group_by_columns])
        
        agg_functions = []
        smalL_count_columns = []
        for col, agg in zip(value_cols, agg_methods):
            if agg == "SMALL_count_mode":
                smalL_count_columns.append(col)
            elif agg == "count_all":
                smalL_count_columns.append(col)
            else:
                if agg == "nunique":
                    agg_functions.append(f"COUNT(DISTINCT `{col}`) AS `{col}_nunique`")
                elif agg == "sum":
                    agg_functions.append(f"SUM(`{col}`) AS `{col}_sum`")
                elif agg == "count":
                    agg_functions.append(f"COUNT(`{col}`) AS `{col}_count`")
                elif agg == "avg":
                    agg_functions.append(f"AVG(`{col}`) AS `{col}_avg`")
                elif agg == "min":
                    agg_functions.append(f"MIN(`{col}`) AS `{col}_min`")
                elif agg == "max":
                    agg_functions.append(f"MAX(`{col}`) AS `{col}_max`")
                else:
                    agg_functions.append(f"COUNT(`{col}`) AS `{col}_count`")

        sql = f"""
            SELECT {', '.join([f'`{col}`' for col in group_by_columns] + agg_functions)}
            FROM `{table_name}`
            GROUP BY {group_by_str}
        """
        self.cursor.execute(sql)

        base_result = self.cursor.fetchall()
        columns = [desc[0] for desc in self.cursor.description]
        df = pd.DataFrame(base_result, columns=columns)

        for col, agg in zip(value_cols, agg_methods):
            if agg == "SMALL_count_mode":
                smalL_result = self.SMALL_count_mode(self.cursor, table_name, col, "", group_by_columns)
            elif agg == "count_all":
                smalL_result = self.SMALL_count_mode(self.cursor, table_name, col, "count_all", group_by_columns)
            else:
                continue

            smalL_df = pd.DataFrame(
                list(smalL_result.items()), columns=["group_key", f"{col}_{agg}"]
            )
            df["group_key"] = df[group_by_columns].astype(str).agg("_".join, axis=1)
            df = pd.merge(df, smalL_df, on="group_key", how="left")
            df.drop("group_key", axis=1, inplace=True)
        df["报表类型"]="{'SQL':"+str(group_by_columns)+"}"    
        df['数据库文件']=self.db_path
        
        return df

    def create_pivot_tool_gui(self):
        """
        创建数据分组和透视工具的图形用户界面 (GUI)。
        """
        root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

        root.title("数据分组工具(DB)")

        sw = root.winfo_screenwidth()
        sh = root.winfo_screenheight()
        ww = 820
        wh = 200
        x = (sw - ww) // 2
        y = (sh - wh) // 2
        root.geometry(f"{ww}x{wh}+{x}+{y}")

        self.cursor.execute(f"PRAGMA table_info(`{self.table_name}`)")
        columns = [row[1] for row in self.cursor.fetchall()]

        row_vars = []
        row_labels = ["行标签1", "行标签2", "行标签3", "行标签4", "行标签5"]
        for i, label in enumerate(row_labels):
            tk.Label(root, text=label).grid(row=0, column=i)
            var = tk.StringVar(root)
            var.set("")
            dropdown = ttk.Combobox(root, textvariable=var, values=columns)
            dropdown.grid(row=1, column=i)
            row_vars.append(var)

        value_vars = []
        agg_method_vars = []
        value_labels = ["值列1", "值列2", "值列3", "值列4", "值列5"]
        for i, label in enumerate(value_labels):
            tk.Label(root, text=label).grid(row=2, column=i)
            var = tk.StringVar(root)
            var.set("")
            dropdown = ttk.Combobox(root, textvariable=var, values=columns)
            dropdown.grid(row=3, column=i)
            value_vars.append(var)

            methods = ["count", "sum", "avg", "min", "max", "nunique", "SMALL_count_mode", "count_all"]
            agg_var = tk.StringVar(root)
            agg_var.set("count")
            agg_dropdown = ttk.Combobox(root, textvariable=agg_var, values=methods)
            agg_dropdown.grid(row=4, column=i)
            agg_method_vars.append(agg_var)

        submit_button = ttk.Button(root, text="提交", command=lambda: self.on_submit(row_vars, value_vars, agg_method_vars))
        submit_button.grid(row=5, column=0, columnspan=5)

        root.mainloop()

    def on_submit(self, row_vars, value_vars, agg_method_vars):
        """
        提交按钮的回调函数，用于处理用户的选择并执行数据透视操作。

        参数:
            row_vars (list): 行标签的下拉菜单变量列表。
            value_vars (list): 值列的下拉菜单变量列表。
            agg_method_vars (list): 聚合方法的下拉菜单变量列表。
        """
        row_labels = [var.get() for var in row_vars if var.get() != ""]
        value_cols = [var.get() for var in value_vars if var.get() != ""]
        agg_methods = [var.get() for var in agg_method_vars if var.get() != ""]

        method = [row_labels, [], value_cols, agg_methods, {}, []]
        result_df = self.create_pivot_tool(self.table_name, method)
        PROGRAM_display_df_in_treeview(result_df,0,0)
        #print("数据透视表结果:")
        #print(result_df)

    def close(self):
        """关闭数据库连接"""
        self.conn.close()
        
def SQL_excels_to_db():
    """将用户选择的一个或多个文件（Excel、CSV、JSON）写入SQLite数据库，并打印行数"""
    def center_window(window, width, height):
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        window.geometry(f'{width}x{height}+{x}+{y}')

    def select_files():
        nonlocal file_paths
        file_types = [("Excel files", "*.xls *.xlsx"), ("CSV files", "*.csv"), ("JSON files", "*.json")]
        file_paths = filedialog.askopenfilenames(title="选择文件", filetypes=file_types)
        if file_paths:
            file_label.config(text=f"已选择 {len(file_paths)} 个文件")
        else:
            file_label.config(text="没有选择文件")

    def process_files():
        if not file_paths:
            messagebox.showwarning("警告", "请先选择文件！")
            return

        # 检查列是否一致
        columns = None
        for file in file_paths:
            if file.endswith('.xls') or file.endswith('.xlsx'):
                df = pd.read_excel(file, sheet_name=0, nrows=0)
            elif file.endswith('.csv'):
                df = pd.read_csv(file, nrows=0)
            elif file.endswith('.json'):
                df = pd.read_json(file, nrows=0)
            else:
                messagebox.showerror("错误", f"不支持的文件类型: {file}")
                return

            if columns is None:
                columns = df.columns
            elif not columns.equals(df.columns):
                messagebox.showerror("错误", "所选文件的列不一致。")
                return

        # 选择保存的数据库文件路径
        db_path = filedialog.asksaveasfilename(
            title="保存数据库文件",
            defaultextension=".db",
            filetypes=[("SQLite数据库文件", "*.db")]
        )

        if not db_path:
            messagebox.showwarning("警告", "未选择保存路径。")
            return

        # 如果数据库已存在，则删除
        if os.path.exists(db_path):
            os.remove(db_path)

        # 连接数据库
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        table_name = 'table1'

        # 创建表
        create_table_sql = "CREATE TABLE {} ({})".format(
            table_name,
            ', '.join(['"{}" TEXT'.format(col) for col in columns])
        )
        cursor.execute(create_table_sql)

        # 逐文件逐块写入数据库
        chunk_size = 10000  # 每次读取和写入的行数
        total_rows = 0

        for file in file_paths:
            print(f"处理文件: {file}")
            if file.endswith('.xls') or file.endswith('.xlsx'):
                excel_file = pd.ExcelFile(file)
                sheet_name = excel_file.sheet_names[0]  # 获取第一个工作表的名称
                print(f"读取工作表: {sheet_name}")
                sheet = excel_file.parse(sheet_name=sheet_name)
                total_sheet_rows = len(sheet)
                for start_row in range(0, total_sheet_rows, chunk_size):
                    end_row = min(start_row + chunk_size, total_sheet_rows)
                    chunk = excel_file.parse(sheet_name=sheet_name, skiprows=start_row, nrows=chunk_size)
                    chunk.to_sql(table_name, conn, if_exists='append', index=False)
                    total_rows += len(chunk)
                    print(f"已写入 {total_rows} 行")
            elif file.endswith('.csv'):
                for chunk in pd.read_csv(file, chunksize=chunk_size):
                    chunk.to_sql(table_name, conn, if_exists='append', index=False)
                    total_rows += len(chunk)
                    print(f"已写入 {total_rows} 行")
            elif file.endswith('.json'):
                df = pd.read_json(file)
                df.to_sql(table_name, conn, if_exists='append', index=False)
                total_rows += len(df)
                print(f"已写入 {total_rows} 行")

        # 关闭数据库连接
        conn.close()

        # 打印总行数
        messagebox.showinfo('构建成功', f"数据库 '{db_path}' 中的表 '{table_name}' 共有 {total_rows} 行。")

    # 创建Tkinter根窗口
    root = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("文件导入SQLite数据库")
    center_window(root, 400, 200)

    file_paths = []

    # 创建GUI组件
    tk.Label(root, text="选择文件类型: Excel, CSV, JSON").pack(pady=10)
    ttk.Button(root, text="选择文件", command=select_files).pack(pady=5)
    file_label = tk.Label(root, text="没有选择文件")
    file_label.pack(pady=5)
    ttk.Button(root, text="开始导入", command=process_files).pack(pady=10)

    # 运行Tkinter事件循环
    root.mainloop()        
  
def SQL_database_deduplication_tool():

    print(
        """
    **********************************************************************    
    SQLite 数据库去重工具

    功能：
    1. 选择 SQLite 数据库文件。
    2. 显示表中的所有列，允许用户多选列。
    3. 提供去重规则选项：保留第一个或最后一个重复项。
    4. 在执行去重操作前，自动备份数据库文件。
    5. 执行去重操作，并显示去重前后的行数。

    使用方法：
    1. 启动程序后，点击“选择数据库文件”按钮，选择一个 SQLite 数据库文件。
    2. 在 Treeview 中选择要去重的列（可多选）。
    3. 选择去重规则（保留第一个或最后一个）。
    4. 点击“确定”按钮，程序会自动备份数据库并执行去重操作。
    5. 操作完成后，程序会弹出去重前后的行数。

    注意：
    - 去重操作会直接修改数据库文件，请确保已备份重要数据。
    *************************************************************************
    """
    )

    def remove_duplicates(db_path, columns, keep_first=True):
        """
        去重函数

        参数：
        - db_path: 数据库文件路径。
        - columns: 需要去重的列名列表。
        - keep_first: 是否保留第一个重复项（默认为 True，保留第一个；False 则保留最后一个）。

        返回值：
        - before_count: 去重前的行数。
        - after_count: 去重后的行数。
        """
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        # 获取去重前的行数
        cursor.execute(f"SELECT COUNT(*) FROM table1")
        before_count = cursor.fetchone()[0]

        # 构建去重SQL语句
        columns_str = ", ".join(columns)
        order = "ROWID" if keep_first else "ROWID DESC"
        sql = f"""
        DELETE FROM table1
        WHERE ROWID NOT IN (
            SELECT MIN(ROWID)
            FROM table1
            GROUP BY {columns_str}
            ORDER BY {order}
        )
        """
        cursor.execute(sql)
        conn.commit()

        # 获取去重后的行数
        cursor.execute(f"SELECT COUNT(*) FROM table1")
        after_count = cursor.fetchone()[0]

        conn.close()

        return before_count, after_count

    def create_backup(db_path):
        """
        创建数据库备份文件

        参数：
        - db_path: 数据库文件路径。

        返回值：
        - backup_path: 备份文件的路径。
        """
        backup_path = db_path + ".backup"
        shutil.copyfile(db_path, backup_path)
        return backup_path

    def on_submit():
        """
        提交按钮的回调函数

        功能：
        1. 获取用户选择的列。
        2. 创建数据库备份。
        3. 执行去重操作。
        4. 显示去重前后的行数。
        """
        selected_columns = [tree.item(item, "text") for item in tree.selection()]
        if not selected_columns:
            messagebox.showerror("错误", "请选择至少一列")
            return

        keep_first = var.get() == 1

        # 创建备份文件
        backup_path = create_backup(db_path)
        messagebox.showinfo("备份成功", f"数据库已备份到: {backup_path}")

        # 执行去重操作
        before_count, after_count = remove_duplicates(db_path, selected_columns, keep_first)

        messagebox.showinfo("去重成功", f"去重前行数: {before_count}\n去重后行数: {after_count}")
        root.destroy()

    def select_database():
        """
        选择数据库文件的回调函数

        功能：
        1. 打开文件选择对话框，选择 SQLite 数据库文件。
        2. 读取表的列信息并显示在 Treeview 中。
        3. 启用提交按钮。
        """
        global db_path
        db_path = filedialog.askopenfilename(filetypes=[("SQLite数据库", "*.db")])
        if not db_path:
            return

        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(table1)")
        columns = [row[1] for row in cursor.fetchall()]
        conn.close()

        for col in columns:
            tree.insert("", "end", text=col)

        submit_button.config(state=tk.NORMAL)

    # 创建主窗口
    root = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("数据库去重工具")

    # 设置窗口大小
    window_width = 400
    window_height = 400
    root.geometry(f"{window_width}x{window_height}")

    # 居中窗口
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    center_x = int((screen_width - window_width) / 2)
    center_y = int((screen_height - window_height) / 2)
    root.geometry(f"+{center_x}+{center_y}")

    # 让窗口成为焦点并显示在最前面
    root.focus_force()
    root.attributes('-topmost', True)

    # 选择数据库文件按钮
    select_button = ttk.Button(root, text="选择数据库文件", command=select_database)
    select_button.pack(pady=10)

    # Treeview显示列
    tree = ttk.Treeview(root, selectmode="extended")
    tree.pack(fill=tk.BOTH, expand=True)

    # 去重规则选择
    var = tk.IntVar(value=0)  # 默认选择“保留最后一个”
    rule_frame = tk.Frame(root)
    rule_frame.pack(pady=10)
    tk.Radiobutton(rule_frame, text="保留第一个", variable=var, value=1).pack(side=tk.LEFT)
    tk.Radiobutton(rule_frame, text="保留最后一个", variable=var, value=0).pack(side=tk.LEFT)

    # 提交按钮
    submit_button = ttk.Button(root, text="确定", command=on_submit, state=tk.DISABLED)
    submit_button.pack(pady=10)

    # 运行主循环
    root.mainloop()



#df转DB
def SQL_df_to_sqlite_db_with_gui(df, table_name='table1'):  
    """
    df转DB
    """ 
    # 创建Tkinter窗口实例  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.withdraw()  # 隐藏主窗口  
  
    # 弹出文件选择对话框让用户选择数据库文件  
    db_path = filedialog.asksaveasfilename(defaultextension=".db",  
                                           filetypes=[("SQLite Database Files", "*.db"), ("All Files", "*.*")])  
  
    # 如果用户选择了文件  
    if db_path:  
        try:  
            # 连接到SQLite数据库（如果文件不存在，它会被创建）  
            conn = sqlite3.connect(db_path)   
            # 显示正在处理的消息  
            print(f"Saving DataFrame to {db_path} as table {table_name}...")   
            df.to_sql(table_name, conn, if_exists='replace', index=False)  
            # 关闭数据库连接  
            conn.close()  
  
            # 显示完成消息  
            messagebox.showinfo(title="提示", message=f"\n数据已保存到数据库文件： {db_path} 内部表单名： {table_name}.")  
  
        except sqlite3.Error as e:  
            messagebox.showinfo(title="提示", message=f"错误: {e}")  
    else:  
        messagebox.showinfo(title="提示", message="未提供文件。") 



def SQL_update_sqlite_db_with_df(df):
    """
    **********************************************************************
    功能：将给定的DataFrame数据追加到用户选择的SQLite数据库的表中。
    
    注意事项：
    1. 该函数会弹出一个文件选择对话框，用户需要选择一个SQLite数据库文件（.db）。
    2. 函数会自动获取数据库中的第一个表，并将DataFrame的数据追加到该表中。
    3. 在追加数据之前，函数会检查DataFrame的列名与数据库表的列名是否一致。
        - 如果列名不一致，函数会提示错误并停止执行。
        - 如果列名一致，数据将被追加到数据库表中。
    4. 追加数据后，数据库表不会自动去重。如果需要去重，请手动执行相关操作。
    5. 该函数不会修改DataFrame中的数据。
    6. 请确保DataFrame中的数据与数据库表的数据类型兼容，以避免插入错误。
    7. 在操作前，函数会自动备份数据库文件。
    **********************************************************************
    """
    print("***************************************************************")
    print("功能：将给定的DataFrame数据追加到用户选择的SQLite数据库的表中。")
    print("注意事项：")
    print("1. 请确保DataFrame的列名与数据库表的列名一致。")
    print("2. 数据追加后，数据库表不会自动去重。")
    print("3. 请确保DataFrame中的数据与数据库表的数据类型兼容。")
    print("4. 如果列名不一致，函数会提示错误并停止执行。")
    print("5. 在操作前，函数会自动备份数据库文件。\n")
    print("***************************************************************")
    # 创建Tkinter窗口实例  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.withdraw()  # 隐藏主窗口  

    # 弹出文件选择对话框让用户选择数据库文件  
    db_path = filedialog.askopenfilename(title="选择SQLite数据库文件", filetypes=[("SQLite数据库文件", "*.db"), ("所有文件", "*.*")])  

    # 如果用户选择了文件  
    if db_path:  
        try:  
            # 创建备份文件
            backup_path = db_path + ".backup"
            shutil.copyfile(db_path, backup_path)
            messagebox.showinfo(title="备份成功", message=f"数据库已备份到: {backup_path}")

            # 连接到SQLite数据库  
            conn = sqlite3.connect(db_path)  

            # 获取数据库第一个表的名称  
            cursor = conn.cursor()  
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name LIMIT 1;")  
            table_name = cursor.fetchone()[0]  

            # 从数据库中读取现有数据的列名  
            db_columns_query = cursor.execute(f"PRAGMA table_info({table_name})")  
            db_columns = [column[1] for column in db_columns_query.fetchall()]  

            # 获取DataFrame的列名  
            df_columns = list(df.columns)  

            # 检查列差异  
            columns_in_df_not_in_db = [col for col in df_columns if col not in db_columns]  
            columns_in_db_not_in_df = [col for col in db_columns if col not in df_columns]  

            # 打印列差异  
            if columns_in_df_not_in_db:  
                print(f"\n表单有的列但数据库没有的列: {columns_in_df_not_in_db}，您需要保证具有相同的列才能够追加。\n")  
            if columns_in_db_not_in_df:  
                print(f"表单没有的列但数据库有的列: {columns_in_db_not_in_df}，您需要保证具有相同的列才能够追加。\n")  
            if columns_in_df_not_in_db or columns_in_db_not_in_df:
                messagebox.showinfo(title="错误", message=f"表单和数据库的列名不统一，无法追加，详情请查看命令提示符。")  
            # 如果列名完全相同，则追加数据并进行去重  
            if not (columns_in_df_not_in_db or columns_in_db_not_in_df):  
                # 将DataFrame数据追加到数据库表  
                df.to_sql(table_name, conn, if_exists='append', index=False)  

                # 提交更改  
                conn.commit()  

                messagebox.showinfo(title="提示", message=f"数据已经追加到 {table_name} 表单 {db_path}。数据库没有去重，如果您需要去重，请另外执行相关功能。")  

            # 关闭数据库连接  
            conn.close()  

        except sqlite3.Error as e:  
            messagebox.showinfo(title="提示", message=f"错误: {e}")  
            if conn:  
                conn.close()  
    else:  
        messagebox.showinfo(title="提示", message="未选择文件。")  

#SQL查询
def SQL_create_query_gui(df, methon=None, conn=None):  
    """SQL查询工具GUI窗口，用于构建和执行SQL查询。"""  
    def add_query_condition(event=None):  
        # 获取选择的列名和输入的查询内容  
        column_name = column_dropdown.get()  
        query_value = query_entry.get()  
        
        # 如果两者都不为空，则构建查询条件并添加到文本框  
        if column_name and query_value:  
            # 对列名进行转义处理（用反引号或方括号括起来）
            escaped_column = f'"{column_name}"'  # SQLite使用双引号或方括号转义列名
            
            # 检查是否输入了算数运算符  
            match = re.match(r"([<>!=]=?)\s*(\d+)", query_value)  
            if match:  
                operator, value = match.groups()  
                # 将用户输入的运算符转换为SQL运算符  
                sql_operator = {  
                    "=": "=",  
                    "!=": "<>",  
                    ">": ">",  
                    "<": "<",  
                    ">=": ">=",  
                    "<=": "<="  
                }.get(operator)  
                if sql_operator:  
                    condition = f"{escaped_column} {sql_operator} {value}"  
                else:  
                    messagebox.showerror("Error", "Invalid operator. Please use >, <, >=, <=, !=, or =.")  
                    return  
            # 检查是否输入了日期范围  
            elif query_value.startswith('time(') and query_value.endswith(')'):  
                try:  
                    # 提取日期范围  
                    start, end = query_value[5:-1].split('-')  
                    # 格式化日期为YYYY-MM-DD  
                    start_date = f"{start[:4]}-{start[4:6]}-{start[6:]}"  
                    end_date = f"{end[:4]}-{end[4:6]}-{end[6:]}"  
                    # 构建日期范围查询条件  
                    condition = f"{escaped_column} BETWEEN '{start_date}' AND '{end_date}'"  
                except ValueError:  
                    # 如果日期格式不正确，显示错误消息  
                    messagebox.showerror("Error", "Invalid date range format. Please use 'time(YYYYMMDD-YYYYMMDD)'.")  
                    return  
            elif query_value.startswith("df['") and query_value.endswith("']"):  
                # 检查是否提供了 DataFrame 和列名  
                if df is None:  
                    messagebox.showerror("Error", "未提供 DataFrame。")  
                    return  
                inner_column_name = query_value[4:-2]  
                if inner_column_name not in df.columns:  
                    messagebox.showerror("Error", f"列名 '{inner_column_name}' 不存在于 DataFrame 中。")  
                    return  
                
                # 获取 DataFrame 中该列的唯一值  
                values_to_match = df[inner_column_name].dropna().unique().tolist()  
                if not values_to_match:  
                    messagebox.showerror("Error", f"列 '{inner_column_name}' 中没有有效数据。")  
                    return  
                
                # 如果值过多，提示性能问题  
                if len(values_to_match) > 900:  
                    print(f"警告：列 '{inner_column_name}' 中有 {len(values_to_match)} 个唯一值，可能会导致性能问题。")  
                    # 构建 SQL IN 子句  
                    in_clause = ', '.join([f"'{value}'" for value in values_to_match])  
                    condition = f"{escaped_column} IN ({in_clause})"  
                else:  
                    # 构建多个 LIKE 条件  
                    conditions = [f"{escaped_column} LIKE '%{value}%'" for value in values_to_match]  
                    condition = " OR ".join(conditions)  
            else:  
                # 构建查询条件  
                values_to_match = query_value.split("|")  
                conditions = [f"{escaped_column} LIKE '%{value}%'" for value in values_to_match]  
                condition = " OR ".join(conditions)  
        
            # 检查文本框是否为空，如果不为空，则添加AND连接符  
            if sql_text.get("1.0", tk.END).strip():  
                sql_text.insert(tk.END, f" AND ({condition})")  
            else:  
                sql_text.insert(tk.END, f"WHERE ({condition})")  
        
            # 清空输入框以便输入下一个条件  
            query_entry.delete(0, tk.END)
      
    def execute_combined_query():  
        # 获取完整的查询语句  
        query = f"SELECT * FROM {first_table} {sql_text.get('1.0', tk.END)}"  
        try:  
            data = pd.read_sql_query(query, conn)  
            PROGRAM_display_df_in_treeview(data, 0, 0)  
        except sqlite3.Error as e:  
            messagebox.showerror("Error", f"An error occurred: {e.args[0]}")  
      
    def return_query():  
        methon.delete("1.0", tk.END)  
        s = sql_text.get('1.0', tk.END)  
        # 查找第一个 "WHERE" 的位置  
        index = s.find("WHERE")  
        # 如果找到了 "WHERE"  
        if index != -1:  
            # 使用切片来删除第一个 "WHERE"  
            s = s[:index] + s[index + len("WHERE"):]  
        methon.insert(tk.END, s.strip())  
        root.destroy()  
      
    def populate_columns():  
        # 获取第一个表的列名并填充到下拉菜单  
        cursor = conn.cursor()  
        cursor.execute(f"PRAGMA table_info({first_table})")  
        columns = [row[1] for row in cursor.fetchall()]  
        column_dropdown['values'] = columns  
      
    def on_closing():  
        conn.close()  
        root.destroy()  
      
    # 打印使用方法  
    print("""
    *************************************************************************************
    SQL查询工具使用方法：
    1. 选择列名和输入查询条件。
    2. 支持的查询条件格式（在值列输入）：
       - 数值比较：>10, <20, >=30, <=40, !=50, =60 
       - 日期范围：time(YYYYMMDD-YYYYMMDD)
       - DataFrame列匹配：df['列名']  (利用目前载入的表单数据上的列，数据库中选择列包含表单特定列的所有数据)
       - 模糊匹配：value1|value2|value3
    3. 点击“添加条件”将条件添加到查询语句中。
    4. 点击“执行”按钮运行查询并显示结果。
    5. 点击“返回”按钮将查询语句返回到主程序。
    *************************************************************************************
    """)  
      
    # 如果未提供数据库连接，弹出文件选择对话框  
    if conn is None:  
        file_path = filedialog.askopenfilename(filetypes=[("SQLite Database Files", "*.db"), ("All Files", "*.*")])  
        if not file_path:  # 用户取消选择文件  
            print("用户取消选择文件。")  
            return default_data  
        conn = sqlite3.connect(file_path)  
      
    # 创建主窗口  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("SQL查询工具")  
    root.protocol("WM_DELETE_WINDOW", on_closing)  # 设置关闭窗口时的回调函数  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 570  # 窗口宽度  
    wh = 300  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
      
    # 获取数据库的第一个表名  
    cursor = conn.cursor()  
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name LIMIT 1;")  
    global first_table  # 使用全局变量以便在多个函数间共享  
    first_table = cursor.fetchone()[0]  
    # 
    # 创建带滚动条的文本框用于显示构建的查询  
    sql_text = scrolledtext.ScrolledText(root, width=40, height=10, wrap=tk.WORD)  
    sql_text.grid(row=0, column=0, columnspan=3, padx=2, pady=2, sticky=tk.W + tk.E)  
      
    # 创建列名下拉菜单  
    column_label = tk.Label(root, text="请选择列:")  
    column_label.grid(row=2, column=0, padx=10, sticky=tk.W)  
    column_dropdown = ttk.Combobox(root)  
    column_dropdown.grid(row=2, column=1, padx=10)  
    populate_columns()  


    qcolumn_label = tk.Label(root, text="值:")  
    qcolumn_label.grid(row=2, column=2, padx=10, sticky=tk.W)  

    # 创建输入框用于输入查询内容  
    query_entry = tk.Entry(root, width=30)  
    query_entry.grid(row=2, column=3, padx=10)  
    

    # 创建“增加”按钮  
    add_button = ttk.Button(root, text="添加条件", command=add_query_condition)  
    add_button.grid(row=3, column=3, padx=10)  
      
    # 创建执行按钮  
    execute_button = ttk.Button(root, text="执行", command=execute_combined_query)  
    execute_button.grid(row=3, column=0, columnspan=3, padx=10, pady=10)  
    

    
    if methon is not None:  
        return_button = ttk.Button(root, text="返回",  command=return_query)  
        return_button.grid(row=3, column=1, columnspan=3, padx=10, pady=10)  
      
    # 运行主循环  
    root.mainloop()


############################################################################################################################
#真实世界研究
############################################################################################################################
class AAA_06_RWS:
    pass


#数据编码
class RWS_data_encoding_gui:
    def __init__(self, df):
        """
        初始化数据编码工具。

        参数:
        df (pd.DataFrame): 需要编码的数据。
        """
        self.df = df
        self.operation_trace = []
        self.encoded_df = None  # 用于存储编码后的数据

        # 创建主窗口并隐藏
        self.root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)

        self.root.withdraw()  # 隐藏主窗口

        # 创建数据编码工具窗口
        self.top = tk.Toplevel()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.top)
        self.top.title("数据编码工具")
        self.setup_gui()

    def setup_gui(self):
        """设置GUI界面"""
        self.center_window(1200, 800)

        # 左侧面板
        left_frame = ttk.Frame(self.top)
        left_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        # 列选择
        column_frame = ttk.LabelFrame(left_frame, text="选择列")
        column_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(column_frame, text="选择需要编码的列:").grid(row=0, column=0, padx=5, pady=5)
        self.column_combobox = ttk.Combobox(column_frame, state="readonly")
        self.column_combobox.grid(row=0, column=1, padx=5, pady=5)
        self.column_combobox['values'] = self.df.columns.tolist()
        self.column_combobox.bind("<<ComboboxSelected>>", self.on_column_select)

        # 编码选项（单选）
        encoding_options_frame = ttk.LabelFrame(left_frame, text="编码选项")
        encoding_options_frame.pack(fill="x", padx=10, pady=5)

        # 使用单选按钮确保互斥
        self.encoding_option_var = tk.StringVar(value="none")  # 默认选择“无编码”（普通编码）

        ttk.Radiobutton(encoding_options_frame, text="普通编码", variable=self.encoding_option_var, value="none").pack(anchor="w")
        ttk.Radiobutton(encoding_options_frame, text="独热编码", variable=self.encoding_option_var, value="one_hot").pack(anchor="w")
        ttk.Radiobutton(encoding_options_frame, text="分段编码", variable=self.encoding_option_var, value="binning").pack(anchor="w")
        ttk.Radiobutton(encoding_options_frame, text="缩放", variable=self.encoding_option_var, value="scaling", command=self.on_scaling_toggle).pack(anchor="w")

        # 缩放方法选择（仅在缩放选项选中时显示）
        self.scaling_method_var = tk.StringVar(value="标准化")
        self.scaling_method_frame = ttk.Frame(encoding_options_frame)
        self.scaling_method_frame.pack(fill="x", padx=10, pady=5)

        ttk.Radiobutton(self.scaling_method_frame, text="标准化", variable=self.scaling_method_var, value="标准化").pack(anchor="w")
        ttk.Radiobutton(self.scaling_method_frame, text="归一化", variable=self.scaling_method_var, value="归一化").pack(anchor="w")

        # 默认隐藏缩放方法选择
        self.scaling_method_frame.pack_forget()

        # 编码规则输入
        self.encoding_frame = ttk.LabelFrame(left_frame, text="输入编码规则")
        self.encoding_frame.pack(fill="x", padx=10, pady=5)

        self.encoding_text = ScrolledText(self.encoding_frame, height=5)
        self.encoding_text.pack(fill="x", padx=5, pady=5)

        # 数据预览
        display_frame = ttk.LabelFrame(left_frame, text="数据预览（前10条）")
        display_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.display_text = ScrolledText(display_frame, height=10)
        self.display_text.pack(fill="both", expand=True, padx=5, pady=5)
        self.display_text.insert(tk.END, self.df.head(10).to_string())

        # 应用按钮
        apply_button = ttk.Button(left_frame, text="应用编码", command=self.apply_encoding)
        apply_button.pack(pady=10)

        # 右侧面板（操作轨迹）
        right_frame = ttk.Frame(self.top)
        right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        # 操作轨迹框
        operation_trace_frame = ttk.LabelFrame(right_frame, text="操作轨迹（JSON格式）")
        operation_trace_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.operation_trace_text = ScrolledText(operation_trace_frame, height=20)
        self.operation_trace_text.pack(fill="both", expand=True, padx=5, pady=5)

        # 操作轨迹按钮
        operation_trace_buttons_frame = ttk.Frame(right_frame)
        operation_trace_buttons_frame.pack(fill="x", padx=10, pady=5)

        save_button = ttk.Button(operation_trace_buttons_frame, text="保存操作轨迹", command=self.save_operation_trace)
        save_button.pack(side="left", padx=5, pady=5)

        load_button = ttk.Button(operation_trace_buttons_frame, text="打开操作轨迹", command=self.load_operation_trace)
        load_button.pack(side="left", padx=5, pady=5)

        replay_button = ttk.Button(operation_trace_buttons_frame, text="复现操作轨迹", command=self.replay_operation_trace)
        replay_button.pack(side="left", padx=5, pady=5)

        # 关闭窗口时触发回调
        self.top.protocol("WM_DELETE_WINDOW", self.on_close)

    def center_window(self, width, height):
        """设置窗口居中"""
        screen_width = self.top.winfo_screenwidth()
        screen_height = self.top.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.top.geometry(f"{width}x{height}+{x}+{y}")

    def on_column_select(self, event):
        """当用户选择列时，更新编码规则输入框"""
        self.update_encoding_rules()

    def on_scaling_toggle(self):
        """当用户选择缩放选项时，显示缩放方法选择框"""
        if self.encoding_option_var.get() == "scaling":
            self.scaling_method_frame.pack(fill="x", padx=10, pady=5)
        else:
            self.scaling_method_frame.pack_forget()

    def update_encoding_rules(self):
        """根据当前选择的列和选项，更新编码规则输入框"""
        selected_col = self.column_combobox.get()
        if not selected_col:
            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, "请先选择需要编码的列。")
            return

        encoding_option = self.encoding_option_var.get()

        if encoding_option == "one_hot":
            # 独热编码提示
            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, 
                "独热编码将自动将选定的列转换为多个二进制列（1 和 0），每个唯一值对应一列。\n"
                "无需手动输入编码规则。"
            )
        elif encoding_option == "scaling" and pd.api.types.is_numeric_dtype(self.df[selected_col]):
            # 缩放提示
            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, 
                f"缩放将对列 '{selected_col}' 进行 {self.scaling_method_var.get()} 处理。\n"
                "无需手动输入编码规则。"
            )
        elif encoding_option == "binning" and pd.api.types.is_numeric_dtype(self.df[selected_col]):
            # 分段编码建议
            min_val = self.df[selected_col].min()
            max_val = self.df[selected_col].max()
            bin_size = (max_val - min_val) / 3
            suggested_rules = "\n".join([
                f"{min_val + i * bin_size}-{min_val + (i + 1) * bin_size}:{i}"
                for i in range(3)
            ])
            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, suggested_rules)
        else:
            # 普通编码建议
            unique_values = self.df[selected_col].unique()
            suggested_rules = "\n".join([f"{value}:{i}" for i, value in enumerate(unique_values)])
            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, suggested_rules)

    def apply_encoding(self, skip_trace=False):
        """应用编码规则并更新数据"""
        try:
            selected_col = self.column_combobox.get()
            if not selected_col:
                messagebox.showerror("错误", "请选择需要编码的列！")
                return

            encoding_option = self.encoding_option_var.get()

            if not skip_trace:
                operation = {
                    "column": selected_col,
                    "encoding_option": encoding_option,
                    "encoding_rules": self.encoding_text.get("1.0", tk.END).strip(),
                    "scaling_method": self.scaling_method_var.get() if encoding_option == "scaling" else None,
                }

            if encoding_option == "scaling" and pd.api.types.is_numeric_dtype(self.df[selected_col]):
                scaling_method = self.scaling_method_var.get()
                if scaling_method == "标准化":
                    scaler = StandardScaler()
                elif scaling_method == "归一化":
                    scaler = MinMaxScaler()
                else:
                    messagebox.showerror("错误", "请选择缩放方法！")
                    return

                new_col_name = f"{selected_col}_缩放后"
                self.df[new_col_name] = scaler.fit_transform(self.df[[selected_col]])
                print(f"列 '{selected_col}' 已进行{scaling_method}处理，新列 '{new_col_name}' 已生成！")
            elif encoding_option == "binning" and pd.api.types.is_numeric_dtype(self.df[selected_col]):
                encoding_rules = self.encoding_text.get("1.0", tk.END).strip()
                if not encoding_rules:
                    messagebox.showerror("错误", "请输入编码规则！")
                    return

                rules = {}
                for line in encoding_rules.split("\n"):
                    if ":" in line:
                        key, value = line.split(":", 1)
                        lower, upper = map(float, key.strip().split("-"))
                        rules[(lower, upper)] = value.strip()

                new_col_name = f"{selected_col}_分段编码后"
                def bin_value(x):
                    for (lower, upper), code in rules.items():
                        if lower <= x <= upper:
                            return code
                    return None
                self.df[new_col_name] = self.df[selected_col].apply(bin_value)
                print(f"列 '{selected_col}' 分段编码完成，新列 '{new_col_name}' 已生成！")
            elif encoding_option == "one_hot":
                unique_values = self.df[selected_col].nunique()
                if unique_values > 2:
                    # 保留源数据列，生成独热编码列
                    self.df = pd.concat([self.df, pd.get_dummies(self.df[selected_col], prefix=selected_col, drop_first=True, dtype=int)], axis=1)
                    print(f"列 '{selected_col}' 已进行独热编码！")
                else:
                    messagebox.showinfo("警告", f"列 '{selected_col}' 只有 {unique_values} 个分类，独热编码可能不适用。")
            else:
                # 普通编码
                encoding_rules = self.encoding_text.get("1.0", tk.END).strip()
                if not encoding_rules:
                    messagebox.showerror("错误", "请输入编码规则！")
                    return

                rules = {}
                for line in encoding_rules.split("\n"):
                    if ":" in line:
                        key, value = line.split(":", 1)
                        if pd.api.types.is_numeric_dtype(self.df[selected_col]):
                            key = float(key.strip()) if "." in key else int(key.strip())
                        else:
                            key = key.strip()
                        rules[key] = value.strip()

                new_col_name = f"{selected_col}_编码后"
                self.df[new_col_name] = self.df[selected_col].map(rules)
                print(f"列 '{selected_col}' 普通编码完成，新列 '{new_col_name}' 已生成！")

            self.display_text.delete(1.0, tk.END)
            self.display_text.insert(tk.END, self.df.head(10).to_string())

            if not skip_trace:
                self.operation_trace.append(operation)
                self.operation_trace_text.insert(tk.END, json.dumps(operation, ensure_ascii=False, indent=4) + "\n")
        except Exception as e:
            messagebox.showerror("错误", f"编码失败: {e}")

    def save_operation_trace(self):
        """保存操作轨迹到文件"""
        if not self.operation_trace:
            messagebox.showerror("错误", "操作轨迹为空！")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")],
            title="保存操作轨迹"
        )
        if file_path:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(self.operation_trace, f, ensure_ascii=False, indent=4)
            messagebox.showinfo("成功", "操作轨迹已保存！")

    def load_operation_trace(self):
        """从文件加载操作轨迹"""
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON文件", "*.json"), ("所有文件", "*.*")],
            title="打开操作轨迹"
        )
        if file_path:
            with open(file_path, "r", encoding="utf-8") as f:
                loaded_trace = json.load(f)
            self.operation_trace.clear()
            self.operation_trace.extend(loaded_trace)
            self.operation_trace_text.delete("1.0", tk.END)
            for op in loaded_trace:
                self.operation_trace_text.insert(tk.END, json.dumps(op, ensure_ascii=False, indent=4) + "\n")
            messagebox.showinfo("成功", "操作轨迹已加载！")

    def replay_operation_trace(self):
        """复现操作轨迹"""
        if not self.operation_trace:
            messagebox.showerror("错误", "操作轨迹为空！")
            return

        for op in self.operation_trace:
            self.column_combobox.set(op["column"])
            self.encoding_option_var.set(op["encoding_option"])
            self.update_encoding_rules()

            self.encoding_text.delete("1.0", tk.END)
            self.encoding_text.insert(tk.END, op["encoding_rules"])

            if op["encoding_option"] == "scaling":
                self.scaling_method_var.set(op["scaling_method"])

            self.apply_encoding(skip_trace=True)

        print("操作轨迹已复现！")

    def on_close(self):
        """关闭窗口时触发回调"""
        self.encoded_df = self.df  # 将编码后的数据存储到属性中
        self.top.destroy()
        self.root.quit()  # 退出主事件循环

    def run(self):
        """运行GUI并返回编码后的数据"""
        self.top.mainloop()  # 启动事件循环
        return self.encoded_df


#多因素回归
class RWS_MultiFactorRegressionGUI:
    def __init__(self, df):
        self.df = df
        self.root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
        self.root.title("回归分析工具")
        self.setup_gui()

    def setup_gui(self):
        self.center_window(800, 700)

        # 模型选择
        model_frame = ttk.LabelFrame(self.root, text="选择回归模型")
        model_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(model_frame, text="选择模型类型:").grid(row=0, column=0, padx=5, pady=5)
        self.model_combobox = ttk.Combobox(model_frame, values=["linear", "logistic", "multinomial", "poisson"], state="readonly")
        self.model_combobox.grid(row=0, column=1, padx=5, pady=5)
        self.model_combobox.current(0)

        # 变量选择
        var_frame = ttk.LabelFrame(self.root, text="选择变量")
        var_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(var_frame, text="结果变量:").grid(row=0, column=0, padx=5, pady=5)
        self.outcome_combobox = ttk.Combobox(var_frame, state="readonly")
        self.outcome_combobox.grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(var_frame, text="暴露变量:").grid(row=1, column=0, padx=5, pady=5)
        self.exposure_combobox = ttk.Combobox(var_frame, state="readonly")
        self.exposure_combobox.grid(row=1, column=1, padx=5, pady=5)
        ttk.Label(var_frame, text="混杂因素:").grid(row=2, column=0, padx=5, pady=5)
        self.covariates_listbox = tk.Listbox(var_frame, selectmode="multiple", width=50, height=6)
        self.covariates_listbox.grid(row=2, column=1, padx=5, pady=5)
        scrollbar = tk.Scrollbar(var_frame, orient="vertical", command=self.covariates_listbox.yview)
        scrollbar.grid(row=2, column=2, sticky="ns")
        self.covariates_listbox.config(yscrollcommand=scrollbar.set)

        # 填充变量名
        columns = self.df.columns.tolist()
        self.outcome_combobox['values'] = columns
        self.exposure_combobox['values'] = columns
        for col in columns:
            self.covariates_listbox.insert(tk.END, col)

        # 参数设置
        param_frame = ttk.LabelFrame(self.root, text="参数设置")
        param_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(param_frame, text="最大迭代次数 (maxiter):").grid(row=0, column=0, padx=5, pady=5)
        self.maxiter_entry = ttk.Entry(param_frame)
        self.maxiter_entry.grid(row=0, column=1, padx=5, pady=5)
        self.maxiter_entry.insert(0, "500")
        ttk.Label(param_frame, text="优化方法 (method):").grid(row=1, column=0, padx=5, pady=5)
        self.method_combobox = ttk.Combobox(param_frame, values=["bfgs", "newton", "lbfgs"], state="readonly")
        self.method_combobox.grid(row=1, column=1, padx=5, pady=5)
        self.method_combobox.current(0)

        # 结果输出
        result_frame = ttk.LabelFrame(self.root, text="终端输出")
        result_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.result_text = ScrolledText(result_frame, height=10)
        self.result_text.pack(fill="both", expand=True, padx=5, pady=5)

        # 运行按钮
        ttk.Button(
            self.root,
            text="运行回归分析",
            command=lambda: self.run_regression(
                self.model_combobox.get(),
                self.outcome_combobox.get(),
                self.exposure_combobox.get(),
                [self.covariates_listbox.get(i) for i in self.covariates_listbox.curselection()],
                self.maxiter_entry.get(),
                self.method_combobox.get()
            )
        ).pack(pady=10)

    def center_window(self, width, height):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def run_regression(self, model_type, outcome_var, exposure_var, covariates, maxiter, method):
        if not outcome_var or not exposure_var:
            messagebox.showerror("错误", "请选择结果变量和暴露变量！")
            return

        try:
            old_stdout = sys.stdout
            sys.stdout = io.StringIO()

            # 检查变量类型
            for var in [outcome_var, exposure_var] + covariates:
                if not pd.api.types.is_numeric_dtype(self.df[var]):
                    raise ValueError(f"变量 '{var}' 不是数值类型！")

            # 构建公式
            formula = f"{outcome_var} ~ {exposure_var}"
            if covariates:
                formula += " + " + " + ".join(covariates)

            print(f"正在运行 {model_type} 回归分析...")
            print(f"公式: {formula}")

            # 选择模型
            if model_type == "linear":
                model = ols(formula, data=self.df).fit()
            elif model_type == "logistic":
                model = logit(formula, data=self.df).fit(maxiter=int(maxiter), method=method)
            elif model_type == "multinomial":
                model = mnlogit(formula, data=self.df).fit(maxiter=int(maxiter), method=method)
            elif model_type == "poisson":
                model = poisson(formula, data=self.df).fit(maxiter=int(maxiter), method=method)
            else:
                raise ValueError("不支持的模型类型")

            # 提取回归结果
            results_summary = model.summary()
            print(results_summary)
            results_df = pd.read_html(results_summary.tables[1].as_html(), header=0, index_col=0)[0]
            results_df.columns = ["系数", "标准误差", "z值", "P值", "[0.025", "0.975]"]

            # 初始化 OR 值列
            results_df["OR值"] = "-"
            results_df["OR值_95%CI_下限"] = "-"
            results_df["OR值_95%CI_上限"] = "-"

            # 计算 OR 值
            results_df["结果归属"] = results_df.index
            for idx, row in results_df.iterrows():
                try:
                    results_df.at[idx, "OR值"] = np.exp(float(row["系数"]))
                    results_df.at[idx, "OR值_95%CI_下限"] = np.exp(float(row["[0.025"]))
                    results_df.at[idx, "OR值_95%CI_上限"] = np.exp(float(row["0.975]"]))
                except Exception as e:
                    print(f'无法计算OR: {e}, 行数据: {row}')
                    continue



            # 打印整合后的结果
            print("\n整合后的回归结果:")
            print(results_df)

            output = sys.stdout.getvalue()
            sys.stdout = old_stdout
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, output)
            PROGRAM_display_df_in_treeview(results_df,0,0)
        except Exception as e:
            sys.stdout = old_stdout
            messagebox.showerror("错误", f"回归分析失败: {str(e)}")

    def run(self):
        self.root.mainloop() 



#倾向性评分
class RWS_PropensityScoreAnalysisGUI:
    def __init__(self, df):
        """
        初始化倾向性评分分析工具。

        参数:
        df (pd.DataFrame): 需要分析的数据。
        """
        self.df = df
        self.root_pro = tk.Tk()  # 使用 root_pro 替代 root
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root_pro)

        self.root_pro.title("倾向性评分分析工具")
        self.setup_gui()

    def setup_gui(self):
        """设置GUI界面"""
        self.center_window(800, 600)

        # 方法选择
        method_frame = ttk.LabelFrame(self.root_pro, text="选择倾向性评分方法")
        method_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(method_frame, text="选择方法:").grid(row=0, column=0, padx=5, pady=5)
        self.method_combobox = ttk.Combobox(method_frame, state="readonly")
        self.method_combobox.grid(row=0, column=1, padx=5, pady=5)
        self.method_combobox['values'] = ["匹配", "分层", "逆概率加权"]
        self.method_combobox.current(0)  # 默认选择第一个选项
        self.method_combobox.bind("<<ComboboxSelected>>", self.toggle_inputs)

        # 匹配方法选择
        self.matching_method_frame = ttk.LabelFrame(self.root_pro, text="选择匹配方法")
        self.matching_method_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(self.matching_method_frame, text="选择匹配方法:").grid(row=0, column=0, padx=5, pady=5)
        self.matching_method_combobox = ttk.Combobox(self.matching_method_frame, state="readonly")
        self.matching_method_combobox.grid(row=0, column=1, padx=5, pady=5)
        self.matching_method_combobox['values'] = ["最近邻匹配", "卡钳匹配"]
        self.matching_method_combobox.current(0)  # 默认选择第一个选项
        self.matching_method_combobox.bind("<<ComboboxSelected>>", self.toggle_inputs)

        # 卡钳值输入框
        self.caliper_frame = ttk.LabelFrame(self.root_pro, text="卡钳值（仅用于卡钳匹配）")
        self.caliper = tk.DoubleVar(value=0.2)  # 调整卡钳值
        ttk.Entry(self.caliper_frame, textvariable=self.caliper).pack(anchor="w")
        self.caliper_frame.pack_forget()  # 默认隐藏

        # 分层数选择
        self.num_strata_frame = ttk.LabelFrame(self.root_pro, text="选择分层数")
        self.num_strata = tk.IntVar(value=5)
        ttk.Entry(self.num_strata_frame, textvariable=self.num_strata).pack(anchor="w")
        self.num_strata_frame.pack_forget()  # 默认隐藏

        # 变量选择
        var_frame = ttk.LabelFrame(self.root_pro, text="选择变量")
        var_frame.pack(fill="x", padx=10, pady=5)

        ttk.Label(var_frame, text="处理变量:").grid(row=0, column=0, padx=5, pady=5)
        self.treatment_combobox = ttk.Combobox(var_frame, state="readonly")
        self.treatment_combobox.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(var_frame, text="结果变量:").grid(row=1, column=0, padx=5, pady=5)
        self.outcome_combobox = ttk.Combobox(var_frame, state="readonly")
        self.outcome_combobox.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(var_frame, text="混杂因素:").grid(row=2, column=0, padx=5, pady=5)

        # 添加滚动条和更宽的 Listbox
        self.covariates_listbox = tk.Listbox(var_frame, selectmode="multiple", width=50, height=6)
        self.covariates_listbox.grid(row=2, column=1, padx=5, pady=5)

        # 添加垂直滚动条
        scrollbar = tk.Scrollbar(var_frame, orient="vertical", command=self.covariates_listbox.yview)
        scrollbar.grid(row=2, column=2, sticky="ns")
        self.covariates_listbox.config(yscrollcommand=scrollbar.set)

        # 更新变量选择框
        self.update_variable_comboboxes()

        # 结果输出
        result_frame = ttk.LabelFrame(self.root_pro, text="分析结果")
        result_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.result_text = ScrolledText(result_frame, height=10)
        self.result_text.pack(fill="both", expand=True, padx=5, pady=5)

        # 运行按钮
        ttk.Button(self.root_pro, text="运行倾向性评分分析", command=self.run_analysis).pack(pady=10)

    def center_window(self, width, height):
        """设置窗口居中"""
        screen_width = self.root_pro.winfo_screenwidth()
        screen_height = self.root_pro.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root_pro.geometry(f"{width}x{height}+{x}+{y}")

    def update_variable_comboboxes(self):
        """更新变量选择框"""
        columns = self.df.columns.tolist()
        self.treatment_combobox['values'] = columns
        self.outcome_combobox['values'] = columns
        self.covariates_listbox.delete(0, tk.END)
        for col in columns:
            self.covariates_listbox.insert(tk.END, col)

    def run_analysis(self):
        """运行倾向性评分分析并显示结果"""
        treatment_col = self.treatment_combobox.get()
        outcome_col = self.outcome_combobox.get()
        covariates = [self.covariates_listbox.get(i) for i in self.covariates_listbox.curselection()]
        method = self.method_combobox.get()

        # 检查用户是否选择了所有变量
        if not treatment_col or not outcome_col or not covariates:
            messagebox.showerror("错误", "请选择所有变量！")
            return

        try:
            # 检查处理变量和结果变量是否为二分类
            if len(self.df[treatment_col].unique()) != 2:
                raise ValueError("处理变量必须是二分类变量（0 或 1）")
            if len(self.df[outcome_col].unique()) != 2:
                raise ValueError("结果变量必须是二分类变量（0 或 1）")

            # 检查混杂因素是否有缺失值
            if self.df[covariates].isnull().any().any():
                raise ValueError("数据中存在缺失值，请处理后再进行分析")

            # 清空结果显示框
            self.result_text.delete(1.0, tk.END)

            # 打印当前选择的方法和变量
            self.result_text.insert(tk.END, f"正在运行倾向性评分分析...\n")
            self.result_text.insert(tk.END, f"处理变量: {treatment_col}\n")
            self.result_text.insert(tk.END, f"结果变量: {outcome_col}\n")
            self.result_text.insert(tk.END, f"混杂因素: {covariates}\n")
            self.result_text.insert(tk.END, f"方法: {method}\n")

            # 生成倾向性评分
            formula = f"{treatment_col} ~ " + " + ".join(covariates)
            self.result_text.insert(tk.END, '正在使用逻辑回归生成倾向性评分\n')
            try:
                model = logit(formula, data=self.df).fit(disp=False)
            except Exception as e:
                raise ValueError(f"逻辑回归拟合失败: {e}")

            self.df['propensity_score'] = model.predict(self.df)

            # 检查逻辑回归拟合质量
            auc = roc_auc_score(self.df[treatment_col], self.df['propensity_score'])
            self.result_text.insert(tk.END, f"\n逻辑回归拟合质量 (AUC): {auc:.3f}\n")

            # 根据选择的方法处理数据
            results_df = None
            if method == "匹配":
                matching_method = self.matching_method_combobox.get()
                caliper = self.caliper.get() if matching_method == "卡钳匹配" else None
                results_df = self.propensity_score_matching(self.df, treatment_col, 'propensity_score', matching_method, caliper)
                self.result_text.insert(tk.END, "\n=== 匹配后的数据 ===\n")
                self.result_text.insert(tk.END, results_df.to_string() + "\n")
                PROGRAM_display_df_in_treeview(results_df,0,0)
            elif method == "分层":
                num_strata = self.num_strata.get()
                results_df = self.propensity_score_stratification(self.df, 'propensity_score', num_strata)
                self.result_text.insert(tk.END, "\n=== 分层后的数据 ===\n")
                self.result_text.insert(tk.END, results_df.to_string() + "\n")

                # 检查每层样本量
                strata_counts = results_df['strata'].value_counts()
                self.result_text.insert(tk.END, "\n=== 每层样本量 ===\n")
                self.result_text.insert(tk.END, strata_counts.to_string() + "\n")
                PROGRAM_display_df_in_treeview(results_df,0,0)
            elif method == "逆概率加权":
                results_df = self.inverse_probability_weighting(self.df, treatment_col, 'propensity_score')
                self.result_text.insert(tk.END, "\n=== 加权后的数据 ===\n")
                self.result_text.insert(tk.END, results_df.to_string() + "\n")

                # 检查权重分布
                self.result_text.insert(tk.END, "\n=== 权重分布 ===\n")
                self.result_text.insert(tk.END, results_df['weight'].describe().to_string() + "\n")
                PROGRAM_display_df_in_treeview(results_df,0,0)
            else:
                raise ValueError("方法必须是 '匹配'、'分层' 或 '逆概率加权'")

        except Exception as e:
            messagebox.showerror("错误", f"倾向性评分分析失败: {e}")
            self.result_text.insert(tk.END, f"错误: {e}\n")

    def propensity_score_matching(self, df, treatment_col, propensity_score_col, method='最近邻匹配', caliper=None):
        """倾向性评分匹配"""
        treated = df[df[treatment_col] == 1]
        control = df[df[treatment_col] == 0]

        print("\n=== 匹配前的样本量 ===")
        print(f"处理组: {len(treated)}, 对照组: {len(control)}")

        if method == '最近邻匹配':
            nbrs = NearestNeighbors(n_neighbors=1, algorithm='auto').fit(control[[propensity_score_col]])
            distances, indices = nbrs.kneighbors(treated[[propensity_score_col]])
            matched_control = control.iloc[indices.flatten()]
            matched_df = pd.concat([treated, matched_control])
        elif method == '卡钳匹配':
            if caliper is None:
                raise ValueError("卡钳匹配需要指定 caliper 值")
            nbrs = NearestNeighbors(n_neighbors=1, algorithm='auto').fit(control[[propensity_score_col]])
            distances, indices = nbrs.kneighbors(treated[[propensity_score_col]])
            within_caliper = distances.flatten() <= caliper
            if not within_caliper.any():
                raise ValueError("卡钳值过小，没有匹配对。请调整卡钳值。")
            matched_control = control.iloc[indices.flatten()[within_caliper]]
            matched_df = pd.concat([treated.iloc[within_caliper], matched_control])
        else:
            raise ValueError("匹配方法必须是 '最近邻匹配' 或 '卡钳匹配'")

        print("\n=== 匹配后的样本量 ===")
        print(f"处理组: {len(matched_df[matched_df[treatment_col] == 1])}, 对照组: {len(matched_df[matched_df[treatment_col] == 0])}")

        return matched_df

    def propensity_score_stratification(self, df, propensity_score_col, num_strata=5):
        """倾向性评分分层"""
        df['strata'] = pd.qcut(df[propensity_score_col], q=num_strata, labels=False)

        print("\n=== 分层后的样本量 ===")
        print(df['strata'].value_counts())

        return df

    def inverse_probability_weighting(self, df, treatment_col, propensity_score_col):
        """逆概率加权"""
        df['weight'] = np.where(df[treatment_col] == 1, 1 / df[propensity_score_col], 1 / (1 - df[propensity_score_col]))
        # 裁剪权重以避免极端值
        lower_bound = df['weight'].quantile(0.01)
        upper_bound = df['weight'].quantile(0.99)
        df['weight'] = df['weight'].clip(lower=lower_bound, upper=upper_bound)
        return df

    def toggle_inputs(self, event=None):
        """根据选择的方法显示或隐藏输入框"""
        method = self.method_combobox.get()
        if method == "匹配":
            self.matching_method_frame.pack(fill="x", padx=10, pady=5)
            matching_method = self.matching_method_combobox.get()
            if matching_method == "卡钳匹配":
                self.caliper_frame.pack(fill="x", padx=10, pady=5)
            else:
                self.caliper_frame.pack_forget()
            self.num_strata_frame.pack_forget()
        elif method == "分层":
            self.matching_method_frame.pack_forget()
            self.caliper_frame.pack_forget()
            self.num_strata_frame.pack(fill="x", padx=10, pady=5)
        elif method == "逆概率加权":
            self.matching_method_frame.pack_forget()
            self.caliper_frame.pack_forget()
            self.num_strata_frame.pack_forget()

    def run(self):
        """运行GUI主循环"""
        self.root_pro.mainloop()


   
def RWS_summarize_missing_values(df):
    """
    RWS质量评估：检测每列空值。
    """
    summary_data = []
    
    # 遍历每一列
    for column in df.columns:
        total_rows = df.shape[0]
        null_rows = df[column].isnull().sum()
        null_percentage = round((null_rows / total_rows) * 100, 2)
        
        # 将结果添加到列表中
        summary_data.append({
            '列名': column,
            '总行数': total_rows,
            '空值行数': null_rows,
            '空值比例': null_percentage
        })
    
    # 将列表转换为DataFrame
    summary_df = pd.DataFrame(summary_data)
    summary_df["报表类型"]="{'group_nan':['"+summary_df["列名"].astype(str)+str("']}") 
    return summary_df




############################################################################################################################
#中途用到的实用小函数
############################################################################################################################
class AAA_07_Small():
    pass

def SMALL_easyreadT(bos):  # 查看表格
    """行列互换查看表格"""
    bos[
        "#####分隔符#########"
    ] = "######################################################################"
    bos2 = bos.stack(dropna=False)
    bos2 = pd.DataFrame(bos2).reset_index()
    bos2.columns = ["序号", "条目", "详细描述T"]
    bos2["逐条查看"] = "逐条查看"
    bos2["报表类型"] = "逐条查看"
    return bos2


def SMALL_merge_dataframes(df1, df1_col, df2, df2_col='PT'):  
    '''文件合并中间函数 '''
    merged_df = pd.merge(df1, df2, left_on=df1_col, right_on=df2_col, how='left')  
      
    return merged_df 

def Small_update_df(df1, df2, key_column):
    '''更新表格 '''
    # 依据列在 df2 中存在的 df1 部分
    df1_existing = df1[df1[key_column].isin(df2[key_column])]
    # 共同列（除依据列外）
    common_columns = [col for col in df1_existing.columns if col != key_column and col in df2.columns]
    # 删除共同列
    df1_existing = df1_existing.drop(columns=common_columns)
    # 与 df2 合并
    merged_existing = pd.merge(df1_existing, df2[common_columns + [key_column]], on=key_column, how='left')
    # 依据列在 df2 中不存在的 df1 部分
    df1_non_existing = df1[~df1[key_column].isin(df2[key_column])]
    # 合并回 df1
    updated_df = pd.concat([merged_existing, df1_non_existing])
    return updated_df
    

    
def SMALL_read_and_merge_files():
    """打开文件选择对话框，读取CSV或Excel文件，并返回合并后的数据。"""
    # 忽略 openpyxl 的 UserWarning
    import warnings
    warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
    file_types = [("Excel files", "*.xls;*.xlsx"), ("CSV files", "*.csv"), ("All files", "*.*")]
    root = tk.Tk()  # 需要一个Tk root实例来运行filedialog
    root.withdraw()  # 隐藏Tk窗口
    file_paths = filedialog.askopenfilenames(title="请选择文件", filetypes=file_types)

    if not file_paths:
        print("用户取消载入文件。")
        return None  # 如果用户取消了对话框，返回 None

    # 读取文件
    all_data = []
    for file_path in file_paths:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        #为医疗器械合并导入做个兼容    
        if '报告编码' in df.columns and '注册证编号'  in df.columns  and  '注册证编号/曾用注册证编号'  not in df.columns:
            df['注册证编号/曾用注册证编号']= df['注册证编号'].copy()
           
        all_data.append(df)

    # 检查所有文件的列是否相同且顺序一致
    if len(all_data) > 1:
        reference_columns = all_data[0].columns
        for idx, df in enumerate(all_data[1:]):
            if not df.columns.equals(reference_columns):
                # 弹出提示框，询问用户是否采用兼容导入
                user_choice = messagebox.askyesno("列名不一致", "导入的文件列名不一致，是否采用兼容导入？")
                if not user_choice:
                    print("用户取消兼容导入。")
                    return None
                else:
                    # 采用兼容导入
                    # 找到列最多的文件，以其列顺序为基准
                    max_columns_index = max(range(len(all_data)), key=lambda i: len(all_data[i].columns))
                    base_columns = all_data[max_columns_index].columns.tolist()  # 基准列顺序

                    # 找到所有文件的列名并集
                    all_columns = set().union(*[set(df.columns) for df in all_data])  # 所有列的并集

                    # 找到新增列（基准列中没有的列），并按字母顺序排序
                    additional_columns = sorted(all_columns - set(base_columns))  # 新增列

                    # 最终列顺序：基准列 + 新增列
                    final_columns = base_columns + additional_columns

                    # 在每个文件中添加缺失的列，并按最终列顺序排序
                    for i, df in enumerate(all_data):
                        # 添加缺失的列
                        for col in all_columns - set(df.columns):
                            df[col] = None  # 添加缺失列，默认值为 None

                        # 添加“源文件”列
                        df['源文件'] = file_paths[i].split('/')[-1]  # 使用文件名填充

                        # 按最终列顺序排序
                        df = df[final_columns + ['源文件']]  # 确保“源文件”列在最后
                        all_data[i] = df  # 更新 DataFrame

                    # 合并所有文件
                    combined_data = pd.concat(all_data, ignore_index=True)

                    # 处理“报告编码”列
                    try:
                        combined_data.loc[:, '报告编码'] = combined_data['报告编码'].astype(str)
                    except KeyError:
                        pass  # 如果“报告编码”列不存在，忽略错误

                    # 删除 Unnamed 相关列
                    combined_data = combined_data.loc[:, ~combined_data.columns.str.contains("^Unnamed", case=False, na=False)]
                    return combined_data

    # 合并数据
    combined_data = pd.concat(all_data, ignore_index=True) if len(all_data) > 1 else all_data[0]

    # 删除 Unnamed 相关列
    combined_data = combined_data.loc[:, ~combined_data.columns.str.contains("^Unnamed", case=False, na=False)]

    # 处理“报告编码”列
    try:
        combined_data.loc[:, '报告编码'] = combined_data['报告编码'].astype(str)
    except KeyError:
        pass  # 如果“报告编码”列不存在，忽略错误

    return combined_data.reset_index(drop=True)
    

class Small_FilterDataWithGUI:
    '''数据筛选 -前K个'''
    def __init__(self, df):
        self.df = df.fillna("")  # 填充空值
        self.ori = df
        self.root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)

        self.root.title("数据处理器")
        self.root.geometry("400x350")  # 调整窗口大小
        self._center_window()
        self._create_widgets()

    def _center_window(self):
        """将窗口居中显示"""
        self.root.update_idletasks()
        w_width = self.root.winfo_width()
        w_height = self.root.winfo_height()
        s_width = self.root.winfo_screenwidth()
        s_height = self.root.winfo_screenheight()
        x = (s_width - 400) // 2
        y = (s_height - 350) // 2  # 调整窗口居中位置
        self.root.geometry(f"400x350+{x}+{y}")

    def _create_widgets(self):
        """创建GUI组件"""
        # 列选择部分
        column_frame = ttk.LabelFrame(self.root, text="选择列", padding=10)
        column_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

        ttk.Label(column_frame, text="选择列：").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # 使用ttk.Combobox替代OptionMenu
        self.column_var = tk.StringVar(value=self.df.columns[0])  # 默认选择第一列
        self.column_combobox = ttk.Combobox(
            column_frame, 
            textvariable=self.column_var,
            values=list(self.df.columns),
            state="readonly",
            width=20
        )
        self.column_combobox.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        # 筛选模式部分
        mode_frame = ttk.LabelFrame(self.root, text="筛选模式", padding=10)
        mode_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

        self.mode_options = ["按阈值筛选（N）", "按前K个筛选"]
        
        ttk.Label(mode_frame, text="筛选模式：").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # 使用ttk.Combobox替代OptionMenu
        self.mode_var = tk.StringVar(value=self.mode_options[0])  # 默认选择第一个选项
        self.mode_combobox = ttk.Combobox(
            mode_frame,
            textvariable=self.mode_var,
            values=self.mode_options,
            state="readonly",
            width=20
        )
        self.mode_combobox.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        # 参数输入部分
        input_frame = ttk.LabelFrame(self.root, text="参数输入", padding=10)
        input_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

        ttk.Label(input_frame, text="阈值（N）：").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.entry_threshold = ttk.Entry(input_frame)
        self.entry_threshold.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        ttk.Label(input_frame, text="前K个：").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.entry_top_k = ttk.Entry(input_frame)
        self.entry_top_k.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        # 处理数据按钮
        ttk.Button(self.root, text="处理数据", command=self.on_submit) \
            .grid(row=3, column=0, columnspan=2, padx=10, pady=10, sticky="ew")

    def process_column(self, cell, threshold=None, top_k=None):
        """处理单个单元格的内容"""
        if not isinstance(cell, str):  # 检查是否为字符串
            return ""  # 如果不是字符串，返回空字符串

        matches = re.findall(r'\s*([^、（）]+)\s*（\s*(\d+)\s*）\s*', cell)
        events = [(text.strip(), int(count.strip())) for text, count in matches if text.strip()]

        # 根据条件筛选
        if threshold is not None:
            events = [event for event in events if event[1] > threshold]
        elif top_k is not None:
            events = sorted(events, key=lambda x: (-x[1], x[0]))[:top_k]

        # 生成字符串（无事件时返回空字符串）
        return '、'.join([f"{text}（{count}）" for text, count in events]) if events else ""

    def process_df(self, df, column, threshold=None, top_k=None):
        """处理整个DataFrame的指定列"""
        new_df = df.copy()
        new_df[column] = new_df[column].apply(lambda x: self.process_column(x, threshold, top_k))
        return new_df

    def on_submit(self):
        """处理提交事件"""
        try:
            threshold = None
            top_k = None

            # 获取用户选择的模式
            selected_mode = self.mode_combobox.get()

            if selected_mode == "按阈值筛选（N）":
                threshold = int(self.entry_threshold.get())
            elif selected_mode == "按前K个筛选":
                top_k = int(self.entry_top_k.get())

            # 获取用户选择的列
            #selected_column = self.column_var.get()
            selected_column = self.column_combobox.get()

            # 处理数据
            result_df = self.process_df(self.df, selected_column, threshold, top_k)
            PROGRAM_display_df_in_treeview(result_df, 0, self.ori)
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数字！")
        except Exception as e:
            messagebox.showerror("错误", str(e))

    def run(self):
        """运行GUI"""
        self.root.mainloop()


        
def SMALL_expand_dict_column(df, column_name, nested_columns_to_preserve=None):
    """
    将指定列中的字典或字典格式的字符串拆分为新的列。
    对于指定的嵌套列（如 '不良反应'），将其视为文本，不进一步展开。
    :param df: 输入的 DataFrame
    :param column_name: 需要拆分的列名
    :param nested_columns_to_preserve: 需要保留为文本的嵌套列名列表（如 ['不良反应']）
    :return: 返回扩展后的 DataFrame
    """
    # 确保列存在
    if column_name not in df.columns:
        raise ValueError(f"列 '{column_name}' 不存在于 DataFrame 中！")

    # 如果没有指定需要保留的嵌套列，默认为空列表
    if nested_columns_to_preserve is None:
        nested_columns_to_preserve = []

    # 定义一个函数来解析字典或字符串
    def parse_cell(cell_value):
        if isinstance(cell_value, str):
            try:
                # 尝试解析为字典
                return ast.literal_eval(cell_value)
            except (ValueError, SyntaxError):
                # 如果解析失败，返回空字典
                return {}
        elif isinstance(cell_value, dict):
            return cell_value
        else:
            return {}

    # 定义一个函数来处理嵌套字典
    def handle_nested_dict(d):
        for key in nested_columns_to_preserve:
            if key in d and isinstance(d[key], (dict, list)):
                # 将嵌套字典或列表转换为字符串
                d[key] = str(d[key])
        return d

    # 解析列中的每个单元格
    parsed_dicts = df[column_name].apply(parse_cell)

    # 处理需要保留为文本的嵌套字典
    parsed_dicts = parsed_dicts.apply(handle_nested_dict)

    # 将解析后的字典转换为 DataFrame
    expanded_df = pd.DataFrame(parsed_dicts.tolist())

    # 将扩展后的 DataFrame 合并到原始 DataFrame
    df = df.drop(column_name, axis=1).join(expanded_df)

    return df
       
def SMALL_merge_dataframes(df1, df1_col, df2_path=None, df2_col='PT'):  
    '''文件合并中间函数-药品特制 '''
      #药品预制-加SOC（PT扩行-MedDra法） 
    # 如果df2_path为None，弹出文件选择对话框让用户选择文件  
    if df2_path is None:  
        root = tk.Tk()  
        root.withdraw()  # 隐藏主窗口  
        df2_path = filedialog.askopenfilename()  # 弹出文件选择对话框  
      
    # 读取df2文件（如果提供了路径） 
    elif isinstance(df2_path, pd.DataFrame):
        df2=df2_path
        #df2.reset_index(inplace=True)
    elif df2_path: 
        script_dir = os.path.dirname(os.path.abspath(__file__))  
        config_file = os.path.join(script_dir, df2_path) 
        df2 = pd.read_excel(config_file)  # 假设df2是CSV文件，如果是其他格式请相应修改  
    
    elif  isinstance(df2_path, pd.DataFrame):
        pass
    else:  
        raise ValueError("No path provided for df2 and no file selected.")  
    #df1.reset_index(inplace=True)    
    try:  
        df1 = df1.drop(['code', 'Chinese', '级别', 'PT','HLT', 'HLGT', 'SOC', '主SOC'], axis=1)  
    except:  
        pass
          
    # 使用merge函数将df2拼接到df1上（基于指定的列名）
    print(df1_col,df2_col)  
    #print(df2)
    merged_df = pd.merge(df1, df2, left_on=df1_col, right_on=df2_col, how='left')  
    try:  
        merged_df['SOC']= merged_df['SOC'].fillna('其他（未规整）')  
    except:  
        pass 
    return merged_df  


def SMALL_save_df_as_txt_files(df, directory):
    """将 DataFrame 的每一行保存为单独的 txt 文件"""
    if not directory:
        print("No directory selected.")
        return


    def generate_random_string(length=5, has_upper=True, has_lower=True, has_digit=True):
        """生成包含大写字母、小写字母和数字的随机字符串"""
        if has_upper:
            characters = string.ascii_uppercase
        else:
            characters = ''
        if has_lower:
            characters += string.ascii_lowercase
        if has_digit:
            characters += string.digits
        
        while True:
            random_string = ''.join(random.choice(characters) for _ in range(length))
            if (any(c in string.ascii_uppercase for c in random_string) and
                any(c in string.ascii_lowercase for c in random_string) and
                any(c in string.digits for c in random_string)):
                return random_string
     
    def sanitize_filename(filename):
        """清理文件名，删除不合规字符和空格"""
        # Windows 文件名不允许的字符集
        invalid_chars = r'\/:*?"<>|'
        # 替换不合规字符为空字符串，并删除多余的空格
        cleaned_filename = ''.join(c for c in filename if c not in invalid_chars).replace(' ', '_')
        # 确保文件名不以空格或点开头（虽然在这个上下文中不太可能出现）
        if cleaned_filename.startswith('_') or cleaned_filename.startswith('.'):
            cleaned_filename = cleaned_filename[1:]
        # 如果文件名变为空字符串，则返回一个默认名（例如 "file_")
        if not cleaned_filename:
            cleaned_filename = "file_"
        return cleaned_filename



    for index, row in df.iterrows():
        if '报告编码' in df.columns and '产品名称' in df.columns:
            file_name = f"{row['报告编码']}_{row['产品名称']}_{generate_random_string()}.txt"
        elif '唯一标识' in df.columns and '通用名称' in df.columns:
            file_name = f"{row['唯一标识']}_{row['通用名称']}_{generate_random_string()}.txt"
        else:
            first_col_value = str(row[df.columns[0]]).strip()  # 转换为文本并去除首尾空格
            file_name = f"{first_col_value}_{generate_random_string()}.txt"
        
        # 清理文件名
        cleaned_file_name = sanitize_filename(file_name)
        file_path = os.path.join(directory, cleaned_file_name)
        file_path=file_path.replace(" ", "")
        with open(file_path, 'w', encoding='utf-8') as file:
            for col in df.columns:
                file.write(f"{col}：{row[col]}\n")
        print(f"Saved {file_path}")
    print('全部文件已保存。')


def SMALL_apply_operation(df, mycols, col_name, operation):
    """
    对DataFrame中的两列应用任意运算。

    参数:
        df (pd.DataFrame): 输入的DataFrame。
        mycols (list): 包含两列名称的列表。
        col_name (str): 结果列的名称。
        operation (str): 运算式子，例如 'a+b' 或 'a-b'。

    返回:
        pd.DataFrame: 包含运算结果的DataFrame。
    """

    def extract_first_date(date_str):
        """
        从包含多个日期的字符串中提取第一个日期。
        例如：'2022-05-30;2022-05-30;2022-05-30;2022-05-30' -> '2022-05-30'
        """
        if pd.isna(date_str):
            return pd.NaT
        dates = date_str.split(';')
        first_date = dates[0].strip()
        return first_date
    def check_time_format(df, column_name):
        # 检查列中的数据是否可以转换为时间格式
        try:
            # 尝试解析第一行数据
            pd.to_datetime(df[column_name].dropna().iloc[0], format='%Y-%m-%d')
            return True
        except:
            try:
                pd.to_datetime(df[column_name].dropna().iloc[0], format='%Y-%m-%d %H:%M:%S')
                return True
            except:
                return False

    col1, col2 = mycols

    # 检测输入列是否为时间格式
    if check_time_format(df, col1) and check_time_format(df, col2):
        # 提取第一个日期
        df[col1] = df[col1].apply(extract_first_date)
        df[col2] = df[col2].apply(extract_first_date)

        # 尝试解析日期，使用更灵活的格式
        time1 = pd.to_datetime(df[col1], errors='coerce', format='mixed')
        time2 = pd.to_datetime(df[col2], errors='coerce', format='mixed')

        if time1.isnull().all() or time2.isnull().all():
            raise ValueError("日期列解析失败，请检查日期格式是否正确。")

        is_time = True
    else:
        is_time = False

    # 根据运算式子执行相应的计算
    if is_time:
        # 如果是时间运算，计算时间差（以天为单位）
        if operation == "a-b":
            time_diff = (time1 - time2).dt.days
        elif operation == "b-a":
            time_diff = (time2 - time1).dt.days
        else:
            raise ValueError("时间只支持相减。")
        df[col_name] = time_diff
    else:
        # 如果是非时间运算，使用eval()函数执行运算
        try:
            df[col_name] = eval(operation, {"__builtins__": None}, {'a': df[col1], 'b': df[col2]})
        except Exception as e:
            raise ValueError(f"无法执行运算：{e}")

    return df



    
def SMALL_find_based_on_expression(df,expression):  
    """
    条件查找，范例："(df['a']>5) | ((3<=df['a']) & (df['a']<10)) | (df['d']!=6) | (df['e']==7) | (df['f']=='r')
    """
    #条件查找，范例："(df['a']>5) | ((3<=df['a']) & (df['a']<10)) | (df['d']!=6) | (df['e']==7) | (df['f']=='r')"  
    try:  
        mask = pd.eval(expression, engine='python', parser='pandas', local_dict={'df': df})  
    except Exception as e:  
        raise ValueError(f"无法解析表达式: {e}")  
          
    # 返回符合条件的行  
    return df[mask]  
  

  
def SMALL_assign_value_based_on_expression(df, target_column, expression, value):  
    """  
    根据给定的表达式给DataFrame的target_column列赋值value。  
    value 可以是数字、文本或者一个关于df的表达式（如 "df['a']" 或 "df['a']/12"）。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        target_column (str): 需要赋值的列名。  
        expression (str): 用于生成布尔掩码的表达式字符串。  
        value: 需要赋给满足表达式条件的行的值，可以是数字、字符串或者关于df的表达式。  
  
    返回:  
        pd.DataFrame: 更新后的DataFrame。  
    """  
    try:  
        # 使用pd.eval来评估表达式并生成布尔掩码  
        mask = pd.eval(expression, engine='python', parser='pandas', local_dict={'df': df})  
        # 确保mask是布尔类型的Series  
        if not isinstance(mask, pd.Series) or not mask.dtype == bool:  
            raise ValueError("表达式返回的不是有效的布尔序列")  
  
        # 如果value是一个字符串并且包含"df", 则尝试将其解析为一个关于df的表达式  
        if isinstance(value, str) and "df" in value:  
            try:  
                value_series = pd.eval(value, engine='python', parser='pandas', local_dict={'df': df})  
                if not isinstance(value_series, pd.Series):  
                    raise ValueError("value 表达式返回的不是有效的pandas Series")  
            except Exception as e:  
                raise ValueError(f"无法解析value表达式: {e}")  
        # 如果value不是字符串或者不包含"df", 则直接将其视为一个常量值  
        else:  
            value_series = pd.Series(value, index=df.index)  # 创建一个常量Series，以便可以广播到目标列  
  
        # 根据掩码给目标列赋值  
        df.loc[mask, target_column] = value_series  
        return df  
  
    except Exception as e:  
        raise ValueError(f"无法解析表达式或赋值: {e}")  

   

    
def SMALL_last_non_null_value(df, columns_list, new_column_name):  
    """  
    在DataFrame df中创建一个新列，其值为columns_list中最后一个非空值。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        columns_list (list of str): 要检查的列名列表，按顺序。  
        new_column_name (str): 新列的名称。  
          
    返回:  
        pd.DataFrame: 包含新列的原始DataFrame。  
    """  
      
    # 遍历每一行，找到最后一个非空值  
    def get_last_non_null(row):  
        for col in reversed(columns_list):  # 从后向前遍历列  
            if pd.notna(row[col]):  
                return row[col]  
        return np.nan  # 如果没有非空值，则返回NaN  
  
    # 应用函数并创建新列  
    df[new_column_name] = df.apply(get_last_non_null, axis=1)  
      
    return df  


def SMALL_calculate_ratios(df, list1, list2):  
    """  
    计算构成比。  
    """  
    # 确保输入列表长度相等  
    if len(list1) != len(list2):  
        raise ValueError("两个列表的长度必须相等")  
  
    # 遍历列表中的列名，并计算比值  
    for i in range(len(list1)):  
        col_name1 = list1[i]  
        col_name2 = list2[i]  
  
        # 确保列名在DataFrame中存在  
        if col_name1 not in df.columns or col_name2 not in df.columns:  
            raise ValueError(f"列名 {col_name1} 或 {col_name2} 不在DataFrame中")  
  
        # 计算比值并保留两位小数  
        df[f"{col_name1}_{col_name2}_ratio"] = df[col_name1] / df[col_name2]*100  
        df[f"{col_name1}_{col_name2}_ratio"] = df[f"{col_name1}_{col_name2}_ratio"].round(2)  
  
    return df 

def SMALL_expand_dict_like_columns(df):   
    """  
    转化字典列。  
    """  
    def format_dict_like_values(value):    
        if isinstance(value, dict):    
            # 使用列表推导式生成格式化后的键值对字符串，然后用' '.join连接  
            formatted_items = [f"{k}（{v}）" for k, v in value.items()]  
            # 使用'、'.join连接字符串，并使用rstrip移除尾部的'、'  
            return '、'.join(formatted_items).rstrip('、')  
        elif isinstance(value, str) and value.startswith('{') and value.endswith('}'):    
            try:    
                data_dict = eval(value)  # 注意：使用eval有风险，应确保数据安全  
                if isinstance(data_dict, dict):    
                    formatted_items = [f"{k}（{v}）" for k, v in data_dict.items()]  
                    return '、'.join(formatted_items).rstrip('、')  
            except:    
                pass    
        return value  # 如果不是字典或类似字典的字符串，则返回原始值    
        
    df_transformed = df.copy()    
    for column in df_transformed.columns:    
        df_transformed[column] = df_transformed[column].apply(format_dict_like_values)    
    return df_transformed  
      
def SMALL_add_composition_ratio(df, column_name, new_column_name='构成比'):  
    """  
    在原始DataFrame（除最后一行合计外）中添加指定列的构成比。  
    这里假设column_name列已经包含了频数，且最后一行是合计。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        column_name (str): 需要计算构成比的列名，该列应包含频数（除最后一行）。  
        new_column_name (str): 新添加的构成比列名，默认为'构成比'。  
          
    返回:  
        pd.DataFrame: 原始DataFrame的拷贝（除最后一行外），其中包含一个新列用于存储构成比（保留两位小数）。  
    """  
    # 检查列名是否在DataFrame中  
    if column_name not in df.columns:  
        raise ValueError(f"Column '{column_name}' not found in DataFrame.")  
      
    # 获取最后一行的合计值，并创建一个不包含最后一行的DataFrame拷贝  
    total = df.iloc[-1][column_name]  

      
    # 计算构成比（以百分比表示）  
    composition_ratio = (df[column_name] / total) * 100  
      
    # 在DataFrame拷贝中添加新列，存储构成比  
    df[new_column_name] = composition_ratio.round(2)  
      
    return df.sort_values('构成比', ascending=False)
          
def SMALL_count_mode(series, mode):  
    """  
    对给定的pandas Series进行计数，并根据mode参数指定的分隔符拆分复合症状并计算每个症状的计数。  
      
    参数:  
        series (pandas.Series): 输入的pandas Series对象，包含要计数的数据。  
        mode (str): 指定用于拆分复合症状的分隔符，可以是单个字符或多个字符。  
      
    返回:  
        dict: 拆分后并计算每个症状计数的新字典。  
    """  
    result = series.value_counts().to_dict()  
    filtered_result = {k: v for k, v in result.items() if v > 0}  
    if mode=="":
        return filtered_result
      
    new_dict = {}  
    for key, value in filtered_result.items():  
        symptoms = re.split(mode, key)  # 使用正则表达式根据指定的分隔符进行拆分  
        for symptom in symptoms:  
            symptom = symptom.strip()  
            if symptom:  # 检查症状是否为非空字符串  
                if symptom in new_dict:  
                    new_dict[symptom] += value  
                else:  
                    new_dict[symptom] = value  
      
    return new_dict
    
def SMALL_get_list(input_str, df, *args):  
    """  
    转成列表。  
    """  
    input_str = str(input_str)  
      
    if pd.isnull(input_str):  
        return []  
  
    result_list = []  
      
    if ("use(" in input_str) and (").file"  in input_str):  
        pattern = r"[(](.*?)[)]"  
        matches = re.findall(pattern, input_str)  
        result_list = df[matches[0]].astype(str).tolist()  
    else:  
        result_list = [input_str]  
      
    # 合并重复的字符串操作  
    result_list = ",".join(result_list)  
    result_list = result_list.replace("（严重）", "").replace("（一般）", "")  
    for separator in ["┋", ";", "；", "、", "，", ",", "|"]:  
        result_list = result_list.replace(separator, ",")  
    #print(result_list)  
    result_list = result_list.split(",") 
    #result_list.sort()  
      
    return result_list    

def SMALL_save_dict(data):  
    """保存文件"""  
    file_path_flhz = filedialog.asksaveasfilename(  
        title="保存文件",  
        initialfile="排序后的原始数据",  
        defaultextension=".xlsx",  
        filetypes=[("Excel 工作簿", "*.xlsx"), ("Excel 97-2003 工作簿", "*.xls")],  
    )  
    if not file_path_flhz:  
        return  # 如果用户点击取消，则退出函数
        
    if "详细描述T" in data.columns:
        data["详细描述T"]=data["详细描述T"].astype(str)

    if "报告编码" in data.columns:
        data["报告编码"]=data["报告编码"].astype(str)     
        
        
    try:  
        with pd.ExcelWriter(file_path_flhz, engine="xlsxwriter") as writer:  
            data.to_excel(writer, sheet_name="导出的数据", index=False)  
        messagebox.showinfo(title="提示", message="文件写入成功。")  
    except Exception as e:  
        messagebox.showerror("错误", f"保存文件时出错: {e}")


    
def SMALL_align_values_to_most_frequent(df, group_col, cols_to_align):  
    """  
    对于DataFrame中的每个group_col的唯一值，找到cols_to_align中对应列的最频繁出现的值，  
    并将该值赋给所有具有相同group_col值的行。不在cols_to_align中的列保持不变。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        group_col (str): 用于分组的列名。  
        cols_to_align (list of str): 需要对齐的列名列表。  
          
    返回:  
        pd.DataFrame: 更新后的DataFrame，包含所有原始列，且列顺序不变。  
    """  
    # 复制原始DataFrame以保持其不变  
    df_aligned = df.copy()  
      
    # 遍历需要对齐的列  
    for col in cols_to_align:  
        # 使用groupby找到每个group_col值的最频繁出现的值  
        mode_series = df.groupby(group_col)[col].transform(lambda x: x.mode()[0])  
        # 将结果对齐到原始DataFrame的相应列  
        df_aligned[col] = mode_series  
      
    return df_aligned  
    
def SMALL_add_count_and_ratio(df, a, agg_col=None, methon=None):  
    """  
    函数说明：  
    该函数接受一个pandas DataFrame（df）、列名（a）、聚合列名（agg_col）以及聚合方法（methon）作为输入。  
    使用groupby方法对列a进行分组，并根据agg_col和methon计算每个组的聚合值。  
    接下来，计算每个组的构成比，并将其转换为百分数形式，保留两位小数。  
    最后，按构成比的大小进行排序，并返回一个新的DataFrame，  
    其中包含列a的唯一值、每个唯一值的计数或聚合值以及构成比。  
  
    参数：  
    df：pandas DataFrame，包含要进行分组和聚合的数据。  
    a：字符串或字符串列表，指定要进行分组的列名。  
    agg_col：字符串，可选，指定要进行聚合的列名。如果为None，则对整个组进行聚合。  
    methon：字符串，指定聚合方法（例如，'count'、'sum'、'nunique'等）。  
  
    返回值：  
    pandas DataFrame，包含列a的唯一值、每个唯一值的计数或聚合值以及构成比（百分数形式，保留两位小数），并按构成比大小排序。  
    """  
  
  
    # 对列a进行分组  
    if agg_col is not None:  
        grouped = df.groupby(a)[agg_col]  
    else:  
        grouped = df.groupby(a).size().reset_index(name='计数')  
      
    # 根据methon计算聚合值  
    if methon == 'count':  
        grouped = grouped.count().reset_index(name='计数')  
    elif methon == 'sum':  
        grouped = grouped.sum().reset_index(name='计数')  
    elif methon == 'nunique':  
        grouped = grouped.nunique().reset_index(name='计数')  
  
    # 计算构成比并转换为百分数形式，保留两位小数  
    total_count = grouped['计数'].sum()  
    grouped['构成比'] = (grouped['计数'] / total_count) * 100  
    grouped['构成比'] = grouped['构成比'].round(2)  
  
    # 按构成比大小进行排序  
    grouped = grouped.sort_values('构成比', ascending=False)  
  
    # 添加报表类型列（注意：此部分逻辑可能需要根据实际需求进行调整）  
    lst = []  
    if isinstance(a, str):  
        lst.append(a)  
    elif isinstance(a, list):  
        lst.extend(a)  
    grouped["报表类型"] = "{'grouped':" + str(lst) + "}"  
  
    return grouped

def SMALL_add_sep(result):
    """  
    溯源支持（关键词法） 
    """  
    result['报表类型']=result['报表类型'].str.replace('grouped','group_sep')
    return result
    
def SMALL_add_count_and_ratio_exp(df, a,sep):  
    """  
    当列扩行透视 
    """  
    df=df.copy()
    df=CLEAN_expand_rows(df,sep,[a])
    df=df.reset_index(drop=True)
    # 对列a进行分组并计算计数  
    grouped = df.groupby(a).size().reset_index(name='计数')  
      
    # 计算构成比并转换为百分数形式，保留两位小数  
    total_count = grouped['计数'].sum()  
    grouped['构成比'] = (grouped['计数'] / total_count) * 100  
    grouped['构成比'] = grouped['构成比'].round(2)  
      
    # 按构成比大小进行排序  
    grouped = grouped.sort_values('构成比', ascending=False)  
    
    lst = []  
    if isinstance(a, str):  
        lst.append(a)  
    elif isinstance(a, list):  
        lst.extend(a)
    
    grouped["报表类型"]="{'group_sep':"+str(lst)+str("}")       
    return grouped    

############################################################################################################################
#预先配置的函数
############################################################################################################################
class AAA_08_CLean():
    pass
def CLEAN_table(df):  
    print(
        """
    **********************************************
    数据清洗工具

    功能：
    1. 提供 GUI 界面供用户选择列并执行多种数据清洗操作。
    2. 支持复制合并、条件赋值、两列运算、重命名列、替换字符、填充空值、删除空值行、转换格式、保留包含、删除包含、众数对齐、扩展多行、还原扩行、扩展多列、按列去重、升序排列、降序排列、删除本列、整体去重、选列保留、重置索引、查看数据、查看轨迹、批量规整、加关键词、赋关键词等操作。

    使用方法：
    1. 启动程序后，选择需要操作的列。
    2. 根据需求选择相应的操作（如复制合并、条件赋值等）。
    3. 输入必要的参数（如新列名、分隔符、查找字符等）。
    4. 点击相应的按钮执行操作。
    5. 结果会显示在 GUI 界面中。
    6. !!!!!!!!!!!!!!!!您可以'查看轨迹'，将操作轨迹另存为文件，轨迹可以通过'批量规整'按钮选择该文件复现。

    注意：
    - 如果未选择任何列或参数不正确，程序会提示错误。
    - 数据清洗结果会显示在 GUI 界面中。
    **********************************************
    """
    )

    # 为验证hash而作。
    df_hash_auto = df.copy()

    operation_log = []  

    def on_confirm(dfsd, method): 
        nonlocal df  # dfsd实际没有作用，没有哪里使用
        
        # 增加记录
        operation = {  
            "方法": "",  
            "作用列": [],  
            "参数": {},  
        } 
        selected_columns = [tree.item(item_id, "text") for item_id in tree.selection()]  # 获取多选的列名    
        operation["作用列"] = selected_columns   
        operation["方法"] = method               
        if not selected_columns and method not in ["重置索引", "查看数据", "查看轨迹", "批量规整", "新建一列", "整体去重"]: 
            messagebox.showinfo("错误", "没有选择任何列")  
            return  # 如果没有选择任何列，则直接返回  
                      
        if method == "复制合并":  
            new_column_name = new_column_entry.get()  # 获取新列名  
            separator = separator_entry.get()  # 获取分隔符             
            # 检查新列名是否为空或与现有列名重复  
            if not new_column_name or new_column_name in df.columns:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return  
            df[new_column_name] = df[selected_columns].apply(lambda row: separator.join(row.astype(str)), axis=1)  
            operation["参数"] = {"new_column_name": new_column_name, "separator": separator}  

        elif method == "条件赋值":  
            compare_value = corr_express.get()  # 获取条件。
            G_value = cor_entry2.get()  # 获取赋值。        
            if len(selected_columns) != 1:  
                messagebox.showinfo("错误", "请选择1个数据列。")  
                return               
            df = SMALL_assign_value_based_on_expression(df, selected_columns[0], compare_value, G_value)
            operation["参数"] = {"compare_value": compare_value, "G_value": G_value}  

        elif method == "两列运算":  
            new_column_name = Enew_column_entry.get()  # 获取新列名  
            separator = Eseparator_entry.get()  # 获取分隔符             
            if len(selected_columns) != 2:  
                messagebox.showinfo("错误", "请选择两个数据列。")  
                return  
            df = SMALL_apply_operation(df, selected_columns, new_column_name, separator) 
            operation["参数"] = {"new_column_name": new_column_name, "separator": separator}  

        elif method == "重命名列":  
            if len(selected_columns) > 1:  
                messagebox.showinfo("提示", "请只选择一列进行重命名")  
                return  
            new_column_name = Anew_column_entry.get().strip()  # 获取新列名并去除首尾空格                
            if not new_column_name or new_column_name in df.columns and new_column_name != selected_columns[0]:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return   
            df.rename(columns={selected_columns[0]: new_column_name}, inplace=True)  
            operation["参数"] = {"new_column_name": new_column_name}  

        elif method == "填充空值":  
            fill_value = fill_entry.get().strip()  # 获取填充值并去除首尾空格  
            fill_method = fill_method_combobox.get()  # 获取填充方式

            if fill_method == "手工填写":
                if not fill_value:
                    messagebox.showinfo("错误", "填充值不能为空")
                    return
                try:
                    # 尝试将填充值转换为适当的数据类型
                    fill_value = df[selected_columns[0]].dtype.type(fill_value)
                except ValueError:
                    messagebox.showinfo("错误", "填充值与列的数据类型不匹配")
                    return
            else:
                # 根据选择的填充方式计算填充值
                if fill_method == "均值填充":
                    fill_value = df[selected_columns[0]].mean()
                elif fill_method == "众数填充":
                    fill_value = df[selected_columns[0]].mode()[0]
                elif fill_method == "50%分位数填充":
                    fill_value = df[selected_columns[0]].quantile(0.5)
                elif fill_method == "最大值填充":
                    fill_value = df[selected_columns[0]].max()
                elif fill_method == "最小值填充":
                    fill_value = df[selected_columns[0]].min()
                else:
                    messagebox.showinfo("错误", "未知的填充方式")
                    return

            for column in selected_columns:
                df[column].fillna(fill_value, inplace=True)  # 使用指定的填充值填充选定列的空值
            operation["参数"] = {"fill_value": fill_value, "fill_method": fill_method}

        elif method == "删空值行":   
            for column in selected_columns:  
                df.dropna(subset=[column], inplace=True)  

        elif method == "转换格式":  
            target_format = format_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not target_format:  
                messagebox.showinfo("错误", "目标格式不能为空")  
                return  
            for column in selected_columns:  
                try: 
                    if target_format == "日期":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m-%d')
                    elif target_format == "年份":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                    elif target_format == "月份":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m')
                    elif target_format == "季度":
                        quarterx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.quarter
                        yearx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                        df[column]  = yearx.astype(str) + "Q" + quarterx.astype(str)
                    elif target_format == "round2":
                        df[column] = round(df[column].astype(float), 2) 
                    else: 
                        # 尝试将选定列的数据类型转换为目标格式  
                        df[column] = df[column].astype(target_format)  
                except ValueError as e:  
                    messagebox.showinfo("错误", f"无法将列 '{column}' 转换为 {target_format} 格式。错误消息：{str(e)}")  
                    return
            operation["参数"] = {"target_format": target_format}  

        elif method == "保留包含" or method == "删除包含":  
            keyword_ct = ct_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not keyword_ct:  
                messagebox.showinfo("错误", "关键词不能为空")  
                return  

            try:  
                if method == "删除包含":
                    df = df[~df[selected_columns[0]].str.contains(keyword_ct, na=False)]  #
                elif method == "保留包含":
                    df = df[df[selected_columns[0]].str.contains(keyword_ct, na=False)]               
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"keyword_ct": keyword_ct}  

        elif method == "众数对齐":  
            group_col = ffd_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not group_col:  
                messagebox.showinfo("错误", "对齐源列不能为空")  
                return  

            try: 
                df = SMALL_align_values_to_most_frequent(df, group_col, selected_columns)             
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"group_col": group_col}  

        elif method == "扩展多行":  
            word_sep = exp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not word_sep:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  

            try:  
                df = CLEAN_expand_rows(df.copy(), word_sep, selected_columns) 
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep": word_sep}           

        elif method == "还原扩行":  
            Rword_sep = Rexp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            depend_row=xformat_entry.get().strip() 
            if not Rword_sep:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  

            try:  
                df = CLEAN_expand_rows_REVERT(df, Rword_sep, selected_columns,depend_row) 
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep": Rword_sep,'depend_row':depend_row}   

        elif method == "扩展多列":  
            word_sep2 = exp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not word_sep2:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  
            try:  
                df = CLEAN_expand_cols(df, word_sep2, selected_columns[0]) 
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep2": word_sep2}           

        elif method == "按列去重":  
            df.drop_duplicates(subset=selected_columns, keep='first', inplace=True)

        elif method == "升序排列":  
            df.sort_values(by=selected_columns, ascending=[True] * len(selected_columns), inplace=True)  

        elif method == "降序排列":  
            df.sort_values(by=selected_columns, ascending=[False] * len(selected_columns), inplace=True)     

        elif method == "删除本列":  
            df.drop(columns=selected_columns, inplace=True)

        elif method == "整体去重":  
            df.drop_duplicates(inplace=True)            

        elif method == "选列保留":  
            all_columns = df.columns.tolist()  # 获取所有列名  
            columns_to_drop = [col for col in all_columns if col not in selected_columns]  # 找出不在selected_columns中的列名  
            df.drop(columns=columns_to_drop, inplace=True)  # 删除这些列

        elif method == "重置索引":  
            df.reset_index(inplace=True) 


        elif method == "替换字符":  
            for col in selected_columns:
                df[col] = df[col].astype(str).str.replace(find_entry.get(), replace_entry.get())

        elif method == "查看数据":  
            pass 

        elif method == "查看轨迹":  
            if operation_log == []:
                return
            PROGRAM_display_df_in_treeview(pd.DataFrame(operation_log), 0, df)
            return 

        elif method == "批量规整":  
            guize = SMALL_read_and_merge_files()
            df = CLEAN_replay_operations(df, guize)
            operation["参数"] = {"guize": guize.to_dict(orient='list')}  

        elif method == "加关键词" or method == "赋关键词":  # XFDD
            mark_column_name = mark_column_entry.get().strip()  # 获取新列名并去除首尾空格  
            predefined_all = fdel_column_entry2.get().strip() 
            parts0 = predefined_all.split("|") 
            predefined_column = parts0[0]
            predefined_objects = parts0[1]  
            keyword_include = keyword_entry.get().strip()  # 获取关键词并去除首尾空格   
            parts1 = keyword_include.split("|")  
            mark_value = parts1[0] 
            include_str = keyword_include
            exclude_str = fdel_column_entry.get()
            df = CLEAN_filter_and_mark_rows(df, selected_columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, method) 
            operation["参数"] = {"mark_column_name": mark_column_name, "include_str": include_str, "exclude_str": exclude_str, "predefined_column": predefined_column, "predefined_objects": predefined_objects, "mark_value": mark_value, "method": method}  

        else:  
            return  

        """更新TreeView以显示DataFrame的列名"""  
        # 清空当前的TreeView内容  
        tree.delete(*tree.get_children())  
        # 根据最新的DataFrame列名重新填充TreeView  
        for col in df.columns:  
            tree.insert("", tk.END, text=str(col), values=())

        operation_log.append(operation)  
        print(operation_log)                
        PROGRAM_display_df_in_treeview(df, 0, 0)  # 更新树状视图  

    title = "数据清洗工具" 
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title(title)  

    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 1230  # 窗口宽度  
    wh = 630  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  

    # 创建主框架，用于放置所有控件  
    main_frame = ttk.Frame(root)  
    main_frame.pack(fill="both", expand=True, padx=10, pady=10)  

    # 创建左侧树状视图，显示df的列名  
    tree_frame = ttk.Frame(main_frame)  
    tree_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)  
    tree_label = ttk.Label(tree_frame, text="请选择要操作的列:")  
    tree_label.pack(anchor='w')  

    # 添加滚动条  
    scrollbar = ttk.Scrollbar(tree_frame)  
    scrollbar.pack(side="right", fill="y")  

    tree = ttk.Treeview(tree_frame, yscrollcommand=scrollbar.set)  
    scrollbar.config(command=tree.yview)  
    tree['columns'] = ('Column',)  
    tree.column('#0', width=200, stretch=tk.NO)  
    tree.heading('#0', text='列名', anchor=tk.W)  

    for col in df.columns:  
        tree.insert("", tk.END, text=str(col), values=())  

    tree.pack(fill="both", expand=True, padx=5, pady=5)  

    # 创建右侧内容区域，使用Grid布局来放置所有控件在同一行  
    right_frame = ttk.Frame(main_frame)  
    right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)  

    # FSM 复制合并   
    separator_label = ttk.Label(right_frame, text="分隔符：")  
    separator_label.grid(row=0, column=0, padx=5, pady=2, sticky='w')  
    separator_entry = tk.Entry(right_frame, width=20)  
    separator_entry.grid(row=0, column=1, padx=5, pady=2, sticky='w')  
    new_column_label = ttk.Label(right_frame, text="新列名：")  
    new_column_label.grid(row=0, column=2, padx=5, pady=2, sticky='w')  
    new_column_entry = tk.Entry(right_frame, width=20)  
    new_column_entry.grid(row=0, column=3, padx=5, pady=2, sticky='w')  
    confirm_button = ttk.Button(right_frame, text="复制合并", command=lambda: on_confirm(df, "复制合并"))  
    confirm_button.grid(row=0, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # 字符替换 
    find_label = ttk.Label(right_frame, text="查找的字符：")  
    find_label.grid(row=1, column=0, padx=5, pady=2, sticky='w')  
    find_entry = tk.Entry(right_frame, width=20)  
    find_entry.grid(row=1, column=1, padx=5, pady=2, sticky='w')  
    replace_label = ttk.Label(right_frame, text="替换值：")  
    replace_label.grid(row=1, column=2, padx=5, pady=2, sticky='w')  
    replace_entry = tk.Entry(right_frame, width=20)  
    replace_entry.grid(row=1, column=3, padx=5, pady=2, sticky='w')  
    replace_button = ttk.Button(right_frame, text="替换字符", command=lambda: on_confirm(df, "替换字符"))  
    replace_button.grid(row=1, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2) 

    # FSM 重命名列   
    Anew_column_label = ttk.Label(right_frame, text="新列名：")  
    Anew_column_label.grid(row=2, column=2, padx=5, pady=2, sticky='w')  
    Anew_column_entry = tk.Entry(right_frame, width=20)  
    Anew_column_entry.grid(row=2, column=3, padx=5, pady=2, sticky='w')  
    rename_button = ttk.Button(right_frame, text="修改列名", command=lambda: on_confirm(df, "重命名列"))  
    rename_button.grid(row=2, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  
    Rrename_button = ttk.Button(right_frame, text="最后非空", command=lambda: on_confirm(df, "最后非空"))  
    Rrename_button.grid(row=2, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 空值填充  删空值行
    fill_label = ttk.Label(right_frame, text="填充值：")  
    fill_label.grid(row=3, column=0, padx=5, pady=2, sticky='w')  
    fill_entry = tk.Entry(right_frame, width=20)  
    fill_entry.grid(row=3, column=1, padx=5, pady=2, sticky='w')  

    # 添加填充方式下拉框
    fill_method_label = ttk.Label(right_frame, text="填充方式：")  
    fill_method_label.grid(row=3, column=2, padx=5, pady=2, sticky='w')  
    fill_method_combobox = ttk.Combobox(right_frame, values=["手工填写", "均值填充", "众数填充", "50%分位数填充", "最大值填充", "最小值填充"], width=20)  
    fill_method_combobox.grid(row=3, column=3, padx=5, pady=2, sticky='w')  
    fill_method_combobox.current(0)  # 默认选择手工填写

    fill_button = ttk.Button(right_frame, text="填充空值", command=lambda: on_confirm(df, "填充空值"))  
    fill_button.grid(row=3, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    delfill_button = ttk.Button(right_frame, text="删空值行", command=lambda: on_confirm(df, "删空值行"))  
    delfill_button.grid(row=3, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 格式转换 
    format_options = ["str", "float", "int", "round2", "日期", "月份", "季度", "年份"]   
    format_label = ttk.Label(right_frame, text="目标格式（如str）：")  
    format_label.grid(row=4, column=2, padx=5, pady=2, sticky='w')  
    format_entry = ttk.Combobox(right_frame, values=format_options, width=20)   
    format_entry.grid(row=4, column=3, padx=5, pady=2, sticky='w')  
    convert_button = ttk.Button(right_frame, text="转换格式", command=lambda: on_confirm(df, "转换格式"))   
    convert_button.grid(row=4, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 扩展多行   扩展多列
       
    exp_row_label = ttk.Label(right_frame, text="分隔符：")  
    exp_row_label.grid(row=5, column=2, padx=5, pady=2, sticky='w')  
    exp_row_entry = tk.Entry(right_frame, width=20)  
    exp_row_entry.grid(row=5, column=3, padx=5, pady=2, sticky='w')  
    exp_row_button = ttk.Button(right_frame, text="扩展多行", command=lambda: on_confirm(df, "扩展多行"))  
    exp_row_button.grid(row=5, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  
    exp_col_button = ttk.Button(right_frame, text="扩展多列", command=lambda: on_confirm(df, "扩展多列"))  
    exp_col_button.grid(row=5, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 含与不含  
    ct_label = ttk.Label(right_frame, text="关键词（仅一列）：")  
    ct_label.grid(row=6, column=2, padx=5, pady=2, sticky='w')  
    ct_entry = tk.Entry(right_frame, width=20)  
    ct_entry.grid(row=6, column=3, padx=5, pady=2, sticky='w')  
    ct_keep_button = ttk.Button(right_frame, text="保留包含", command=lambda: on_confirm(df, "保留包含"))  
    ct_keep_button.grid(row=6, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  
    ct_del_button = ttk.Button(right_frame, text="删除包含", command=lambda: on_confirm(df, "删除包含"))  
    ct_del_button.grid(row=6, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 众数对齐  
    ffd_label = ttk.Label(right_frame, text="众数源列：")  
    ffd_label.grid(row=7, column=2, padx=5, pady=2, sticky='w')  
    ffd_entry = tk.Entry(right_frame, width=20)  
    ffd_entry.grid(row=7, column=3, padx=5, pady=2, sticky='w')  
    ffd_button = ttk.Button(right_frame, text="众数对齐", command=lambda: on_confirm(df, "众数对齐"))  
    ffd_button.grid(row=7, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 还原扩行 
    xformat_options = df.columns.tolist()  
    xformat_label = ttk.Label(right_frame, text='依据列：')  
    xformat_label.grid(row=11, column=0, padx=5, pady=2, sticky='w')  
    xformat_entry = ttk.Combobox(right_frame, values=xformat_options, width=20)   
    xformat_entry.grid(row=11, column=1, padx=5, pady=2, sticky='w') 
     
    Rexp_row_label = ttk.Label(right_frame, text="分隔符：")  
    Rexp_row_label.grid(row=11, column=2, padx=5, pady=2, sticky='w')  
    Rexp_row_entry = tk.Entry(right_frame, width=20)  
    Rexp_row_entry.grid(row=11, column=3, padx=5, pady=2, sticky='w')  
    Rexp_row_button = ttk.Button(right_frame, text="还原扩行", command=lambda: on_confirm(df, "还原扩行"))  
    Rexp_row_button.grid(row=11, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 两列运算   
    Eseparator_label = ttk.Label(right_frame, text="运算式(a+b)：")  
    Eseparator_label.grid(row=12, column=0, padx=5, pady=2, sticky='w')  
    Eseparator_entry = tk.Entry(right_frame, width=20)  
    Eseparator_entry.grid(row=12, column=1, padx=5, pady=2, sticky='w')  
    Enew_column_label = ttk.Label(right_frame, text="新列名：")  
    Enew_column_label.grid(row=12, column=2, padx=5, pady=2, sticky='w')  
    Enew_column_entry = tk.Entry(right_frame, width=20)  
    Enew_column_entry.grid(row=12, column=3, padx=5, pady=2, sticky='w')  
    Econfirm_button = ttk.Button(right_frame, text="两列运算", command=lambda: on_confirm(df, "两列运算"))  
    Econfirm_button.grid(row=12, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 新建列    new_column_valuex
    AAnew_column_label = ttk.Label(right_frame, text="新列名：")  
    AAnew_column_label.grid(row=13, column=2, padx=5, pady=2, sticky='w')  
    AAnew_column_entry = tk.Entry(right_frame, width=20)  
    AAnew_column_entry.grid(row=13, column=3, padx=5, pady=2, sticky='w')  
    AAnew_column_label2 = ttk.Label(right_frame, text="赋值：")  
    AAnew_column_label2.grid(row=13, column=0, padx=5, pady=2, sticky='w')  
    new_column_valuex_entry = tk.Entry(right_frame, width=20)  
    new_column_valuex_entry.grid(row=13, column=1, padx=5, pady=2, sticky='w')
    creatnew_button = ttk.Button(right_frame, text="新建一列", command=lambda: on_confirm(df, "新建一列"))  
    creatnew_button.grid(row=13, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 并用追加
    SAAnew_column_label = ttk.Label(right_frame, text="依据列：")  
    SAAnew_column_label.grid(row=14, column=2, padx=5, pady=2, sticky='w')  
    SAAnew_column_entry = tk.Entry(right_frame, width=20)  
    SAAnew_column_entry.grid(row=14, column=3, padx=5, pady=2, sticky='w')  
    SAAnew_column_label2 = ttk.Label(right_frame, text="新列名：")  
    SAAnew_column_label2.grid(row=14, column=0, padx=5, pady=2, sticky='w')  
    Snew_column_valuex_entry = tk.Entry(right_frame, width=20)  
    Snew_column_valuex_entry.grid(row=14, column=1, padx=5, pady=2, sticky='w')
    Screatnew_button = ttk.Button(right_frame, text="并用追加", command=lambda: on_confirm(df, "并用追加"))  
    Screatnew_button.grid(row=14, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # 条件赋值 
    cor_value_label = ttk.Label(right_frame, text="表达式df['a']：")  
    cor_value_label.grid(row=15, column=0, padx=5, pady=2, sticky='w')  
    corr_express = tk.Entry(right_frame, width=20)  
    corr_express.grid(row=15, column=1, padx=5, pady=2, sticky='w')  
    cor_value_labels = ttk.Label(right_frame, text="赋值：")  
    cor_value_labels.grid(row=15, column=2, padx=5, pady=2, sticky='w')  
    cor_entry2 = tk.Entry(right_frame, width=20)  
    cor_entry2.grid(row=15, column=3, padx=5, pady=2, sticky='w')  
    Xcor_button = ttk.Button(right_frame, text="条件赋值", command=lambda: on_confirm(df, "条件赋值"))  
    Xcor_button.grid(row=15, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2) 

    # 找词标记  
    column_label = ttk.Label(right_frame, text="被标记列：")  
    column_label.grid(row=17, column=0, padx=5, pady=2, sticky='w')  
    mark_column_entry = tk.Entry(right_frame, width=20)  
    mark_column_entry.grid(row=17, column=1, padx=5, pady=2, sticky='w')     
    keyword_label = ttk.Label(right_frame, text="标记|词1|词2...")  
    keyword_label.grid(row=17, column=2, padx=5, pady=2, sticky='w')      
    keyword_entry = tk.Entry(right_frame, width=20)  
    keyword_entry.grid(row=17, column=3, padx=5, pady=2, sticky='w')  
    fdel_column_label = ttk.Label(right_frame, text="排除值：")  
    fdel_column_label.grid(row=18, column=0, padx=5, pady=2, sticky='w')  
    fdel_column_entry = tk.Entry(right_frame, width=20)  
    fdel_column_entry.grid(row=18, column=1, padx=5, pady=2, sticky='w') 
    fdel_column_label2 = ttk.Label(right_frame, text="作用对象|列内对象：")  
    fdel_column_label2.grid(row=18, column=2, padx=5, pady=2, sticky='w')  
    fdel_column_entry2 = tk.Entry(right_frame, width=20)  
    fdel_column_entry2.insert(0, "所有列|所有对象")  # 在这里插入默认值
    fdel_column_entry2.grid(row=18, column=3, padx=5, pady=2, sticky='w')            
    add_keyword_button = ttk.Button(right_frame, text="找词加标", command=lambda: on_confirm(df, "加关键词"))  
    add_keyword_button.grid(row=18, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)
    add_keyword_button2 = ttk.Button(right_frame, text="找词赋标", command=lambda: on_confirm(df, "赋关键词"))  
    add_keyword_button2.grid(row=18, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)

    # FSM 升序排列 ## 
    sort_value_up_button = ttk.Button(right_frame, text="升序排列", command=lambda: on_confirm(df, "升序排列"))  
    sort_value_up_button.grid(row=19, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 降序排列 ## 
    sort_value_down_button = ttk.Button(right_frame, text="降序排列", command=lambda: on_confirm(df, "降序排列"))  
    sort_value_down_button.grid(row=20, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 按列去重 
    drop_duplicates_button = ttk.Button(right_frame, text="按列去重", command=lambda: on_confirm(df, "按列去重"))  
    drop_duplicates_button.grid(row=19, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 删除本列 
    drop_button = ttk.Button(right_frame, text="删除本列", command=lambda: on_confirm(df, "删除本列"))  
    drop_button.grid(row=19, column=2, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 选列保留 
    keep_button = ttk.Button(right_frame, text="选列保留", command=lambda: on_confirm(df, "选列保留"))  
    keep_button.grid(row=19, column=0, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 整体去重 
    dropdu_button = ttk.Button(right_frame, text="整体去重", command=lambda: on_confirm(df, "整体去重"))  
    dropdu_button.grid(row=19, column=1, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 重置索引 
    reset_button = ttk.Button(right_frame, text="重置索引", command=lambda: on_confirm(df, "重置索引"))  
    reset_button.grid(row=20, column=0, padx=(0, 10), pady=2, sticky='e', columnspan=2) 

    # FSM 批量规整 
    sll_button = ttk.Button(right_frame, text="批量规整", command=lambda: on_confirm(df, "批量规整"))  
    sll_button.grid(row=20, column=1, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 查看轨迹 
    rec_button = ttk.Button(right_frame, text="查看轨迹", command=lambda: on_confirm(df, "查看轨迹"))  
    rec_button.grid(row=20, column=2, padx=(0, 10), pady=2, sticky='e', columnspan=2)  

    # FSM 查看数据 
    adint_button = ttk.Button(right_frame, text="查看数据", command=lambda: on_confirm(df, "查看数据"))  
    adint_button.grid(row=20, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2) 

    root.lift()
    root.attributes("-topmost", True)
    root.attributes("-topmost", False)

def CLEAN_replay_operations(df, operation_log):  
    """  
    根据operation_log中的记录，在df上复现操作。  
    """ 
    for ids, operation in operation_log.iterrows():  
        try:
            method = operation["方法"]  
            columns = ast.literal_eval(operation["作用列"])
            selected_columns = columns
            params = eval(str(operation["参数"]))   

            if method == "复制合并":  
                new_column_name = params["new_column_name"] 
                separator = params["separator"]  
                if new_column_name in df.columns:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df[new_column_name] = df[columns].apply(lambda row: separator.join(row.astype(str)), axis=1)  
                print(f"已合并列，并创建新列：{new_column_name}")  

            elif method == "两列运算":  
                new_column_name = params["new_column_name"] 
                separator = params["separator"]  
                if len(columns) != 2: 
                    print("错误", "请选择两个数据列。")  
                    return  
                df = SMALL_apply_operation(df, columns, new_column_name, separator)                
                print(f"已开展相关运算：{new_column_name}")  

            elif method == "众数对齐":  
                group_col = params["group_col"]  
                if not group_col:  
                    print("错误，对齐源列不能为空")  
                    return  
                try:  
                    df = SMALL_align_values_to_most_frequent(df, group_col, columns)             
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    return
                print(f"已对齐列 {columns} 到众数列 {group_col}")  

            elif method == "重命名列":  
                if len(columns) > 1:  
                    print("错误：重命名操作只能选择一列")  
                    continue  # 跳过当前操作，继续下一个  
                new_column_name = params["new_column_name"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df.rename(columns={columns[0]: new_column_name}, inplace=True)  
                print(f"已将列 {columns[0]} 重命名为 {new_column_name}")  

            elif method == "最后非空":  
                new_column_name = params["new_column_name"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df = SMALL_last_non_null_value(df, columns, new_column_name) 
                print(f"已设置最后非空为 {new_column_name}")  

            elif method == "并用追加":  
                new_column_name = params["new_column_name"]  
                new_column_valuex = params["new_column_valuex"]  
                df = CLEAN_fill_column_c_based_on_a_and_b(df, new_column_name, columns[0], new_column_valuex)     
                print(f"已并用追加列 {columns[0]} 到新列 {new_column_name}")  

            elif method == "新建一列":  
                if len(columns) > 1:  
                    print("错误：新建列操作只能选择一列")  
                    continue  # 跳过当前操作，继续下一个  
                new_column_name = params["new_column_name"]
                new_column_valuex = params["new_column_valuex"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df[new_column_name] = new_column_valuex 
                print(f"已新建列 {new_column_name}")  

            elif method == "替换字符":  
                find_text = params["find_text"]  
                replace_text = params["replace_text"]  
                if not find_text:  
                    print("错误：请输入要查找的字符")  
                    continue  
                for col in columns:  
                    df[col] = df[col].astype(str).str.replace(find_text, replace_text)  
                print(f"已在列 {columns} 中替换字符：{find_text} -> {replace_text}")  

            elif method == "条件赋值":  
                compare_value = params["compare_value"]  
                G_value = params["G_value"]  
                if len(columns) != 1:  
                    print("错误", "请选择1个数据列。")  
                    return               
                df = SMALL_assign_value_based_on_expression(df, columns[0], compare_value, G_value) 
                print(f"已完成条件赋值")  

            elif method == "填充空值":  
                fill_value = params["fill_value"]  
                fill_method = params["fill_method"]  

                if fill_method == "手工填写":
                    if not fill_value:
                        print("错误：填充值不能为空")
                        continue
                    try:
                        fill_value = df[columns[0]].dtype.type(fill_value)
                    except ValueError:
                        print("错误：填充值与列的数据类型不匹配")
                        continue
                else:
                    if fill_method == "均值填充":
                        fill_value = df[columns[0]].mean()
                    elif fill_method == "众数填充":
                        fill_value = df[columns[0]].mode()[0]
                    elif fill_method == "50%分位数填充":
                        fill_value = df[columns[0]].quantile(0.5)
                    elif fill_method == "最大值填充":
                        fill_value = df[columns[0]].max()
                    elif fill_method == "最小值填充":
                        fill_value = df[columns[0]].min()
                    else:
                        print("错误：未知的填充方式")
                        continue

                for column in columns:
                    df[column].fillna(fill_value, inplace=True)  
                print(f"已在列 {columns} 中填充空值：{fill_value}")  

            elif method == "删空值行":   
                for column in columns:  
                    df.dropna(subset=[column], inplace=True)  
                print(f"已删除列 {columns} 中的空值行")  

            elif method == "转换格式":  
                target_format = params["target_format"]  
                if not target_format:  
                    print("错误：目标格式不能为空")  
                    continue  
                for column in columns:  
                    try: 
                        if target_format == "日期":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m-%d')
                        elif target_format == "年份":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                        elif target_format == "月份":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m')
                        elif target_format == "季度":
                            quarterx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.quarter
                            yearx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                            df[column]  = yearx.astype(str) + "Q" + quarterx.astype(str)
                        elif target_format == "round2":
                            df[column] = round(df[column].astype(float), 2) 
                        else: 
                            df[column] = df[column].astype(target_format)  
                    except ValueError as e:  
                        print(f"错误：无法将列 '{column}' 转换为 {target_format} 格式。错误消息：{str(e)}")  
                        continue  
                print(f"已将列 {columns} 的格式转换为 {target_format}")  

            elif method == "保留包含" or method == "删除包含":  
                keyword_ct = params["keyword_ct"]  
                if not keyword_ct:  
                    print("错误", "关键词不能为空")  
                    return  
                try:  
                    if method == "删除包含":
                        df = df[~df[columns[0]].str.contains(keyword_ct, na=False)]
                    elif method == "保留包含":
                        df = df[df[columns[0]].str.contains(keyword_ct, na=False)]               
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    return
                print(f"已根据关键词 {keyword_ct} 完成 {method} 操作")  

            elif method == "扩展多行":  
                word_sep = params["word_sep"]  
                if not word_sep:  
                    print("错误：分隔符不能为空")  
                    continue  
                try:  
                    df = CLEAN_expand_rows(df, word_sep, columns)  
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    continue  
                print(f"已使用分隔符 {word_sep} 在列 {columns} 中扩展多行")  

            elif method == "还原扩行":  
                word_sep = params["word_sep"] 
                depend_row= params["depend_row"] 
                if not word_sep:  
                    print("错误：分隔符不能为空")  
                    continue  
                try:  
                    df = CLEAN_expand_rows_REVERT(df, word_sep, columns,depend_row)  
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    continue  
                print(f"已使用分隔符 {word_sep} 在列 {columns} 中还原扩展多行")  

            elif method == "扩展多列":  
                word_sep2 = params["word_sep2"]  
                if not word_sep2:  
                    print("错误：分隔符不能为空")  
                    return  
                try:  
                    df = CLEAN_expand_cols(df, word_sep2, columns[0])
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    return  
                print(f"已使用分隔符 {word_sep2} 在列 {columns[0]} 中扩展多列")  

            elif method == "按列去重":  
                df.drop_duplicates(subset=columns, keep='first', inplace=True)  
                print(f"已按列 {columns} 去重")  

            elif method == "升序排列":  
                df.sort_values(by=columns, ascending=[True] * len(columns), inplace=True)  
                print(f"已按列 {columns} 升序排列")  

            elif method == "降序排列":  
                df.sort_values(by=columns, ascending=[False] * len(columns), inplace=True)  
                print(f"已按列 {columns} 降序排列")  

            elif method == "删除本列":  
                df.drop(columns=columns, inplace=True)  
                print(f"已删除列 {columns}")  

            elif method == "选列保留":  
                all_columns = df.columns.tolist()  
                columns_to_drop = [col for col in all_columns if col not in columns]  
                df.drop(columns=columns_to_drop, inplace=True)  
                print(f"已保留列 {columns}")  

            elif method == "整体去重":  
                df.drop_duplicates(inplace=True)  
                print("已完成整体去重")  

            elif method == "重置索引":  
                df.reset_index(inplace=True)  
                print("已重置索引")  

            elif method == "查看数据":  
                pass  # 不执行任何操作，仅用于查看数据  

            elif method == "批量规整":  
                guize = params["guize"]
                print(guize)
                df = CLEAN_replay_operations(df, pd.DataFrame(eval(str(guize))))
                print("已完成批量规整")  

            elif method == "加关键词" or method == "赋关键词":  
                mark_column_name = params["mark_column_name"]   
                include_str = params["include_str"] 
                exclude_str = params["exclude_str"]   
                predefined_column = params["predefined_column"]  
                predefined_objects = params["predefined_objects"]   
                mark_value = params["mark_value"]    
                df = CLEAN_filter_and_mark_rows(df, columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, method) 
                print(f"已完成 {method} 操作")  

            else:  
                print(f"不支持的操作：{method}")  # 对于不支持的操作类型，打印一条消息并跳过  

            print(f"操作完成：{ids} {method}")  # 打印操作完成的消息  

        except Exception as e:
            print(f"操作失败：{ids} {method} 原因：{str(e)}")

    print("所有批量规整工作完成。")
    return df


def CLEAN_fill_column_c_based_on_a_and_b(df, col_a, col_b, col_c):  
    """  
    对于DataFrame df中col_a列具有相同值的行，将相关行中col_b列的值（除了所在行的值）  
    填写到col_c列中，以逗号分隔。  
      
    参数:  
    df (DataFrame): 输入的DataFrame。  
    col_a (str): 用于比较的列名。  
    col_b (str): 需要被复制值的列名。  
    col_c (str): 目标列名，用于存储合并后的值。  
      
    返回值:  
    DataFrame: 处理后的DataFrame。  
    """  
    # 创建一个空的list用于存储结果  
    df[col_c] = ''  
      
    # 对col_a列中的每个唯一值进行处理  
    for value_a in df[col_a].unique():  
        # 筛选出col_a列等于当前唯一值的所有行  
        sub_df = df[df[col_a] == value_a]  
          
        # 如果筛选出的子DataFrame有多于一行  
        if len(sub_df) > 1:  
            # 遍历子DataFrame的每一行  
            for index, row in sub_df.iterrows():  
                # 从子DataFrame中去掉当前行，然后获取col_b列的值  
                other_values_b = sub_df.drop(index)[col_b].tolist()  
                # 将这些值用逗号连接成一个字符串，并赋值给当前行的col_c列  
                df.at[index, col_c] = ','.join(map(str, other_values_b))  
      
    return df  

def CLEAN_easystat(df, guize,method,*myselect_columns):  
    """  
    根据operation_log中的记录，在df上复现操作。  
    """
    #print(guize)
    mark_column_name=guize['结果列'][0]
    df[mark_column_name]=""

    guize=guize.fillna("XXX---XXXXXXXXX---XXXXX")
    allx=len(guize)
    for ids,cols in guize.iterrows(): 
        print(ids, allx)
        if len(myselect_columns)==0:  
            selected_columns=str(cols["查找位置"]).split("|")
        else:
            selected_columns=myselect_columns[0]
        include_str=cols["值"]
        exclude_str=cols['排除值']
        predefined_column=cols['适用范围列']
        predefined_objects=cols['适用范围']
        
        mark_value=str(cols["值"]).split("|")[0]
        
        CLEAN_filter_and_mark_rows(df, selected_columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, method) 
    df[mark_column_name] = df[mark_column_name].replace("", "其他") 
    df[mark_column_name] = df[mark_column_name].str.replace("nan;", "")
    df[mark_column_name] = df[mark_column_name].str.replace("nan", "其他") 
    return df


def CLEAN_filter_and_mark_rows(df, columns, include_str, exclude_str, predefined_column=None, predefined_objects=None, mark_column_name='mark', mark_value='marked', method_name='add'):  
    """  
    对DataFrame中的行进行筛选和标记。  
      
    参数：  
        df (pandas.DataFrame): 需要处理的DataFrame。  
        columns (list of str): 要检查的列名列表。  
        include_str (str): 需要包含的字符串。  
        exclude_str (str): 需要排除的字符串。  
        predefined_column (str, optional): 预定义的列名，用于进一步筛选行。默认为None。  
        predefined_objects (list, optional): 预定义的对象列表，与predefined_column配合使用。默认为None。  
        mark_column_name (str, optional): 标记列的列名。默认为'mark'。  
        mark_value (str, optional): 标记值。默认为'marked'。  
        method_name (str, optional): 处理标记列的方法，'加关键词'表示追加标记值，'赋关键词'表示替换为标记值。默认为'加关键词'。  
      
    返回：  
        pandas.DataFrame: 处理后的DataFrame，包含新增或更新的标记列。  
      
    步骤：  
    1. 创建一个布尔掩码，初始化为True，以便保留所有行。  
    2. 如果提供了预定义列和对象，则基于它们更新掩码，仅包括预定义对象所在的行。  
    3. 使用布尔掩码筛选包含include_str且不包含exclude_str的行。  
    4. 如果需要标记的列不存在于DataFrame中，则新建该列，并初始化为空字符串。  
    5. 根据method_name的值处理标记列：  
        - 如果method_name为'add'，则将标记值追加到筛选结果为True的行的标记列中（用分号隔开）。  
        - 如果method_name为'replace'，则将筛选结果为True的行的标记列替换为标记值。  
    6. 返回处理后的DataFrame，注意行数保持不变，未筛选出的行不会被删除。  
    """  

    # 创建一个布尔掩码，初始化为True，以便保留所有行  
    mask = pd.Series(True, index=df.index)  
    
    if predefined_column=="所有列" and predefined_objects=="所有对象":
        predefined_column=None
        predefined_objects=None
      
    # 如果提供了预定义列和对象，则基于它们更新掩码  
    if predefined_column and predefined_objects:  
        if predefined_column not in df.columns:  
            raise ValueError(f"Predefined column '{predefined_column}' not found in DataFrame.")  
        #mask &= df[predefined_column] == predefined_objects  # 更新掩码以仅包括预定义对象所在的行  
        mask &= df[predefined_column].str.contains(predefined_objects, na=False)
    
    
    #print( columns, include_str, exclude_str, predefined_column, predefined_objects)
    
      
    # 使用布尔掩码筛选包含include_str且不包含exclude_str的行  
    # 检查exclude_str是否为空  
    if exclude_str:  
        mask &= df[columns].apply(lambda col: col.str.contains(include_str, na=False) & ~col.str.contains(exclude_str, na=False)).any(axis=1)  
    else:  
        mask &= df[columns].apply(lambda col: col.str.contains(include_str, na=False)).any(axis=1)  
           
    # 如果需要标记的列不存在于DataFrame中，则新建该列，并初始化为空字符串  
    if mark_column_name not in df.columns:  
        df[mark_column_name] = ''  # 初始化为空字符串，确保后续可以正确添加或替换标记值  
    # 根据method_name的值处理标记列  
    if method_name == '加关键词':  
        # 追加标记值（用分号隔开），仅对筛选结果为True的行操作  
        df.loc[mask, mark_column_name] = df.loc[mask, mark_column_name].astype(str) + ';' + mark_value  
    elif method_name == '赋关键词':  
        # 替换为标记值，仅对筛选结果为True的行操作  
        df.loc[mask, mark_column_name] = mark_value  
    else:  
        raise ValueError(f"Invalid method_name '{method_name}'. Only 'add' or 'replace' are accepted.")  
    df[mark_column_name] = df[mark_column_name].str.lstrip(';')

    # 返回处理后的DataFrame，注意行数保持不变，未筛选出的行不会被删除  
    return df

def CLEAN_expand_cols(dfs, delimiter, col):  
    """  
    扩展指定列的函数。  只支持单列。
  
    参数:  
    df (DataFrame): 输入的DataFrame。  
    delimiter (str): 用于拆分指定列的分隔符。  
    col (str): 需要拆分的列名。  
  
    返回值:  
    DataFrame: 扩展后的DataFrame，包含原始列和新的拆分列。  
  
    步骤说明:  
    1. 使用分隔符拆分指定的列，并创建虚拟/指示列。  
    2. 通过将原始列名和唯一值附加到一起来重命名列。  
    3. 将原始DataFrame与虚拟DataFrame沿列方向（axis=1）连接。  
    """  
    # Split the column by delimiter and create dummy columns    
    dummy_df = dfs[col].str.get_dummies(sep=delimiter)    
        
    # Rename the columns by appending the original column name and the unique values    
    dummy_df.columns = [f'{col}_{value}' for value in dummy_df.columns]    
        
    # Concatenate the original dataframe with the dummy dataframe    
    result_df = pd.concat([dfs, dummy_df], axis=1)    
   
    return result_df
  
class CLEAN_DataFrameUpdateApp:
    '''更新表格 '''
    def __init__(self, root, df1, df2):
        self.root = root
        self.root.title("更新表格")
        self.center_window(1024, 768)
        self.df1 = df1


        # 获取 df1 和 df2 共同的列
        common_columns = list(set(df1.columns) & set(df2))
        self.df2 = df2[common_columns]
        if not common_columns:
            messagebox.showerror("错误", "两个 DataFrame 没有共同的列！")
            self.root.destroy()
            return

        # 创建下拉菜单选择 key_column
        ttk.Label(root, text="请选择关键列:").pack(pady=10)
        self.key_column_var = tk.StringVar(root)
        self.key_column_var.set(common_columns[0])  # 默认选择第一个共同列
        self.key_column_dropdown = ttk.Combobox(root, textvariable=self.key_column_var, values=common_columns)
        self.key_column_dropdown.pack()

        # 创建按钮执行更新操作
        self.update_button = ttk.Button(root, text="更新表格", command=self.update_dataframe)
        self.update_button.pack(pady=20)

    def center_window(self, width, height):
        """窗口居中显示"""
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def update_dataframe(self):
        """执行更新操作"""
        key_column = self.key_column_var.get()
        updated_df = Small_update_df(self.df1, self.df2, key_column)

        PROGRAM_display_df_in_treeview(updated_df,0,0)


#行选择
def CLEAN_select_rows_interactively(df):
    """ 交互式行选择函数（完整版） """
    
    def parse_top_n(input_str, max_rows):
        """ 解析前N行选择 """
        if not input_str.isdigit():
            raise ValueError("请输入有效数字")
        
        n = min(int(input_str), max_rows)
        return list(range(n)) if n > 0 else []

    def parse_range_selection(input_str, max_rows):
        """ 解析范围选择 """
        if not input_str:
            return []
        
        selected = set()
        parts = [p.strip() for p in input_str.split(',') if p.strip()]
        
        for part in parts:
            if '-' in part:
                start, end = map(int, part.split('-', 1))
                start = max(1, start)
                end = min(end, max_rows)
                if start > end:
                    start, end = end, start
                selected.update(range(start-1, end))
            else:
                row = int(part)
                if 1 <= row <= max_rows:
                    selected.add(row-1)
        
        return sorted(selected)

    def parse_random_selection(input_str, max_rows):
        """ 解析随机选择 """
        if not input_str.isdigit():
            raise ValueError("请输入有效数字")
        
        n = int(input_str)
        if n <= 0:
            return []
        
        n_samples = min(n, max_rows)
        return random.sample(range(max_rows), n_samples)

    
    df.reset_index(drop=True)
    root = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("行选择器")
    
    # 设置窗口居中
    window_width = 450
    window_height = 250
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width - window_width) // 2
    y = (screen_height - window_height) // 2
    root.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    # 创建模式选择变量
    mode_var = tk.StringVar(value="前N行")
    
    # 创建主界面组件
    frame = ttk.Frame(root, padding=20)
    frame.pack(expand=True, fill=tk.BOTH)
    
    # 模式选择下拉菜单
    ttk.Label(frame, text="选择模式:").grid(row=0, column=0, sticky=tk.W)
    mode_menu = ttk.OptionMenu(
        frame, mode_var,
        "前N行",
        "前N行",
        "按条件选择",
        "随机抽取"
    )
    mode_menu.grid(row=0, column=1, sticky=tk.EW, padx=10, pady=5)
    
    # 输入框
    input_label = ttk.Label(frame, text="输入参数:")
    input_label.grid(row=1, column=0, sticky=tk.W, pady=5)
    
    input_entry = ttk.Entry(frame, width=35)
    input_entry.grid(row=1, column=1, padx=10, pady=5)
    
    # 动态示例提示
    example_text = tk.Text(frame, height=5, width=45, wrap=tk.WORD)
    example_text.grid(row=2, column=0, columnspan=2, pady=10)
    
    def update_example():
        example_text.configure(state=tk.NORMAL)
        example_text.delete(1.0, tk.END)
        mode = mode_var.get()
        if mode == "前N行":
            example_text.insert(tk.END, "格式说明：\n• 输入数字（如100）：返回前100行\n• 自动处理超出行数情况")
        elif mode == "按条件选择":
            example_text.insert(tk.END, "格式说明：\n• 2-5：返回2到5行\n• 1,3,5：返回指定行\n• 3-9,11-15：多范围选择")
        elif mode == "随机抽取":
            example_text.insert(tk.END, "格式说明：\n• 输入数字（如50）：随机抽取50行\n• 自动去重且不重复抽样")
        example_text.configure(state=tk.DISABLED)
    
    mode_var.trace_add("write", lambda *_: update_example())
    update_example()
    
    # 确认按钮
    def on_confirm():
        try:
            user_input = input_entry.get().strip()
            mode = mode_var.get()
            
            if mode == "前N行":
                indices = parse_top_n(user_input, len(df))
            elif mode == "按条件选择":
                indices = parse_range_selection(user_input, len(df))
            elif mode == "随机抽取":
                indices = parse_random_selection(user_input, len(df))
                
            if not indices:
                messagebox.showwarning("警告", "没有选择有效行！")
                root.result = pd.DataFrame()
            else:
                root.result = df.iloc[indices]
            
            root.destroy()
            PROGRAM_display_df_in_treeview(root.result ,0,0)
        except Exception as e:
            messagebox.showerror("错误", f"输入错误: {str(e)}")
    
    confirm_btn = ttk.Button(frame, text="确认", command=on_confirm)
    confirm_btn.grid(row=3, column=1, sticky=tk.E, pady=10)
    
    # 取消按钮
    def on_cancel():
        root.result = pd.DataFrame()
        root.destroy()
    
    cancel_btn = ttk.Button(frame, text="取消", command=on_cancel)
    cancel_btn.grid(row=3, column=0, sticky=tk.W, pady=10)
    
    root.mainloop()
    return getattr(root, 'result', pd.DataFrame())

   
def CLEAN_expand_rows(df, sep, cols):  
    """  
    拆分成行的函数，sep是分隔符，cols是需要拆分的列的列表。 支持多列。
    """  
    # 重置索引  
    df = df.reset_index(drop=True)  
      
    # 拆分指定的列  
    frames = []  
    for col in cols:  
        frame = df[col].str.split(sep, expand=True).stack().reset_index(level=1, drop=True).to_frame(col)  
        frames.append(frame)  
      
    # 合并拆分的列  
    result = pd.concat(frames, axis=1)  
      
    # 获取原始数据框中非拆分列的部分  
    other_cols = [col for col in df.columns if col not in cols]  
    remaining_df = df[other_cols]  
      
    # 如果有非拆分列，则合并拆分的列与非拆分列的部分  
    if other_cols:  
        result = pd.merge(result, remaining_df, left_index=True, right_index=True)  
    result['groupby_column_0'] = result.index.copy()  
    # 清理结果并返回  
    result = result.reset_index(drop=True)  
    return result    

def CLEAN_expand_rows_REVERT(df, delimiter, merge_columns, groupby_column="groupby_column_0"):
    """
    根据指定的列合并df的行，并将特定列的值用分号连接起来，同时保留其他列的第一个值。
    
    参数:
    df (DataFrame): 要合并的DataFrame。
    delimiter (str): 用于合并字符串值的分隔符。
    merge_columns (list of str): 需要合并值的列名列表（允许传入多列）。
    groupby_column (str): 用于分组的列名，默认为"groupby_column_0"。
    
    返回值:
    merged_df (DataFrame): 合并后的DataFrame。
    """
    # 检查merge_columns是否为空
    if not merge_columns:
        raise ValueError("merge_columns不能为空")
    
    # 检查merge_columns中的列是否都在df中
    if not all(col in df.columns for col in merge_columns):
        raise ValueError("merge_columns中的某些列不在df中")
    
    # 检查groupby_column是否在df中
    if groupby_column not in df.columns:
        raise ValueError(f"groupby_column '{groupby_column}' 不在df中")
    
    # 对于需要合并的字符串列，使用groupby和agg来合并它们的值
    merged_values = df.groupby(groupby_column)[merge_columns].agg(lambda x: delimiter.join(x.astype(str)))
    
    # 对于其他列，使用first()保留每个分组的第一个值
    other_cols = [col for col in df.columns if col not in merge_columns and col != groupby_column]
    first_values = df.groupby(groupby_column)[other_cols].first()
    
    # 将合并后的值和其他列的第一个值合并到一个DataFrame中
    merged_df = pd.concat([first_values, merged_values], axis=1)
    
    # 重置索引
    merged_df.reset_index(inplace=True)
    
    return merged_df
   
class CLEAN_DataMaskingApp:
    """
    数据脱敏函数。
    """
    def __init__(self, rs,df):

        self.root = tk.Toplevel()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)
        self.root.title("数据脱敏工具")
        self.center_window(600, 800)
        self.data = df.copy()  # 通过 DataFrame 传入数据
        self.mapping_keys = {}  # 存储脱敏映射关系
        self.original_data = df.copy()  # 保存原始数据用于恢复

        # GUI 布局
        self.frame = ttk.Frame(self.root)
        self.frame.pack(fill=tk.BOTH, expand=True)

        # Treeview 显示列名
        self.tree_frame = ttk.Frame(self.frame)
        self.tree_frame.pack(fill=tk.BOTH, expand=True)
        self.tree = ttk.Treeview(self.tree_frame, columns=("列名",), show="headings")
        self.tree.heading("列名", text="列名")
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 添加 Treeview 滚动条
        self.tree_scroll = ttk.Scrollbar(self.tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=self.tree_scroll.set)

        # 显示列名
        for col in self.data.columns:
            self.tree.insert("", tk.END, values=(col,))

        # 进度条
        self.progress = ttk.Progressbar(self.frame, orient=tk.HORIZONTAL, length=400, mode='determinate')
        self.progress.pack(pady=10)

        # 更多功能下拉框（加宽）
        self.extra_options_label = ttk.Label(self.frame, text="更多功能:")
        self.extra_options_label.pack(pady=5)
        
        self.extra_options = ttk.Combobox(
            self.frame, 
            values=["无", "去除非选中列中的选中列信息", "去除非选中列中的选中列信息和医患信息"], 
            state="readonly",
            width=40
        )
        self.extra_options.current(0)
        self.extra_options.pack(pady=5)
        self.extra_options.bind("<<ComboboxSelected>>", self.toggle_extra_fields)

        # 姓氏和关键词输入框
        self.surnames_frame = ttk.Frame(self.frame)
        self.surnames_label = ttk.Label(self.surnames_frame, text="姓氏（以|隔开）:")
        self.surnames_label.grid(row=0, column=0, sticky="w", pady=5)

        self.surnames_text = tk.Text(self.surnames_frame, height=5, width=50)
        self.surnames_scroll = ttk.Scrollbar(self.surnames_frame, orient=tk.VERTICAL, command=self.surnames_text.yview)
        self.surnames_text.configure(yscrollcommand=self.surnames_scroll.set)
        self.surnames_text.grid(row=1, column=0, sticky="nsew")
        self.surnames_scroll.grid(row=1, column=1, sticky="ns")

        self.keywords_frame = ttk.Frame(self.frame)
        self.keywords_label = ttk.Label(self.keywords_frame, text="医患识别关键词（以|隔开）:")
        self.keywords_label.grid(row=0, column=0, sticky="w", pady=5)

        self.keywords_text = tk.Text(self.keywords_frame, height=5, width=50)
        self.keywords_scroll = ttk.Scrollbar(self.keywords_frame, orient=tk.VERTICAL, command=self.keywords_text.yview)
        self.keywords_text.configure(yscrollcommand=self.keywords_scroll.set)
        self.keywords_text.grid(row=1, column=0, sticky="nsew")
        self.keywords_scroll.grid(row=1, column=1, sticky="ns")

        # 恢复默认值按钮
        self.btn_reset = ttk.Button(self.frame, text="恢复默认值", command=self.reset_to_default)

        # 脱敏按钮
        self.btn_mask = ttk.Button(self.frame, text="开始脱敏", command=self.mask_data)
        self.btn_mask.pack(pady=10)
        self.init_surnames_and_keywords()


    def center_window(self, width, height):
        """窗口居中显示"""
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def update_treeview(self):
        """更新 Treeview 显示列名"""
        self.tree.delete(*self.tree.get_children())
        for col in self.data.columns:
            self.tree.insert("", tk.END, values=(col,))

    def toggle_extra_fields(self, event=None):
        """切换显示姓氏和关键词输入框"""
        if self.extra_options.get() == "去除非选中列中的选中列信息和医患信息":
            self.surnames_frame.pack(pady=10)
            self.keywords_frame.pack(pady=10)
            self.btn_reset.pack(pady=10)
        else:
            self.surnames_frame.pack_forget()
            self.keywords_frame.pack_forget()
            self.btn_reset.pack_forget()


    def replace_related_values_with_xxx(self, df, columns):
        """替换相关值为XXX，但保留空值不变"""
        df_copy = df.copy()
        total_steps = len(columns) * len(df_copy)
        current_step = 0

        for col in columns:
            for index, row in df_copy.iterrows():
                value_to_replace = row[col]
                if pd.notna(value_to_replace):
                    for other_col in df_copy.columns:
                        if other_col != col and pd.notna(row[other_col]):
                            df_copy.at[index, other_col] = str(row[other_col]).replace(str(value_to_replace), 'XXX')
                current_step += 1
                self.progress['value'] = (current_step / total_steps) * 100
                self.root.update_idletasks()
        return df_copy

    def process_dataframe(self, df):
        """去除特定情境下的人名"""
        surnames = self.get_surnames_from_xml()
        keywords = self.get_keywords_from_xml()

        surnames_pattern = '|'.join(surnames)
        keywords_pattern = '|'.join(keywords)
        pattern = rf'({keywords_pattern})\s*([{surnames_pattern}][\u4e00-\u9fa5]{{0,2}})|([{surnames_pattern}][\u4e00-\u9fa5]{{0,2}})\s*({keywords_pattern})'

        def remove_names(text):
            if not isinstance(text, str):
                return text
            return re.sub(pattern, lambda match: f'{match.group(1) or ""}XXX{match.group(4) or ""}', text)

        total_steps = len(df.columns) * len(df)
        current_step = 0
        for col in df.columns:
            df[col] = df[col].apply(remove_names)
            current_step += len(df)
            self.progress['value'] = (current_step / total_steps) * 100
            self.root.update_idletasks()
        return df

    def mask_data(self):
        """对选择的列进行脱敏并保存"""
        if self.data is None:
            messagebox.showwarning("提示", "数据为空！")
            return

        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("提示", "请选择需要脱敏的列！")
            return

        selected_columns = [self.tree.item(item, "values")[0] for item in selected_items]
        self.original_data = self.data.copy()

        # 处理额外功能
        extra_option = self.extra_options.get()
        if extra_option == "去除非选中列中的选中列信息":
            self.progress['value'] = 0
            self.data = self.replace_related_values_with_xxx(self.data, selected_columns)
        elif extra_option == "去除非选中列中的选中列信息和医患信息":
            self.progress['value'] = 0
            self.data = self.replace_related_values_with_xxx(self.data, selected_columns)
            self.progress['value'] = 0
            self.data = self.process_dataframe(self.data)

        # 生成脱敏映射关系
        self.mapping_keys = {}
        for col in selected_columns:
            unique_values = self.data[col].unique()
            mapping = {value: str(uuid.uuid4()) for value in unique_values}
            self.mapping_keys[col] = mapping
            self.data[col] = self.data[col].map(mapping)



        # 强制保存
        output_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx")],
            title="保存脱敏后的文件"
        )
        if not output_path:
            return

        self.data.to_excel(output_path, index=False)

        # 保存密钥
        key_path = os.path.join(
            os.path.dirname(output_path),
            f"masking_keys_{os.path.splitext(os.path.basename(output_path))[0]}.json"
        )
        with open(key_path, "w", encoding="utf-8") as f:
            json.dump(self.mapping_keys, f, ensure_ascii=False, indent=4)

        messagebox.showinfo("完成", f"文件已保存至:\n{output_path}\n密钥已保存至:\n{key_path}")
        # 假设这里有一个函数用于显示数据
        PROGRAM_display_df_in_treeview(self.data, 0, 0)


    def init_surnames_and_keywords(self):
        """初始化姓氏和关键词"""
        global csdir
        if not os.path.exists(os.path.join(csdir, "maskname.xml")):
            self.create_default_maskname_xml()
        self.update_text_fields_from_xml()

    def create_default_maskname_xml(self):
        """创建默认的 maskname.xml 文件"""
        global csdir
        root = ET.Element("maskname")
        surnames = ET.SubElement(root, "surnames")
        surnames.text = "王|李|张|刘|陈|杨|赵|吴|徐|孙|胡|朱|林|何|郭|罗|梁|宋|郑|谢|韩|唐|冯|董|萧|曹|袁|邓|许|傅|沈|曾|彭|吕|卢|蒋|蔡|贾|丁|魏|薛|叶|阎"
        keywords = ET.SubElement(root, "keywords")
        keywords.text = "患者|病人|医生|护士|护工|护士长|护长|主任|主治医生|医师"

        tree = ET.ElementTree(root)
        tree.write(os.path.join(csdir, "maskname.xml"), encoding="utf-8", xml_declaration=True)

    def update_text_fields_from_xml(self):
        """从 XML 文件更新输入框内容"""
        global csdir
        tree = ET.parse(os.path.join(csdir, "maskname.xml"))
        root = tree.getroot()

        surnames = root.find("surnames").text
        keywords = root.find("keywords").text

        self.surnames_text.delete("1.0", tk.END)
        self.surnames_text.insert("1.0", surnames)
        self.keywords_text.delete("1.0", tk.END)
        self.keywords_text.insert("1.0", keywords)

    def save_config(self, event=None):
        """实时保存配置到 maskname.xml"""
        global csdir
        root = ET.Element("maskname")
        surnames = ET.SubElement(root, "surnames")
        surnames.text = self.surnames_text.get("1.0", tk.END).strip()
        keywords = ET.SubElement(root, "keywords")
        keywords.text = self.keywords_text.get("1.0", tk.END).strip()

        tree = ET.ElementTree(root)
        tree.write(os.path.join(csdir, "maskname.xml"), encoding="utf-8", xml_declaration=True)

    def get_surnames_from_xml(self):
        """从 XML 文件中获取姓氏"""
        global csdir
        tree = ET.parse(os.path.join(csdir, "maskname.xml"))
        root = tree.getroot()
        return root.find("surnames").text.split("|")

    def get_keywords_from_xml(self):
        """从 XML 文件中获取关键词"""
        global csdir
        tree = ET.parse(os.path.join(csdir, "maskname.xml"))
        root = tree.getroot()
        return root.find("keywords").text.split("|")

    def reset_to_default(self):
        """恢复默认值"""
        confirm = messagebox.askyesno("确认操作", "是否要恢复默认值？该操作会导致自定义配置丢失！")
        if confirm:
            self.create_default_maskname_xml()
            self.update_text_fields_from_xml()
            messagebox.showinfo("提示", "已恢复默认值！")


def CLEAN_restore_data(data):
    """恢复数据"""
    if 1==1:
        key_path = filedialog.askopenfilename(
            title="选择密钥文件",
            filetypes=[("JSON文件", "*.json")]
        )
        if not key_path:
            return

        try:
            with open(key_path, "r", encoding="utf-8") as f:
                keys = json.load(f)
        except Exception as e:
            messagebox.showerror("错误", f"密钥文件读取失败:\n{str(e)}")
            return

        data2 = data.copy()
        for col, mapping in keys.items():
            reverse_mapping = {v: k for k, v in mapping.items()}
            data2[col] = data2[col].map(reverse_mapping)

        print("完成", "数据恢复成功！")
        PROGRAM_display_df_in_treeview(data2, 0, 0)
  
def CLEAN_data_masking_qixie(df, log_callback=None, progress_callback=None):
    """医疗器械脱敏函数"""

    def update_progress(value):
        if progress_callback:
            progress_callback(value)
    
    def remove_medical_names(df):
        """Remove doctor/patient names from the dataframe"""
        print("开始移除医生/患者姓名...")
        # Try to load surnames and keywords from XML
        try:
            surnames = get_surnames_from_xml()
            keywords = get_keywords_from_xml()
            print("成功从XML加载姓氏和关键词")
        except Exception as e:
            print(f"无法从XML加载配置，使用默认值: {str(e)}")
            # Use defaults if XML not available
            surnames = "王|李|张|刘|陈|杨|赵|吴|徐|孙|胡|朱|林|何|郭|罗|梁|宋|郑|谢|韩|唐|冯|董|萧|曹|袁|邓|许|傅|沈|曾|彭|吕|卢|蒋|蔡|贾|丁|魏|薛|叶|阎".split("|")
            keywords = "患者|病人|医生|护士|护工|护士长|护长|主任|主治医生|医师".split("|")
        
        # Build regex pattern
        surnames_pattern = '|'.join(surnames)
        keywords_pattern = '|'.join(keywords)
        pattern = rf'({keywords_pattern})\s*([{surnames_pattern}][\u4e00-\u9fa5]{{0,2}})|([{surnames_pattern}][\u4e00-\u9fa5]{{0,2}})\s*({keywords_pattern})'
        
        def remove_names(text):
            """Remove names from text based on pattern"""
            if not isinstance(text, str):
                return text
            return re.sub(pattern, lambda match: f'{match.group(1) or ""}XXX{match.group(4) or ""}', text)
        
        # Apply to all columns
        for col in df.columns:
            df[col] = df[col].apply(remove_names)
        
        print("医生/患者姓名移除完成")
        return df

    def replace_related_values_with_xxx(df, columns):
        """替换相关值为XXX，但保留空值不变"""
        print("开始替换相关值为XXX...")
        df_copy = df.copy()
        
        for col in columns:
            print(f"正在处理列: {col}")
            for index, row in df_copy.iterrows():
                value_to_replace = row[col]
                # 只有当值不是空值时才进行处理
                if pd.notna(value_to_replace):
                    for other_col in df_copy.columns:
                        if other_col != col:
                            # 确保目标列的值也不是空值
                            if pd.notna(row[other_col]):
                                df_copy.at[index, other_col] = str(row[other_col]).replace(str(value_to_replace), 'XXX')
        
        print("相关值替换完成")
        return df_copy

    def get_surnames_from_xml():
        """Get surnames from XML configuration"""
        csdir = os.path.dirname(__file__)
        tree = ET.parse(os.path.join(csdir, "maskname.xml"))
        root = tree.getroot()
        return root.find("surnames").text.split("|")

    def get_keywords_from_xml():
        """Get keywords from XML configuration"""
        csdir = os.path.dirname(__file__)
        tree = ET.parse(os.path.join(csdir, "maskname.xml"))
        root = tree.getroot()
        return root.find("keywords").text.split("|")
    
    current_progress = 0
    if progress_callback:
        progress_callback(current_progress)
    
    print("开始医疗器械数据脱敏处理")
    
    # Step 1: Keep only the specified columns
    required_columns = [
        "报告编码", "报告日期", "单位名称", "产品名称", "注册证编号/曾用注册证编号", 
        "型号", "规格", "产品批号", "产品编号", "生产日期", "有效期至", 
        "上市许可持有人名称", "事件发生日期", "发现或获知日期", "伤害", 
        "伤害表现", "年龄", "年龄类型", "性别", "器械故障表现", 
        "预期治疗疾病或作用", "器械使用日期", "使用场所", "场所名称", 
        "使用过程", "合并用药/械情况说明", "事件原因分析", "事件原因分析描述", 
        "初步处置情况", "经营企业使用单位报告状态", "报告类型", 
        "是否开展了调查", "调查情况", "关联性评价", "事件原因分析.1", 
        "是否需要开展产品风险评价", "计划提交时间", "是否已采取控制措施", 
        "具体控制措施", "未采取控制措施原因", "是否为错报误报报告", 
        "错报误报说明", "是否合并报告", "合并报告编码", "持有人报告状态"
    ]
    
    print("过滤数据列，仅保留所需列...")
    # Filter the dataframe to keep only required columns
    df = df[required_columns]
    current_progress += 10
    if progress_callback:
        progress_callback(current_progress)
    
    # Step 2: Perform column-based masking (removing sensitive info)
    columns_to_mask = [
        "单位名称", "注册证编号/曾用注册证编号", "型号", "规格", 
        "产品批号", "产品编号", "上市许可持有人名称", "场所名称"
    ]
    
    # First try to remove doctor/patient names
    print("开始移除医疗相关姓名...")
    df = remove_medical_names(df)
    current_progress += 20
    if progress_callback:
        progress_callback(current_progress)
    
    # Then perform column-based masking (replace related values with XXX)
    print("开始列基础脱敏...")
    df = replace_related_values_with_xxx(df, columns_to_mask)
    current_progress += 20
    if progress_callback:
        progress_callback(current_progress)
    
    # Step 3: Perform selective masking (with recovery capability)
    print("开始选择性脱敏(可恢复)...")
    selective_columns = columns_to_mask  # Same columns as above
    mapping_keys = {}
    
    for col in selective_columns:
        print(f"处理列: {col}")
        unique_values = df[col].unique()
        mapping = {value: str(uuid.uuid4()) for value in unique_values if pd.notna(value)}
        mapping_keys[col] = mapping
        df[col] = df[col].map(mapping)
    
    current_progress += 20
    if progress_callback:
        progress_callback(current_progress)
    
    # Step 4: Save the files
    print("准备保存脱敏文件...")
    output_path = filedialog.asksaveasfilename(
        defaultextension=".xlsx", 
        filetypes=[("Excel 文件", "*.xlsx")],
        title="保存脱敏后的文件"
    )
    
    if not output_path:
        print("用户取消了保存操作")
        return None, None  # User cancelled
    
    # Save the masked data
    print("正在保存脱敏数据...")
    df.to_excel(output_path, index=False)
    
    # Save the recovery file (only for selective masking)
    output_dir = os.path.dirname(output_path)
    output_filename = os.path.splitext(os.path.basename(output_path))[0]
    recovery_path = os.path.join(output_dir, f"{output_filename}_recovery.json")
    
    print("正在保存恢复文件...")
    with open(recovery_path, "w", encoding="utf-8") as f:
        json.dump(mapping_keys, f, ensure_ascii=False, indent=4)
    
    current_progress += 20
    if progress_callback:
        progress_callback(current_progress)
    
    print(f"脱敏完成!\n脱敏文件: {output_path}\n恢复文件: {recovery_path}")
    messagebox.showinfo(
        "成功", 
        f"脱敏完成！\n"
        f"脱敏数据已保存到: {output_path}\n"
        f"恢复文件已保存到: {recovery_path}"
    )
    
    PROGRAM_display_df_in_treeview(df,0,0)





############################################################################################################################
#实用工具函数
############################################################################################################################
class AAA_09_Tools:
    pass


#秩和检验，包括TOOLS_rank_sum_test_cout，TOOLS_rank_sum_test两个函数#######################
def TOOLS_rank_sum_test_cout(group1, group2):
    # 将两个样本组合并
    combined = np.concatenate((group1, group2))
    
    # 对合并后的样本进行排序，并获取排序后的索引
    sorted_indices = np.argsort(combined)
    
    # 根据排序后的索引，计算每个样本的秩次
    ranks = np.empty_like(sorted_indices)
    ranks[sorted_indices] = np.arange(1, len(combined) + 1)
    
    # 分别提取两个样本组的秩次
    ranks_group1 = ranks[:len(group1)]
    ranks_group2 = ranks[len(group1):]
    
    # 计算两个样本组的秩和
    rank_sum_group1 = np.sum(ranks_group1)
    rank_sum_group2 = np.sum(ranks_group2)
    
    # 计算检验统计量U
    U = min(rank_sum_group1, rank_sum_group2)
    
    # 计算p值
    p_value = ranksums(ranks_group1, ranks_group2)[1]
    
    return U, p_value
 
#秩和检验附属函数
def TOOLS_rank_sum_test(df):  
    print(
        """
    **********************************************
    秩和检验工具

    功能：
    1. 对两组数据进行秩和检验，计算统计量 U 和 p 值。
    2. 提供 GUI 界面供用户选择数据列和比较对象。

    使用方法：
    1. 启动程序后，选择比较对象所处的列（如：通用名称）。
    2. 选择比较对象列中的两个不同元素。
    3. 选择比较目标列（如：年龄段）。
    4. 选择比较目标的数值列（如：计数）。
    5. 点击“确定”按钮，程序会计算并显示秩和检验结果。

    注意：
    - DataFrame 必须至少包含 3 列。
    - 比较对象列和目标列不能相同。
    **********************************************
    """
    )
    
    if len(df.columns) < 3:  
        print("DataFrame must have at least 3 columns.")  
        return  
  
    def print_selected_values():  
        # 获取选择的列名和目标列名  
        selected_column = column_var.get()  
        target_column = target_var.get()  
        value_column = value_var.get()  
  
        # 检查选择的列是否有效  
        if selected_column == target_column:  
            print("Selected column and target column cannot be the same.")  
            return  
  
        # 从DataFrame中提取选定的列和目标列的值  
        column_values = df[selected_column].unique()  
        target_values = df[target_column].unique()  
  
        # 获取选择的两个不同元素  
        element1 = element1_var.get()  
        element2 = element2_var.get()  
  
        # 找到与选定元素对应的行，并提取目标数值列的值  
        selected_rows1 = df[df[selected_column] == element1]  
        selected_rows2 = df[df[selected_column] == element2]  
        values1 = selected_rows1[value_column].values  
        values2 = selected_rows2[value_column].values  
        result=TOOLS_rank_sum_test_cout(values1, values2)  
        # 打印提取的数值列  
        data1=f"Values for {element1} in {target_column}: {values1}\n"
        data2=f"Values for {element2} in {target_column}: {values2}\n"  
        data3="U:"+str(result[0])+"\n"
        data4="P:"+str(result[1])+"\n"       

        PROGRAM_display_content_in_textbox("秩和检验结果：\n"+data1+data2+data3+data4)
  
    def update_elements(event=None):    
        selected_column = column_var.get()    
        element1_dropdown['values'] = sorted(df[selected_column].unique())    
        element2_dropdown['values'] = sorted(df[selected_column].unique()) 
  
    # 创建GUI窗口和组件    
    root = tk.Tk()   
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title('秩和检验工具')    
  
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）          
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 710  # 窗口宽度          
    wh = 200  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")   
  
    # 第一行：选择比较对象所处的列    
    ttk.Label(root, text="请选择比较对象所处的列(如：通用名称)").grid(row=0, column=0, padx=5, pady=5)    
    column_var = tk.StringVar(root)    
    column_var.set(df.columns[0])  # 默认选择第一列    
    column_dropdown = ttk.Combobox(root, textvariable=column_var, values=df.columns.tolist())    
    column_dropdown.grid(row=0, column=1, padx=5, pady=5)    
    column_dropdown.bind('<<ComboboxSelected>>', update_elements)    
    
    # 第二行：选择第一行选择结果列内两个不同元素    
    ttk.Label(root, text="请选择第一个比较对象:").grid(row=1, column=0, padx=5, pady=5)    
    element1_var = tk.StringVar(root)    
    element1_dropdown = ttk.Combobox(root, textvariable=element1_var)    
    element1_dropdown.grid(row=1, column=1, padx=5, pady=5)    
    ttk.Label(root, text="请选择第二个比较对象:").grid(row=1, column=2, padx=5, pady=5)    
    element2_var = tk.StringVar(root)    
    element2_dropdown = ttk.Combobox(root, textvariable=element2_var)    
    element2_dropdown.grid(row=1, column=3, padx=5, pady=5)    
    
    # 第三行：选择比较目标列    
    ttk.Label(root, text="请选择比较目标列(如：年龄段)").grid(row=2, column=0, padx=5, pady=5)    
    target_var = tk.StringVar(root)    
    target_var.set(df.columns[0])  # 默认选择第二列    
    target_dropdown = ttk.Combobox(root, textvariable=target_var, values=df.columns.tolist())    
    target_dropdown.grid(row=2, column=1, padx=5, pady=5)    
    
    # 第四行：选择比较数值列    
    ttk.Label(root, text="请选择比较目标的数值列（如：计数）:").grid(row=3, column=0, padx=5, pady=5)    
    value_var = tk.StringVar(root)    
    value_var.set(df.columns[0])  # 默认选择第三列    
    value_dropdown = ttk.Combobox(root, textvariable=value_var, values=df.columns.tolist())    
    value_dropdown.grid(row=3, column=1, padx=5, pady=5)    
    
    # 第五行：打印按钮    
    print_button = ttk.Button(root, text="确定", command=print_selected_values)    
    print_button.grid(row=4, column=3, padx=5, pady=5)    
    
    update_elements()  # 初始化时更新一次    
    root.mainloop()  # 运行GUI窗口


#ROR计算-DB
def TOOLS_ROR_from_DB_get_abcd(df, name_char, drug_col='drugname', pt_col='pt', soc_col='soc_name', a_col='a'):  
    # 筛选出drugname包含name_char的行  
    df_filtered = df[df[drug_col].str.contains(name_char, na=False)].copy()  
  
    # 创建用于存储结果的列  
    df_filtered['b'] = 0  
    df_filtered['c'] = 0  
    df_filtered['d'] = 0  
  
    # 创建一个数据框，包含每种药物名称下 'a' 值的总和，不考虑患者类型和 SOC 名称  
    drug_totals = df.groupby(drug_col)[a_col].sum().reset_index()  
  
    # 合并药物名称的 'a' 值总和到筛选后的数据框  
    df_filtered = df_filtered.merge(drug_totals, on=drug_col, how='left', suffixes=('', '_total'))  
  
    # 计算 'b' 值  
    df_filtered['b'] = df_filtered[f'{a_col}_total'] - df_filtered[a_col]  
  
    # 计算 'c' 值  
    # 首先创建一个辅助数据框，包含每个患者类型和 SOC 名称下不同药物的 'a' 值总和  
    group_totals = df.groupby([pt_col, soc_col])[a_col].sum().reset_index()  
  
    # 然后将这个辅助数据框与筛选后的数据框合并  
    df_filtered = df_filtered.merge(group_totals, on=[pt_col, soc_col], how='left', suffixes=('', '_group_total'))  
  
    df_filtered['c'] = df_filtered[f'{a_col}_group_total'] - df_filtered[a_col]  
  
    # 确保 'c' 值不为负（根据业务逻辑调整）  
    df_filtered['c'] = df_filtered['c'].clip(lower=0)  
  
    # 计算 'd' 值  
    total_a = df[a_col].sum()  
    df_filtered['d'] = total_a - (df_filtered[a_col] + df_filtered['b'] + df_filtered['c'])  
  
    return df_filtered     

#ROR计算-DB    
def TOOLS_ROR_from_DB(conn, name_field, pt_field, soc_field, pid, name_char, result_mode, display_mode=None, count_mode=None,additional_where_clause=None,abcd_value=None):  
    """    
    从SQLite数据库中高效获取数据并计算相关统计信息。    
    
    参数:    
    conn (sqlite3.Connection): 数据库连接对象。    
    name_field (str): 名称字段名。    
    pt_field (str): pt字段名。    
    soc_field (str): soc字段名。    
    pid (str): 患者ID字段名。  
    name_char (str): 要筛选的名称字符。    
    additional_where_clause (str, optional): 额外的WHERE子句条件。默认为None。  
    
    返回:    
    DataFrame: 包含name, pt, soc, a, b, c, d列的DataFrame。    
    """    
    # 构建基本的WHERE子句  
    where_clause = f"{name_field} LIKE '%{name_char}%'"  
    
  
    # 如果提供了额外的WHERE子句，则追加它  
    if additional_where_clause and additional_where_clause.strip(): 
        where_clause += f" AND ({additional_where_clause})"  
    print(where_clause)
    # 第一步：获取包含name_char的所有行，作为目标数据，用于查看原始数据   
    print("正在读取目标对象的数据作为溯源。")
    time1=time.time()
    query_target = f"""    
    SELECT    
        *    
    FROM    
        table1    
    WHERE    
        {where_clause}    
    """  

    dfs = pd.read_sql_query(query_target, conn)  
    
  
    # 第二步：计算所有药品的a（这部分代码没有变化）
    print(time.time() -time1) 
    print("正在读取数据库计算所有品种的a值...。")
    if abcd_value==None or  abcd_value=="":
        if additional_where_clause and additional_where_clause.strip():
            print("限定条件："+str(additional_where_clause)) 
            query_a = f"""    
            SELECT     
                {name_field},     
                {soc_field},     
                {pt_field},     
                COUNT(DISTINCT {pid}) AS a    
            FROM     
                table1  
            WHERE    
                {additional_where_clause}     
            GROUP BY     
                {name_field},     
                {soc_field},     
                {pt_field}    
            """  
        else:
            print("限定条件：无。") 
            query_a = f"""    
            SELECT     
                {name_field},     
                {soc_field},     
                {pt_field},     
                COUNT(DISTINCT {pid}) AS a    
            FROM     
                table1    
            GROUP BY     
                {name_field},     
                {soc_field},     
                {pt_field}    
            """       
        list_all = pd.read_sql_query(query_a, conn)    
        print(time.time() -time1) 
        # 假设这两个函数是在其他地方定义的 
        print("正在计算abcd值...")  
        result = TOOLS_ROR_from_DB_get_abcd(list_all, name_char, name_field, pt_field, soc_field)  
        print(time.time() -time1)     
    else:
        pass
    result = result.fillna(0) 
    print("正在计算ROR值...")  
    ROR_result = TOOLS_ROR_STAT_0(result)  
    print(time.time() -time1) 
    if result_mode=="a>=3&ROR_CI_95_low>1":
       ROR_result=ROR_result[(ROR_result["a"]>=3)&(ROR_result["ROR_CI_95_low"]>1)] 
       ROR_result=ROR_result.sort_values(by="ROR", ascending=False).reset_index(drop=True)
    ROR_result["报表类型"]="{'grouped':"+str([name_field, pt_field, soc_field])+str("}")
    if display_mode=="对比表":
       ROR_result=TOOLS_create_pivot_tool(ROR_result,[[soc_field, pt_field], [name_field], ['a', 'b', 'c', 'd', 'ROR', 'ROR_CI_95_low', 'PRR', 'PRR_CI_95_low'], ['sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum'], '', ''])
       del ROR_result["PRR合计"]       
       del ROR_result["ROR合计"]
    return ROR_result,dfs

#ROR计算-DB      
def TOOLS_ROR_from_DB_GUI(df,table_name='table1'):  
    """  
    弹出一个对话框让用户选择数据库文件，并返回选择的文件路径以及过滤参数。  
    """  
    print("************************************************************************************")  
    print("功能说明：")
    print("1. 该工具用于从SQLite数据库中读取数据，并计算ROR（报告比值比）和PRR（比例报告比）。")
    print("2. 用户需要选择一个SQLite数据库文件，并设置目标字段、事件字段、SOC字段和患者ID字段。")
    print("3. 支持模糊查询条件筛选目标数据。")
    print("4. 计算结果包括ROR、PRR及其95%置信区间。")
    print("5. 用户可以选择信号标准和结果展示模式（对比表或详细表）。")
    print("\n使用方法：")
    print("1. 点击“选择数据库文件”按钮，选择一个SQLite数据库文件。")
    print("2. 在下拉框中选择目标字段、事件字段、SOC字段和患者ID字段。")
    print("3. 输入模糊查询条件（可选）。")
    print("4. 点击“执行”按钮，查看计算结果。")
    print("************************************************************************************")  
    # 选择数据库文件  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.withdraw()  
    file_path = filedialog.askopenfilename(title="Select Database File", filetypes=[("SQLite Database", "*.db;*.sqlite;*.sqlite3")])  
    if not file_path:  
        return None, None, None, None  
    root.destroy()  
  
    # 连接到SQLite数据库  
    conn = sqlite3.connect(file_path)  
  
    # 获取字段名列表  
    cursor = conn.cursor()  
    cursor.execute(f"PRAGMA table_info({table_name})")  
    columns = [row[1] for row in cursor.fetchall()]  
  
    # 创建新窗口以选择字段和输入字符  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("ROR和PRR计算器（from数据库）")  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 450  # 窗口宽度  
    wh = 500  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
  
    var_name = tk.StringVar(root)  
    var_pt = tk.StringVar(root)  
    var_soc = tk.StringVar(root)  
    var_char = tk.StringVar(root)  
    var_pid = tk.StringVar(root)  
    def on_ok():  
        name_field = var_name.get()  
        pt_field = var_pt.get()  
        soc_field = var_soc.get()  
        name_char = var_char.get() 
        pid=var_pid.get() 
        
        result_mode=xcode_var.get() 
        display_mode=dcode_var.get()
        count_mode=  Mcode_var.get() 
        abcd_value=abcd_var.get()
        additional_where_clause= text_box.get("1.0", tk.END)  
        # 这里调用你的函数，我假设它存在并可以接受这些参数  
        result,dfs = TOOLS_ROR_from_DB(conn, name_field, pt_field, soc_field, pid,name_char,result_mode,display_mode,count_mode,additional_where_clause,abcd_value) 
        
        PROGRAM_display_df_in_treeview(result,1,dfs)
    def add_ok(event=None):  
        condition = SQL_create_query_gui(df,text_box,conn) 

        
    # 创建并布局标签和下拉框  
    def create_labeled_dropdown(root, label_text, var, values, row):  
        label = ttk.Label(root, text=label_text)  
        label.grid(row=row, column=0, sticky=tk.W, padx=5, pady=5)  
          
        dropdown = ttk.Combobox(root, textvariable=var, values=values)  
        dropdown.grid(row=row, column=1, sticky=tk.W+tk.E, padx=5, pady=5)  
          
        return dropdown  
      
    # 参数设置标签  
    tk.Label(root, text="参数设置").grid(row=0, column=0, columnspan=2, pady=10)  
      
    # 列设置  
    var_name = tk.StringVar(root)  
    name_field_dropdown = create_labeled_dropdown(root, "计算对象所在的列（比如产品名称列）", var_name, columns, 1)  
    name_field_dropdown.set("NAME")  # 设置默认选项  
      
    var_pt = tk.StringVar(root)  
    pt_field_dropdown = create_labeled_dropdown(root, "事件列（如PT）:", var_pt, columns, 2)  
    pt_field_dropdown.set("PT")  # 设置默认选项  
      
    var_soc = tk.StringVar(root)  
    soc_field_dropdown = create_labeled_dropdown(root, "额外事件列（如SOC）:", var_soc, columns, 3)  
    soc_field_dropdown.set("SOC")  # 设置默认选项  
      
    var_pid = tk.StringVar(root)  
    pid_field_dropdown = create_labeled_dropdown(root, "计数列（如报告编码PID）:", var_pid, columns, 4)  
    pid_field_dropdown.set("PID")  # 设置默认选项  
      
    # 计算模式  
    Mcode_var = tk.StringVar(root)  
    Mcode_var.set("count")  
    Mcode_dropdown = create_labeled_dropdown(root, "计算模式:", Mcode_var, ["count"], 5)  
      
    # 信号标准  
    xcode_var = tk.StringVar(root)  
    xcode_var.set("全部")  
    xcode_dropdown = create_labeled_dropdown(root, "信号标准:", xcode_var, ["a>=3&ROR_CI_95_low>1", "全部"], 6)  
      
    # 结果展示  
    dcode_var = tk.StringVar(root)  
    dcode_var.set("详细表")  
    dcode_dropdown = create_labeled_dropdown(root, "结果展示:", dcode_var, ["对比表", "详细表"], 7)  
      
    # 模糊查询  
    tk.Label(root, text="对象范围(模糊查询):", pady=10).grid(row=8, column=0, columnspan=2, sticky=tk.W)  
    var_char = tk.StringVar(root)  
    name_char_entry = tk.Entry(root, textvariable=var_char)  
    name_char_entry.grid(row=8, column=1, columnspan=2, padx=5, pady=5)  
    # where  
    # 全局范围标签  
    #tk.Label(root, text="a+b+c+d:", pady=10).grid(row=9, column=0, sticky=tk.W)  
    abcd_var = tk.StringVar(root)  
    abcd_char_entry = tk.Entry(root, textvariable=abcd_var)  
    #abcd_char_entry.grid(row=9, column=1, columnspan=2, padx=5, pady=5)  
            
    # 创建一个带滚动条的文本框  
    where_char = tk.StringVar(root)  
    text_frame = tk.Frame(root)  # 创建一个框架来容纳文本框和滚动条  
    text_frame.grid(row=10, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W+tk.E+tk.N+tk.S)  
      
    text_box = tk.Text(text_frame, height=5, width=20, wrap=tk.WORD)  # 创建一个文本框，默认3行  
    text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)  
    text_box.insert(tk.END,"(role_cod LIKE '%PS%')")    
    # 创建一个垂直滚动条，并与文本框关联  
    scroll_bar = ttk.Scrollbar(text_frame, command=text_box.yview)  
    scroll_bar.pack(side=tk.RIGHT, fill=tk.Y)  
      
    # 将文本框的滚动与滚动条关联  
    text_box['yscrollcommand'] = scroll_bar.set  
    # 执行按钮  
    ok_button = ttk.Button(root, text="执行", command=on_ok)  
    ok_button.grid(row=11, column=1, columnspan=2, pady=10)  
    add_button = ttk.Button(root, text="生成前置", command=add_ok)  
    add_button.grid(row=11, column=0, columnspan=2, pady=10)        
    root.mainloop()


#ROR计算-DF
def TOOLS_ROR_STAT_0(df):  
    """  
    根据DataFrame的四个列a, b, c, d，计算TOOLS_ROR_STAT_0及置信区间。  
    直接在原始的df上添加ROR, ROR_CI_95_low, ROR_CI_95_high, PRR, PRR_CI_95_low, PRR_CI_95_high, X2, chi2, p, dof列。  
    """  
      
    def calculate_chi2(row):  
        contingency_table = [[row['a'], row['b']], [row['c'], row['d']]]  
        chi2, p, dof, expected = stats.chi2_contingency(contingency_table)  
        return chi2, p, dof  
      
    # 验证输入是否是DataFrame且包含a, b, c, d四列  
    if not isinstance(df, pd.DataFrame) or not all(col in df.columns for col in ['a', 'b', 'c', 'd']):  
        raise ValueError("输入必须是一个包含a, b, c, d四列的DataFrame")  
      
    # 计算PRR和PRR的标准误  
    df['PRR'] = (df['a'] / (df['a'] + df['b'])) / (df['c'] / (df['c'] + df['d']))  
    df['PRR_SE'] = np.sqrt(1 / df['a'] - 1 / (df['a'] + df['b']) + 1 / df['c'] - 1 / (df['c'] + df['d']))  
    df['PRR_CI_95_low'] = df['PRR'] * np.exp(-1.96 * df['PRR_SE'])  
    df['PRR_CI_95_high'] = df['PRR'] * np.exp(1.96 * df['PRR_SE'])  
      
    # 计算ROR和ROR的标准误  
    df['ROR'] = (df['a'] / df['c']) / (df['b'] / df['d'])  
    df['ROR_SE'] = np.sqrt(1 / df['a'] + 1 / df['b'] + 1 / df['c'] + 1 / df['d'])  
    df['ROR_CI_95_low'] = df['ROR'] * np.exp(-1.96 * df['ROR_SE'])  
    df['ROR_CI_95_high'] = df['ROR'] * np.exp(1.96 * df['ROR_SE'])  
    
    # 计算卡方值  
    #df['X2'] = ((df['a'] * df['d'] - df['b'] * df['c']) ** 2 * (df['a'] + df['b'] + df['c'] + df['d'])) / (  
    #    (df['a'] + df['b']) * (df['c'] + df['d']) * (df['a'] + df['c']) * (df['b'] + df['d'])  
    #)  
        
    # 使用calculate_chi2函数计算卡方值、p值和自由度，并添加到DataFrame中  
    #df[['chi2', 'chi2-p', 'chi2-dof']] = df.apply(calculate_chi2, axis=1, result_type='expand')  
      
    return df.round(2)


#ROR计算-DF
def TOOLS_ROR_from_df(dfs, target_column, event_column, extra_event_column, code_column,result_mode,display_mode,count_mode): 
      
    if extra_event_column=="":
        dfs["事件列分类"]="无分类"
        extra_event_column ="事件列分类" 
    df=dfs.copy()
    #df.drop_duplicates(subset=[target_column, event_column, extra_event_column, code_column], keep='first', inplace=True)
    # 生成List1  目标药品报告数量
    if count_mode=="count":
        print(count_mode)
        list1 = df.groupby(target_column)[code_column].count().reset_index()          
    else:
        list1 = df.groupby(target_column)[code_column].nunique().reset_index()  
    list1.columns = [target_column, 'a_plus_b']  
    list1['abcd'] = list1['a_plus_b'].sum()  # abcd在返回后其他地方使用  
      
    # 生成List2  所有药品目标adr报告数量
    if count_mode=="count":
        list2 = df.groupby([extra_event_column,event_column])[code_column].count().reset_index()          
    else:
        list2 = df.groupby([extra_event_column,event_column])[code_column].nunique().reset_index()  
    list2.columns = [extra_event_column,event_column,'a_plus_c']  
      
    # 生成List3  目标药品目标adr报告数量
    if count_mode=="count":
        list3 = df.groupby([target_column, extra_event_column, event_column])[code_column].count().reset_index()          
    else:
        list3 = df.groupby([target_column, extra_event_column, event_column])[code_column].nunique().reset_index()  
    list3.columns = [target_column, extra_event_column, event_column, 'a']  
      
    # 拼接List1和List3  
    list3 = pd.merge(list3, list1, on=target_column, how='left')  
      
    # 拼接List2和List3  
    list3 = pd.merge(list3, list2, on=[extra_event_column, event_column], how='left')  
      
    list3["b"] = list3['a_plus_b'] - list3["a"]  
    list3["c"] = list3['a_plus_c'] - list3["a"]      
    list3["d"] = list3['abcd'] - list3["a"] - list3["b"] - list3["c"]   
    del list3['a_plus_b']  
    del list3['a_plus_c']   
    del list3['abcd']  
    ROR_result= TOOLS_ROR_STAT_0(list3)
    if result_mode=="a>=3&ROR_CI_95_low>1":
       ROR_result=ROR_result[(ROR_result["a"]>=3)&(ROR_result["ROR_CI_95_low"]>1)] 
       ROR_result=ROR_result.sort_values(by="ROR", ascending=False).reset_index(drop=True)
    ROR_result["报表类型"]="{'grouped':"+str([target_column, event_column, extra_event_column])+str("}")
    if display_mode=="对比表":
       ROR_result=TOOLS_create_pivot_tool(ROR_result,[[extra_event_column, event_column], [target_column], ['a', 'b', 'c', 'd', 'ROR', 'ROR_CI_95_low', 'PRR', 'PRR_CI_95_low'], ['sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum'], '', ''])
       del ROR_result["PRR合计"]       
       del ROR_result["ROR合计"]
   
    PROGRAM_display_df_in_treeview(ROR_result, 1, dfs)  # 这个函数在原始代码中没有提供，因此注释掉  


#ROR计算-DF
def TOOLS_ROR_from_df_with_gui(df):  
    print("************************************************************************************")  
    print("功能说明：")
    print("1. 该工具用于从DataFrame中计算ROR（报告比值比）和PRR（比例报告比）。")
    print("2. 用户需要选择目标字段（如药品名称）、事件字段（如不良反应）、额外事件字段（如器官系统）和计数字段（如报告编码）。")
    print("3. 支持两种计算模式：基于唯一值的计数（nunique）或基于重复值的计数（count）。")
    print("4. 计算结果包括ROR、PRR及其95%置信区间。")
    print("5. 用户可以选择信号标准和结果展示模式（对比表或详细表）。")
    print("\n使用方法：")
    print("1. 在下拉框中选择目标字段、事件字段、额外事件字段和计数字段。")
    print("2. 选择计算模式（nunique或count）。")
    print("3. 选择信号标准和结果展示模式。")
    print("4. 点击“确定”按钮，查看计算结果。")
    print("************************************************************************************")  
    print("ROR计算器：您应该使用一份规整过的原始数据工作。a:目标对象目标事件报告数，b：目标对象不出现目标事件报告数，c:非目标对象目标事件报告数，d：非目标对象不出现目标事件报告数") 
    def on_submits():    
        target_col = target_var.get()    
        event_col = event_var.get()    
        extra_event_col = extra_event_var.get()  # 获取额外的事件列选择  
        code_col = code_var.get()    
        result_mode=xcode_var.get() 
        display_mode=dcode_var.get()
        count_mode=  Mcode_var.get()         
        # 在这里添加对默认选项的检查，确保用户已选择有效选项    
        if any([col == "请选择" for col in [target_var.get(), event_var.get(), extra_event_var.get(), code_var.get()]]):    
            print("请选择所有选项后再提交")    
            return    
    
        result = TOOLS_ROR_from_df(df, target_col, event_col, extra_event_col, code_col,result_mode,display_mode,count_mode)  # 调用处理函数，并传入额外的事件列参数  
            
        # 显示结果（这里只是简单地打印到控制台，您可以根据需要修改）    
        print(result)    
    
    root = tk.Tk()    
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("ROR和PRR计算器")    
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 400  # 窗口宽度          
    wh = 450  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    columns = list(df.columns)    
    target_options = ["请选择"] + columns  # 添加默认选项    
    event_options = ["请选择"] + columns  # 添加默认选项    
    extra_event_options = [""] + columns  # 添加默认选项，用于额外的事件列下拉列表  
    code_options = ["请选择"] + columns  # 添加默认选项    

    
    # 目标列下拉列表    
    target_label = tk.Label(root, text="目标对象列(如产品名称):")    
    target_label.pack()    
    target_var = tk.StringVar(root)    
    target_var.set("")   
    target_dropdown = ttk.Combobox(root, textvariable=target_var, values=target_options)    
    target_dropdown.pack()    
    
    # 事件列下拉列表    
    event_label = tk.Label(root, text="事件列(比如不良反应,建议先扩行):")    
    event_label.pack()    
    event_var = tk.StringVar(root)  # 初始化为默认选项    
    event_var.set("")   
    event_dropdown = ttk.Combobox(root, textvariable=event_var, values=event_options)    
    event_dropdown.pack()    
    
    # 额外的事件列下拉列表（新增）  
    extra_event_label = tk.Label(root, text="额外事件列(非必填，比如器官系统):")  # 新增标签和下拉列表变量等。此列与事件列在同一行。  
    extra_event_label.pack()  # 新增的列与事件列在同一行。
    extra_event_var = tk.StringVar(root)  # 初始化为默认选项    
    extra_event_var.set("")   
    extra_event_dropdown = ttk.Combobox(root, textvariable=extra_event_var, values=extra_event_options)    
    extra_event_dropdown.pack()    
    # 代码列下拉列表  
    code_label = tk.Label(root, text="计数列（比如报告编码）:")  
    code_label.pack()  
    code_var = tk.StringVar(root)  # 初始化为默认选项  
    code_var.set("报告编码")     
    code_dropdown = ttk.Combobox(root, textvariable=code_var, values=code_options)  
    code_dropdown.pack()  
    # 计算方法  
    Mcode_label = tk.Label(root, text="计算模式:")  
    Mcode_label.pack()  
    Mcode_var = tk.StringVar(root)  # 初始化为默认选项  
    Mcode_var.set("nunique")     
    Mcode_dropdown = ttk.Combobox(root, textvariable=Mcode_var, values=["nunique","count"])  
    Mcode_dropdown.pack()  
    # 信号标准 
    xcode_label = tk.Label(root, text="信号标准:")  
    xcode_label.pack()  
    xcode_var = tk.StringVar(root)  # 初始化为默认选项  
    xcode_var.set("a>=3&ROR_CI_95_low>1")     
    xcode_dropdown = ttk.Combobox(root, textvariable=xcode_var, values=["a>=3&ROR_CI_95_low>1","全部"])  
    xcode_dropdown.pack()  

    # 结果展示 
    dcode_label = tk.Label(root, text="结果展示:")  
    dcode_label.pack()  
    dcode_var = tk.StringVar(root)  # 初始化为默认选项  
    dcode_var.set("详细表")     
    dcode_dropdown = ttk.Combobox(root, textvariable=dcode_var, values=["对比表","详细表"])  
    dcode_dropdown.pack()  

    # 提交按钮  
    submit_button = ttk.Button(root, text="确定", command=on_submits)  
    submit_button.pack(pady=20)  

    separator_label = ttk.Label(root, text="采用报告数为基准。 如：(目标对象出现该事件的报告数量/其他对象出现该事件不良反应的报告数量)\n/ (目标对象不出现该事件的报告数量/其他对象不出现该事件的报告数量)")  
    separator_label.pack() 
  
    root.mainloop()


#卡方检验
def TOOLS_drug_reaction_CH2(df, a_drug_freq_col, b_drug_freq_col, a_drug_users_col, b_drug_users_col):  
    #a_drug_users_col=8400
    #b_drug_users_col=1760
    """  
    对DataFrame中的两种药物不良反应进行卡方检验和Fisher精确检验。  
      
    参数:  
        df (pd.DataFrame): 包含不良反应数据的DataFrame。  
        a_drug_freq_col (str): a药发生频次的列名。  
        b_drug_freq_col (str): b药发生频次的列名。  
        a_drug_users_col (str): a药使用人数的列名（或固定值，如果对所有反应都一样）。  
        b_drug_users_col (str): b药使用人数的列名（或固定值，如果对所有反应都一样）。  
            
    返回:  
        pd.DataFrame: 在原DataFrame基础上增加卡方检验和Fisher精确检验的结果列，以及期望计数相关信息。  
    """  
    # 检查传入列名是否存在  
    if not all(col in df.columns for col in [a_drug_freq_col, b_drug_freq_col]):  
        raise ValueError("One or more column names are not found in the input DataFrame.")  
        
    # 尝试获取固定的用户数量，如果失败则假定是列名  
    try:  
        a_drug_users = int(df[a_drug_users_col].iloc[0]) if a_drug_users_col in df.columns else int(a_drug_users_col)  
        b_drug_users = int(df[b_drug_users_col].iloc[0]) if b_drug_users_col in df.columns else int(b_drug_users_col)  
    except ValueError:  
        raise ValueError("a_drug_users_col and b_drug_users_col must be either column names or fixed integer values.")  
        
    # 创建一个新的DataFrame，防止修改原始数据  
    new_df = df.copy()  
    new_df = new_df.fillna(0)  
    new_df[a_drug_freq_col]= new_df[a_drug_freq_col].astype(int)
    new_df[b_drug_freq_col]= new_df[b_drug_freq_col].astype(int)    
      
    # 初始化结果列表  
    results = []  
      
    # 对每一行（每一种不良反应）进行卡方检验和Fisher精确检验  
    for index, row in new_df.iterrows():  
        # 提取当前行的不良反应频次和使用人数  
        a_freq = int(row[a_drug_freq_col])  
        b_freq = int(row[b_drug_freq_col])  
            
        # 假设未使用药物的人没有不良反应  
        a_no_reaction = a_drug_users - a_freq  
        b_no_reaction = b_drug_users - b_freq  
            
        # 构建2x2的频数表  
        contingency_table = [[a_freq, a_no_reaction],  
                             [b_freq, b_no_reaction]]  
            

          
        # 卡方检验  
        chi2, p, dof, exp = chi2_contingency(contingency_table,correction=False)  
        
        
        # 计算期望计数  
        expected_counts = exp  
        expected_counts = expected_counts.reshape(2, 2)  # Ensure it's a 2x2 array  
          
        # 计算小于5的期望计数数量和百分比  
        less_than_5 = (expected_counts < 5).sum()  
        percent_less_than_5 = (less_than_5 / expected_counts.size) * 100  
          
        # 找到最小期望计数  
        min_expected_count = expected_counts.min()  

        # 卡方检验  
        Fix_chi2, Fix_p, Fix_dof, Fix_exp = chi2_contingency(contingency_table)  
  
        # Fisher精确检验  
        oddsratio, p_value = fisher_exact(contingency_table, alternative='two-sided')  
            
        # 收集结果  
        result = {  
            'chi2': chi2,  
            'chi2_p': p,  
            'chi2_dof': dof,  
            'chi2_exp':exp,
            'Fix_chi2': Fix_chi2,  
            'Fix_p': Fix_p,  
            'Fix_dof': Fix_dof,  
            'Fix_exp':Fix_exp,            
            'fisher_oddsratio': oddsratio,  
            'fisher_p_value': p_value,  
            'less_than_5_count': less_than_5,  
            'percent_less_than_5': percent_less_than_5,  
            'min_expected_count': min_expected_count  
        }  
        results.append(result)  
        
    # 将结果添加到DataFrame中  
    new_df = pd.concat([new_df, pd.DataFrame(results, index=new_df.index)], axis=1)  
        
    return new_df

#卡方检验
def TOOLS_drug_reaction_CH2_create_gui(df):

    print(
        """
    **********************************************
    卡方检验和Fisher精确检验工具

    功能：
    1. 对两种药物的不良反应数据进行卡方检验和Fisher精确检验。
    2. 提供GUI界面供用户选择数据列和输入基数。

    使用方法：
    1. 启动程序后，选择事件列、对象A阳性频次列、对象B阳性频次列。
    2. 输入对象A和对象B的基数（固定值）。
    3. 点击“统计”按钮，程序会计算并显示检验结果。

    注意：
    - 若每个条件下的期望频数均大于5，且总样本量大于40，则可使用Pearson卡方检验。
    - 若存在某个条件的期望频数小于5但大于1，应使用Yates矫正卡方检验。
    - 若样本量小于40，或者期望频数存在小于1的情况，应使用Fisher精确检验。
    **********************************************
    """
    )

    # 检查DataFrame是否有足够的列
    if len(df.columns) < 2:
        raise ValueError("DataFrame必须至少包含2列。")

    # 创建主窗口
    root = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("卡方检验和Fisher精确检验")

    sw = root.winfo_screenwidth()
    sh = root.winfo_screenheight()
    ww = 300  # 窗口宽度
    wh = 220  # 窗口高度
    x = (sw - ww) // 2
    y = (sh - wh) // 2
    root.geometry(f"{ww}x{wh}+{x}+{y}")

    # 创建用于选择列名的下拉菜单
    reaction_label = ttk.Label(root, text="事件列:")
    reaction_label.grid(row=0, column=0, padx=5, pady=5)
    reaction_var = tk.StringVar(root)
    reaction_var.set(df.columns[0] if len(df.columns) > 0 else "")  # 设置默认选项
    reaction_dropdown = ttk.Combobox(root, textvariable=reaction_var, state='readonly')
    reaction_dropdown['values'] = df.columns.tolist()
    reaction_dropdown.grid(row=0, column=1, padx=5, pady=5)

    a_drug_label = ttk.Label(root, text="对象A阳性频次:")
    a_drug_label.grid(row=1, column=0, padx=5, pady=5)
    a_drug_var = tk.StringVar(root)
    a_drug_var.set(df.columns[1] if len(df.columns) > 1 else "")  # 设置默认选项
    a_drug_dropdown = ttk.Combobox(root, textvariable=a_drug_var, state='readonly')
    a_drug_dropdown['values'] = df.columns.tolist()
    a_drug_dropdown.grid(row=1, column=1, padx=5, pady=5)

    b_drug_label = ttk.Label(root, text="对象B阳性频次:")
    b_drug_label.grid(row=2, column=0, padx=5, pady=5)
    b_drug_var = tk.StringVar(root)
    b_drug_var.set(df.columns[0] if len(df.columns) > 0 else "")  # 设置默认选项
    b_drug_dropdown = ttk.Combobox(root, textvariable=b_drug_var, state='readonly')
    b_drug_dropdown['values'] = df.columns.tolist()
    b_drug_dropdown.grid(row=2, column=1, padx=5, pady=5)

    # 用户数量是固定的，所以我们使用Entry来输入
    a_users_label = ttk.Label(root, text="对象A基数:")
    a_users_label.grid(row=3, column=0, padx=5, pady=5)
    a_users_entry = ttk.Entry(root)
    a_users_entry.grid(row=3, column=1, padx=5, pady=5)

    b_users_label = ttk.Label(root, text="对象B基数:")
    b_users_label.grid(row=4, column=0, padx=5, pady=5)
    b_users_entry = ttk.Entry(root)
    b_users_entry.grid(row=4, column=1, padx=5, pady=5)

    # 创建一个按钮来执行统计函数
    calculate_button = ttk.Button(
        root,
        text="统计",
        command=lambda: PROGRAM_display_df_in_treeview(
            TOOLS_drug_reaction_CH2(
                df,
                a_drug_var.get(),
                b_drug_var.get(),
                a_users_entry.get(),
                b_users_entry.get()
            ),
            0,
            0
        )
    )
    calculate_button.grid(row=5, column=0, columnspan=2, padx=5, pady=10)

    # 运行主循环
    root.mainloop()


#批量透视  
def TOOLS_stat_all_gui(dfs):  
    print(
        """
    **********************************************
    批量统计和透视工具

    功能：
    1. 对给定的DataFrame进行批量统计和透视操作。
    2. 支持选择计数方法（如nunique、sum、count）和计数列。
    3. 支持选择透视列（可选）。
    4. 支持多选列进行批量操作，并将结果拼接在一起。

    使用方法：
    1. 启动程序后，选择计数方法和计数列。
    2. 在Treeview中选择需要统计的列（可多选）。
    3. 如果需要透视，选择透视列；否则选择“无需透视”。
    4. 点击“执行”按钮，程序会计算并显示结果。
    **********************************************
    """
    )
    df=dfs.fillna(0)
     
    # 创建主窗口  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("批量透视")  
      
    # 定义窗口的大小和位置  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 400  # 窗口宽度  
    wh = 650  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
      
    # 定义下拉框的选项  
    aggregation_methods = ['nunique', 'sum', 'count']  
      
    # 创建标签和下拉框选择计数方法和a列  
    ttk.Label(root, text="请选择方法:").pack()  
    aggregation_var = tk.StringVar(root)  
    aggregation_var.set(aggregation_methods[0])  
    aggregation_dropdown = ttk.Combobox(root, textvariable=aggregation_var, values=aggregation_methods)  
    aggregation_dropdown.pack()  
      
    ttk.Label(root, text="请选择计数列（如：报告编码）:").pack()  
    a_column_var = tk.StringVar(root)  
    a_column_var.set(df.columns.tolist()[0] if not df.columns.empty else '')  
    a_column_dropdown = ttk.Combobox(root, textvariable=a_column_var)  
    a_column_dropdown['values'] = df.columns.tolist()  
    a_column_dropdown.pack()  

    ttk.Label(root, text="透视列:").pack()  
    pvot_column_var = tk.StringVar(root)  
    pvot_column_var.set('无需透视')  
    pvot_column_dropdown = ttk.Combobox(root, textvariable=pvot_column_var)  
    pvot_column_dropdown['values'] = df.columns.tolist()  
    pvot_column_dropdown.pack()  
   
    # 创建Treeview来显示列名及滚动条  
    tree_frame = ttk.Frame(root)  
    tree_frame.pack(side='top', fill='both', expand=True)  
      
    # Treeview  
    tree = ttk.Treeview(tree_frame, columns=("column_name",), show='headings')  
    vertical_scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=tree.yview)  
    tree.configure(yscrollcommand=vertical_scrollbar.set)  
  
    tree.heading("#0", text="列名")  
    tree.column("#0", width=100, anchor='w')  
    tree.column("column_name", width=100, anchor='w')  
      
    # 不再需要这一行，因为我们已经在ttk.Treeview中定义了列  
    # tree["columns"] = ("column_name",)  
      
    # Pack Treeview and Scrollbar  
    tree.pack(side='left', fill='both', expand=True)  
    vertical_scrollbar.pack(side='right', fill='y')  
      
    for column in df.columns:  
        tree.insert("", 'end', text=column, values=(column,))  
  
    def calculate():  
        selected_items = tree.selection()  
        if not selected_items:  
            return  
  
        selected_columns = [tree.set(item, "column_name") for item in selected_items]  
        aggregation_method = aggregation_var.get()  
        a_column = a_column_var.get()  
        
        #add
        pvot_column=pvot_column_var.get()
        ratio_col=str(a_column)+'合计'
        
  
        if not all(col in df.columns for col in selected_columns) or a_column not in df.columns or aggregation_method not in aggregation_methods:  
            return  
  
        result_dfs = []  
        for selected_column in selected_columns: 
            df[selected_column]=df[selected_column].astype(str)
                       
            if pvot_column=="无需透视":
                result_temp=SMALL_add_count_and_ratio(df,selected_column,a_column_var.get(),aggregation_var.get())
            else:
                df[pvot_column]=df[pvot_column].astype(str) 
                result_temp=TOOLS_create_pivot_tool(df,[[selected_column], [pvot_column], [a_column_var.get()], [aggregation_var.get()], '', [ratio_col]])
            
            
            result_temp.rename(columns={selected_column: "项目"}, inplace=True)
            result_temp["列名"]=selected_column
            result_dfs.append(result_temp.copy())
        # 合并所有结果 DataFrame  
        if result_dfs:  
 
            final_result_df = pd.concat(result_dfs, ignore_index=True).reset_index(drop=True)

            PROGRAM_display_df_in_treeview(final_result_df,0,0)  
  
    # 创建按钮来执行计算  
    calculate_button = ttk.Button(root, text="执行", command=calculate)  
    calculate_button.pack()  
  
    root.mainloop()                                             

    
#数据分组和透视附属函数#######################    
def TOOLS_create_pivot_create_multiselect_pivot_gui(df, result_text):  
    """  
    数据分组和透视函数。
    """  
    def on_confirm():  
        """  
        确认按钮的回调函数，用于获取用户的选择并返回  
        """  
        # 获取用户选择的列名  
        selected_ids = tree.selection()  
        selected_col_names = [tree.item(item)["text"] for item in selected_ids]  
        selected_method = method_var.get()  
  
        # 创建一个字典，其中key是列名，value是选定的方法  
        result_dict = {col_name: selected_method for col_name in selected_col_names}  
  
        # 将字典转换为字符串并插入到result_text中  
        result_str = str(result_dict)  
        result_text.insert(tk.END, result_str)  
  
        # 关闭当前窗口  
        top.destroy()  
  
    # 创建新的顶层窗口  
    top = tk.Toplevel()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(top)
    top.title("多选值透视工具")  
    # 设置窗口位置和大小（这里可以自定义）    
    sw = top.winfo_screenwidth()    
    sh = top.winfo_screenheight()    
    ww = 320  # 窗口宽度    
    wh = 300  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    top.geometry(f"{ww}x{wh}+{x}+{y}")   
    # 在窗口中创建Treeview，显示df的所有列  
    tree = ttk.Treeview(top)  
    ysb = ttk.Scrollbar(top, orient='vertical', command=tree.yview)  # 创建垂直滚动条  
    xsb = ttk.Scrollbar(top, orient='horizontal', command=tree.xview)  # 创建水平滚动条  
    tree.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)  # 将滚动条的视图与Treeview绑定  
  
    tree.column("#0", width=200, minwidth=200, stretch=tk.NO)  
    tree.heading("#0", text="列名", anchor=tk.W)  
    for i, col in enumerate(df.columns):  
        tree.insert("", i, text=col)  # 删除values参数以修复Treeview显示问题  
  
    # 安排Treeview和滚动条的位置  
    tree.grid(row=0, column=0, sticky='nsew')  
    ysb.grid(row=0, column=1, sticky='ns')  
    xsb.grid(row=1, column=0, sticky='ew')  
  
    top.grid_rowconfigure(0, weight=1)  
    top.grid_columnconfigure(0, weight=1)  
  
    # 创建方法选择的下拉菜单  
    methods = ["sum", "count", "nunique", "unique", "mean", "median", "std", "size", "cumsum", "SMALL_count_mode", "count_CH_semicolon", "count_EN_semicolon", "count_CH_comma", "count_EN_comma", "count_CH_commas", "count_ALL"]  
    method_var = tk.StringVar(top)  
    method_var.set(methods[0])  # 设置默认方法  
    method_dropdown = ttk.Combobox(top, textvariable=method_var, values=methods)  
    method_dropdown.grid(row=2, column=0 ) 
  
    # 创建确认按钮，并绑定回调函数on_confirm  
    confirm_button = ttk.Button(top, text="确认", command=on_confirm)  
    confirm_button.grid(row=3, column=0)
  
    # 运行新窗口的主循环，等待用户操作  
    top.wait_window()


#数据分组和透视附属函数      
def TOOLS_create_pivot_tool(df,methon): 
    """    
    数据透视工具的的功能函数   
    """
    row_labels=methon[0]
    col_labels=methon[1] 
    value_cols=methon[2] 
    agg_methods=methon[3] 
    text_content=methon[4] 
    all_ratio=methon[5] 


    
    try:
        for i, value in enumerate(agg_methods):  
            if value == "SMALL_count_mode":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "")  
            elif value == "count_CH_semicolon":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "；")  
            elif value == "count_EN_semicolon":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, ";")  
            elif value == "count_CH_comma":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "，")  
            elif value == "count_EN_comma":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, ",")  
            elif value == "count_CH_commas":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "、")  
            elif value == "count_ALL":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "。|；|，|、|,|;|/")
    except:
        print("SMALL_count_mode error")
        
    try:
        for i, value in text_content.items():  
            if value == "SMALL_count_mode":  
                text_content[i] = lambda x: SMALL_count_mode(x, "")  
            elif value == "count_CH_semicolon":  
                text_content[i] = lambda x: SMALL_count_mode(x, "；")  
            elif value == "count_EN_semicolon":  
                text_content[i] = lambda x: SMALL_count_mode(x, ";")  
            elif value == "count_CH_comma":  
                text_content[i] = lambda x: SMALL_count_mode(x, "，")  
            elif value == "count_EN_comma":  
                text_content[i] = lambda x: SMALL_count_mode(x, ",")  
            elif value == "count_CH_commas":  
                text_content[i] = lambda x: SMALL_count_mode(x, "、")  
            elif value == "count_ALL":  
                text_content[i] = lambda x: SMALL_count_mode(x, "。|；|，|、|,|;|/")
    except:
        pass 
 
    if len(col_labels)!=0:
        pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)},margins=True, margins_name='合计') .reset_index() #,dropna=False, fill_value=0
    else:
        pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)}) .reset_index() #,dropna=False, fill_value=0

    #解除多级索引  
    pivot_table.columns = [''.join(col).strip() for col in pivot_table.columns.values]                  
    if len(text_content)>=1:
        try:
            if len(col_labels)!=0:
                pivot_table1 = pd.pivot_table(df, values=text_content.keys(), index=row_labels, columns=[], aggfunc=text_content,margins=True, margins_name='合计') .reset_index() #,dropna=False, fill_value=0
                pivot_table1.columns = [''.join(col).strip() for col in pivot_table1.columns.values] 
                pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)},margins=True, margins_name='合计') .reset_index() #,dropna=False, fill_value=0
            else:
                pivot_table1 = pd.pivot_table(df, values=text_content.keys(), index=row_labels, columns=[], aggfunc=text_content) .reset_index() #,dropna=False, fill_value=0
                pivot_table1.columns = [''.join(col).strip() for col in pivot_table1.columns.values] 
                pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)}) .reset_index() #,dropna=False, fill_value=0
            pivot_table.columns = [''.join(col).strip() for col in pivot_table.columns.values] 
            #del pivot_table1["合计"]
            pivot_table=pd.merge(pivot_table,pivot_table1, on=row_labels, how='left')
   

        except:
            print("多选值列配置错误。") 
    #增加构成比
    if len(all_ratio)!=0: 
        pivot_table=SMALL_add_composition_ratio(pivot_table,all_ratio[0])  #SMALL_calculate_ratios
    #增加严重比、超时比等。
    if len(all_ratio)==3: 
        pivot_table=SMALL_calculate_ratios(pivot_table,all_ratio[1],all_ratio[2])                           
    try:
        pivot_table=SMALL_expand_dict_like_columns(pivot_table)
    except:
        pass
    pivot_table["报表类型"]="{'grouped':"+str(row_labels)+str("}")  

    return pivot_table

#数据分组和透视         
def TOOLS_create_pivot_tool_gui(df, ori):    
    """    
    创建数据透视工具的图形用户界面(GUI)    
    
    参数:    
        df (DataFrame): 输入的数据框    
        ori (object): 原始对象，用于后续处理结果    
    """    
    print(
        """
    **********************************************
    数据分组和透视工具

    功能：
    1. 提供 GUI 界面供用户选择行标签、列标签、值列和聚合方法。
    2. 支持多选值透视功能。
    3. 生成数据透视表并显示结果。

    使用方法：
    1. 启动程序后，选择行标签、列标签、值列和聚合方法。
    2. 点击“提交”按钮生成数据透视表。
    3. 点击“拼接数据分组”按钮进行多选值透视。
    4. 结果会显示在 GUI 界面中。

    注意：
    - 如果未选择任何列，程序会提示警告。
    - 数据透视表会显示在 GUI 界面中。
    **********************************************
    """
    )
    def on_submit():    
        """    
        提交按钮的回调函数，用于处理用户的选择并执行数据透视操作    
        """    
        # 获取用户选择的行标签、列标签、值列和聚合方法    
        row_labels = [var.get() for var in row_vars if var.get() != ""]    
        col_labels = [var.get() for var in col_vars if var.get() != ""]    
        value_cols = [var.get() for var in value_vars if var.get() != ""]    
        agg_methods = [var.get() for var in agg_method_vars if var.get() != ""]        
        

   
        # 创建数据透视表，使用字典指定每个值列的聚合方法    
        
        text_content=""
        try:
           text_content = eval(result_text.get("1.0", tk.END)) 
        except:
            pass
        methon=[row_labels,col_labels,value_cols,agg_methods,text_content,""]
        print(methon)
        pivot_table=TOOLS_create_pivot_tool(df,methon)
        PROGRAM_display_df_in_treeview(pivot_table, 0, df)    

  
    df=df.fillna(0)
    # 创建主窗口    
    root = tk.Tk()   
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("数据分组和透视工具")  # 设置窗口标题为“数据透视工具”    
    
    # 设置窗口位置和大小（这里可以自定义）    
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 820  # 窗口宽度    
    wh = 300  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    
    # 获取数据框的列名列表    
    columnslist = df.columns.tolist()    
    
    # 创建行标签的下拉菜单    
    row_vars = []    
    row_labels = ["行标签1", "行标签2", "行标签3", "行标签4", "行标签5"]    
    for i, label in enumerate(row_labels):    
        tk.Label(root, text=label).grid(row=0, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=1, column=i)    
        row_vars.append(var)    
    
    # 创建列标签的下拉菜单    
    col_vars = []    
    col_labels = ["透视列1", "透视列2", "透视列3", "透视列4", "透视列5"]    
    for i, label in enumerate(col_labels):    
        tk.Label(root, text=label).grid(row=2, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=3, column=i)    
        col_vars.append(var)    
    
    # 创建值列的下拉菜单和对应的聚合方法的下拉菜单    
    value_vars = []    
    agg_method_vars = []    
    value_labels = ["值列1", "值列2", "值列3", "值列4", "值列5"]    
    for i, label in enumerate(value_labels):    
        tk.Label(root, text=label).grid(row=4, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=5, column=i)    
        value_vars.append(var)    
        # 创建聚合方法的下拉菜单    
        methods = ["sum", "count", "nunique", "unique", "mean", "median", "std", "size", "cumsum", "SMALL_count_mode", "count_CH_semicolon", "count_EN_semicolon", "count_CH_comma", "count_EN_comma", "count_CH_commas", "count_ALL"]  # 可根据需要添加其他聚合方法    

        agg_var = tk.StringVar(root)    
        agg_var.set("")  # 设置默认值为第一个聚合方法    
        agg_dropdown = ttk.Combobox(root, textvariable=agg_var, values=methods)    
        agg_dropdown.grid(row=6, column=i)  # 将下拉菜单放置到网格中，与对应的值列下拉菜单对齐显示    
        agg_method_vars.append(agg_var)    
    
    # 创建提交按钮，并绑定回调函数on_submit，用于处理用户的选择和执行数据透视操作    
    submit_button = ttk.Button(root, text="提交", command=on_submit)  # 创建提交按钮并设置文本和回调函数    
    submit_button.grid(row=7, column=0, columnspan=3)  # 将提交按钮放置到网格中，并跨3列显示    
    multiselect_button = ttk.Button(root, text="拼接数据分组", command=lambda: TOOLS_create_pivot_create_multiselect_pivot_gui(df,result_text))  # 创建按钮并绑定功能函数create_multiselect_pivot_gui()      
    multiselect_button.grid(row=7, column=3, columnspan=3)  # 将按钮放置到网格中，并跨3列显示      
    # 运行主循环，显示窗口并等待用户操作  
    # 创建文本框，用于显示多选值透视的结果  
    result_text = tk.Text(root, height=4, width=100, wrap=tk.WORD)  
    # 创建垂直滚动条  
    scrollbar = tk.Scrollbar(root, command=result_text.yview)     
    # 将滚动条与文本框关联  
    result_text['yscrollcommand'] = scrollbar.set  
    # 将文本框放置到网格中，并跨5列显示  
    result_text.grid(row=8, column=0, columnspan=5, sticky='nsew')  
    # 将滚动条放置在文本框右侧，与文本框同高  
    scrollbar.grid(row=8, column=5, rowspan=4, sticky='ns')  
    root.mainloop() #启动Tkinter事件循环，显示窗口并等待用户操作或关闭窗口时结束程序
    

#趋势分析函数，TOOLS_trend_analysis_GUI，,TOOLS_trend_analysis_with_3_sd,TOOLS_trend_analysis_with_3_sd#######################    
def TOOLS_trend_analysis_GUI(df):  
    
    print(
        """
    **********************************************
    趋势分析工具

    功能：
    1. 提供 GUI 界面供用户选择日期列、事件列、目标筛选列、分析对象、分析窗口、方法和频率。
    2. 支持基于标准差或 IQR 的控制限进行趋势分析。
    3. 支持生成趋势图或趋势表。

    使用方法：
    1. 启动程序后，选择日期列、事件列、目标筛选列、分析对象、分析窗口、方法和频率。
    2. 点击“趋势图”按钮生成趋势图，或点击“趋势表”按钮生成趋势表。
    3. 趋势图和趋势表会显示在 GUI 界面中。

    注意：
    - 如果未选择任何列，程序会提示警告。
    - 趋势图和趋势表会显示在 GUI 界面中。
    **********************************************
    """
    )
    
    # 创建Tkinter窗口  
    # 更新分析对象下拉菜单的去重清单  
    def update_unique_values(event):  
        selected_column = target_column_var.get()  
        if selected_column == "不筛选":  
            analysis_object_var.set("不筛选")  
            analysis_object_dropdown['values'] = []  
            return  
        df[selected_column]=df[selected_column].fillna('-未填写-')  
        unique_values = df[selected_column].unique()  
        sorted_values = sorted(unique_values, key=lambda x: df[selected_column].value_counts()[x], reverse=True)  
          
        # 格式化所有的唯一值，包括值、计数和百分比  
        formatted_values = [str([value, df[selected_column].value_counts()[value], f"{df[selected_column].value_counts(normalize=True)[value]*100:.2f}%"]) for value in sorted_values]


          
        # 更新分析对象变量和下拉菜单的值  
        analysis_object_var.set(formatted_values[0] if formatted_values else [])  
        analysis_object_dropdown['values'] = formatted_values  

     
    
    
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("趋势分析预处理")  
      
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 450  # 窗口宽度  
    wh = 250  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
      
    # 获取df的列名以供选择  
    columns = df.columns.tolist()  
      
    # 设置默认参数  
    default_date_column = "事件发生日期" if "事件发生日期" in columns else '报告日期'  
    default_event_column = "报告编码" if "报告编码" in columns else columns[1]  
    default_windows = 12  
    default_method = "nunique"  
    default_freq = "M"  
      
    # 创建标签和输入框  
    mainframe = ttk.Frame(root, padding="10")  
    mainframe.grid(column=0, row=0, sticky=(tk.N, tk.W, tk.E, tk.S))  
      
    date_label = ttk.Label(mainframe, text="日期列:")  
    date_label.grid(column=1, row=1, sticky=tk.W)  
    date_column_var = tk.StringVar(root)  
    date_column_var.set(default_date_column)  # 初始设置  
    date_column_dropdown = ttk.Combobox(mainframe, textvariable=date_column_var, values=columns, width=40)  # 设置宽度为30个字符  
    date_column_dropdown.grid(column=2, row=1, sticky=(tk.W, tk.E))  
      
    event_label = ttk.Label(mainframe, text="报告编号列:")  
    event_label.grid(column=1, row=2, sticky=tk.W)  
    event_column_var = tk.StringVar(root)  
    event_column_var.set(default_event_column)  # 初始设置  
    event_column_dropdown = ttk.Combobox(mainframe, textvariable=event_column_var, values=columns, width=40)  # 设置宽度为30个字符  
    event_column_dropdown.grid(column=2, row=2, sticky=(tk.W, tk.E))  
      
    # 分析目标列下拉菜单  
    target_label = ttk.Label(mainframe, text="目标筛选列:")  
    target_label.grid(column=1, row=3, sticky=tk.W)  
    target_column_var = tk.StringVar(root)  
    target_column_var.set("不筛选")  # 设置默认值为“不筛选”  
    target_column_dropdown = ttk.Combobox(mainframe, textvariable=target_column_var, values=["不筛选"] + columns, width=40)  # 设置宽度为30个字符  
    target_column_dropdown.grid(column=2, row=3, sticky=(tk.W, tk.E))  
    target_column_dropdown.bind("<<ComboboxSelected>>", update_unique_values)  # 绑定事件，当选定列时更新分析对象下拉菜单  
      
    # 分析对象下拉菜单（显示选定列的去重清单）  
    analysis_object_label = ttk.Label(mainframe, text="目标对象:")  
    analysis_object_label.grid(column=1, row=4, sticky=tk.W)  
    analysis_object_var = tk.StringVar(root)  
    analysis_object_dropdown = ttk.Combobox(mainframe, textvariable=analysis_object_var, width=40)  # 设置宽度为30个字符  
    analysis_object_dropdown.grid(column=2, row=4, sticky=(tk.W, tk.E))  
    update_unique_values(None)  # 初始化时更新分析对象下拉菜单  


    #分析窗口  
    windows_label = ttk.Label(mainframe, text="分析窗口:")  
    windows_label.grid(column=1, row=5, sticky=tk.W)  
    windows_entry = ttk.Entry(mainframe)  
    windows_entry.insert(0, str(default_windows))  
    windows_entry.grid(column=2, row=5, sticky=(tk.W, tk.E))  
  
    method_label = ttk.Label(mainframe, text="方法:")  
    method_label.grid(column=1, row=6, sticky=tk.W)  
    method_var = tk.StringVar(root)  
    method_var.set(default_method)  # 初始设置  
    method_dropdown = ttk.Combobox(mainframe, textvariable=method_var, values=["nunique", "count", "sum"])  # 下拉菜单选择方法  
    method_dropdown.grid(column=2, row=6, sticky=(tk.W, tk.E))  
  
    freq_label = ttk.Label(mainframe, text="频率:")  
    freq_label.grid(column=1, row=7, sticky=tk.W)  
    freq_var = tk.StringVar(root)  
    freq_var.set(default_freq)  # 初始设置  
    freq_dropdown = ttk.Combobox(mainframe, textvariable=freq_var, values=["M", "Q"])  # 下拉菜单选择频率  
    freq_dropdown.grid(column=2, row=7, sticky=(tk.W, tk.E))  

    # 添加控制限选择下拉框  
    control_limit_label = ttk.Label(mainframe, text="控制限:")  
    control_limit_label.grid(column=1, row=8, sticky=tk.W)  
    control_limit_var = tk.StringVar(root)  
    control_limit_var.set("标准差")  # 默认设置为标准差  
    control_limit_dropdown = ttk.Combobox(mainframe, textvariable=control_limit_var, values=["标准差", "IQR"])  
    control_limit_dropdown.grid(column=2, row=8, sticky=(tk.W, tk.E))  
  
    # 创建确认按钮，用于启动趋势分析函数 
     
    #趋势图
    def confirm_1():  
        # 获取用户输入的参数值，并进行趋势分析函数调用  
        date_column = date_column_var.get()  # 获取日期列名称  
        event_column = event_column_var.get()  # 获取事件列名称  
        windows = int(windows_entry.get()) or default_windows  # 获取分析周期，若为空则使用默认值（需转换为整数）  
        method = method_var.get()  # 获取方法（计数或求和）  
        freq = freq_var.get()  # 获取数据聚合频率（月或季度）  
        target_column = target_column_var.get()  
        analysis_object=analysis_object_var.get()

        # 根据target_column的值决定传递哪个df  
        if target_column == "不筛选":  
            df_to_use = df  
        else: 
            
            pattern = r"'(.*?)'"  # 匹配单引号内的任意文本 
            analysis_object = re.findall(pattern,analysis_object) 
            df_to_use = df[df[target_column] == analysis_object[0]]  # 假设这是筛选逻辑  


      
        if control_limit_var.get() == "标准差":  
            TOOLS_trend_analysis_with_3_sd(df_to_use, date_column, event_column, windows, method, "draw", freq)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
        elif control_limit_var.get() == "IQR":  
            TOOLS_trend_analysis_with_1_5IQR(df_to_use, date_column, event_column, windows, method, "draw", freq)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
      
        root.destroy()  # 关闭窗口  
      
    # 趋势表  
    def confirm_2():  
        # 获取用户输入的参数值，并进行趋势分析函数调用  
        date_column = date_column_var.get()  # 获取日期列名称  
        event_column = event_column_var.get()  # 获取事件列名称  
        windows = int(windows_entry.get()) or default_windows  # 获取分析周期，若为空则使用默认值（需转换为整数）  
        method = method_var.get()  # 获取方法（计数或求和）  
        freq = freq_var.get()  # 获取数据聚合频率（月或季度）  
        target_column = target_column_var.get()  
        analysis_object=analysis_object_var.get()

        # 根据target_column的值决定传递哪个df  
        if target_column == "不筛选":  
            df_to_use = df  
        else: 
            pattern = r"'(.*?)'"  # 匹配单引号内的任意文本 
            analysis_object = re.findall(pattern,analysis_object)  

                            
            df_to_use = df[df[target_column] == analysis_object[0]]  # 假设这是筛选逻辑  

      
        if control_limit_var.get() == "标准差":  
            PROGRAM_display_df_in_treeview(TOOLS_trend_analysis_with_3_sd(df_to_use, date_column, event_column, windows, method, "data", freq), 1, 0)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
        elif control_limit_var.get() == "IQR":  
            PROGRAM_display_df_in_treeview(TOOLS_trend_analysis_with_1_5IQR(df_to_use, date_column, event_column, windows, method, "data", freq), 1, 0)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
      
        root.destroy()  # 关闭窗口   
    confirm_button = ttk.Button(mainframe, text="趋势图", command=confirm_1)  # 创建确认按钮，用于启动趋势分析函数，看图  
    confirm_button.grid(column=2, row=9, sticky=(tk.W, tk.E))  
    confirm_button = ttk.Button(mainframe, text="趋势表", command=confirm_2)  # 创建确认按钮，用于启动趋势分析函数，看表  
    confirm_button.grid(column=1, row=9, sticky=(tk.W, tk.E))    
    # 运行Tkinter主循环，等待用户操作  
    root.mainloop()

#趋势分析函数    
def TOOLS_trend_analysis_with_3_sd(df, date_column, event_column, windows, method,draw_or_data, freq='M'):  
    """  
    df: pandas DataFrame, 包含日期和事件的数据  
    date_column: str, 日期列的名称  
    event_column: str, 事件列的名称  ，比如报告表编码
    windows: 分析周期  
    method: 'count' 或 'sum'  或 'unique'
    freq: 数据聚合频率，'M' 表示按月，'Q' 表示按季度  
    draw_or_data:draw绘图，data返回数据
    """  
    # 确保日期列为datetime，并按日期排序  
    df.loc[:, date_column] = pd.to_datetime(df[date_column]) 
    df = df.sort_values(by=date_column)  
  
    # 按指定频率对事件进行分组和计数/求和，即使数据为0也要纳入  
    if method == "count":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].count()  
    elif method == "sum":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].sum()  
    elif method == "nunique":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].nunique()  
    else:  
        raise ValueError("Invalid method. Please use 'count', 'sum', or 'unique'.")  
  
    monthly_counts = monthly_counts.resample(freq).asfreq().fillna(0)  # Resample to ensure monthly/quarterly data, fill NaNs with 0  
  
    # 计算移动平均线和标准差控制限  
    rolling_mean = monthly_counts.rolling(window=windows).mean()  
    rolling_std = monthly_counts.rolling(window=windows).std()  
    UCL_2sd = rolling_mean + 2 * rolling_std  
    LCL_2sd = rolling_mean - 2 * rolling_std  
    UCL_3sd = rolling_mean + 3 * rolling_std  
    LCL_3sd = rolling_mean - 3 * rolling_std  
  
    result_df = pd.DataFrame({  
        'Date': monthly_counts.index,  
        'Monthly Counts': monthly_counts.values,  
        'Rolling Mean': rolling_mean.values,  
        'UCL (2 SD)': UCL_2sd.values,  
        'LCL (2 SD)': LCL_2sd.values,  
        'UCL (3 SD)': UCL_3sd.values,  
        'LCL (3 SD)': LCL_3sd.values  
    })  
    if draw_or_data=="data": 
        return result_df    

    # 创建Tkinter窗口和matplotlib图形    
    root = tk.Tk()   
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("Trend Analysis")    
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）      
    sw = root.winfo_screenwidth()      
    sh = root.winfo_screenheight()      
    ww = 1300  # 窗口宽度      
    wh = 700  # 窗口高度      
    x = (sw - ww) // 2      
    y = (sh - wh) // 2      
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    fig, ax = plt.subplots(figsize=(10, 6))    
        
    # 解决汉字乱码问题    
    plt.rcParams["font.sans-serif"] = [my_font_ch]  # 使用指定的汉字字体类型（此处为黑体）    
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号    
        
    # 绘制趋势图    
    ax.plot(result_df['Date'], result_df['Monthly Counts'], marker='o', label='Monthly Counts')    
    ax.plot(result_df['Date'], result_df['Rolling Mean'], color='red', label='Rolling Mean ({} Months)'.format(windows))    
    ax.plot(result_df['Date'], result_df['UCL (2 SD)'], color='blue', linestyle='--', label='UCL (2 SD)')  # 2 SD line in blue    
    ax.plot(result_df['Date'], result_df['LCL (2 SD)'], color='blue', linestyle='--')  # 2 SD line in blue    
    ax.plot(result_df['Date'], result_df['UCL (3 SD)'], color='green', linestyle='--', label='UCL (3 SD)')  # 3 SD line in green    
    ax.plot(result_df['Date'], result_df['LCL (3 SD)'], color='green', linestyle='--')  # 3 SD line in green    
    ax.set_xlabel('Date')    
    ax.set_ylabel('Event Counts')    
    ax.set_title('Trend Analysis with Control Limits (2 and 3 Standard Deviations)')    
    ax.legend()    
    ax.grid(True)    
        
    # 使用FigureCanvasTkAgg将图形嵌入到Tkinter窗口中    
    canvas = FigureCanvasTkAgg(fig, master=root)    
    canvas.draw()    
    canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)    
        
    # 创建和显示工具条  
    toolbar = NavigationToolbar2Tk(canvas, root)  
    toolbar.update()  
    canvas.get_tk_widget().pack()  

    # 运行Tkinter事件循环  
    root.mainloop()

#趋势分析函数
def TOOLS_trend_analysis_with_1_5IQR(df, date_column, event_column, windows, method, draw_or_data, freq='M'):  
    """  
    df: pandas DataFrame, 包含日期和事件的数据  
    date_column: str, 日期列的名称  
    event_column: str, 事件列的名称  
    windows: 分析周期  
    method: 'count' 或 'sum'  或 'unique'  
    freq: 数据聚合频率，'M' 表示按月，'Q' 表示按季度  
    draw_or_data:draw绘图，data返回数据  
    """  
    # 确保日期列为datetime，并按日期排序  
    df.loc[:, date_column] = pd.to_datetime(df[date_column])  
    df = df.sort_values(by=date_column)  
  
    # 按指定频率对事件进行分组和计数/求和，即使数据为0也要纳入  
    if method == "count":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].count()  
    elif method == "sum":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].sum()  
    elif method == "nunique":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].nunique()  
    else:  
        raise ValueError("Invalid method. Please use 'count', 'sum', or 'unique'.")  
  
    monthly_counts = monthly_counts.resample(freq).asfreq().fillna(0)  # Resample to ensure monthly/quarterly data, fill NaNs with 0  
  
    # 计算移动平均线和IQR控制限  
    rolling_mean = monthly_counts.rolling(window=windows).mean()  
    rolling_25p = monthly_counts.rolling(window=windows).quantile(0.25)  # 25th percentile  
    rolling_50p = monthly_counts.rolling(window=windows).quantile(0.5)  # 50th percentile (median)  
    rolling_75p = monthly_counts.rolling(window=windows).quantile(0.75)  # 75th percentile  
    IQR = rolling_75p - rolling_25p  # Interquartile range  
    UCL = rolling_75p + 1.5 * IQR  # Upper control limit (1.5 IQR above 75th percentile)  
    LCL = rolling_25p - 1.5 * IQR  # Lower control limit (1.5 IQR below 25th percentile)  
  
    result_df = pd.DataFrame({  
        'Date': monthly_counts.index,  
        'Monthly Counts': monthly_counts.values,  
        'Rolling Mean': rolling_mean.values,  
        'Rolling 25th Percentile': rolling_25p.values,  
        'Rolling Median': rolling_50p.values,  
        'Rolling 75th Percentile': rolling_75p.values,  
        'UCL (1.5 IQR)': UCL.values,  
        'LCL (1.5 IQR)': LCL.values  
    })  
    if draw_or_data=="data":   
        return result_df      
  
    # 创建Tkinter窗口和matplotlib图形      
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("Trend Analysis")      
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）        
    sw = root.winfo_screenwidth()        
    sh = root.winfo_screenheight()        
    ww = 1300  # 窗口宽度        
    wh = 700  # 窗口高度        
    x = (sw - ww) // 2        
    y = (sh - wh) // 2        
    root.geometry(f"{ww}x{wh}+{x}+{y}")      
    fig, ax = plt.subplots(figsize=(10, 6))      
          
    # 解决汉字乱码问题      
    plt.rcParams["font.sans-serif"] = [my_font_ch]  # 使用指定的汉字字体类型（此处为黑体）      
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号      
          
    # 绘制趋势图      
    ax.plot(result_df['Date'], result_df['Monthly Counts'], marker='o', label='Monthly Counts')      
    ax.plot(result_df['Date'], result_df['Rolling Mean'], color='red', label='Rolling Mean ({} Months)'.format(windows))      
    ax.plot(result_df['Date'], result_df['Rolling Median'], color='orange', linestyle='--', label='Rolling Median ({} Months)'.format(windows))  # Median line in orange      
    ax.plot(result_df['Date'], result_df['Rolling 25th Percentile'], color='green', linestyle='--', label='Rolling 25th Percentile ({} Months)'.format(windows))  # 25th percentile line in green      
    ax.plot(result_df['Date'], result_df['Rolling 75th Percentile'], color='blue', linestyle='--', label='Rolling 75th Percentile ({} Months)'.format(windows))  # 75th percentile line in blue       
    ax.plot(result_df['Date'], result_df['UCL (1.5 IQR)'], color='purple', linestyle='--', label='UCL (1.5 IQR)')  # UCL line in purple      
    ax.plot(result_df['Date'], result_df['LCL (1.5 IQR)'], color='purple', linestyle='--')  # LCL line in purple       
    ax.set_xlabel('Date')      
    ax.set_ylabel('Event Counts')      
    ax.set_title('Trend Analysis with Control Limits (1.5 IQR)')      
    ax.legend()      
    ax.grid(True)      
          
    # 使用FigureCanvasTkAgg将图形嵌入到Tkinter窗口中      
    canvas = FigureCanvasTkAgg(fig, master=root)      
    canvas.draw()      
    canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)      
          
    # 创建和显示工具条    
    toolbar = NavigationToolbar2Tk(canvas, root)    
    toolbar.update()    
    canvas.get_tk_widget().pack()    
  
    # 运行Tkinter事件循环    
    root.mainloop() 





############################################################################################################################
#绘图函数
############################################################################################################################
class AAA_10_DRAW():
    pass

#数据看板    
class DRAW_show_analysis_gui:
    """
    数据可视化面板-引导函数。
    """
    def __init__(self, df):
        self.df = df
        self.root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(self.root)

        self.root.title("选择列名")
        self._center_window(400, 300)
        self._setup_ui()
        
    def _center_window(self, width, height):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        self.root.geometry(f"{width}x{height}+{x}+{y}")
        
    def _setup_ui(self):
        columns = list(self.df.columns)
        
        ttk.Label(self.root, text="请选择需要透视查看的相关列").pack(pady=10)
        
        frame = ttk.Frame(self.root)
        frame.pack(pady=10)
        
        # 直接创建Combobox，不绑定StringVar
        ttk.Label(frame, text="分组列:").grid(row=0, column=0, sticky="e")
        self.combo1 = ttk.Combobox(frame, values=columns, state="readonly")
        self.combo1.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(frame, text="透视列:").grid(row=1, column=0, sticky="e")
        self.combo2 = ttk.Combobox(frame, values=columns, state="readonly")
        self.combo2.grid(row=1, column=1, padx=5, pady=5)
        ttk.Label(frame, text="时间列:").grid(row=2, column=0, sticky="e")
        self.combo3 = ttk.Combobox(frame, values=columns, state="readonly")
        self.combo3.grid(row=2, column=1, padx=5, pady=5)
        
        btn_frame = ttk.Frame(self.root)
        btn_frame.pack(pady=20)
        
        ttk.Button(btn_frame, text="确认", command=self._on_confirm).grid(row=0, column=0, padx=10)
        ttk.Button(btn_frame, text="取消", command=self.root.destroy).grid(row=0, column=1, padx=10)
    
    def _on_confirm(self):
        # 直接获取Combobox当前值
        v1 = self.combo1.get()
        v2 = self.combo2.get()
        v3 = self.combo3.get()
        
        print(f"当前选择: {v1}, {v2}, {v3}")  # 调试输出
        
        if not all([v1, v2, v3]):
            messagebox.showwarning("警告", "请完整选择三个列")
            return
            
        DRAW_show_analysis(self.df, v1, v2, v3)
        self.root.destroy()
        
    def run(self):
        self.root.mainloop()

#数据看板附属函数 
def DRAW_show_analysis(df, target1, column, event_date_col):
    """数据可视化面板-主分析函数，显示医疗器械数据分析界面"""
    if 1==1:
        # 设置中文字体
        plt.rcParams['font.sans-serif'] = [my_font_ch]
        plt.rcParams['axes.unicode_minus'] = False

        # 定义实际业务列名
        report_code_col = '报告编码'
        fault_performance_col = column
        registration_col = target1
        
        # 检查必要列是否存在
        required_cols = [report_code_col, fault_performance_col, event_date_col, registration_col]
        missing_cols = [col for col in required_cols if col not in df.columns]

        if missing_cols:
            raise ValueError(f"缺少必要列: {', '.join(missing_cols)}")
        
        df[[report_code_col, fault_performance_col, event_date_col, registration_col]]= df[[report_code_col, fault_performance_col, event_date_col, registration_col]].fillna('-未填写-')
        
        # 创建主窗口
        root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

        root.title("可视化面板")
        root.geometry("1400x800")
        
        # 当前选中的数据
        current_reg = '(全部)'
        current_pivot_df = pd.DataFrame()
        current_filtered_df = pd.DataFrame()
        current_display_data = pd.DataFrame()
        current_chart = 'pie'
        current_time_granularity = 'month'

        # --------------------------------------------
        # 内部函数定义
        # --------------------------------------------

        class RightClickMenu:
            def __init__(self, parent, treeview):
                self.menu = tk.Menu(parent, tearoff=0)
                self.treeview = treeview
                self.menu.add_command(label="导出整个表格", command=self.export_csv)
                self.menu.add_command(label="大模型辅助分析", command=self.print_table)
                self.menu.add_command(label="大模型辅助分析(原始报告）", command=self.find_original_reports)
            
            def show(self, event):
                self.menu.post(event.x_root, event.y_root)
            
            def export_csv(self):
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")],
                    title="保存表格为Excel"
                )
                if file_path:
                    try:
                        import openpyxl
                        columns = self.treeview["columns"]
                        data = []
                        for item in self.treeview.get_children():
                            values = self.treeview.item(item, 'values')
                            data.append(values)
                        
                        df = pd.DataFrame(data, columns=columns)
                        df.to_excel(file_path, index=False, engine='openpyxl')
                        messagebox.showinfo("成功", "表格已成功导出为Excel文件！")
                    except ImportError:
                        messagebox.showerror("依赖缺失", "导出Excel需要openpyxl库，请执行 pip install openpyxl 安装")
                    except Exception as e:
                        messagebox.showerror("错误", f"导出失败: {str(e)}")
            
            def print_table(self):
                columns = self.treeview["columns"]
                data = []
                for item in self.treeview.get_children():
                    values = self.treeview.item(item, 'values')
                    data.append(values)
                
                df = pd.DataFrame(data, columns=columns)
                LLM_Send_Form(df)
            
            def find_original_reports(self):
                """找出原始报告中与当前表格所有报告编码一致的行"""
                # 检查当前treeview是否有报告编码列
                columns = self.treeview["columns"]
                if '报告编码' not in columns:
                    messagebox.showwarning("警告", "当前表格中没有'报告编码'列！")
                    return
                
                # 获取当前表格中所有报告编码
                report_codes = []
                for item in self.treeview.get_children():
                    values = self.treeview.item(item, 'values')
                    col_index = columns.index('报告编码')
                    report_codes.append(values[col_index])
                
                if not report_codes:
                    messagebox.showinfo("提示", "当前表格中没有数据")
                    return
                
                # 在原始数据中查找这些报告编码
                try:
                    original_reports = df[df['报告编码'].isin(report_codes)]
                    if original_reports.empty:
                        messagebox.showinfo("结果", "未找到匹配的原始报告")
                    else:
                        LLM_Send_Form(original_reports)
                except Exception as e:
                    messagebox.showerror("错误", f"查找失败: {str(e)}")
                    
        def update_chart(pivot_data):
            """更新右上角图表"""
            nonlocal current_chart
            ax.clear()
            ax.set_facecolor('white')
            fig.set_tight_layout(False)
            fig.subplots_adjust(left=0, right=1, bottom=0, top=1)

            sorted_data = pivot_data.sort_values('报告数量', ascending=False)
            
            if current_chart == 'bar':
                fig.set_size_inches(10, 6)
                display_data = sorted_data.head(15)
                y_pos = np.arange(len(display_data))
                bars = ax.barh(
                    y_pos,
                    display_data['报告数量'],
                    height=0.6,
                    align='center',
                    color='#1f77b4'
                )
                ax.invert_yaxis()
                
                ax.set_yticks(y_pos)
                ax.set_yticklabels(display_data[fault_performance_col], fontsize=10)
                ax.set_xlabel('报告数量', fontsize=10)
                ax.set_title('统计对象报告数量TOP15（降序排列）', pad=20, fontsize=12)
                
                max_value = display_data['报告数量'].max()
                for i, value in enumerate(display_data['报告数量']):
                    ax.text(
                        value + max_value * 0.02,
                        i,
                        f'{int(value)}',
                        va='center',
                        ha='left',
                        fontsize=9
                    )
                
                ax.grid(axis='x', linestyle=':', alpha=0.4)
                ax.margins(y=0.03)
                ax.set_xlim(right=max_value * 1.2)
                
            else:
                fig.set_size_inches(6, 4)
                if len(sorted_data) > 15:
                    top_data = sorted_data.head(15)
                    other_sum = sorted_data['报告数量'][15:].sum()
                    combined_data = pd.concat([
                        top_data,
                        pd.DataFrame({fault_performance_col: ['其他'], '报告数量': [other_sum]})
                    ])
                else:
                    combined_data = sorted_data
                    
                wedges, texts, autotexts = ax.pie(
                    combined_data['报告数量'],
                    labels=None,
                    autopct=lambda p: f'{p:.1f}%' if p >= 3 else '',
                    startangle=90,
                    counterclock=False,
                    pctdistance=0.8,
                    textprops={'fontsize': 8}
                )
                
                if ax.get_legend():
                    ax.get_legend().remove()
                    
                ax.legend(
                    wedges,
                    combined_data[fault_performance_col],
                    title="统计对象",
                    loc="center left",
                    bbox_to_anchor=(1, 0, 0.5, 1),
                    fontsize=8
                )
                ax.set_title('统计对象占比分布', fontsize=12)
            
            fig.tight_layout()
            canvas.draw_idle()

        def update_trend(data):
            """更新右下角趋势图"""
            trend_ax.clear()
            if not data.empty:
                try:
                    data = data.copy()
                    
                    if event_date_col not in data.columns:
                        trend_ax.set_title(f"错误: 缺少日期列 '{event_date_col}'")
                        trend_canvas.draw()
                        return
                        
                    try:
                        date_series = pd.to_datetime(
                            data[event_date_col].astype(str), 
                            errors='coerce'
                        )
                        
                        if date_series.isnull().all():
                            trend_ax.set_title(f"错误: 无法解析日期数据")
                            trend_canvas.draw()
                            return
                            
                        data = data.assign(_datetime_=date_series).dropna(subset=['_datetime_'])
                        
                        if data.empty:
                            trend_ax.set_title("无有效日期数据")
                            trend_canvas.draw()
                            return
                            
                        min_date = data['_datetime_'].min()
                        max_date = data['_datetime_'].max()
                        
                        if current_time_granularity == 'day':
                            full_range = pd.date_range(min_date.floor('D'), max_date.ceil('D'), freq='D')
                            time_series = data.groupby(data['_datetime_'].dt.floor('D'))[report_code_col].nunique()
                            x_label = '日期'
                        elif current_time_granularity == 'month':
                            full_range = pd.date_range(min_date.floor('D'), max_date.ceil('D'), freq='MS')
                            time_series = data.groupby(data['_datetime_'].dt.to_period('M'))[report_code_col].nunique()
                            x_label = '月份'
                        elif current_time_granularity == 'quarter':
                            full_range = pd.date_range(min_date.floor('D'), max_date.ceil('D'), freq='QS')
                            time_series = data.groupby(data['_datetime_'].dt.to_period('Q'))[report_code_col].nunique()
                            x_label = '季度'
                        elif current_time_granularity == 'year':
                            full_range = pd.date_range(min_date.floor('D'), max_date.ceil('D'), freq='YS')
                            time_series = data.groupby(data['_datetime_'].dt.year)[report_code_col].nunique()
                            x_label = '年份'
                        
                        time_series_df = time_series.reset_index()
                        time_series_df.columns = ['time_period', 'count']
                        
                        if current_time_granularity == 'day':
                            full_df = pd.DataFrame({'time_period': full_range})
                        elif current_time_granularity == 'month':
                            full_df = pd.DataFrame({'time_period': full_range.to_period('M')})
                        elif current_time_granularity == 'quarter':
                            full_df = pd.DataFrame({'time_period': full_range.to_period('Q')})
                        elif current_time_granularity == 'year':
                            full_df = pd.DataFrame({'time_period': full_range.year})
                        
                        merged = pd.merge(full_df, time_series_df, on='time_period', how='left').fillna(0)
                        merged = merged.sort_values('time_period')
                        
                        x_labels = [
                            str(period) if count > 0 else '' 
                            for period, count in zip(merged['time_period'], merged['count'])
                        ]
                        x_values = range(len(x_labels))
                        y_values = merged['count'].astype(int)
                        
                        trend_ax.plot(x_values, y_values, marker='o', label='报告数量')
                        trend_ax.set_xticks(x_values)
                        trend_ax.set_xticklabels(x_labels)
                        trend_ax.set_title('时间趋势图')
                        trend_ax.set_xlabel(x_label)
                        trend_ax.set_ylabel('报告数量')
                        trend_ax.legend()
                        
                        if len(x_labels) > 5:
                            trend_ax.tick_params(axis='x', rotation=45)
                        
                        trend_ax.grid(True, linestyle='--', alpha=0.6)
                        trend_fig.tight_layout()
                        
                    except Exception as e:
                        trend_ax.set_title(f"处理错误: {str(e)}")
                        
                except Exception as e:
                    trend_ax.set_title(f"严重错误: {str(e)}")
            else:
                trend_ax.set_title("无数据")
            
            trend_canvas.draw()

        def update_bottom_left(data):
            """更新左下角数据表格"""
            nonlocal current_display_data
            current_display_data = data.copy()
            bl_tree.delete(*bl_tree.get_children())
            for _, row in data.iterrows():
                bl_tree.insert('', 'end', values=list(row[display_cols]))
            update_trend(data)

        def update_views(reg_number='(全部)'):
            """更新所有视图"""
            nonlocal current_reg, current_pivot_df, current_filtered_df
            
            current_reg = reg_number
            if reg_number == '(全部)':
                filtered = df.copy()
            else:
                filtered = df[df[registration_col] == reg_number].copy()
            
            current_filtered_df = filtered
            
            pivot_data = filtered.groupby(fault_performance_col)[report_code_col].nunique().reset_index()
            pivot_data.columns = [fault_performance_col, '报告数量']
            pivot_data = pivot_data.sort_values('报告数量', ascending=False)
            current_pivot_df = pivot_data
            
            tl_tree.delete(*tl_tree.get_children())
            for _, row in pivot_data.iterrows():
                tl_tree.insert('', 'end', values=list(row))
            
            update_chart(pivot_data)
            update_bottom_left(filtered)

        def on_reg_select(event):
            """注册证编号选择事件"""
            selected = reg_tree.selection()
            if selected:
                value = reg_tree.item(selected[0], 'values')[0]
                update_views(value)

            
        def on_pivot_select(event):
            """透视表选择事件（单击）"""
            selected = tl_tree.selection()
            if selected and not current_filtered_df.empty:
                value = tl_tree.item(selected[0], 'values')[0]
                filtered = current_filtered_df[current_filtered_df[fault_performance_col] == value].copy()
                update_bottom_left(filtered)
            else:
                update_bottom_left(current_filtered_df)

        def on_pivot_double_click(event):
            """透视表双击事件"""
            selected = tl_tree.selection()
            if selected and not current_filtered_df.empty:
                value = tl_tree.item(selected[0], 'values')[0]
                filtered = current_filtered_df[current_filtered_df[fault_performance_col] == value].copy()
                PROGRAM_display_df_in_treeview(filtered, 'psur', filtered)


        def print_data_row(event):
            """打印数据表格当前行对应的原始数据"""
            selected = bl_tree.selection()
            if selected:
                item = bl_tree.selection()[0]
                values = bl_tree.item(item, 'values')
                report_code = values[0]  # 假设第一列是报告编码
                
                # 在原始数据中查找该报告编码
                original_data = df[df[report_code_col] == report_code]
                PROGRAM_display_content_in_textbox(original_data.to_dict())



        def toggle_chart():
            """切换图表类型"""
            nonlocal current_chart
            current_chart = 'pie' if current_chart == 'bar' else 'bar'
            update_chart(current_pivot_df)

        def export_chart():
            """导出图表"""
            file_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG图片", "*.png"), ("所有文件", "*.*")],
                title="保存图表"
            )
            if file_path:
                try:
                    fig.savefig(file_path, dpi=300, bbox_inches='tight')
                    messagebox.showinfo("成功", "图表已成功导出！")
                except Exception as e:
                    messagebox.showerror("错误", f"导出失败: {str(e)}")

        def export_trend_chart():
            """导出趋势图"""
            file_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG图片", "*.png"), ("所有文件", "*.*")],
                title="保存趋势图"
            )
            if file_path:
                try:
                    trend_fig.savefig(file_path, dpi=300, bbox_inches='tight')
                    messagebox.showinfo("成功", "趋势图已成功导出！")
                except Exception as e:
                    messagebox.showerror("错误", f"导出失败: {str(e)}")

        def set_time_granularity(granularity):
            """设置时间粒度"""
            nonlocal current_time_granularity
            current_time_granularity = granularity
            update_trend(current_display_data)

        def print_row(event):
            """打印当前行"""
            item = reg_tree.selection()[0]
            values = reg_tree.item(item, 'values')
            print(f"\n表格当前行: {values}")

        def print_pivot_row(event):
            """打印透视表当前行"""
            item = tl_tree.selection()[0]
            values = tl_tree.item(item, 'values')
            print(f"\n透视表当前行: {values}")



        # --------------------------------------------
        # 界面布局
        # --------------------------------------------

        # 创建左侧注册证编号/曾用注册证编号列表框架
        left_frame = ttk.Frame(root, width=210)
        left_frame.pack(side='left', fill='y')
        left_frame.pack_propagate(False)

        # 注册证编号/曾用注册证编号列表
        reg_df = df.groupby(registration_col)[report_code_col].nunique().reset_index()
        reg_df = reg_df.sort_values('报告编码', ascending=False)
        reg_df.columns = [registration_col, '报告数量']

        reg_container = ttk.Frame(left_frame)
        reg_container.pack(fill='both', expand=True, padx=5, pady=5)

        reg_tree = ttk.Treeview(
            reg_container,
            columns=[registration_col, '报告数量'],
            show='headings',
            selectmode='browse'
        )

        reg_vsb = ttk.Scrollbar(reg_container, orient="vertical", command=reg_tree.yview)
        reg_hsb = ttk.Scrollbar(reg_container, orient="horizontal", command=reg_tree.xview)
        reg_tree.configure(yscrollcommand=reg_vsb.set, xscrollcommand=reg_hsb.set)

        reg_tree.heading(registration_col, text=target1)
        reg_tree.column(registration_col, width=120)
        reg_tree.heading('报告数量', text='报告数量')
        reg_tree.column('报告数量', width=60, anchor='e')

        reg_tree.grid(row=0, column=0, sticky="nsew")
        reg_vsb.grid(row=0, column=1, sticky="ns")
        reg_hsb.grid(row=1, column=0, sticky="ew")

        reg_container.grid_rowconfigure(0, weight=1)
        reg_container.grid_columnconfigure(0, weight=1)

        for _, row in reg_df.iterrows():
            reg_tree.insert('', 'end', values=list(row))

        reg_tree.insert('', 0, values=['(全部)', df[report_code_col].nunique()], tags=('all',))
        reg_tree.tag_configure('all', background='#f0f0ff')
        
        reg_tree_menu = RightClickMenu(root, reg_tree)
        reg_tree.bind("<Button-3>", reg_tree_menu.show)
        reg_tree.bind("<Double-1>", print_row)
        reg_tree.bind('<<TreeviewSelect>>', on_reg_select)

        # 创建主内容框架（四个象限）
        main_frame = ttk.Frame(root)
        main_frame.pack(side='right', fill='both', expand=True)
        
        top_left = ttk.Frame(main_frame)
        top_right = ttk.Frame(main_frame)
        bottom_left = ttk.Frame(main_frame)
        bottom_right = ttk.Frame(main_frame)

        for i in range(2):
            main_frame.rowconfigure(i, weight=1)
            main_frame.columnconfigure(i, weight=1)
        top_left.grid(row=0, column=0, sticky="nsew")
        top_right.grid(row=0, column=1, sticky="nsew")
        bottom_left.grid(row=1, column=0, sticky="nsew")
        bottom_right.grid(row=1, column=1, sticky="nsew")

        # 左上：数据透视表
        tl_frame = ttk.Frame(top_left)
        tl_frame.pack(fill='both', expand=True)
        
        tl_tree = ttk.Treeview(
            tl_frame,
            columns=[fault_performance_col, '报告数量'],
            show='headings',
            selectmode='browse'
        )
        
        vsb = ttk.Scrollbar(tl_frame, orient="vertical", command=tl_tree.yview)
        hsb = ttk.Scrollbar(tl_frame, orient="horizontal", command=tl_tree.xview)
        tl_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        tl_tree.heading(fault_performance_col, text=fault_performance_col)
        tl_tree.column(fault_performance_col, width=150, anchor='center')
        tl_tree.heading('报告数量', text='报告数量')
        tl_tree.column('报告数量', width=80, anchor='center')
        
        tl_tree_menu = RightClickMenu(root, tl_tree)
        tl_tree.bind("<Button-3>", tl_tree_menu.show)
        tl_tree.bind('<<TreeviewSelect>>', on_pivot_select)  # 单击事件
        tl_tree.bind("<Double-1>", on_pivot_double_click)    # 双击事件
        
        
        tl_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        tl_frame.grid_rowconfigure(0, weight=1)
        tl_frame.grid_columnconfigure(0, weight=1)

        # 右上：图表区域
        tr_frame = ttk.Frame(top_right)
        tr_frame.pack(fill='both', expand=True)
        
        fig, ax = plt.subplots(figsize=(6, 4))
        canvas = FigureCanvasTkAgg(fig, master=tr_frame)
        
        chart_btn_frame = ttk.Frame(tr_frame)
        chart_btn_frame.pack(side='bottom', fill='x')
        
        ttk.Button(chart_btn_frame, text="切换图表类型", command=toggle_chart).pack(side='left', padx=2)
        ttk.Button(chart_btn_frame, text="导出图表", command=export_chart).pack(side='right', padx=2)
        
        canvas.get_tk_widget().pack(fill='both', expand=True)

        # 左下：原始数据表格
        bl_frame = ttk.Frame(bottom_left)
        bl_frame.pack(fill='both', expand=True)
        
        display_cols = [report_code_col, registration_col, fault_performance_col, event_date_col]
        bl_tree = ttk.Treeview(
            bl_frame,
            columns=display_cols,
            show='headings'
        )
        
        bl_vsb = ttk.Scrollbar(bl_frame, orient="vertical", command=bl_tree.yview)
        bl_hsb = ttk.Scrollbar(bl_frame, orient="horizontal", command=bl_tree.xview)
        bl_tree.configure(yscrollcommand=bl_vsb.set, xscrollcommand=bl_hsb.set)
        
        for col in display_cols:
            bl_tree.heading(col, text=col)
            bl_tree.column(col, width=120)
        
        bl_tree_menu = RightClickMenu(root, bl_tree)
        bl_tree.bind("<Button-3>", bl_tree_menu.show)
        bl_tree.bind("<Double-1>", print_data_row)
        
        bl_tree.grid(row=0, column=0, sticky='nsew')
        bl_vsb.grid(row=0, column=1, sticky='ns')
        bl_hsb.grid(row=1, column=0, sticky='ew')
        bl_frame.grid_rowconfigure(0, weight=1)
        bl_frame.grid_columnconfigure(0, weight=1)



        # 右下：趋势分析图
        br_frame = ttk.Frame(bottom_right)
        br_frame.pack(fill='both', expand=True)
        
        trend_fig, trend_ax = plt.subplots(figsize=(6, 4))
        trend_canvas = FigureCanvasTkAgg(trend_fig, master=br_frame)
        
        time_btn_frame = ttk.Frame(br_frame)
        time_btn_frame.pack(side='bottom', fill='x')
        
        ttk.Button(time_btn_frame, text="按日", command=lambda: set_time_granularity('day')).pack(side='left', padx=2)
        ttk.Button(time_btn_frame, text="按月", command=lambda: set_time_granularity('month')).pack(side='left', padx=2)
        ttk.Button(time_btn_frame, text="按季", command=lambda: set_time_granularity('quarter')).pack(side='left', padx=2)
        ttk.Button(time_btn_frame, text="按年", command=lambda: set_time_granularity('year')).pack(side='left', padx=2)
        ttk.Button(time_btn_frame, text="导出图表", command=export_trend_chart).pack(side='right', padx=2)
        
        trend_canvas.get_tk_widget().pack(fill='both', expand=True)

        # 初始显示所有数据
        update_views()
        
        root.mainloop()


#绘图函数  
def DRAW_plot_df(ori_df):
    print("************************************************************************************")  
    print("功能说明：")
    print("1. 该工具用于对输入的DataFrame进行数据可视化，支持多种图表类型。")
    print("2. 支持的图表类型包括散点图、柱状图、折线图、饼图、帕累托图、热力图等。")
    print("3. 用户可以选择X轴、Y轴、Z分层、M补充等字段，动态生成图表。")
    print("4. 支持自动生成X轴数值、动态调整图例位置、清除图表等功能。")
    print("\n使用方法：")
    print("1. 在界面中选择X轴标签、Y轴标签、Y轴数值等字段（最简单需要这三个）。")
    print("2. 选择Z分层字段（可选）和M补充字段（可选）。")
    print("3. 选择图表类型（如散点图、柱状图等）。")
    print("4. 点击“绘图”按钮生成图表。")
    print("5. 使用工具栏进行图表操作（如缩放、保存等）。")
    print("6. 点击“图例”按钮调整图例显示方式。")
    print("7. 如需清除当前图表，请点击“清除”按钮。")
    print("************************************************************************************")  
    
    df=ori_df.copy() 
    df=df.fillna(0)
    print("绘图工具使用您的一份副本工作。请确保您的数据是经过汇总整理、且相关绘图要素无重复项。")
    cycle_iterator = itertools.cycle([1, 2, 3]) 

    def create_dropdown(frame, label_text, variable, options):  
        """Create a dropdown menu."""  
        label = ttk.Label(frame, text=label_text)  
        label.pack(side=tk.LEFT, pady=5)  
        dropdown = ttk.Combobox(frame, textvariable=variable, values=options, state='readonly',width= 15)  
        dropdown.pack(side=tk.LEFT,padx=3)  
        return dropdown  

    def clear_plot():  
        """Clears the current plot and resets axis properties."""  
        ax.clear()  # 清除轴上的所有内容  
        # 重置轴的限制和视图（可能需要根据具体情况调整）  
        ax.relim()  
        ax.autoscale_view()  
        # 可能还需要其他设置来确保下一个图能够正确显示  clear_legend
        canvas.draw()  # 重新绘制画布以更新显示 

    def clear_legend():  
        """Clears the current plot and resets axis properties."""  
         
        iss=next(cycle_iterator)
        if iss==1:
            ax.legend_ = None
        elif iss==2:
            ax.legend(loc='upper right', bbox_to_anchor=(1.11, 1.0), fontsize=8, borderaxespad=0.0) 
        elif iss==3:  
            ax.legend()   
                    
        # 清除轴上的所有内容  
        # 重置轴的限制和视图（可能需要根据具体情况调整）  
        ax.relim()  
        ax.autoscale_view()  
        # 可能还需要其他设置来确保下一个图能够正确显示  clear_legend
        canvas.draw()  # 重新绘制画布以更新显示 
        
    def draw_text(ax,texts):
        # 遍历每个条形并添加文本  
        for xxbar in ax.patches: 
            print(xxbar) 
            # 获取条形的x位置、宽度和高度  
            x = xxbar.get_x()  
            width = xxbar.get_width()  
            height = xxbar.get_height()  
            # 计算文本的位置（在条形的顶部中心）  
            text_x = x + width / 2  
            text_y = height     
            # 获取当前条形的索引（用于从texts中获取相应的文本）  
            index = int(xxbar.get_x()+1)    
            # 在条形顶部添加文本  
            ax.text(text_x, text_y, texts[index], ha='center', va='bottom', fontsize=8)     
        # 调整y轴的上限，以确保文本完全可见  
        ax.relim()  
        ax.autoscale_view()
  
 

  
    def draw_plot(x_label,x_value, y_label, y_value, z_column, m_column, plot_type, my_colors): 
        """Draw the plot based on user's selection."""  
        dfs=df.copy()
        #print(x_label,x_value, y_label, y_value, z_column, m_column, plot_type, my_colors)
        if x_value=="-自动生成-":
            dfs[x_value] = pd.factorize(dfs[x_label])[0]+1 
            dfs[x_value] =dfs[x_value].astype(int)
        if y_label=="-不显示-":
            dfs[y_label] = None
        next_color = next(my_colors)  # 获取下一个颜色
        if plot_type == "---------":
            return

 
        #################################################          

        if plot_type == "帕累托图(XY)":  
            clear_plot()  # 清除之前的图形  
            ax2 = ax.twinx()  # 创建第二个坐标轴对象用于累计百分比  
      
            # 确保使用传入的dfs而不是未定义的df  
            dfs = dfs.sort_values(by=y_value, ascending=False)  # 按y_value列的值降序排序  

            dfs['Cumulative Percentage'] = dfs[y_value].cumsum() / dfs[y_value].sum() * 100  # 计算累计百分比  
      
            # 绘制条形图  
            dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color, legend=False)  
            next_color = next(my_colors)  # 获取下一个颜色  
      
            # 绘制累计百分比曲线  
            dfs.plot(kind='line',  x=x_label, y='Cumulative Percentage', ax=ax2, color=next_color, marker='D', ms=4)  

            draw_text(ax,dfs[y_label])

            # 设置坐标轴格式和标签  
            ax.set_xlabel(x_label)  # 设置X轴标签  
   
            ax.set_ylabel('count')  # 设置左侧Y轴标签  
            ax2.set_ylabel('Cumulative Percentage (%)')  # 设置右侧Y轴标签  
            ax2.yaxis.set_ticks_position('right')  # 将累计百分比的坐标轴放在右边  
            ax.set_xticklabels(dfs[x_label], rotation=90)  # 设置X轴标签  

            ax2.set_ylim(0, 100)  # 设置累计百分比的y轴范围  
            ax.set_title('Pareto Chart')  # 设置图表标题  

        elif plot_type == "热力图":  
            clear_plot()
            def SMALL_pre_process_dataframe(df):  
                """  
                处理DataFrame中的非数字列和缺失值。  
                  
                参数:  
                    df (pd.DataFrame): 输入的DataFrame。  
                      
                返回:  
                    pd.DataFrame: 处理后的DataFrame。  
                """  
                # 复制原始DataFrame以防修改原始数据  
                df_processed = df.copy()  
                  
                # 遍历每一列  
                for col in df_processed.columns:  
                    # 检查列的数据类型  
                    if pd.api.types.is_string_dtype(df_processed[col]):  
                        # 如果是字符串类型，尝试将其转换为类别型并进行编码  
                        df_processed[col] = df_processed[col].astype('category').cat.codes  
                    elif pd.api.types.is_numeric_dtype(df_processed[col]):  
                        # 如果是数字类型但有缺失值，使用均值填充  
                        df_processed[col].fillna(df_processed[col].mean(), inplace=True)  
                    else:  
                        # 对于其他类型，可以考虑删除列或转换为适当的类型（根据具体情况）  
                        # 这里我们选择删除列作为示例（根据实际情况调整）  
                        df_processed.drop(columns=col, inplace=True)  
                        print(f"Column '{col}' has been dropped because its data type is not supported.")  
                  
                # 额外处理：检查是否有任何非数字值（可能是由于之前的步骤中未处理的特殊情况）  
                # 如果有，将这些值替换为NaN，并使用该列的均值填充  
                for col in df_processed.columns:  
                    df_processed[col] = pd.to_numeric(df_processed[col], errors='coerce')  
                    if df_processed[col].isnull().sum() > 0:  
                        df_processed[col].fillna(df_processed[col].mean(), inplace=True)  
                  
                return df_processed  
            dfs=SMALL_pre_process_dataframe(dfs.copy())
            # 计算相关性矩阵  
            corr = dfs.corr()    
            # 生成一个mask用于隐藏上半部分的热力图（因为相关性矩阵是对称的）  
            mask = np.triu(np.ones_like(corr, dtype=bool))
            sns.heatmap(corr, mask=mask, cmap='coolwarm', annot=True, fmt=".2f")  
        

              
        ######################################            
        elif plot_type == "散点图(XY-Z分层-M大小)":  #分层散点图(XYZ)
            clear_plot()
            # 设置点的大小
            if m_column !="-不选择-":
                min_size = dfs[m_column].min()  
                max_size = dfs[m_column].max()  
                if min_size < 0 or max_size > 1:  
                    dfs['normalized_size'] = (dfs[m_column] - min_size) / (max_size - min_size)*100  
                    size_column = 'normalized_size'
            else:
                dfs['normalized_size']=100
                size_column = 'normalized_size'
            
            #如果不分层
            if z_column =="-不选择-":
                dfs.plots=dfs.plot(kind='scatter', x=x_value, y=y_value, ax=ax,s=size_column, label=x_value) 
                jittered_x=x_value
                jittered_y=y_value

            #如果分层
            else:
                clear_plot()

  
                # 添加抖动和透明度参数  
                jitter_amount = 0.2  # 抖动量，根据需要调整  
                alpha_value = 0.7  # 透明度值，根据需要调整 
                jittered_x="X("+str(x_value)+")"
                jittered_y="X("+str(y_value)+")"
                dfs[jittered_x]=0
                dfs[jittered_y]=0
                dfs[jittered_x] = dfs[x_value] + np.random.uniform(-jitter_amount, jitter_amount, size=len(dfs))  
                dfs[jittered_y] = dfs[y_value] + np.random.uniform(-jitter_amount, jitter_amount, size=len(dfs))
                
                
                for i in dfs[z_column].unique():  
                    product_data = dfs[dfs[z_column] == i]  
  
                    product_data.plot(kind='scatter', x=jittered_x, y=jittered_y, ax=ax, c=next_color, alpha=0.7, label=i,s=size_column) 
                    next_color = next(my_colors) 


            labels = dfs[y_label]  
            texts = [ax.text(x00, y00, z00, color='black', size=8) for x00, y00, z00 in zip(dfs[jittered_x], dfs[jittered_y], labels)] 
        
            ax.set_xticks(dfs[x_value])  # 设置X轴刻度为index列的值  
            if x_value!="-自动生成-":
                ax.set_xticklabels(dfs[x_value])  # 设置X轴标签为x_label_column列的值  
            else:
                ax.set_xticklabels(dfs[x_label])
            ax.legend()


        elif plot_type == "折线图(XY-Z分层)": 
            clear_plot()   
            if z_column =="-不选择-":  
                dfs.plot(kind='line', x=x_value, y=y_value, ax=ax, color=next_color, label=x_label, marker='D', ms=3)  
            else:
                clear_plot() 
                for i in df[z_column].unique():  
                    product_data = dfs[dfs[z_column] == i].sort_values(by=x_value)  
                    product_data.plot(kind='line', x=x_value, y=y_value, ax=ax, color=next_color, alpha=0.7, label=i, marker='D', ms=3)
                    next_color = next(my_colors) 
            ax.set_xticks(dfs[x_value])  # 设置X轴刻度为index列的值  
            if x_value!="-自动生成-":
                ax.set_xticklabels(dfs[x_value])  # 设置X轴标签为x_label_column列的值  
            else:
                ax.set_xticklabels(dfs[x_label])
            ax.legend()

            dfs.plot(kind='scatter', x=x_value, y=y_value, ax=ax, s=0.01) 
            labels = dfs[y_label]  
            texts = [ax.text(x00, y00, z00, color='black', size=8) for x00, y00, z00 in zip(dfs[x_value], dfs[y_value], labels)]   

                      
            
        elif plot_type == "左右比对条形图(XYM)":         
            clear_plot()
            # 确保数据是numpy数组以便于计算  
            data1 = dfs[y_value].values  
            data2 = dfs[m_column].values  
            labels= dfs[x_value].values  
              
            # 计算每侧数据的总和，用于缩放条形图的高度以保持总面积不变  
            total1 = data1.sum()  
            total2 = data2.sum()  
            max_total = max(total1, total2)  # 用于确定条形的最大高度  
            bar_width = 0.4  # 条形宽度  
            gap = 0.05  # 间隔宽度  
            n = len(labels)  # 标签的数量  
              
            # 计算条形的y轴位置  
            y_pos = np.arange(n) - gap / 2  # 条形图左侧的位置  
            y_pos_right = y_pos + bar_width + gap  # 
              
            # 绘制左侧的条形图（负方向）    
            ax.barh(y_pos, -data1 / max_total * 100, height=bar_width, align='center', edgecolor='white', label=y_value)    
            # 绘制右侧的条形图（正方向），注意这里应该使用 m_column 作为标签    
            ax.barh(y_pos, data2 / max_total * 100, height=bar_width, align='center', edgecolor='white', label=m_column)  
              
            # 设置y轴的刻度位置和标签，以及网格线等属性  
            ax.set_yticks(y_pos)  # 将y轴刻度设置在条形的中间位置  
            ax.set_yticklabels(labels)  # 设置y轴的刻度标签为传入的标签列表  
            ax.invert_yaxis()  # 反转y轴，使得标签从底部开始向上排列  
            ax.axvline(x=0, linewidth=1)  # 添加中心线  
            ax.set_title("横向对比图")  # 图表标题 
            
            
            figure.tight_layout()  # 自动调整子图参数，使之填充整个图像区域（可选） 
            # 添加图例，并设置其位置和属性  
            ax.legend()    
        

                    
        elif plot_type == "添加网格":         
            ax.grid(True) 
                                    
        
        elif plot_type == "柱状图(XY-Z分层)":  
            clear_plot() 
            if x_value!="-自动生成-":
                messagebox.showinfo("提示", "X数值仅支持-自动生成-")
                return
                
            if z_column=="-不选择-": 
                dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color)  
            else:
                pivot_dfs = dfs.pivot(index=z_column, columns=x_label, values=y_value)  
                pivot_dfs.plot(kind='bar', width=0.8, ax=ax)   
                   

            draw_text(ax,dfs[y_label])
            
        elif plot_type == "堆叠柱状图(XYM-小到大)":  
            # 绘制堆叠柱状图 
            clear_plot() 
            dfs.plot(kind='bar', x=x_label, y=m_column, ax=ax, color=next_color) 
            #draw_text(ax,dfs[y_label])
            next_color = next(my_colors) 
            dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color) 


            
        elif plot_type == "横向条形图(XY-Z分层)":  
            clear_plot() 
            if x_value!="-自动生成-":
                messagebox.showinfo("提示", "X数值仅支持-自动生成-")
                return
                
            if z_column=="-不选择-": 
                dfs.plot(kind='barh', x=x_label, y=y_value, ax=ax, color=next_color)  
                
                
            else:
                pivot_dfs = dfs.pivot(index=z_column, columns=x_label, values=y_value)  
                pivot_dfs.plot(kind='barh', width=0.8, ax=ax)           
                  
                # 为每个条形添加数值标签  
            if y_label!="-不显示-":
                for p in ax.patches:  
                    ax.annotate(str(p.get_width()), (p.get_width(), p.get_y() + p.get_height() / 2), ha='left', va='center')  



      
            
            
        elif plot_type == "横向堆叠条形图(XYM-小到大)":  
            # 绘制横向条形图 
            clear_plot() 
            dfs.plot(kind='barh', x=x_label, y=m_column, ax=ax, color=next_color)    
            next_color = next(my_colors)
            dfs.plot(kind='barh', x=x_label, y=y_value, ax=ax, color=next_color) 
            if y_label!="-不显示-":
                for p in ax.patches:  
                    ax.annotate(str(p.get_width()), (p.get_width(), p.get_y() + p.get_height() / 2), ha='left', va='center')              



            
        ################################################# 
        elif plot_type == "饼图(XY)":  
            clear_plot() 
            ax.pie(dfs[y_value], labels=dfs[x_label], autopct='%1.1f%%', startangle=90)  
            ax.axis('equal')  
        #################################################    
                        
        else:  
            print("不支持的图表类型。请检查您的输入。")  
        
            #格式设置
       
        labels = df[x_label].tolist() 
        fontsize = 8  # 初始字体大小  
        rotation = 90  # 初始旋转角度  
        plt.xticks(fontsize=fontsize)  # 更新字体大小  
        plt.xticks(rotation=rotation)  # 设置标签旋转角度  
      
          
        ax.yaxis.tick_left()  # Ensure y-axis is on the left (reset if needed)   
        ax.relim()  
        ax.autoscale_view()  
        if ax.legend_ and plot_type!="左右比对条形图(XYM)":  # 只在有图例时才显示图例  
            ax.legend(loc='upper right', bbox_to_anchor=(1.11, 1.0), fontsize=8, borderaxespad=0.0) 
        canvas.draw()  
  
    # Create the main window  
    root = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)

    root.title("数据可视化工具")  
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 1400  # 窗口宽度          
    wh = 600  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")   
    main_frame = ttk.Frame(root, padding="0")  
    main_frame.pack(fill=tk.BOTH, expand=True)  
    
    # Set up matplotlib with Chinese font support  
    plt.rcParams["font.sans-serif"] = [my_font_ch]  # Use specified Chinese font type (here: SimHei)  
    plt.rcParams['axes.unicode_minus'] = False  # Display negative signs properly  
    my_colors = itertools.cycle(plt.rcParams['axes.prop_cycle'].by_key()['color'])  # 创建一个颜色循环器 
    # Create the figure and axes (only once)  
    figure, ax = plt.subplots(figsize=(10, 4))  # Adjust figure size as needed  
    canvas = FigureCanvasTkAgg(figure, master=main_frame)  
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, pady=10)  
    toolbar = NavigationToolbar2Tk(canvas, main_frame)  
    toolbar.update()  
    canvas._tkcanvas.pack(fill=tk.BOTH, expand=True)  # Re-pack the canvas to include the toolbar space  
    toolbar.pack(fill=tk.X)  # Place the toolbar at the top or appropriate location  
  
    cols = ["-不选择-"]+df.columns.tolist()  # Get a list of column names  
  
    # Create labels and drop-down menus (using the create_dropdown function)  
    x_label_vars = tk.StringVar(root)  
    #x_label_vars.set("单位名称")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "X标签:", x_label_vars,cols)  # Use column names list as selectable values  
    
    x_value_vars = tk.StringVar(root)  
    x_value_vars.set("-自动生成-")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "X数值:", x_value_vars, ["-自动生成-"]+df.columns.tolist())  # Use column names list as selectable values  

    y_label_vars = tk.StringVar(root)  
    y_label_vars.set("-不显示-")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "Y标签:", y_label_vars, ["-不显示-"]+df.columns.tolist())  # Use column names list as selectable values  

    y_value_vars = tk.StringVar(root)  
    #y_value_vars.set("报告编码其他")  # Default value for "Y-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "Y数值:", y_value_vars, cols)  # Use column names list as selectable values  
     
    z_column_var = tk.StringVar(root)    
    z_column_var.set("-不选择-")  # 假设数据框中有一个名为"风险级别"的列  
    create_dropdown(main_frame, "Z分层:", z_column_var, cols)     
    
    m_column_var = tk.StringVar(root)  
    m_column_var.set("-不选择-")  # 
    create_dropdown(main_frame, "M补充:", m_column_var, cols)  
  

    plot_type_var = tk.StringVar(root)  
    plot_type_var.set("散点图(XY-Z分层-M大小)")  # Default plot type is scatter plot  
    plot_type_dropdown = create_dropdown(main_frame, "绘图类型:", plot_type_var, ["散点图(XY-Z分层-M大小)","---------", "柱状图(XY-Z分层)","堆叠柱状图(XYM-小到大)","横向条形图(XY-Z分层)","左右比对条形图(XYM)","横向堆叠条形图(XYM-小到大)","---------","折线图(XY-Z分层)","帕累托图(XY)","---------","饼图(XY)","---------", "热力图","添加网格"])  # Allow selecting plot type from a list of options  

    
    # 创建一个清除按钮   
    #clear_button = ttk.Button(main_frame, text="清除", command=clear_plot)  
    #clear_button.pack(side=tk.LEFT, pady=5) 

    # 创建一个清除按钮   
    Mclear_button = ttk.Button(main_frame, text="图例", command=clear_legend)  
    Mclear_button.pack(side=tk.LEFT, pady=5)  

    # Create a button to trigger plotting and lambda function to call draw_plot with selected options and pass Y-axis position value correctly  
    plot_button = ttk.Button(main_frame, text="绘图", command=lambda: draw_plot(x_label_vars.get(), x_value_vars.get(),y_label_vars.get(),y_value_vars.get(), z_column_var.get(), m_column_var.get(), plot_type_var.get(),my_colors))  
    plot_button.pack(pady=5) 
     

    root.lift()
    root.attributes("-topmost", True)
    root.attributes("-topmost", False)
      
    # 定义一个trace方法，当z_column_var的值发生变化时更新Label的内容  
    def update_risk_label(*args):  
        text=z_column_var.get()
        print(text)  
      
    # 设置trace方法，使得当z_column_var的值发生变化时调用update_risk_label函数  
    z_column_var.trace("w", update_risk_label) 
    root.mainloop()  # Start the main event loop to run the GUI application

############################################################################################################################
#通用程序和界面
############################################################################################################################
class AAA_11_PROGRAM():
    pass
######################################
#暂存合并
def PROGRAM__temp_save_df(df):
    print("************************************************************************************")
    print("功能说明：")
    print("1. 该工具用于临时保存DataFrame，并支持对已保存的DataFrame进行拼接和合并操作。")
    print("2. 支持以下功能：")
    print("   - 临时保存DataFrame，并为每个DataFrame命名。")
    print("   - 对已保存的DataFrame进行垂直拼接（追加行）。")
    print("   - 对已保存的DataFrame进行水平合并（追加列）。")
    print("   - 双击已保存的DataFrame名称可以查看其内容。")
    print("   - 右键已保存的表单，支持重命名和移除操作。")
    print("3. 垂直拼接要求所有DataFrame的列数相同。")
    print("4. 水平合并要求DataFrame具有相同的行索引或某些共同列。")
    print("\n使用方法：")
    print("1. 在输入框中输入文件名，点击“暂存”按钮保存当前DataFrame。")
    print("2. 在已保存的文件列表中选择一个或多个文件：")
    print("   - 点击“按行拼接”按钮，将选中的文件垂直拼接（追加行）。")
    print("   - 点击“按列拼接”按钮，将选中的文件水平合并（追加列）。")
    print("3. 双击已保存的文件名称，可以查看其内容。")
    print("4. 右键已保存的表单，可以重命名或移除。")
    print("5. 输入新文件名后，点击“暂存”按钮保存拼接或合并后的DataFrame。")
    print("************************************************************************************")

    def save_file():
        filename = entry.get().strip()
        if filename and filename not in global_dfs:
            global_dfs[filename] = df
            refresh_tree()
            entry.delete(0, tk.END)  # 清空输入框
        elif filename in global_dfs:
            messagebox.showwarning("警告", "该名称已存在，请另选名称。")
            mroot.lift()  # 将表单暂存器提到最前面

    def refresh_tree():
        tree.delete(*tree.get_children())
        for filename in global_dfs:
            tree.insert("", tk.END, text=filename)

    def check_filename(new_filename):
        """检查表单名是否填写或重复"""
        if not new_filename:
            messagebox.showwarning("警告", "表单名不能为空，请填写表单名。")
            mroot.lift()  # 将表单暂存器提到最前面
            return False
        if new_filename in global_dfs:
            messagebox.showwarning("警告", "该表单名已存在，请另选名称。")
            mroot.lift()  # 将表单暂存器提到最前面
            return False
        return True

    def concat_dfs_with_same_number_of_columns(dfs):
        """
        拼接具有相同数量列的多个DataFrame，使用第一个DataFrame的列名。

        参数:
            dfs (list of DataFrame): 要拼接的DataFrame列表。

        返回:
            DataFrame: 拼接后的DataFrame，使用第一个DataFrame的列名。
        """
        if not isinstance(dfs, list) or len(dfs) == 0:
            raise ValueError("输入必须是一个包含至少一个DataFrame的非空列表。")

        # 检查所有DataFrame的列数是否相同
        num_columns = dfs[0].shape[1]
        for df in dfs:
            if df.shape[1] != num_columns:
                raise ValueError("所有DataFrame必须具有相同数量的列。")

        # 使用第一个DataFrame的列名作为模板
        first_columns = dfs[0].columns

        # 将每个DataFrame的列名更改为第一个DataFrame的列名
        aligned_dfs = []
        for df in dfs:
            df.columns = first_columns
            aligned_dfs.append(df)

        # 拼接DataFrame列表，由于列名已经对齐，这里不需要忽略索引或列名
        concatenated = pd.concat(aligned_dfs)

        return concatenated

    def concat_files():
        # 检查表单名是否填写或重复
        if not check_filename(entry.get().strip()):
            return

        selected_items = tree.selection()
        if selected_items:  # 检查是否有选中的项目
            filenames = []
            new_dfs = []
            for item in selected_items:
                filename = tree.item(item, "text")
                filenames.append(filename)
                new_dfs.append(global_dfs[filename])

            new_df = concat_dfs_with_same_number_of_columns(new_dfs)  # 合并所有选中的DataFrame
            new_filename = entry.get().strip()

            if new_filename and new_filename not in global_dfs:
                global_dfs[new_filename] = new_df
                refresh_tree()
            else:
                messagebox.showwarning("警告", "该新名称已存在，请另选名称。")
                mroot.lift()  # 将表单暂存器提到最前面
        else:
            messagebox.showinfo("提示", "请选择要合并的行。")
            mroot.lift()  # 将表单暂存器提到最前面

    def merge_files():
        # 检查表单名是否填写或重复
        if not check_filename(entry.get().strip()):
            return

        selected_items = tree.selection()
        if len(selected_items) < 2:
            messagebox.showwarning("警告", "请选择至少两个文件进行合并。")
            mroot.lift()  # 将表单暂存器提到最前面
            return

        # 获取选中的文件名
        filenames = [tree.item(item, "text") for item in selected_items]
        selected_dfs = [global_dfs[filename] for filename in filenames]

        # 创建主窗口
        merge_root = tk.Tk()
        bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(merge_root)

        merge_root.title("追列合并")
        sw = merge_root.winfo_screenwidth()
        sh = merge_root.winfo_screenheight()
        ww = 600  # 窗口宽度
        wh = 230  # 窗口高度
        x = (sw - ww) // 2
        y = (sh - wh) // 2
        merge_root.geometry(f"{ww}x{wh}+{x}+{y}")

        # 选择主文件和追加文件
        label_main = tk.Label(merge_root, text="请选择主文件:")
        label_main.grid(row=0, column=0, padx=10, pady=10)
        main_file_var = tk.StringVar(merge_root)
        main_file_var.set(filenames[0])  # 默认选择第一个文件
        main_file_menu = ttk.Combobox(merge_root, textvariable=main_file_var, values=filenames)
        main_file_menu.grid(row=0, column=1, padx=10)

        label_append = tk.Label(merge_root, text="请选择追加文件:")
        label_append.grid(row=1, column=0, padx=10, pady=10)
        append_file_var = tk.StringVar(merge_root)
        append_file_var.set(filenames[1])  # 默认选择第二个文件
        append_file_menu = ttk.Combobox(merge_root, textvariable=append_file_var, values=filenames)
        append_file_menu.grid(row=1, column=1, padx=10)

        # 选择合并依据列
        label_left_on = tk.Label(merge_root, text="请选择主文件的合并依据列 (left_on):")
        label_left_on.grid(row=2, column=0, padx=10, pady=10)
        left_on_var = tk.StringVar(merge_root)
        left_on_menu = ttk.Combobox(merge_root, textvariable=left_on_var, values=selected_dfs[0].columns.tolist())
        left_on_menu.grid(row=2, column=1, padx=10)

        label_right_on = tk.Label(merge_root, text="请选择追加文件的合并依据列 (right_on):")
        label_right_on.grid(row=3, column=0, padx=10, pady=10)
        right_on_var = tk.StringVar(merge_root)
        right_on_menu = ttk.Combobox(merge_root, textvariable=right_on_var, values=selected_dfs[1].columns.tolist())
        right_on_menu.grid(row=3, column=1, padx=10)

        # 更新列选择菜单
        def update_columns(*args):
            main_file = main_file_var.get()
            append_file = append_file_var.get()
            main_df = global_dfs[main_file]
            append_df = global_dfs[append_file]
            left_on_menu['values'] = main_df.columns.tolist()
            right_on_menu['values'] = append_df.columns.tolist()

        main_file_var.trace_add('write', update_columns)
        append_file_var.trace_add('write', update_columns)

        # 确定按钮
        def on_ok():
            main_file = main_file_var.get()
            append_file = append_file_var.get()
            left_on = left_on_var.get()
            right_on = right_on_var.get()

            if not left_on or not right_on:
                messagebox.showerror("错误", "请选择合并依据列。")
                merge_root.lift()  # 将合并窗口提到最前面
                return


            main_df = global_dfs[main_file]
            append_df = global_dfs[append_file]
            
            main_df[left_on]=main_df[left_on].astype(str)
            #print(main_df[left_on].head(10))
            append_df[right_on]=append_df[right_on].astype(str)
            #print(append_df[right_on].head(10))
            
            # 检查列名是否存在于各自的DataFrame中
            if left_on not in main_df.columns or right_on not in append_df.columns:
                messagebox.showerror("错误", "选择的列名不存在。")
                merge_root.lift()  # 将合并窗口提到最前面
                return

            # 检查 left_on 和 right_on 是否有相同项目
            if not set(main_df[left_on]).intersection(set(append_df[right_on])):
                messagebox.showwarning("警告", f"{left_on} 和 {right_on} 列没有相同项目，不予合并。")
                merge_root.lift()  # 将合并窗口提到最前面
                return

            # 检查是否有重复列（除了 left_on 和 right_on 之外）
            common_columns = set(main_df.columns).intersection(set(append_df.columns))
            common_columns.discard(left_on)
            common_columns.discard(right_on)

            if common_columns:
                # 如果有重复列，弹出提示框
                msg = f"发现 {len(common_columns)} 个重复列:\n\n"
                msg += "\n".join(common_columns)
                msg += "\n\n这些列将被重命名为 '列名_x' 和 '列名_y'。\n是否继续合并？"

                if not messagebox.askyesno("重复列警告", msg):
                    return  # 用户取消合并

            try:
                # 合并DataFrame
                merged_df = pd.merge(main_df, append_df, left_on=left_on, right_on=right_on, how='left', suffixes=('_x', '_y'))
                # 保存合并后的DataFrame
                new_filename = entry.get().strip()
                if new_filename and new_filename not in global_dfs:
                    global_dfs[new_filename] = merged_df
                    refresh_tree()
                else:
                    messagebox.showwarning("警告", "该新名称已存在，请另选名称。")
                    mroot.lift()  # 将表单暂存器提到最前面
                # 关闭窗口
                merge_root.destroy()
            except Exception as e:
                messagebox.showerror("错误", f"合并失败，原因: {e}")
                merge_root.lift()  # 将合并窗口提到最前面

        button_ok = ttk.Button(merge_root, text="确定", command=on_ok)
        button_ok.grid(row=4, column=0, columnspan=2, pady=10)

        # 取消按钮
        def on_cancel():
            merge_root.destroy()

        button_cancel = ttk.Button(merge_root, text="取消", command=on_cancel)
        button_cancel.grid(row=4, column=1, columnspan=2, pady=10)

        # 运行主循环
        merge_root.mainloop()

    def on_double_click(event):
        selected_item = tree.selection()[0]
        filename = tree.item(selected_item, "text")
        # 假设您有一个函数来显示DataFrame (这里应该是实际逻辑，而不是print)
        # PROGRAM_display_df_in_treeview(global_dfs[filename], 1, 0)
        PROGRAM_display_df_in_treeview(global_dfs[filename], 0, 0)  # 此处仅为演示

    def show_context_menu(event):
        """显示右键菜单"""
        item = tree.identify_row(event.y)
        if item:
            tree.selection_set(item)
            menu.post(event.x_root, event.y_root)

    def rename_file():
        """重命名选中的表单"""
        selected_item = tree.selection()[0]
        old_filename = tree.item(selected_item, "text")
        new_filename = simpledialog.askstring("重命名", "请输入新的表单名:", parent=mroot)
        if new_filename and new_filename != old_filename:
            if new_filename in global_dfs:
                messagebox.showwarning("警告", "该表单名已存在，请另选名称。")
                mroot.lift()  # 将表单暂存器提到最前面
            else:
                global_dfs[new_filename] = global_dfs.pop(old_filename)
                refresh_tree()

    def remove_file():
        """移除选中的表单"""
        selected_item = tree.selection()[0]
        filename = tree.item(selected_item, "text")
        if messagebox.askyesno("确认", f"确定要移除表单 '{filename}' 吗？"):
            global_dfs.pop(filename)
            mroot.lift() 
            refresh_tree()

    # 创建GUI窗口
    mroot = tk.Tk()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(mroot)

    mroot.title("表单暂存和合并")
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）
    sw = mroot.winfo_screenwidth()
    sh = mroot.winfo_screenheight()
    ww = 500  # 窗口宽度
    wh = 400  # 窗口高度
    x = (sw - ww) // 2
    y = (sh - wh) // 2
    mroot.geometry(f"{ww}x{wh}+{x}+{y}")

    # 使用Frame来更好地组织控件
    top_frame = ttk.Frame(mroot, padding="10")
    top_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
    bottom_frame = ttk.Frame(mroot)
    bottom_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
    mroot.grid_rowconfigure(1, weight=1)  # 让TreeView占据多余空间
    mroot.grid_columnconfigure(0, weight=1)  # 让窗口在水平方向上可伸缩

    # 输入框用于填写文件名或新文件名（合并时）
    ttk.Label(top_frame, text="表单名/新表单名:").grid(row=0, column=0, sticky=tk.W, pady=5)
    entry = ttk.Entry(top_frame)
    entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=10)

    # concat按钮（垂直合并）
    concat_button = ttk.Button(top_frame, text="按行拼接", command=concat_files)
    concat_button.grid(row=1, column=4, columnspan=2, pady=5)
    # merge按钮（这里只是打印提示，实际应添加merge逻辑）
    merge_button = ttk.Button(top_frame, text="按列拼接", command=merge_files)
    merge_button.grid(row=1, column=2, columnspan=2, pady=5)

    # 保存按钮（暂存）
    save_button = ttk.Button(top_frame, text="暂存", command=save_file)
    save_button.grid(row=1, column=0, columnspan=2, pady=5)

    # TreeView控件显示已保存的文件列表
    ttk.Label(bottom_frame, text="已保存的表单:").grid(row=0, column=0, sticky=tk.W, pady=5)
    tree = ttk.Treeview(bottom_frame)
    tree_scrollbar = ttk.Scrollbar(bottom_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=tree_scrollbar.set)
    tree.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
    tree_scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))
    bottom_frame.grid_rowconfigure(1, weight=1)  # 让TreeView在Frame内垂直方向上可伸缩
    bottom_frame.grid_columnconfigure(0, weight=1)  # 让TreeView在Frame内水平方向上可伸缩
    tree.bind("<Double-1>", on_double_click)  # 绑定双击事件以显示DataFrame内容（需实现显示逻辑）
    tree.bind("<<TreeviewSelect>>", lambda event: entry.delete(0, tk.END))  # 清除输入框当选择树中的项时（为新名称做准备）

    # 右键菜单
    menu = tk.Menu(mroot, tearoff=0)
    menu.add_command(label="重命名", command=rename_file)
    menu.add_command(label="移除", command=remove_file)
    tree.bind("<Button-3>", show_context_menu)  # 绑定右键事件

    refresh_tree()  # 初始化TreeView控件
    mroot.mainloop()

#表格追加合并
def PROGRAM_merge_dataframes(df1):  
    # 创建主窗口  
    print(
        """
    **********************************************
    追加合并工具

    功能：
    1. 提供 GUI 界面供用户选择主文件和需要追加的文件。
    2. 支持选择合并依据列（left_on 和 right_on）。
    3. 执行左连接合并操作，并显示合并后的结果。

    使用方法：
    1. 启动程序后，选择需要追加的 Excel 文件。
    2. 选择主文件的合并依据列（left_on）和需要追加文件的合并依据列（right_on）。
    3. 点击“确定”按钮执行合并操作。
    4. 合并结果会显示在 GUI 界面中。

    注意：
    - 如果未选择文件或列，程序会提示错误。
    - 合并结果会显示在 GUI 界面中。
    **********************************************
    """
    )
    
    # 确定按钮  
    def on_ok(df1, df2, left_on, right_on):  
        if not left_on or not right_on:  
            messagebox.showerror("错误", '没有正确选择列。')  
            return  
  
        # 检查列名是否存在于各自的DataFrame中  
        if left_on not in df1.columns or right_on not in df2.columns:  
            messagebox.showerror("错误", "没有正确选择列。")  
            return  

        df1[left_on]=df1[left_on].astype(str)
        df2[right_on]=df2[right_on].astype(str)

        # 检查 A1 和 B1 是否有相同项目
        if not set(df1[left_on]).intersection(set(df2[right_on])):
            messagebox.showwarning("警告", f"{left_on} 和 {right_on} 列没有相同项目，不予合并。")
            return
  
        # 检查是否有重复列（除了 left_on 和 right_on 之外）
        common_columns = set(df1.columns).intersection(set(df2.columns))
        common_columns.discard(left_on)
        common_columns.discard(right_on)
        
        if common_columns:
            # 如果有重复列，弹出提示框
            msg = f"发现 {len(common_columns)} 个重复列:\n\n"
            msg += "\n".join(common_columns)
            msg += "\n\n这些列将被重命名为 '列名_x' 和 '列名_y'。\n是否继续合并？"
            
            if not messagebox.askyesno("重复列警告", msg):
                return  # 用户取消合并
        
        try:  
            # 合并DataFrame  
            merged_df = pd.merge(df1, df2, left_on=left_on, right_on=right_on, how='left', suffixes=('_x', '_y'))  
            # 打印合并后的DataFrame（您可以替换为其他功能）  
            PROGRAM_display_df_in_treeview(merged_df, 0, 0)  
            # 关闭窗口  
            bsroot.destroy()  
        except Exception as e:  
            messagebox.showerror("错误", f"合并失败，原因: {e}")  
    
    bsroot = tk.Tk()  
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(bsroot)

    bsroot.title("追加合并")    
    sw = bsroot.winfo_screenwidth()          
    sh = bsroot.winfo_screenheight()          
    ww = 600  # 窗口宽度          
    wh = 150  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    bsroot.geometry(f"{ww}x{wh}+{x}+{y}")   
    # 隐藏主窗口直到文件对话框弹出  
    bsroot.withdraw()  
  
    # 弹出文件选择对话框，选择xlsx或xls文件  
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])  
    if not file_path:  
        messagebox.showerror("错误", "未提供文件名!")  
        return df1  
  
    # 读取选择的文件为DataFrame  
    try:  
        df2 = pd.read_excel(file_path)  
    except Exception as e:  
        messagebox.showerror("错误", f"打开文件失败: {e}")  
        return df1  
  
    # 显示主窗口  
    bsroot.deiconify()  
  
    # 获取df1和df2的列名  
    columns_df1 = df1.columns.tolist()  
    columns_df2 = df2.columns.tolist()  
  
    # 变量来存储选中的列名  
    selected_col_df1 = tk.StringVar(bsroot)  
    selected_col_df2 = tk.StringVar(bsroot)  
  
    # 创建并布局界面元素  
    label_df1 = tk.Label(bsroot, text="请选择主文件的合并依据列 (left_on):")  
    label_df1.grid(row=0, column=0, padx=10, pady=10)  
  
    combo_df1 = ttk.Combobox(bsroot, textvariable=selected_col_df1, values=columns_df1)  
    combo_df1.grid(row=0, column=1, padx=10)  
  
    label_df2 = tk.Label(bsroot, text="请选择需要追加文件的合并依据列 (right_on):")  
    label_df2.grid(row=1, column=0, padx=10, pady=10)  
  
    combo_df2 = ttk.Combobox(bsroot, textvariable=selected_col_df2, values=columns_df2)  
    combo_df2.grid(row=1, column=1, padx=10)  
  
    button_ok = ttk.Button(bsroot, text="确定", command=lambda: on_ok(df1, df2.drop_duplicates(selected_col_df2.get()), selected_col_df1.get(), selected_col_df2.get()))  
    button_ok.grid(row=2, column=0, columnspan=2, pady=20)  
  
    # 取消按钮  
    def on_cancel():  
        bsroot.destroy()  
  
    button_cancel = ttk.Button(bsroot, text="取消", command=on_cancel)  
    button_cancel.grid(row=2, column=1, columnspan=2, pady=10)  
  
    # 运行主循环  
    bsroot.mainloop()  


def PROGRAM_update_progress_bar(win,win_progressbar,now_schedule, all_schedule):  
    """更新进度条"""  
    win_progressbar['value']  = min(now_schedule / all_schedule * 100, 100)  # 限制在0-100%范围内  
    win.update()  # 更新界面  
    
def PROGRAM_thread_it(func, *args):
    """将函数打包进线程"""
    # 创建
    t = threading.Thread(target=func, args=args)
    # 守护 !!!
    t.setDaemon(True)
    # 启动
    t.start() 
       
def PROGRAM_show_report_window(df, report_code_column):
    """报告查看"""
    # 创建一个Toplevel窗口
    root = tk.Toplevel()
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)
    root.title("报告查看器")
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 1024 # 窗口宽度    
    wh = 768  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")   
    
    # 创建一个Frame来包含左右两列
    frame = tk.Frame(root)
    frame.pack(fill=tk.BOTH, expand=True)
    
    # 确保报告编码列是字符串类型
    df[report_code_column] = df[report_code_column].astype(str)
    
    # 左边：报告编码列表框
    left_frame = tk.Frame(frame)
    left_frame.pack(side=tk.LEFT, fill=tk.Y, expand=False)
    
    scrollbar_left = tk.Scrollbar(left_frame, orient=tk.VERTICAL)
    scrollbar_left.pack(side=tk.RIGHT, fill=tk.Y)
    
    report_code_listbox = tk.Listbox(left_frame, yscrollcommand=scrollbar_left.set)
    report_code_listbox.pack(fill=tk.BOTH, expand=True)
    scrollbar_left.config(command=report_code_listbox.yview)
    
    report_codes = df[report_code_column].unique()
    for code in report_codes:
        report_code_listbox.insert(tk.END, code)
    
    # 右边：报告内容显示区域
    right_frame = tk.Frame(frame)
    right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
    
    text_area = scrolledtext.ScrolledText(right_frame, wrap=tk.WORD, width=70, height=20)
    text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    def prevent_key_input(event):
        return "break"  # 阻止默认行为
    text_area.bind("<Key>", prevent_key_input)
    
    # 定义右键菜单
    def show_context_menu(event):
        if text_area.tag_ranges("sel"):
            context_menu.post(event.x_root, event.y_root)
    
    context_menu = tk.Menu(root, tearoff=0)
    context_menu.add_command(label="复制", command=lambda: text_area.event_generate("<<Copy>>"))
    text_area.bind("<Button-3>", show_context_menu)
    
    # 绑定列表框选择事件来更新文本区域
    def update_text_area(event=None):
        selected_index = report_code_listbox.curselection()
        if selected_index:
            selected_code = report_code_listbox.get(selected_index[0])
            selected_rows = df[df[report_code_column] == selected_code]
            text_area.delete(1.0, tk.END)  # 清空文本区域
            if not selected_rows.empty:
                for _, row in selected_rows.iterrows():
                    for column, value in row.items():
                        text_area.insert(tk.END, f"{column}: {value}\n")
                    text_area.insert(tk.END, "\n")  # 每个报告条目后添加空行
                    # 下面的分隔符可以视需要添加或删除
                    # text_area.insert(tk.END, "########################################################\n")
            else:
                text_area.insert(tk.END, "未找到匹配的报告编码。\n")
        else:
            # 默认显示第一条数据
            if not report_codes.empty:
                update_text_area(None, report_codes[0])  # 这里需要修改update_text_area以接受可选的默认代码
    
    # 注意：我们需要稍微修改这个函数来接受一个可选的默认代码参数
    def updated_update_text_area(event=None, default_code=None):
        if default_code:
            selected_code = default_code
            selected_rows = df[df[report_code_column] == selected_code]
            if not selected_rows.empty:
                report_code_listbox.select_set(0)  # 选中列表中的第一项
                update_text_area_internal(selected_rows)
            else:
                text_area.insert(tk.END, "未找到匹配的报告编码。\n")
        else:
            selected_index = report_code_listbox.curselection()
            if selected_index:
                selected_code = report_code_listbox.get(selected_index[0])
                selected_rows = df[df[report_code_column] == selected_code]
                update_text_area_internal(selected_rows)
    
    def update_text_area_internal(selected_rows):
        text_area.delete(1.0, tk.END)  # 清空文本区域
        for _, row in selected_rows.iterrows():
            for column, value in row.items():
                text_area.insert(tk.END, f"{column}: {value}\n")
            text_area.insert(tk.END, "\n")  # 每个报告条目后添加空行
            text_area.insert(tk.END, "###############################################################")  # 每个报告条目后添加空行
            text_area.insert(tk.END, "\n")  # 每个报告条目后添加空行    
    # 绑定列表框选择事件
    report_code_listbox.bind("<<ListboxSelect>>", updated_update_text_area)
    
    # 默认显示第一条数据
    updated_update_text_area(default_code=report_codes[0])
    
    # 允许文本区域复制
    def copy_to_clipboard(event=None):
        try:
            selected_text = text_area.selection_get()
            root.clipboard_clear()
            root.clipboard_append(selected_text)
        except tk.TclError:
            pass
    
    text_area.bind("<<Copy>>", copy_to_clipboard)
    
    root.mainloop()
    
def PROGRAM_display_content_in_textbox(content):
    """
    在弹出窗口中显示给定内容，并增加导出为TXT的功能。
    
    参数:
        content (any): 要显示的内容，可以是字典、整数、浮点数或字符串。
    """
    def create_popup_menu(event):
        """
        为文本框创建右键菜单，并添加复制和导出为TXT功能。
        
        参数:
            event (tkinter.Event): 触发此函数的事件对象。
        """
        popup_menu = tk.Menu(text_widget, tearoff=0)  # 创建右键菜单
        popup_menu.add_command(label="复制", command=lambda: text_widget.event_generate("<<Copy>>"))  # 添加复制功能
        popup_menu.add_command(label="导出为TXT", command=export_to_txt)  # 添加导出为TXT功能
        popup_menu.tk_popup(event.x_root, event.y_root)  # 在指定位置显示菜单
 
    def export_to_txt():
        # 获取文本框内容
        text_content = text_widget.get("1.0", tk.END)
        
        # 保存TXT文件
        txt_file_path = filedialog.asksaveasfilename(
            title="保存文件",
            initialfile="output",
            defaultextension=".txt",
            filetypes=[("Text file", "*.txt")],
        )
        
        if txt_file_path:  # 确保用户选择了文件路径
            with open(txt_file_path, "w", encoding="utf-8") as f:
                f.write(text_content)
        
        messagebox.showinfo("成功", f"内容已导出到 {txt_file_path}")
 
    root = tk.Tk()  # 创建主窗口
    
    root.withdraw()  # 隐藏主窗口
 
    popup = tk.Toplevel(root)  # 创建弹出窗口
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(popup)
    popup.title("查看内容")  # 设置弹出窗口的标题
 
    # 设置弹出窗口大小和位置
    sw_treeQ = popup.winfo_screenwidth()
    sh_treeQ = popup.winfo_screenheight()
    ww_treeQ = 800
    wh_treeQ = 600
    x_treeQ = (sw_treeQ - ww_treeQ) // 2
    y_treeQ = (sh_treeQ - wh_treeQ) // 2
    popup.geometry(f"{ww_treeQ}x{wh_treeQ}+{x_treeQ}+{y_treeQ}")
 
    text_widget = ScrolledText(popup, height=40, width=80, bg="#FFFFFF", wrap=tk.WORD)
    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)  # 设置文本框的布局和填充方式

    # 根据内容类型显示内容
    if isinstance(content, dict):  # 如果内容是字典类型
        for key, value in content.items():  # 遍历字典并显示其内容
            text_widget.insert(tk.END, f"{key}：{value}\n")  # 在文本框末尾插入内容
    elif isinstance(content, (int, float, str)):  # 如果内容是整数、浮点数或字符串类型
        text_widget.insert(tk.END, str(content))  # 将内容转换为字符串并显示在文本框末尾
    else:  # 如果内容是其他类型
        messagebox.showerror("错误", "不支持的内容类型")  # 显示错误消息框
        popup.destroy()  # 销毁弹出窗口
        return  # 结束函数执行
 
    text_widget.config(state=tk.DISABLED)  # 设置文本框为不可编辑状态
    text_widget.bind("<Button-3>", create_popup_menu)  # 绑定鼠标右键单击事件到创建右键菜单的函数
    popup.mainloop()


              
def PROGRAM_display_df_in_treeview(ori_owercount_easyread, methon, ori, page_size=100):    
    """    
    在Treeview中分页展示DataFrame的内容，并添加点击列标题进行排序的功能。       
    参数:    
        ori_owercount_easyread: 工作文件，一个Pandas DataFrame。    
        ori:源文件。
        page_size (int): 每页显示的行数，默认为10000。    
    """ 
    def PSRU_ori_owercount_easyread(df):
        if '报表类型' in df.columns:
            m=df.loc[0,'报表类型']
            df=SMALL_expand_dict_like_columns(df)
            df['报表类型']=m
            new_columns = [col.replace('报告表编码', '') if '报告表编码' in col else col for col in df.columns]  
            df.columns = new_columns
            for col in df.columns:  
                try:  
                    df[col] = df[col].fillna(0).astype(int)  
                except ValueError:  
                    pass
        return df
    if ori_owercount_easyread is None:
        print("传入的数据为空!")
        return
    
    #兼容PSUR
    if methon=="psur":
        ori_owercount_easyread=PSRU_ori_owercount_easyread(ori_owercount_easyread)


    
    if isinstance(ori, pd.DataFrame):  
        pass  
    else:  
        ori=ori_owercount_easyread
     
    
      
    def sort_column(col, tree):   
        nonlocal current_page  # 声明current_page为非局部变量 
        nonlocal current_sort_order
        nonlocal ori_owercount_easyread           
        current_page = 0  
        ori_owercount_easyread = ori_owercount_easyread.sort_values(by=col, ascending=current_sort_order)    
        update_treeview(tree, ori_owercount_easyread)    
        # 切换排序顺序  
        current_sort_order = not current_sort_order   
    def update_treeview(tree, df):    
        tree.delete(*tree.get_children())  # 清除Treeview中的所有数据    
        start = current_page * page_size    
        end = start + page_size    
        for index, row in df[start:end].iterrows():    
            tree.insert("", "end", text=str(index), values=list(row))    
      
    def go_to_page(BF,page):    
        nonlocal current_page  # 声明current_page为非局部变量 
        
        total_pages = len(ori_owercount_easyread) // page_size + (len(ori_owercount_easyread) % page_size > 0) 
        current_page = page  
        update_treeview(tree, ori_owercount_easyread)    
        update_textbox(entry_text, str(len(ori_owercount_easyread))+"，"+str(current_page+1)+"/"+str(total_pages))  
        print("当前页：",current_page+1,"/",total_pages)

    def copy_to_clipboard():    
        try:    
            selected_item = tree.selection()[0]  # 获取选定的行    
            selected_values = tree.item(selected_item, "values")  # 获取选定的行的值    
            selected_dict = dict(zip(ori_owercount_easyread.columns, selected_values))  # 转换为字典格式    
            root.clipboard_clear()  # 清除Tk的剪贴板内容    
            root.clipboard_append(str(selected_dict))  # 将字典字符串复制到Tk的剪贴板中    
        except IndexError:  # 如果没有选定任何行，则不执行复制操作    
            pass    

    def view_to_clipboard():    
        try:    
            selected_item = tree.selection()[0]  # 获取选定的行    
            selected_values = tree.item(selected_item, "values")  # 获取选定的行的值    
            selected_dict = dict(zip(ori_owercount_easyread.columns, selected_values))  # 转换为字典格式    
            PROGRAM_display_content_in_textbox(selected_dict) 
        except IndexError:  # 如果没有选定任何行，则不执行复制操作    
            pass  
    
    def show_popup_menu(event):    
        try:    
            tree.selection_set(tree.identify_row(event.y))  # 选定点击的行    
            popup_menu.tk_popup(event.x_root, event.y_root)  # 弹出右键菜单    
        except AttributeError:  # 如果没有点击在行上，则不弹出菜单    
            pass    
    
    def handle_double_click(event):    
        df = ori_owercount_easyread  # 在这里明确df的来源是ori_owercount_easyread  
        if "报表类型" in df.columns:  # 如果双击的是“报表类型”列，显示警告框    
            for item in tree.selection():    
                selection = tree.item(item, "values")    
                content_dict = dict(zip(df.columns, selection))    
                SETTING_ReportType(content_dict.copy(),methon,ori) 
        else:  # 如果双击的不是“报表类型”列，弹出名为VIEWDICT的文本窗口展示“xxxx”    
            for item in tree.selection():    
                selection = tree.item(item, "values")    
                content_dict = dict(zip(df.columns, selection))    
                PROGRAM_display_content_in_textbox(content_dict)  # 同上，确保这个函数在其他地方已经定义  

    def PROGRAM_create_button(frame, text, command):  
        """创建一个按钮并添加到frame"""  
        ttk.Button(  
            frame,  
            text=text,  
            command=command,  
        ).pack(side=tk.LEFT) 
        
           
    def update_textbox(textbox, text_to_display):  
        textbox.config(text=text_to_display)  # 更新标签的文本内容   
           
    # 初始化Tk窗口    
    root = tk.Tk()   
     
    root.title(mytitle)   
  
    
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）    
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 1310  # 窗口宽度    
    wh = 600  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    bg_color,frame_color,button_color,text_color,highlight_color=SETTING_style_UI(root)
    # 创建Treeview并设置列    
    tree = ttk.Treeview(root)    
    tree["columns"] = list(ori_owercount_easyread.columns)    
    current_sort_column = None    
    current_sort_order = True  # True表示升序，False表示降序    
    current_page = 0  # 当前页数，从0开始计数。这里初始化了current_page变量。  
    total_pages = len(ori_owercount_easyread) // page_size + (len(ori_owercount_easyread) % page_size > 0)  # 总页数    



    
    # 配置列宽度和标题，并绑定点击事件和双击事件    
    for i, col in enumerate(ori_owercount_easyread.columns):    
        tree.column(col, width=200, stretch=tk.NO)    
        tree.heading(col, text=col, command=lambda c=col: sort_column(c, tree))    
    tree.bind("<Double-1>", handle_double_click)  # 绑定双击事件    


    # 个性化设置列宽 
     
    column_widths=SETTING_get_width()
    for column, width in column_widths.items():  
        try:  
            tree.column(column, minwidth=0, width=width, stretch=tk.NO)  
        except:   
            pass 

    
    update_treeview(tree, ori_owercount_easyread)  # 初次加载DataFrame数据到Treeview中（第一页）    
    

    # 创建分页控件（使用ttk.Label和ttk.Button）    
    frame0 = ttk.Frame(root)    
    ttk.Button(frame0, text="<<", command=lambda: go_to_page("B",max(0, current_page - 1))).pack(side="left") 
    entry_text = tk.Label(frame0, text=str(len(ori_owercount_easyread))+"，"+str(current_page+1)+"/"+str(total_pages),width=20)  
    entry_text.pack(side=tk.LEFT)  # 确保文本框可见    
    ttk.Button(frame0, text=">>", command=lambda: go_to_page("F",min(total_pages - 1, current_page + 1))).pack(side="left")        
    frame0.pack(side="bottom", fill="x")  # 将分页控件放置在底部  
 
    framecanvas0 = ttk.Frame(root, width=1310, height=20)
    framecanvas0.pack(side="bottom")
    framecanvas = ttk.Frame(root, width=1310, height=20)
    framecanvas.pack(side="bottom")            
    # 右键菜单的实现      
    popup_menu = tk.Menu(root, tearoff=0)  # 创建弹出菜单实例      
    popup_menu.add_command(label="复制", command=copy_to_clipboard)  # 添加复制选项到菜单中      
    popup_menu.add_command(label="查看", command=view_to_clipboard)  # 添加复制选项到菜单中      
    tree.bind("<Button-3>", show_popup_menu)  # 绑定右键点击事件到弹出菜单函数上    
    tree.bind("<Double-1>", handle_double_click)  # 绑定双击事件     
    # 创建垂直滚动条  
    vertical_scrollbar = ttk.Scrollbar(root, orient="vertical", command=tree.yview)  
    tree.configure(yscrollcommand=vertical_scrollbar.set)  
  
    # 创建水平滚动条  
    horizontal_scrollbar = ttk.Scrollbar(root, orient="horizontal", command=tree.xview)  
    tree.configure(xscrollcommand=horizontal_scrollbar.set)    
          
    vertical_scrollbar.pack(side="right", fill="y")  # 垂直滚动条放置在右侧  
    horizontal_scrollbar.pack(side="bottom", fill="x")  # 水平滚动条放置在底部  
    
    #创建进度条
    win_progress_frame = ttk.Frame(frame0)  
    win_progress_frame.pack(side="left",pady=1)  
    win_progressbar = ttk.Progressbar(win_progress_frame, orient='horizontal', mode='determinate', length=100)  
    #win_progressbar.pack(fill=tk.X, expand=True, pady=5) 
    #创建选择组件
    xt11 = tk.StringVar()  
    xt11.set("列名")  
     
    import_se1 = tk.Label(frame0, text="位置：")  
    import_se1.pack(side=tk.LEFT)      
    comvalue = tk.StringVar()  # 窗体自带的文本，新建一个值  
    comboxlist = ttk.Combobox(  
        frame0, width=12, height=30, state="readonly", textvariable=comvalue  
    )   
    comboxlist["values"] = ori_owercount_easyread.columns.tolist()  
    #comboxlist.current(0)  # 选择第一个  
    comboxlist.bind("<<ComboboxSelected>>", lambda *arg: xt11.set(comboxlist.get()))  # 绑定事件,(下拉列表框被选中时，绑定XT11SET函数)  
    comboxlist.pack(side=tk.LEFT)    
    import_se3 = tk.Label(frame0, text="文本：")  
    import_se3.pack(side=tk.LEFT)  
    
      # 使用 Text widget 替换 Entry widget  
    xentry_t22 = tk.Text(frame0, width=20, height=1)  # 设置宽度和高度以适应多行文本  
    xentry_t22.pack(side=tk.LEFT)  
      # 插入原始内容到 Text widget  
    xentry_t22.insert(tk.END,"关键词1|关键词2")  
      
    import_se4 = tk.Label(frame0, text="操作：") 
    import_se4.pack(side=tk.LEFT) 
    #创建通用按钮  
    #PROGRAM_create_button(framecanvas, "测试专用", lambda:PROGRAM_display_df_in_treeview(TOOLS_ROR_STAT_0(ori_owercount_easyread),0,0)) 
      

            
                  
    PROGRAM_create_button(frame0, "含", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str).str.contains(str(xentry_t22.get("1.0", tk.END).strip()), na=False)], methon, ori_owercount_easyread)) 
    PROGRAM_create_button(frame0, "无", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[~ori_owercount_easyread[xt11.get()].astype(str).str.contains(str(xentry_t22.get("1.0", tk.END).strip()), na=False)], methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "是", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str)==str(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "大", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(float)>float(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "小", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(float)<float(xentry_t22.get("1.0", tk.END).strip())],methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "等", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str)==str(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "升", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.sort_values(by=(xt11.get()),ascending=[True],na_position="last") , methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "降", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.sort_values(by=(xt11.get()),ascending=[False],na_position="last") , methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "扩", lambda: PROGRAM_display_df_in_treeview(CLEAN_expand_rows(ori_owercount_easyread.copy(),str(xentry_t22.get("1.0", tk.END).strip()), [xt11.get()]),0,0)) 
    PROGRAM_create_button(frame0, "浏", lambda: PROGRAM_show_report_window(ori_owercount_easyread.copy(), xt11.get())) 
    PROGRAM_create_button(frame0, "源", lambda: PROGRAM_display_df_in_treeview(ori,0,ori)) 
    PROGRAM_create_button(frame0, "列", lambda: PROGRAM_display_df_in_treeview(SMALL_expand_dict_column(ori_owercount_easyread.copy(),xt11.get()),0,0)) 




    #增加PSUR兼容性
    if psur==0:

        PROGRAM_create_button(frame0, "找", lambda: PROGRAM_display_df_in_treeview(SMALL_find_based_on_expression(ori_owercount_easyread,str(xentry_t22.get("1.0", tk.END).strip())), methon, ori_owercount_easyread))
        PROGRAM_create_button(frame0, "透", lambda: PROGRAM_display_df_in_treeview(SMALL_add_count_and_ratio(ori_owercount_easyread,xt11.get()),1,ori_owercount_easyread) )
        PROGRAM_create_button(frame0, "拆", lambda: PROGRAM_display_df_in_treeview(SMALL_add_count_and_ratio_exp(ori_owercount_easyread,str(xt11.get()),xentry_t22.get("1.0", tk.END).strip()),methon,ori_owercount_easyread) )
        PROGRAM_create_button(frame0, "存", lambda: PROGRAM__temp_save_df(ori_owercount_easyread.copy()) )

        PROGRAM_create_button(framecanvas0, "智能导入", lambda: PSUR_read_files_and_clean())
        PROGRAM_create_button(framecanvas0, "数据清洗", lambda: CLEAN_table(ori_owercount_easyread.copy()))
        PROGRAM_create_button(framecanvas0, "分组透视", lambda:TOOLS_create_pivot_tool_gui(ori_owercount_easyread.copy(),ori_owercount_easyread))       
        PROGRAM_create_button(framecanvas0, "暂存合并", lambda: PROGRAM__temp_save_df(ori_owercount_easyread.copy()) )
        PROGRAM_create_button(framecanvas0, "图形绘制", lambda:DRAW_plot_df(ori_owercount_easyread))  
    if '报告编码' in ori_owercount_easyread.columns:
        PROGRAM_create_button(framecanvas0, "可视面板", lambda:DRAW_show_analysis_gui(ori_owercount_easyread).run())       
 
        #创建菜单栏
    SETTING_create_menu({"windows":root,"win_progressbar":win_progressbar,"ori_owercount_easyread":ori_owercount_easyread,"ori":ori})  
    
    #创建PSUR专属按钮
    if  '是否已清洗' in ori_owercount_easyread.columns    and '报告表编码' in ori_owercount_easyread.columns    and '---药品数据规整---' in ori_owercount_easyread.columns:
        PROGRAM_create_button(framecanvas, "导入说明书", lambda:PSUR_check_adr_in_word(ori_owercount_easyread))       
        PROGRAM_create_button(framecanvas, "报告类型", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-伤害'], ['怀疑/并用'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "不良反应（SOC）",  lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['SOC'], ['-伤害'], ['报告表编码'], ['count'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'SOC', '不良反应名称（规整）', '不良反应发生时间'))
        PROGRAM_create_button(framecanvas, "不良反应（PT）", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['SOC', '不良反应名称（规整）'], ['-伤害'], ['报告表编码'], ['count'], {'疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'通用名称', '不良反应名称（规整）', '不良反应发生时间'))

        PROGRAM_create_button(framecanvas, "年龄段", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['年龄段'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "肝肾孕哺", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['特殊人群'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "性别", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['性别'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "时隔", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['不良反应发生时间减用药开始时间'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "报告年份", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['报告年份'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "用法用量", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['用法用量'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "批号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['生产批号'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'批准文号', '生产批号', '不良反应发生时间'))
        PROGRAM_create_button(framecanvas, "批准文号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','批准文号','通用名称'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori) )
        PROGRAM_create_button(framecanvas, "通用名称", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','通用名称'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori) )
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'通用名称', 'SOC', '不良反应发生时间'))
        PROGRAM_create_button(framecanvas, "ROR", lambda:TOOLS_ROR_from_df_with_gui(ori_owercount_easyread.copy()) )
        PROGRAM_create_button(framecanvas, "新ADR检测", lambda:PROGRAM_display_df_in_treeview(PSUR_get_new_GUI(ori_owercount_easyread),0,0) )
        PROGRAM_create_button(framecanvas, "活性成分", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','最主要的一个活性成分'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori) )
        PROGRAM_create_button(framecanvas, "药品分类", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.fillna('-未填写-'),[['药品分类'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori))
        
        PROGRAM_create_button(framecanvas, "报告单位", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','医院名称'], ['-伤害'], ['报告表编码'], ['nunique'], {},  ['报告表编码合计']]),'psur',ori) )
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'-监测机构', '医院名称', '报告日期'))
        PROGRAM_create_button(frame0, "PT计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], ['-伤害'], ['报告表编码'], ['count'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(frame0, "编码唯一值计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       

    #创建PSUR专属按钮
    if  '是否已清洗' in ori_owercount_easyread.columns    and  '报告编码' in ori_owercount_easyread.columns  and '器械故障表现' in ori_owercount_easyread.columns   and '-持有人-品名-证号列' in ori_owercount_easyread.columns:
        PROGRAM_create_button(framecanvas, "严重程度", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['伤害'], [], ['报告编码'], ['nunique'], {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "故障", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['器械故障表现'], [], ['报告编码'], ['nunique'], {'伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),'psur',ori))       
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'-注册证备案证', '器械故障表现', '事件发生日期'))
     

        PROGRAM_create_button(framecanvas, "伤害", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['伤害表现'], [], ['报告编码'], ['nunique'], {'器械故障表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),'psur',ori))       
        PROGRAM_create_button(framecanvas, "敏感词", lambda:PROGRAM_display_df_in_treeview(SMALL_add_sep(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；|;',['敏感词']),[['敏感词'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True)),'psur',ori))       
        PROGRAM_create_button(framecanvas, "通用故障", lambda:PROGRAM_display_df_in_treeview(SMALL_add_sep(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；|;',['通用故障表现']),[['通用故障表现'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True)),'psur',ori))       
        PROGRAM_create_button(framecanvas, "产品名称", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['产品类别', '产品名称'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "证号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "批号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-批号'], [], ['报告编码'], ['nunique'], {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'-注册证备案证', '产品批号', '事件发生日期'))
        PROGRAM_create_button(framecanvas, "型号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-型号'], [], ['报告编码'], ['nunique'], {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "规格", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-规格'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 

        if  '二级类别' in ori_owercount_easyread.columns:
            PROGRAM_create_button(framecanvas, "大类", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['产品大类'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '产品名称': 'SMALL_count_mode'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
            PROGRAM_create_button(framecanvas, "中类", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['一级类别'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '产品名称': 'SMALL_count_mode'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
            PROGRAM_create_button(framecanvas, "小类", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['二级类别'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '产品名称': 'SMALL_count_mode'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 


        PROGRAM_create_button(framecanvas, "报告单位", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['单位名称'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "图", lambda:DRAW_show_analysis(ori_owercount_easyread,'-监测机构', '单位名称', '报告日期'))
        PROGRAM_create_button(framecanvas, "注册人", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告编码', ascending=False).reset_index(drop=True),0,ori)) 
        PROGRAM_create_button(framecanvas, "报告月份", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['报告月份'], [], ['报告编码'], ['nunique'],  {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]).sort_values(by='报告月份', ascending=False).reset_index(drop=True),0,ori))
        PROGRAM_create_button(framecanvas, "故障标记", lambda:PSUR_get_guize2(ori_owercount_easyread)) 
        PROGRAM_create_button(framecanvas, "同类比较", lambda:TOOLS_ROR_from_df_with_gui(ori_owercount_easyread.copy()) )
        PROGRAM_create_button(framecanvas, "趋势分析", lambda:TOOLS_trend_analysis_GUI(ori_owercount_easyread.copy()))   
        PROGRAM_create_button(framecanvas, "导出", lambda:SMALL_save_dict(ori_owercount_easyread)) 

        #一级敏感词
        #二级敏感词
        #故障表现（规整）
        #故障表现
        #伤害表现
        #型号、规格
        PROGRAM_create_button(frame0, "PT计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], [], ['报告编码'], ['count'], {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]),'psur',ori))       
        PROGRAM_create_button(frame0, "编码唯一值计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], [], ['报告编码'], ['nunique'], {'器械故障表现': 'SMALL_count_mode', '伤害表现': 'SMALL_count_mode', '敏感词': 'count_ALL','通用故障表现': 'count_ALL'}, []]),'psur',ori))       

           
    # 在Tk窗口中放置Treeview并运行主循环      
    tree.pack(side="top", fill="both", expand=True) # Treeview填充窗口的剩余空间      

    root.lift()
    #root.attributes("-topmost", True)
    #root.attributes("-topmost", False)  
    root.mainloop() # 启动Tk窗口的主事件循环```python
        


############################################################################################################################
#入口函数
############################################################################################################################

class AAA_00_main():

    global version_now
    global usergroup
    global setting_cfg
    global csdir
    global peizhidir
    global biaozhun
    global global_dfs
    global psur
    global mytitle 
    global db_ver_i
    global memory_db



    password = "my_secret_password"
    try:
        memory_db = SQL_decrypt_db_and_print_tables(password)
        db_ver = SQL_gettable(memory_db, 'version')['version'][0]
        print('资源库版本：',db_ver)
        
        if float(db_ver)<float(db_ver_i):
            print('资源库版本不是最新的，正在更新...')
            script_dir = os.path.dirname(os.path.abspath(__file__))
            # 获取程序目录下的 sor.zip（实际是 zip 文件）
            def_zip_path=sor_path
            #def_zip_path = os.path.join(script_dir, "sor.zip")
            # 获取程序目录下的 sor.zip（实际是 zip 文件）
            if os.path.exists(def_zip_path):
                try:
                    # 解压 sor.zip 到资源库目录，处理中文文件名
                    with zipfile.ZipFile(def_zip_path, 'r') as zip_ref:
                        for file_info in zip_ref.infolist():
                            # 处理中文文件名
                            file_name = file_info.filename.encode('cp437').decode('gbk')  # 从 cp437 转换为 gbk
                            target_path = os.path.join(peizhidir, file_name)
                            # 确保目标目录存在
                            os.makedirs(os.path.dirname(target_path), exist_ok=True)
                            # 解压文件
                            with open(target_path, 'wb') as target_file:
                                target_file.write(zip_ref.read(file_info.filename))
                    print(f"资源库更新完成: {peizhidir}")
                    db_ver = SQL_gettable(memory_db, 'version')['version'][0]
                    print('当前资源库版本：',db_ver)
                except Exception as e:
                    print(f"解压资源库文件失败,资源库更新失败: {e}")
        
        
    except:
        print('未发现内置资源库文件。')
    #os.path.join(peizhidir, file_name)
    try:
        biaozhun["药品清洗"]= SQL_gettable(memory_db, 'easy_药品规整-基础清洗').reset_index(drop=True) 
        biaozhun["药品清洗品种评价"]= SQL_gettable(memory_db, 'easy_药品规整-基础清洗（品种评价）').reset_index(drop=True) 

        biaozhun["药品清洗-扩行后"]= SQL_gettable(memory_db,'easy_药品规整-基础清洗-扩行后').reset_index(drop=True)         

        biaozhun["药品PT清洗"]=SQL_gettable(memory_db, 'easy_药品规整-PT标准化').reset_index(drop=True)
        biaozhun["药品不良反应名称规整-AI"]=SQL_gettable(memory_db, 'easy_药品规整-不良反应名称规整-AI').reset_index(drop=True)  
        biaozhun["药品关键词"]=SQL_gettable(memory_db, 'easy_药品规整-SOC-关键词').reset_index(drop=True) 
        biaozhun["药品重点关注库"]=SQL_gettable(memory_db, 'easy_药品规整-重点关注').reset_index(drop=True)  
        biaozhun["药品信息库"]=SQL_gettable(memory_db, 'easy_药品规整-药品信息库').reset_index(drop=True)  

        biaozhun["药品分类库"]=SQL_gettable(memory_db, 'easy_药品规整-药品分类-关键词').reset_index(drop=True)   

        biaozhun["器械清洗"]=SQL_gettable(memory_db, 'easy_器械规整-基础清洗').reset_index(drop=True) 
        
        biaozhun["器械关键词"]=SQL_gettable(memory_db, 'easy_器械规整-SOC-关键词').reset_index(drop=True) 
        
        biaozhun["器械关键词（仅故障表现）"]= biaozhun["器械关键词"].copy()
        biaozhun["器械关键词（仅故障表现）"]['查找位置']="器械故障表现"
        try:
            biaozhun["器械品种库"]=SQL_gettable(memory_db, 'easy_器械规整-品种信息库-测试').reset_index(drop=True)            
        except:
            pass
        
        biaozhun["meddra"]=0   

        print("已载入标准库。\n\n")
    except:
        if psur!=1:
            print("未载入资源库中的标准库,部分功能可能无法使用。可能资源库未正确配置，如需标准库，请联系开发者索取。")

           
    app = PROGRAM_display_df_in_treeview(default_data,0,default_data)  # 创建应用实例  
    #app=DRAW_plot_df(default_data)




       
if __name__ == "__main__":  
    AAA_00_main()
    print("done.")
