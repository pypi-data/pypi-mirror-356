import pathlib
import io 
from PIL import Image
import pdfplumber
import docx
import openpyxl
import pytesseract
import datetime
from oletools.olevba import VBA_Parser
from openpyxl.utils import get_column_letter
from typing import Union, BinaryIO, IO, List, Dict
from synergech_doc_extractor.base import BaseExtractor
from synergech_doc_extractor.exception_handling import ExtractionError,FileReadError,UnsupportedFormatError


import io
import pathlib
import base64
from typing import List, Dict
import pdfplumber
import pytesseract
from PIL import Image
import io
import pathlib
import base64
from typing import List, Dict
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image

class PDFExtractor(BaseExtractor):
    def extract_text(self) -> List[Dict[str, str]]:
        try:
            if isinstance(self.source, pathlib.Path):
                with open(self.source, 'rb') as f:
                    file_bytes = f.read()
            elif isinstance(self.source, bytes):
                file_bytes = self.source
            else:
                file_bytes = self.source.read()

            result = []

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text(
                        x_tolerance=1.5,
                        x_tolerance_ratio=None,
                        y_tolerance=1.5,
                        layout=True,
                        x_density=7.25,
                        y_density=13,
                        line_dir_render=None,
                        char_dir_render=None
                    ) or ""

                    images_dict = {}
                    image_counter = 1

                    # Check if page has detectable images via pdfplumber
                    if page.images:
                        for img in page.images:
                            try:
                                bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                                cropped = page.crop(bbox).to_image(resolution=300)
                                img_pil = cropped.original

                                # OCR
                                ocr_text = pytesseract.image_to_string(img_pil).strip()
                                if ocr_text:
                                    text += f"\n\n[Image OCR {image_counter}]: {ocr_text}"

                                # Base64
                                buffer = io.BytesIO()
                                img_pil.save(buffer, format="PNG")
                                encoded_image = base64.b64encode(buffer.getvalue()).decode("utf-8")
                                images_dict[f"image_{image_counter}"] = encoded_image

                                image_counter += 1
                            except Exception:
                                continue
                    else:
                        # Fallback: rasterize entire page using pdf2image
                        fallback_img = convert_from_bytes(file_bytes, dpi=300, first_page=i + 1, last_page=i + 1)[0]

                        # OCR
                        ocr_text = pytesseract.image_to_string(fallback_img).strip()
                        if ocr_text:
                            text += f"\n\n[Page OCR]: {ocr_text}"

                        # Base64
                        buffer = io.BytesIO()
                        fallback_img.save(buffer, format="PNG")
                        encoded_image = base64.b64encode(buffer.getvalue()).decode("utf-8")
                        images_dict[f"image_full_page_{i + 1}"] = encoded_image

                    result.append({
                        "page": i + 1,
                        "text": text.strip(),
                        "images": images_dict
                    })

            return result

        except Exception as e:
            raise FileReadError(f"Failed to extract PDF: {str(e)}")

class DocxExtractor(BaseExtractor):
    def extract_text(self) -> List[Dict[str, str]]:
        try:
            if isinstance(self.source, (str, pathlib.Path)):
                doc = docx.Document(str(self.source))
            elif isinstance(self.source, bytes):
                doc = docx.Document(io.BytesIO(self.source))
            else:
                doc = docx.Document(self.source)
            text = "\n".join([p.text for p in doc.paragraphs])
            return [{"page": "1", "text": text.strip()}]
        except Exception as e:
            raise FileReadError(f"Failed to extract DOCX: {str(e)}")

class ExcelExtractor(BaseExtractor):
    def extract_text(self) -> List[Dict[str, str]]:
        try:
            print("")
            if isinstance(self.source, pathlib.Path):
                wb = openpyxl.load_workbook(self.source, data_only=False, keep_vba=True)
            elif isinstance(self.source, bytes):
                wb = openpyxl.load_workbook(io.BytesIO(self.source), data_only=False, keep_vba=True)
            else:
                wb = openpyxl.load_workbook(self.source, data_only=False, keep_vba=True)

            # Named Ranges
            named_ranges = []
            for defn in wb.defined_names:
                name = defn.name
                if not name:
                    continue
                dests = []
                for title, coord in defn.destinations:
                    dests.append({"sheet": title, "range": coord})
                named_ranges.append({"name": name, "destinations": dests})

            # Macros
            macros_list = []
            if isinstance(self.source, (str, pathlib.Path)) and str(self.source).lower().endswith(('.xlsm', '.xlsb')):
                parser = VBA_Parser(str(self.source))
                try:
                    if parser.detect_vba_macros():
                        for (_, _, name, vba_code) in parser.extract_all_macros():
                            if name:
                                macros_list.append({"module": name, "code": vba_code})
                finally:
                    parser.close()

            result = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]

                # Text layout
                max_row, max_col = ws.max_row, ws.max_column
                col_widths = [0] * max_col
                all_values = []
                for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                    values = []
                    for i, cell in enumerate(row):
                        val = "" if cell.value is None else str(cell.value)
                        values.append(val)
                        col_widths[i] = max(col_widths[i], len(val))
                    all_values.append(values)
                col_widths = [w + 2 for w in col_widths]

                text_lines = [" " * 5 + "".join([get_column_letter(i+1).ljust(col_widths[i]) for i in range(max_col)])]
                for idx, row_vals in enumerate(all_values, 1):
                    row_text = str(idx).rjust(4) + " " + "".join([val.ljust(col_widths[i]) for i, val in enumerate(row_vals)])
                    text_lines.append(row_text)

                # Cell metadata
                sheet_data = []
                merged_ranges = [str(r) for r in ws.merged_cells.ranges]
                for row in ws.iter_rows(values_only=False):
                    for cell in row:
                        coord = cell.coordinate
                        value = cell.value
                        formula = value if isinstance(value, str) and value.startswith("=") else None
                        value = None if formula else value

                        font = cell.font
                        font_info = {
                            "name": font.name,
                            "size": font.sz,
                            "bold": font.b,
                            "italic": font.i,
                            "color": font.color.rgb if font.color and font.color.type == "rgb" else None
                        }

                        fill = cell.fill
                        fill_color = fill.fgColor.rgb if fill and fill.fgColor and fill.fgColor.type == "rgb" else None

                        merged_range = next((rng for rng in merged_ranges if coord in rng), None)

                        entry = {
                            "coordinate": coord,
                            "value": value.isoformat() if isinstance(value, datetime.datetime) else value,
                            "formula": formula,
                            "font": font_info,
                            "fill_color": fill_color,
                            "merged_range": merged_range
                        }
                        sheet_data.append(entry)

                # Charts
                chart_objs = []
                if hasattr(ws, "_charts"):
                    for chart in ws._charts:
                        chart_info = {"type": type(chart).__name__}
                        try:
                            chart_info["title"] = chart.title.tx.rich.p[0].r[0].t.value
                        except Exception:
                            chart_info["title"] = None
                        series_info = []
                        for series in chart.series:
                            info = {}
                            if series.val and series.val.numRef:
                                info["values"] = series.val.numRef.f
                            if series.cat and series.cat.strRef:
                                info["categories"] = series.cat.strRef.f
                            series_info.append(info)
                        chart_info["series"] = series_info
                        chart_objs.append(chart_info)

                # Formatted data (header -> {value, cell, formula})
                formatted = []
                if all_values:
                    headers = all_values[0]
                    for r_idx, row in enumerate(all_values[1:], start=2):  # Start from 2nd row (Excel row 2)
                        row_dict = {}
                        for c_idx, cell_value in enumerate(row):
                            if c_idx < len(headers) and headers[c_idx]:
                                header = headers[c_idx]
                                cell_coord = f"{get_column_letter(c_idx + 1)}{r_idx}"
                                cell_obj = ws.cell(row=r_idx, column=c_idx + 1)

                                entry = {
                                    "value": cell_obj.value,
                                    "cell": cell_coord
                                }
                                

                                if cell_obj.data_type == "f" or (isinstance(cell_obj.value, str) and cell_obj.value.startswith("=")):
                                    entry["formula"] = cell_obj.value

                                row_dict[header] = entry
                        if any(v.get("value") is not None for v in row_dict.values()):
                            formatted.append(row_dict)


                result.append({
                    "sheet": sheet,
                    "text": "\n".join(text_lines),
                    "cell_data": sheet_data,
                    "charts": chart_objs,
                    "macros": macros_list,
                    "named_ranges": named_ranges,
                    "formatted": formatted
                })

            return result
        except Exception as e:
            raise FileReadError(f"Failed to extract Excel: {str(e)}")
